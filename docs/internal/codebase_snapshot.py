# CODEBASE SNAPSHOT - HFMM Thesis (13 May 2026)
# This file is a read-only upload-context artifact; do not execute it.
# Project: High-Frequency Market Making via Reinforcement Learning under Different Volatility Regimes.
# Contents: active Python pipeline/support/audit files plus important JSON configs.
# Excludes: scripts/legacy, generated result artifacts, virtualenv/cache/build files, raw data, and obsolete manuscript outputs.
# Format: each block records FILE, PURPOSE, and STATUS before full file contents.

# ============================================================
# FILE: run.py
# PURPOSE: Main config-driven dispatcher, including resume support and WP5.5/WP6 jobs.
# STATUS: active
# ============================================================

"""
Ana Çalıştırıcı (Dispatcher)
-----------------------------
Config JSON dosyasını okur ve ilgili iş paketini (WP) çalıştırır.
Kullanım: python run.py --config config/<config_dosyasi>.json

İş paketleri (job değeri):
  w0_smoke          → WP0 smoke test
  w1_naive_sweep    → WP1 naive sweep
  w1_as_baseline    → WP1 AS baseline
  w1_compare        → WP1 karşılaştırma
  w2_synth          → WP2 sentetik veri üretimi
  w3_sanity         → WP3 ortam doğrulama
  w4_ppo            → WP4 PPO eğitimi
  w5_eval           → WP5 OOS değerlendirme
  w5_ablation_eta   → WP5 η ablasyon
  w5_ablation_skew  → WP5 skew penalty ablasyon
  w5_detector_compare → WP5 detector robustness
"""

import argparse
import json

from src.run_context import setup_run, finalize_run
from src.wp0.w0_smoke import wp0_smoke
from src.wp1.w1_naive_sweep import job_entry as w1_naive_sweep
from src.wp1.w1_as_baseline import job_entry as w1_as_baseline
from src.wp1.w1_compare import job_entry as w1_compare
from src.wp2.job_w2_synth import job_entry as w2_synth
from src.wp3.w3_sanity import job_entry as w3_sanity
from src.wp4.job_w4_ppo import job_entry as w4_ppo
from src.wp5.job_w5_eval import job_entry as w5_eval
from src.wp5.job_w5_ablation_eta import job_entry as w5_ablation_eta
from src.wp5.job_w5_ablation_skew import job_entry as w5_ablation_skew
from src.wp5.job_w5_detector_compare import job_entry as w5_detector_compare



def parse_args():
    p = argparse.ArgumentParser()
    p.add_argument("--config", type=str, default="config/base.json")
    p.add_argument(
        "--resume",
        type=str,
        default=None,
        help=(
            "Resume an existing run dir by ID (e.g. "
            "20260426-012241_seed42_wp6-sweep-pilot_6508b4f). The job will reuse "
            "that run dir, skip completed model files, and append metrics rows "
            "for newly trained cells only."
        ),
    )
    return p.parse_args()


def main():
    args = parse_args()
    ctx = setup_run(args.config, resume_run_id=args.resume)

    try:
        # setup_run bazı yapılarda cfg'yi ctx içine koyar; yoksa dosyadan okuruz
        if hasattr(ctx, "cfg") and isinstance(ctx.cfg, dict):
            cfg = ctx.cfg
        else:
            with open(args.config, "r", encoding="utf-8") as f:
                cfg = json.load(f)

        job = cfg.get("job", "w0_smoke")

        if job == "w0_smoke":
            wp0_smoke(ctx)  # WP0
        elif job == "w1_naive_sweep":
            w1_naive_sweep(cfg, ctx)  # WP1
        elif job == "w1_as_baseline":
            w1_as_baseline(cfg, ctx)
        elif job == "w1_compare":
            w1_compare(cfg, ctx)
        elif job == "w2_synth":
            w2_synth(cfg, ctx)
        elif job == "w3_sanity":
            w3_sanity(cfg, ctx)
        elif job == "w4_ppo":
            w4_ppo(cfg, ctx)
        elif job == "w5_eval":
            w5_eval(cfg, ctx)
        elif job == "w5_ablation_eta":
            w5_ablation_eta(cfg, ctx)
        elif job == "w5_ablation_skew":
            w5_ablation_skew(cfg, ctx)
        elif job == "w5_detector_compare":
            w5_detector_compare(cfg, ctx)
        elif job == "w55_audit":
            from src.wp5_5.job_w55_audit import run as run_w55_audit
            run_w55_audit(cfg, ctx)
        elif job == "w55_runtime":
            from src.wp5_5.job_w55_runtime import run as run_w55_runtime
            run_w55_runtime(cfg, ctx)
        elif job == "w55_calibration":
            from src.wp5_5.job_w55_calibration import run as run_w55_calibration
            run_w55_calibration(cfg, ctx)
        elif job == "w6_sweep_pilot":
            from src.wp6.job_w6_sweep_pilot import run as run_w6_sweep_pilot
            run_w6_sweep_pilot(cfg, ctx)
        elif job == "w6_sweep_full":
            from src.wp6.job_w6_sweep_full import run as run_w6_sweep_full
            run_w6_sweep_full(cfg, ctx)
        else:
            raise ValueError(f"Unknown job: {job}")

        finalize_run(ctx, "success")

    except Exception as e:
        ctx.logger.exception("Run crashed.")
        finalize_run(ctx, "failed", error=str(e))
        raise


if __name__ == "__main__":
    main()

# ============================================================
# FILE: src/__init__.py
# PURPOSE: Package marker for the thesis source tree.
# STATUS: support
# ============================================================

# makes src a package

# ============================================================
# FILE: src/run_context.py
# PURPOSE: Run lifecycle, reproducibility metadata, config snapshots, CSV metrics, and resume validation.
# STATUS: active
# ============================================================

"""Run lifecycle management: context, seeding, logging, and metrics."""
# Çalıştırma Yaşam Döngüsü (WP0)
# ----------------------------------
# Her deneyi benzersiz run_id ile izler, config snapshot ve git bilgisi kaydeder.
# setup_run() → RunContext oluşturur, finalize_run() → durumu yazar.
# CSVMetricLogger ile metrikler satır satır CSV'ye eklenir.

from __future__ import annotations

import json
import logging
import os
import platform
import subprocess
import sys
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, Optional

import random

try:
    import numpy as np
except Exception:
    np = None


def _utc_timestamp() -> str:
    return datetime.now(timezone.utc).strftime("%Y%m%d-%H%M%S")


def get_git_commit_short() -> str:
    """Return short git commit hash if available, otherwise 'nogit'."""
    try:
        out = subprocess.check_output(
            ["git", "rev-parse", "--short", "HEAD"],
            stderr=subprocess.DEVNULL
        )
        return out.decode("utf-8").strip()
    except Exception:
        return "nogit"


# Benzersiz çalıştırma kimliği üretir: YYYYMMDD-HHMMSS_seed<S>_<tag>_<commit>
# Tekrarlanabilirlik için seed ve git commit hash'i içerir.
def make_run_id(seed: int, run_tag: Optional[str] = None) -> str:
    """
    run_id format: YYYYMMDD-HHMMSS_seed<seed>_<tag>_<commit>
    """
    ts = _utc_timestamp()
    commit = get_git_commit_short()
    tag = (run_tag or "run").replace(" ", "-")
    return f"{ts}_seed{seed}_{tag}_{commit}"


def load_json(path: Path) -> Dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def save_json(path: Path, obj: Dict[str, Any]) -> None:
    path.write_text(json.dumps(obj, indent=2, ensure_ascii=False), encoding="utf-8")


def set_global_seed(seed: int) -> None:
    """
    Best-effort reproducibility across stdlib + numpy.
    """
    os.environ["PYTHONHASHSEED"] = str(seed)
    random.seed(seed)
    if np is not None:
        np.random.seed(seed)


def build_logger(run_dir: Path, level: str = "INFO") -> logging.Logger:
    logger = logging.getLogger("thesis")
    logger.setLevel(level.upper())
    logger.handlers.clear()
    logger.propagate = False

    fmt = logging.Formatter("%(asctime)s | %(levelname)s | %(message)s")

    ch = logging.StreamHandler(stream=sys.stdout)
    ch.setFormatter(fmt)
    logger.addHandler(ch)

    fh = logging.FileHandler(run_dir / "run.log", encoding="utf-8")
    fh.setFormatter(fmt)
    logger.addHandler(fh)

    return logger


class CSVMetricLogger:
    """Append-only CSV metric writer with schema enforcement.

    The schema (fieldnames) may be declared at construction time. If not,
    it is inferred from the first log() call and frozen thereafter. Every
    subsequent log() call must use the same set of keys, otherwise a
    ValueError is raised — silent column-mismatch corruption is never
    written to disk.
    """

    def __init__(self, path: Path, fieldnames: Optional[list] = None):
        self.path = path
        self._header_written = False
        self._fieldnames: Optional[list] = (
            list(fieldnames) if fieldnames is not None else None
        )

    def log(self, row: Dict[str, Any]) -> None:
        import csv

        if self._fieldnames is None:
            self._fieldnames = list(row.keys())
        else:
            row_keys = set(row.keys())
            schema_keys = set(self._fieldnames)
            if row_keys != schema_keys:
                missing = sorted(schema_keys - row_keys)
                extra = sorted(row_keys - schema_keys)
                raise ValueError(
                    f"CSVMetricLogger schema mismatch for {self.path.name}: "
                    f"expected {sorted(schema_keys)}, got {sorted(row_keys)}; "
                    f"missing={missing} extra={extra}"
                )

        self.path.parent.mkdir(parents=True, exist_ok=True)
        write_header = (not self._header_written) and (not self.path.exists())

        with self.path.open("a", newline="", encoding="utf-8") as f:
            writer = csv.DictWriter(f, fieldnames=self._fieldnames)
            if write_header:
                writer.writeheader()
                self._header_written = True
            writer.writerow(row)


def _diff_config(saved: Dict[str, Any], current: Dict[str, Any], prefix: str = "") -> list:
    """Return dotted-path diff lines for keys that differ between saved and current."""
    diffs = []
    for k in sorted(set(saved.keys()) | set(current.keys())):
        path = f"{prefix}.{k}" if prefix else k
        if k not in saved:
            diffs.append(f"{path}: <missing> != {current[k]!r}")
        elif k not in current:
            diffs.append(f"{path}: {saved[k]!r} != <missing>")
        elif isinstance(saved[k], dict) and isinstance(current[k], dict):
            diffs.extend(_diff_config(saved[k], current[k], prefix=path))
        elif saved[k] != current[k]:
            diffs.append(f"{path}: {saved[k]!r} != {current[k]!r}")
    return diffs


@dataclass
class RunContext:
    run_id: str
    run_dir: Path
    plots_dir: Path
    config: Dict[str, Any]
    logger: logging.Logger
    metrics: CSVMetricLogger
    resume_run_id: Optional[str] = None


# Yeni bir deney çalıştırmasını başlatır: dizin oluşturur, seed ayarlar,
# config snapshot ve meta bilgilerini kaydeder, logger ve metrics nesnelerini döner.
def setup_run(
    config_path: str | Path,
    results_root: str | Path = "results/runs",
    resume_run_id: Optional[str] = None,
    resume_force: bool = False,
) -> RunContext:
    config_path = Path(config_path)
    cfg = load_json(config_path)

    if "seed" not in cfg or not isinstance(cfg["seed"], int):
        raise ValueError("Config must include an integer field: 'seed'")

    seed = int(cfg["seed"])
    results_root = Path(results_root)
    if resume_run_id is not None:
        run_id = resume_run_id
        run_dir = results_root / run_id
        if not run_dir.exists():
            raise FileNotFoundError(
                f"--resume target does not exist: {run_dir.as_posix()}"
            )
        plots_dir = run_dir / "plots"
        plots_dir.mkdir(parents=True, exist_ok=True)

        # Validate the current config matches the snapshot saved by the
        # original non-resume run. A mismatched resume can silently produce
        # orphan-config runs (audit issue B9). Refuse on mismatch unless
        # the caller explicitly opts in via resume_force=True. A missing
        # snapshot (legacy runs predating snapshot saving) is permitted.
        snapshot_path = run_dir / "config_snapshot.json"
        if snapshot_path.exists():
            saved = load_json(snapshot_path)
            saved_cmp = {k: v for k, v in saved.items() if k != "run_id"}
            current_cmp = {k: v for k, v in cfg.items() if k != "run_id"}
            diffs = _diff_config(saved_cmp, current_cmp)
            if diffs:
                msg = (
                    f"Resume config mismatch for run_id={resume_run_id}.\n"
                    f"  Saved snapshot: {snapshot_path.as_posix()}\n"
                    f"  Current config: {Path(config_path).as_posix()}\n"
                    f"  Differing keys ({len(diffs)}):\n    - "
                    + "\n    - ".join(diffs)
                    + "\n  Pass resume_force=True to override."
                )
                if not resume_force:
                    raise ValueError(msg)
                print(f"WARNING: {msg}\nProceeding because resume_force=True.")
    else:
        run_id = make_run_id(seed=seed, run_tag=cfg.get("run_tag"))
        run_dir = results_root / run_id
        plots_dir = run_dir / "plots"
        run_dir.mkdir(parents=True, exist_ok=False)
        plots_dir.mkdir(parents=True, exist_ok=True)

    set_global_seed(seed)

    logger = build_logger(run_dir)

    cfg = {**cfg, "run_id": run_id}
    if resume_run_id is None:
        save_json(run_dir / "config_snapshot.json", cfg)
        meta = {
            "run_id": run_id,
            "started_utc": datetime.now(timezone.utc).isoformat(),
            "config_path": str(config_path.as_posix()),
            "git_commit": get_git_commit_short(),
            "python": sys.version,
            "platform": platform.platform(),
        }
        save_json(run_dir / "meta.json", meta)
        logger.info(f"Run started: {run_id}")
    else:
        logger.info(f"Run RESUMED: {run_id}")
    logger.info(f"Run dir: {run_dir.as_posix()}")

    metrics = CSVMetricLogger(run_dir / "metrics.csv")
    return RunContext(
        run_id=run_id,
        run_dir=run_dir,
        plots_dir=plots_dir,
        config=cfg,
        logger=logger,
        metrics=metrics,
        resume_run_id=resume_run_id,
    )

# Çalıştırmayı sonlandırır: status.json dosyasına başarı/hata durumunu yazar.
def finalize_run(ctx: RunContext, status: str = "success", error: str | None = None) -> None:
    payload = {
        "run_id": ctx.run_id,
        "status": status,
        "finished_utc": datetime.now(timezone.utc).isoformat(),
        "error": error,
    }
    save_json(ctx.run_dir / "status.json", payload)

    if status == "success":
        ctx.logger.info("Run finished successfully.")
    else:
        ctx.logger.error(f"Run finished with status={status}. error={error}")

# ============================================================
# FILE: src/wp0/__init__.py
# PURPOSE: Package marker for WP0 smoke-test code.
# STATUS: support
# ============================================================



# ============================================================
# FILE: src/wp0/w0_smoke.py
# PURPOSE: WP0 smoke test used to validate run-context plumbing.
# STATUS: active
# ============================================================

"""
WP0 Smoke Test
--------------
Projenin temel kurulumunu doğrular: env import, config okuma,
tek adım simülasyon. CI/CD benzeri basit sağlık kontrolü.
"""

from __future__ import annotations

try:
    import numpy as np
except Exception:
    np = None

import matplotlib.pyplot as plt

from src.run_context import RunContext, save_json


def wp0_smoke(ctx: RunContext) -> None:
    n_steps = int(ctx.config.get("n_steps", 200))
    seed = int(ctx.config["seed"])

    if np is None:
        raise RuntimeError("numpy is required for smoke test. Install numpy and rerun.")

    rng = np.random.default_rng(seed)
    rets = rng.normal(loc=0.0, scale=0.01, size=n_steps)

    equity_series = np.cumprod(1.0 + rets)

    for t in range(n_steps):
        ctx.metrics.log({"step": t, "ret": float(rets[t]), "equity": float(equity_series[t])})

    plt.figure()
    plt.plot(equity_series)
    plt.title("WP0 Smoke: Equity Curve")
    plt.xlabel("step")
    plt.ylabel("equity")
    out_path = ctx.plots_dir / "equity_curve.png"
    plt.savefig(out_path, dpi=150, bbox_inches="tight")
    plt.close()

    summary = {
        "n_steps": n_steps,
        "seed": seed,
        "final_equity": float(equity_series[-1]),
        "mean_ret": float(np.mean(rets)),
        "vol_ret": float(np.std(rets)),
    }
    save_json(ctx.run_dir / "summary.json", summary)

    ctx.logger.info(f"Saved plot: {out_path.as_posix()}")
    ctx.logger.info(f"Final equity: {summary['final_equity']:.6f}")

# ============================================================
# FILE: src/wp1/__init__.py
# PURPOSE: Package marker for WP1 simulator and baseline code.
# STATUS: support
# ============================================================



# ============================================================
# FILE: src/wp1/sim.py
# PURPOSE: Core market-making simulator with mid-price dynamics, Poisson fills, fees, latency, cash, and inventory state.
# STATUS: active
# ============================================================

"""
Piyasa Simülatörü (WP1)
-----------------------
Bu modül, piyasa yapıcı (market maker) simülasyonunun çekirdeğini oluşturur.
Aritmetik Brownian hareketi ile fiyat üretir, Poisson süreci ile emir dolumlarını (fills) simüle eder.
"""

from __future__ import annotations

from dataclasses import dataclass
from math import exp, sqrt
from typing import Dict, Tuple

import numpy as np


# Piyasa parametrelerini tutan değişmez (frozen) veri sınıfı.
# mid0: başlangıç fiyatı, tick_size: en küçük fiyat adımı,
# dt: zaman adımı (saniye), sigma_mid_ticks: volatilite (tick/√saniye)
@dataclass(frozen=True)
class MarketParams:
    mid0: float
    tick_size: float
    dt: float
    sigma_mid_ticks: float  # mid hareketi: "tick / sqrt(sec)"


# Emir çalıştırma parametrelerini tutan değişmez veri sınıfı.
# A: delta=0 anındaki dolum yoğunluğu, k: yoğunluğun uzaklıkla azalma hızı,
# fee_bps: komisyon (baz puan), latency_steps: gecikme adım sayısı
@dataclass
class ExecParams:
    A: float               # delta=0 iken intensity (fills/sec)
    k: float               # decay (per tick)
    fee_bps: float         # bps (ör: 0.2 = 0.2 bps)
    latency_steps: int     # quote için kaç adım stale mid


# Simülasyonun her adımdaki anlık durumunu tutan veri sınıfı.
# t: adım sayısı, mid: orta fiyat, cash: nakit, inv: envanter (net pozisyon)
@dataclass
class MMState:
    t: int
    mid: float
    cash: float
    inv: int

    @property
    def equity(self) -> float:
        return self.cash + self.inv * self.mid


# Poisson dolum yoğunluğunu hesaplar: λ(δ) = A * e^(-k*δ)
# δ (delta): kotasyonun mid-price'tan tick cinsinden uzaklığı
# Ne kadar uzak kotasyon verirsen, o kadar az dolum olasılığı.
def lambda_intensity(delta_ticks: float, A: float, k: float) -> float:
    """λ(δ)=A e^{-kδ}, δ tick cinsinden, clamp ile negatifleri engeller."""
    d = max(0.0, float(delta_ticks))
    return A * exp(-k * d)


# Bir zaman adımında en az bir dolumun gerçekleşme olasılığını hesaplar.
# Poisson sürecinde: P(N >= 1) = 1 - exp(-λ * dt)
def fill_prob(lmbda: float, dt: float) -> float:
    """Poisson(λ dt): P(N>=1)=1-exp(-λ dt)."""
    l = max(0.0, float(lmbda))
    return 1.0 - exp(-l * float(dt))


# Ana simülatör sınıfı. Her adımda fiyatı günceller ve dolumları simüle eder.
class MMSimulator:
    def __init__(self, market: MarketParams, execp: ExecParams, seed: int):
        self.m = market
        self.e = execp
        self.rng = np.random.default_rng(seed)
        self._mid_hist = []  # mid history for latency

    # Simülasyonu başlangıç durumuna sıfırlar. Her episode başında çağrılır.
    def reset(self) -> MMState:
        self._mid_hist = [self.m.mid0]
        return MMState(t=0, mid=self.m.mid0, cash=0.0, inv=0)

    # Aritmetik Brownian hareketi ile bir sonraki mid-price değerini üretir.
    # d_ticks = sigma * sqrt(dt) * z, z ~ N(0,1)
    def _evolve_mid(self, mid: float) -> float:
        # Arithmetic BM in ticks: dMid_ticks = sigma * sqrt(dt) * z
        z = self.rng.standard_normal()
        d_ticks = self.m.sigma_mid_ticks * sqrt(self.m.dt) * z
        return mid + d_ticks * self.m.tick_size

    # Simülasyonun bir adımını çalıştırır:
    # 1) Fiyatı güncelle (ABM)
    # 2) Gecikmeli fiyatla kotasyon fiyatlarını hesapla (latency adverse selection)
    # 3) Bid ve ask için Poisson dolum olasılıklarını hesapla
    # 4) Rastgele dolumları gerçekleştir, nakit ve envanteri güncelle
    def step(self, s: MMState, delta_bid_ticks: int, delta_ask_ticks: int) -> Tuple[MMState, Dict]:
        # 1) mid moves
        mid_new = self._evolve_mid(s.mid)
        self._mid_hist.append(mid_new)

        # 2) choose stale mid for quoting
        L = int(self.e.latency_steps)
        if L <= 0 or len(self._mid_hist) <= L:
            mid_for_quote = s.mid
        else:
            mid_for_quote = self._mid_hist[-(L + 1)]  # L steps stale

        ts = self.m.tick_size

        # 3) quote prices (based on stale mid)
        bid_px = mid_for_quote - float(delta_bid_ticks) * ts
        ask_px = mid_for_quote + float(delta_ask_ticks) * ts

        # 4) effective deltas vs CURRENT mid (latency adverse selection shows up here)
        # bid: delta_eff = (current_mid - bid_px)/tick
        # ask: delta_eff = (ask_px - current_mid)/tick
        delta_eff_bid = (mid_new - bid_px) / ts
        delta_eff_ask = (ask_px - mid_new) / ts

        lmbda_bid = lambda_intensity(delta_eff_bid, self.e.A, self.e.k)
        lmbda_ask = lambda_intensity(delta_eff_ask, self.e.A, self.e.k)

        p_bid = fill_prob(lmbda_bid, self.m.dt)
        p_ask = fill_prob(lmbda_ask, self.m.dt)

        fill_bid = (self.rng.random() < p_bid)
        fill_ask = (self.rng.random() < p_ask)

        cash = s.cash
        inv = s.inv
        fee_total = 0.0
        fills = 0

        def fee(notional: float) -> float:
            return abs(notional) * (self.e.fee_bps * 1e-4)

        if fill_bid:
            f = fee(bid_px)
            cash -= (bid_px + f)
            inv += 1
            fee_total += f
            fills += 1

        if fill_ask:
            f = fee(ask_px)
            cash += (ask_px - f)
            inv -= 1
            fee_total += f
            fills += 1

        s2 = MMState(t=s.t + 1, mid=mid_new, cash=cash, inv=inv)

        info = {
            "bid_px": bid_px,
            "ask_px": ask_px,
            "delta_eff_bid": float(delta_eff_bid),
            "delta_eff_ask": float(delta_eff_ask),
            "lambda_bid": lmbda_bid,
            "lambda_ask": lmbda_ask,
            "p_bid": p_bid,
            "p_ask": p_ask,
            "fill_bid": bool(fill_bid),
            "fill_ask": bool(fill_ask),
            "fills": fills,
            "fee_total": fee_total,
            "equity": s2.equity,
            "inv": inv,
            "mid": mid_new
        }
        return s2, info

# ============================================================
# FILE: src/wp1/w1_naive_sweep.py
# PURPOSE: WP1 fixed-spread naive baseline sweep.
# STATUS: active
# ============================================================

from __future__ import annotations

from dataclasses import asdict
from pathlib import Path
import json

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt

from src.wp1.sim import MarketParams, ExecParams, MMSimulator


def compute_metrics(equity: np.ndarray, inv: np.ndarray, fills: np.ndarray, dt: float) -> dict:
    rets = np.diff(equity)
    # “Sharpe” burada kaba: mean/std of per-step PnL increments
    mu = rets.mean() if rets.size else 0.0
    sd = rets.std(ddof=1) if rets.size > 1 else 0.0
    sharpe = (mu / sd) * np.sqrt(1.0 / dt) if sd > 0 else 0.0  # scale ~ sqrt(steps/sec)

    # max drawdown
    peak = np.maximum.accumulate(equity)
    dd = equity - peak
    max_dd = dd.min() if dd.size else 0.0

    return {
        "final_equity": float(equity[-1]),
        "mean_step_pnl": float(mu),
        "std_step_pnl": float(sd),
        "sharpe_like": float(sharpe),
        "max_drawdown": float(max_dd),
        "fill_rate": float(fills.sum() / len(fills)),
        "turnover": int(fills.sum()),
        "inv_mean": float(inv.mean()),
        "inv_p95": float(np.quantile(np.abs(inv), 0.95)),
        "inv_p99": float(np.quantile(np.abs(inv), 0.99)),
    }


def save_plot(path: Path, x: np.ndarray, y: np.ndarray, title: str, xlabel: str, ylabel: str) -> None:
    plt.figure()
    plt.plot(x, y)
    plt.title(title)
    plt.xlabel(xlabel)
    plt.ylabel(ylabel)
    plt.tight_layout()
    plt.savefig(path, dpi=150)
    plt.close()


def run(cfg: dict, ctx=None) -> None:
    # output dir (WP0 context varsa onun içine yaz)
    if ctx is not None and hasattr(ctx, "run_dir"):
        out_dir = Path(ctx.run_dir)
    elif ctx is not None and hasattr(ctx, "paths") and "run_dir" in ctx.paths:
        out_dir = Path(ctx.paths["run_dir"])
    else:
        out_dir = Path("results/runs/manual_w1")
        out_dir.mkdir(parents=True, exist_ok=True)

    (out_dir / "plots").mkdir(exist_ok=True)

    market = MarketParams(**cfg["market"])
    execp = ExecParams(**cfg["exec"])

    # snapshot
    (out_dir / "config_snapshot.json").write_text(json.dumps(cfg, indent=2), encoding="utf-8")

    rows = []
    for h in cfg["sweep"]["half_spreads_ticks"]:
        sim = MMSimulator(market, execp, seed=int(cfg["seed"]))
        s = sim.reset()

        n = int(cfg["episode"]["n_steps"])
        equity = np.zeros(n + 1)
        inv = np.zeros(n + 1, dtype=int)
        fills = np.zeros(n, dtype=int)

        equity[0] = s.equity
        inv[0] = s.inv

        for t in range(n):
            # naive: m=0 => delta_bid=delta_ask=h
            s, info = sim.step(s, delta_bid_ticks=int(h), delta_ask_ticks=int(h))
            equity[t + 1] = s.equity
            inv[t + 1] = s.inv
            fills[t] = int(info["fills"])

        m = compute_metrics(equity, inv, fills, dt=market.dt)
        row = {"half_spread_ticks": int(h), **m}
        rows.append(row)

        # save equity curve + plots per h
        df_curve = pd.DataFrame({
            "t": np.arange(n + 1),
            "equity": equity,
            "inv": inv,
        })
        df_curve.to_csv(out_dir / f"equity_curve_h{h}.csv", index=False)

        save_plot(out_dir / "plots" / f"equity_h{h}.png",
                  x=np.arange(n + 1), y=equity,
                  title=f"Equity Curve (h={h} ticks)", xlabel="step", ylabel="equity")

        save_plot(out_dir / "plots" / f"inventory_h{h}.png",
                  x=np.arange(n + 1), y=inv,
                  title=f"Inventory (h={h} ticks)", xlabel="step", ylabel="inventory")

    df = pd.DataFrame(rows).sort_values("half_spread_ticks")
    df.to_csv(out_dir / "metrics_sweep.csv", index=False)

    # quick summary print (logger yoksa bile gör)
    print(df[["half_spread_ticks", "final_equity", "fill_rate", "inv_p99", "max_drawdown", "sharpe_like"]])


# Eğer WP0 run.py bu modülü job olarak çağıracaksa:
def job_entry(cfg: dict, ctx) -> None:
    run(cfg, ctx)

# ============================================================
# FILE: src/wp1/w1_as_baseline.py
# PURPOSE: WP1 Avellaneda-Stoikov baseline and shared metric computation utilities.
# STATUS: active
# ============================================================

"""Avellaneda-Stoikov baseline strategy for WP1."""
# Avellaneda-Stoikov Baz Stratejisi (WP1)
# -----------------------------------------
# Analitik piyasa yapıcılığı formülüne dayanan baz strateji.
# Rezervasyon fiyatı (r) ve yarı-spread (d) hesaplayarak kotasyon üretir.
# compute_metrics() fonksiyonu tüm stratejiler tarafından ortak kullanılır.

from __future__ import annotations

import json
from pathlib import Path

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt

from src.wp1.sim import MarketParams, ExecParams, MMSimulator


def save_plot(path: Path, x: np.ndarray, y: np.ndarray, title: str, xlabel: str, ylabel: str) -> None:
    plt.figure()
    plt.plot(x, y)
    plt.title(title)
    plt.xlabel(xlabel)
    plt.ylabel(ylabel)
    plt.tight_layout()
    plt.savefig(path, dpi=150)
    plt.close()


# Avellaneda-Stoikov kotasyon hesabı: envantere göre asimetrik bid/ask delta üretir.
# r = mid - inv * gamma * sigma² * tau (rezervasyon fiyatı)
# d = 0.5 * gamma * sigma² * tau + (1/gamma) * ln(1 + gamma/k) (yarı-spread)
def as_deltas_ticks(mid: float, inv: int, t: int, cfg: dict, market: MarketParams, execp: ExecParams) -> tuple[int, int]:
    """
    Avellaneda–Stoikov (basit) quote:
    reservation price: r = mid - inv * gamma * sigma^2 * tau
    half-spread:      d = 0.5*gamma*sigma^2*tau + (1/gamma)*log(1 + gamma/k_price)

    Burada k_price = k_ticks / tick_size (çünkü intensity exp(-k_ticks * delta_ticks)).
    """
    ts = market.tick_size
    dt = market.dt

    gamma = float(cfg["as"]["gamma"])
    H = int(cfg["as"]["horizon_steps"])
    tau = max(0.0, (H - t) * dt)  # seconds

    sigma_price = market.sigma_mid_ticks * ts
    sigma2 = sigma_price ** 2  # price^2 / sec

    k_price = execp.k / ts  # 1/price

    r = mid - inv * gamma * sigma2 * tau
    d = 0.5 * gamma * sigma2 * tau + (1.0 / gamma) * np.log(1.0 + gamma / k_price)

    bid = r - d
    ask = r + d

    # delta ticks relative to mid (latency_steps=1 iken simulator içinde mid_for_quote == s.mid)
    delta_bid = int(np.ceil((mid - bid) / ts))
    delta_ask = int(np.ceil((ask - mid) / ts))

    min_d = int(cfg["as"].get("min_delta_ticks", 1))
    max_d = int(cfg["as"].get("max_delta_ticks", 25))
    delta_bid = int(np.clip(delta_bid, min_d, max_d))
    delta_ask = int(np.clip(delta_ask, min_d, max_d))

    return delta_bid, delta_ask


# Performans metriklerini hesaplar: Sharpe oranı, final equity, dolum oranı,
# envanter istatistikleri (ortalama, p95, p99) ve maksimum drawdown.
# Tüm stratejiler (naive, AS, PPO) bu ortak fonksiyonu kullanır.
def compute_metrics(equity: np.ndarray, inv: np.ndarray, fills: np.ndarray, dt: float) -> dict:
    rets = np.diff(equity)
    mu = rets.mean() if rets.size else 0.0
    sd = rets.std(ddof=1) if rets.size > 1 else 0.0
    sharpe = (mu / sd) * np.sqrt(1.0 / dt) if sd > 0 else 0.0

    peak = np.maximum.accumulate(equity)
    dd = equity - peak
    max_dd = dd.min() if dd.size else 0.0

    return {
        "final_equity": float(equity[-1]),
        "fill_rate": float(fills.sum() / len(fills)),
        "turnover": int(fills.sum()),
        "inv_mean": float(inv.mean()),
        "inv_p95": float(np.quantile(np.abs(inv), 0.95)),
        "inv_p99": float(np.quantile(np.abs(inv), 0.99)),
        "max_drawdown": float(max_dd),
        "sharpe_like": float(sharpe),
    }


# AS stratejisini tek episode boyunca çalıştırır.
# Her adımda as_deltas_ticks() ile kotasyon hesaplar, simülatörü ilerletir.
# Equity eğrisi, delta geçmişi ve metrikleri dosyaya yazar.
def run(cfg: dict, ctx=None) -> None:
    if ctx is not None and hasattr(ctx, "run_dir"):
        out_dir = Path(ctx.run_dir)
    elif ctx is not None and hasattr(ctx, "paths") and "run_dir" in ctx.paths:
        out_dir = Path(ctx.paths["run_dir"])
    else:
        out_dir = Path("results/runs/manual_w1_as")
        out_dir.mkdir(parents=True, exist_ok=True)

    (out_dir / "plots").mkdir(exist_ok=True)

    market = MarketParams(**cfg["market"])
    execp = ExecParams(**cfg["exec"])

    (out_dir / "config_snapshot.json").write_text(json.dumps(cfg, indent=2), encoding="utf-8")

    sim = MMSimulator(market, execp, seed=int(cfg["seed"]))
    s = sim.reset()

    n = int(cfg["episode"]["n_steps"])
    equity = np.zeros(n + 1)
    inv = np.zeros(n + 1, dtype=int)
    fills = np.zeros(n, dtype=int)
    deltab = np.zeros(n, dtype=int)
    deltaa = np.zeros(n, dtype=int)

    equity[0] = s.equity
    inv[0] = s.inv

    for t in range(n):
        db, da = as_deltas_ticks(s.mid, s.inv, t, cfg, market, execp)
        deltab[t] = db
        deltaa[t] = da

        s, info = sim.step(s, delta_bid_ticks=db, delta_ask_ticks=da)
        equity[t + 1] = s.equity
        inv[t + 1] = s.inv
        fills[t] = int(info["fills"])

    df_curve = pd.DataFrame({
        "t": np.arange(n + 1),
        "equity": equity,
        "inv": inv
    })
    df_curve.to_csv(out_dir / "equity_curve_as.csv", index=False)

    df_deltas = pd.DataFrame({
        "t": np.arange(n),
        "delta_bid_ticks": deltab,
        "delta_ask_ticks": deltaa
    })
    df_deltas.to_csv(out_dir / "deltas_as.csv", index=False)

    m = compute_metrics(equity, inv, fills, dt=market.dt)
    pd.DataFrame([m]).to_csv(out_dir / "metrics_as.csv", index=False)

    save_plot(out_dir / "plots" / "equity_as.png", np.arange(n + 1), equity, "Equity Curve (AS)", "step", "equity")
    save_plot(out_dir / "plots" / "inventory_as.png", np.arange(n + 1), inv, "Inventory (AS)", "step", "inventory")

    print(pd.DataFrame([m]))


def job_entry(cfg: dict, ctx) -> None:
    run(cfg, ctx)

# ============================================================
# FILE: src/wp1/w1_compare.py
# PURPOSE: WP1 same-seed comparison between naive and Avellaneda-Stoikov baselines.
# STATUS: active
# ============================================================

from __future__ import annotations

import json
from pathlib import Path

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt

from src.wp1.sim import MarketParams, ExecParams, MMSimulator


def save_plot(path: Path, x: np.ndarray, y: np.ndarray, title: str, xlabel: str, ylabel: str) -> None:
    plt.figure()
    plt.plot(x, y)
    plt.title(title)
    plt.xlabel(xlabel)
    plt.ylabel(ylabel)
    plt.tight_layout()
    plt.savefig(path, dpi=150)
    plt.close()


def compute_metrics(equity: np.ndarray, inv: np.ndarray, fills: np.ndarray, dt: float) -> dict:
    rets = np.diff(equity)
    mu = rets.mean() if rets.size else 0.0
    sd = rets.std(ddof=1) if rets.size > 1 else 0.0
    sharpe = (mu / sd) * np.sqrt(1.0 / dt) if sd > 0 else 0.0

    peak = np.maximum.accumulate(equity)
    dd = equity - peak
    max_dd = dd.min() if dd.size else 0.0

    return {
        "final_equity": float(equity[-1]),
        "fill_rate": float(fills.sum() / len(fills)),
        "turnover": int(fills.sum()),
        "inv_mean": float(inv.mean()),
        "inv_p95": float(np.quantile(np.abs(inv), 0.95)),
        "inv_p99": float(np.quantile(np.abs(inv), 0.99)),
        "max_drawdown": float(max_dd),
        "sharpe_like": float(sharpe),
    }


def as_deltas_ticks(mid: float, inv: int, t: int, cfg: dict, market: MarketParams, execp: ExecParams) -> tuple[int, int]:
    ts = market.tick_size
    dt = market.dt

    gamma = float(cfg["as"]["gamma"])
    H = int(cfg["as"]["horizon_steps"])
    tau = max(0.0, (H - t) * dt)  # seconds remaining

    sigma_price = market.sigma_mid_ticks * ts
    sigma2 = sigma_price ** 2

    k_price = execp.k / ts  # 1/price

    r = mid - inv * gamma * sigma2 * tau
    d = 0.5 * gamma * sigma2 * tau + (1.0 / gamma) * np.log(1.0 + gamma / k_price)

    bid = r - d
    ask = r + d

    delta_bid = int(np.ceil((mid - bid) / ts))
    delta_ask = int(np.ceil((ask - mid) / ts))

    min_d = int(cfg["as"].get("min_delta_ticks", 1))
    max_d = int(cfg["as"].get("max_delta_ticks", 25))
    delta_bid = int(np.clip(delta_bid, min_d, max_d))
    delta_ask = int(np.clip(delta_ask, min_d, max_d))

    return delta_bid, delta_ask


def run(cfg: dict, ctx=None) -> None:
    if ctx is not None and hasattr(ctx, "run_dir"):
        out_dir = Path(ctx.run_dir)
    elif ctx is not None and hasattr(ctx, "paths") and "run_dir" in ctx.paths:
        out_dir = Path(ctx.paths["run_dir"])
    else:
        out_dir = Path("results/runs/manual_w1_compare")
        out_dir.mkdir(parents=True, exist_ok=True)

    (out_dir / "plots").mkdir(exist_ok=True)
    (out_dir / "curves").mkdir(exist_ok=True)

    (out_dir / "config_snapshot.json").write_text(json.dumps(cfg, indent=2), encoding="utf-8")

    market = MarketParams(**cfg["market"])
    execp = ExecParams(**cfg["exec"])
    n = int(cfg["episode"]["n_steps"])
    seed = int(cfg["seed"])

    rows = []

    # ---- Naive sweep ----
    for h in cfg["sweep"]["half_spreads_ticks"]:
        sim = MMSimulator(market, execp, seed=seed)
        s = sim.reset()

        equity = np.zeros(n + 1)
        inv = np.zeros(n + 1, dtype=int)
        fills = np.zeros(n, dtype=int)

        equity[0] = s.equity
        inv[0] = s.inv

        for t in range(n):
            s, info = sim.step(s, delta_bid_ticks=int(h), delta_ask_ticks=int(h))
            equity[t + 1] = s.equity
            inv[t + 1] = s.inv
            fills[t] = int(info["fills"])

        m = compute_metrics(equity, inv, fills, dt=market.dt)
        rows.append({"strategy": f"naive_h{h}", **m})

        pd.DataFrame({"t": np.arange(n + 1), "equity": equity, "inv": inv}).to_csv(
            out_dir / "curves" / f"curve_naive_h{h}.csv", index=False
        )
        save_plot(out_dir / "plots" / f"equity_naive_h{h}.png", np.arange(n + 1), equity,
                  f"Equity (naive h={h})", "step", "equity")
        save_plot(out_dir / "plots" / f"inv_naive_h{h}.png", np.arange(n + 1), inv,
                  f"Inventory (naive h={h})", "step", "inventory")

    df_naive = pd.DataFrame(rows).copy()
    df_naive.to_csv(out_dir / "metrics_naive.csv", index=False)

    # ---- AS baseline ----
    sim = MMSimulator(market, execp, seed=seed)
    s = sim.reset()

    equity = np.zeros(n + 1)
    inv = np.zeros(n + 1, dtype=int)
    fills = np.zeros(n, dtype=int)

    equity[0] = s.equity
    inv[0] = s.inv

    for t in range(n):
        db, da = as_deltas_ticks(s.mid, s.inv, t, cfg, market, execp)
        s, info = sim.step(s, delta_bid_ticks=db, delta_ask_ticks=da)
        equity[t + 1] = s.equity
        inv[t + 1] = s.inv
        fills[t] = int(info["fills"])

    m_as = compute_metrics(equity, inv, fills, dt=market.dt)
    row_as = {"strategy": "AS", **m_as}

    pd.DataFrame({"t": np.arange(n + 1), "equity": equity, "inv": inv}).to_csv(
        out_dir / "curves" / "curve_as.csv", index=False
    )
    save_plot(out_dir / "plots" / "equity_as.png", np.arange(n + 1), equity, "Equity (AS)", "step", "equity")
    save_plot(out_dir / "plots" / "inv_as.png", np.arange(n + 1), inv, "Inventory (AS)", "step", "inventory")

    # ---- Compare table ----
    df_compare = pd.concat([df_naive, pd.DataFrame([row_as])], ignore_index=True)
    df_compare = df_compare.sort_values("final_equity", ascending=False)
    df_compare.to_csv(out_dir / "metrics_compare.csv", index=False)

    # quick bar plot: final equity
    plt.figure()
    plt.bar(df_compare["strategy"], df_compare["final_equity"].to_numpy())
    plt.xticks(rotation=45, ha="right")
    plt.title("Final Equity: naive sweep vs AS")
    plt.tight_layout()
    plt.savefig(out_dir / "plots" / "compare_final_equity.png", dpi=150)
    plt.close()

    print(df_compare[["strategy", "final_equity", "fill_rate", "inv_p99", "max_drawdown", "sharpe_like"]])


def job_entry(cfg: dict, ctx) -> None:
    run(cfg, ctx)

# ============================================================
# FILE: src/wp2/__init__.py
# PURPOSE: Package marker for WP2 regime-generation and detector code.
# STATUS: support
# ============================================================



# ============================================================
# FILE: src/wp2/synth_regime.py
# PURPOSE: Synthetic volatility-regime generation, causal/offline detector variants, and WP2 artifact writing.
# STATUS: active
# ============================================================

"""Synthetic regime generation and detection for WP2."""
# Sentetik Rejim Üretimi ve Tespiti (WP2)
# ----------------------------------------
# 3 durumlu (L/M/H) Markov zinciri ile volatilite rejimi üretir.
# Rolling realized volatility (RV) tabanlı dedektörler ve HMM ile rejim tespiti yapar.
# Look-ahead yok: rejim etiketi yalnızca geçmiş veriye dayanır.

from __future__ import annotations

from pathlib import Path

import numpy as np
import pandas as pd

REGIME_LABELS = {0: "L", 1: "M", 2: "H"}

# Varsayılan geçiş matrisi: yüksek köşegen değerleri "yapışkan" rejimler üretir.
# Satır i, sütun j: P(state_t+1 = j | state_t = i)
DEFAULT_TRANS_MATRIX = np.array([
    [0.9967, 0.0023, 0.0010],
    [0.0042, 0.9917, 0.0041],
    [0.0010, 0.0030, 0.9960],
])


# Markov zinciri ile rejim dizisi üretir. Başlangıç durumu M (state=1).
# Her adımda geçiş matrisinden olasılıksal örnekleme yapılır.
def generate_regime_series(
    n_steps: int,
    seed: int,
    cfg: dict | None = None,
    trans_matrix: np.ndarray | None = None,
    rng: np.random.Generator | None = None,
) -> np.ndarray:
    if rng is None:
        rng = np.random.default_rng(seed)
    if trans_matrix is None:
        if cfg is not None and "trans_matrix" in cfg.get("regime", {}):
            trans_matrix = np.array(cfg["regime"]["trans_matrix"])
        else:
            trans_matrix = DEFAULT_TRANS_MATRIX

    regime = np.empty(n_steps, dtype=int)
    state = 1  # start at M
    for t in range(n_steps):
        regime[t] = state
        state = rng.choice(3, p=trans_matrix[state])
    return regime


# Rejim dizisine bağlı olarak sentetik mid-price serisi üretir.
# Her rejimde farklı volatilite çarpanı (sigma_mult) kullanılır.
# Aritmetik Brownian hareket: dMid = sigma_regime * sqrt(dt) * z * tick_size
def generate_mid_series(
    regime_true: np.ndarray,
    cfg: dict,
    rng: np.random.Generator,
) -> tuple[np.ndarray, np.ndarray]:
    mid0 = cfg["market"]["mid0"]
    tick_size = cfg["market"]["tick_size"]
    dt = cfg["market"]["dt"]
    sigma_base = cfg["regime"]["sigma_mid_ticks_base"]
    sigma_mult = cfg["regime"]["sigma_mult"]

    sigma_per_regime = [sigma_base * m for m in sigma_mult]

    n = len(regime_true)
    mid = np.empty(n + 1)
    mid[0] = mid0

    sqrt_dt = np.sqrt(dt)
    z = rng.standard_normal(n)

    for t in range(n):
        state = regime_true[t]
        d_ticks = sigma_per_regime[state] * sqrt_dt * z[t]
        mid[t + 1] = max(tick_size, mid[t] + d_ticks * tick_size)

    ret = np.diff(mid)  # mid[t] - mid[t-1], length n
    return mid, ret


# Kayan pencere ile gerçekleşmiş volatilite (RV) hesaplar.
# İlk `window` adımda NaN döner (yetersiz veri).
# sigma_hat = RV / tick_size (tick cinsinden normalleştirilmiş volatilite)
def compute_rolling_rv(
    mid: np.ndarray,
    window: int,
    tick_size: float,
) -> tuple[np.ndarray, np.ndarray]:
    ret = np.diff(mid, prepend=mid[0])
    n = len(ret)
    rv = np.full(n, np.nan)

    for t in range(window, n):
        rv[t] = np.std(ret[t - window + 1 : t + 1], ddof=1)

    sigma_hat = rv / tick_size
    return rv, sigma_hat


# Warmup dönemindeki sigma_hat dağılımından eşik değerlerini belirler.
# 33. ve 66. persantil: L/M sınırı (thresh_LM) ve M/H sınırı (thresh_MH)
def calibrate_thresholds(
    sigma_hat: np.ndarray,
    warmup_end: int,
) -> tuple[float, float]:
    vals = sigma_hat[:warmup_end]
    vals = vals[~np.isnan(vals)]
    thresh_LM = float(np.percentile(vals, 33))
    thresh_MH = float(np.percentile(vals, 66))
    return thresh_LM, thresh_MH


# Eşik tabanlı rejim tespiti (rv_baseline dedektör).
# Warmup dönemi "warmup" etiketi alır; sonrası sigma_hat'a göre L/M/H.
def assign_regime_hat(
    sigma_hat: np.ndarray,
    thresh_LM: float,
    thresh_MH: float,
    warmup_end: int,
) -> list[str]:
    n = len(sigma_hat)
    regime_hat: list[str] = []
    for t in range(n):
        if t < warmup_end:
            regime_hat.append("warmup")
        elif sigma_hat[t] < thresh_LM:
            regime_hat.append("L")
        elif sigma_hat[t] < thresh_MH:
            regime_hat.append("M")
        else:
            regime_hat.append("H")
    return regime_hat


# Bekleme süresi filtresi: min_dwell adımdan kısa rejim geçişlerini bastırır.
# Kısa geçişleri bir önceki rejim etiketiyle değiştirir (gürültü azaltma).
def apply_dwell_filter(
    regime_labels: list[str],
    min_dwell: int = 5,
) -> list[str]:
    """Offline dwell smoothing for regime labels.

    NOT CAUSAL. To decide the label at index i, this function scans
    forward up to min_dwell positions to compute the run length starting
    at i, then overwrites short runs with the previous label. This
    introduces a bounded forward look (up to min_dwell steps) and must
    not be used in any pipeline that claims online or causal detection.

    Intended use: offline robustness comparison only (see compare_detectors.py
    and stats_detector_robustness.py). All main WP4/WP5/WP6 pipelines use
    the causal baseline detector assign_regime_hat.

    Args:
        regime_labels: list of labels (including "warmup")
        min_dwell: minimum run length; shorter runs are smoothed

    Returns:
        smoothed labels, same length
    """
    out = list(regime_labels)
    n = len(out)

    # Identify contiguous runs of non-warmup labels
    i = 0
    while i < n:
        if out[i] == "warmup":
            i += 1
            continue
        # start of a run
        j = i + 1
        while j < n and out[j] == out[i]:
            j += 1
        run_len = j - i
        if run_len < min_dwell:
            # find the previous non-warmup label (fall back to current if none)
            prev = out[i]
            for k in range(i - 1, -1, -1):
                if out[k] != "warmup":
                    prev = out[k]
                    break
            for k in range(i, j):
                out[k] = prev
        i = j
    return out


def assign_regime_hat_dwell(
    sigma_hat: np.ndarray,
    thresh_LM: float,
    thresh_MH: float,
    warmup_end: int,
    min_dwell: int = 5,
) -> list[str]:
    """Threshold-based regime detection followed by a dwell filter."""
    raw = assign_regime_hat(sigma_hat, thresh_LM, thresh_MH, warmup_end)
    return apply_dwell_filter(raw, min_dwell)


# GaussianHMM tabanlı rejim tespiti (hmm dedektör).
# Warmup verisinde eğitilir, sonrası nedensel (causal) tahmin yapar.
# HMM durumları varyansa göre L/M/H'ye eşlenir (düşük varyans → L).
def assign_regime_hat_hmm(
    sigma_hat: np.ndarray,
    warmup_end: int,
    n_states: int = 3,
    seed: int = 0,
) -> list[str]:
    """HMM-based regime detection (GaussianHMM on rolling sigma_hat).

    * Fits on non-NaN sigma_hat values up to *warmup_end* (no look-ahead).
    * Predicts causally: at each step t >= warmup_end the model uses only
      sigma_hat[warmup_end:t+1], dropping any leading NaNs.
    * Maps HMM states to L/M/H by emission variance (lowest → L, highest → H).
    """
    from hmmlearn.hmm import GaussianHMM

    n = len(sigma_hat)
    labels: list[str] = ["warmup"] * n

    # --- Fit on warmup data, dropping NaNs ---
    train = sigma_hat[:warmup_end]
    train = train[~np.isnan(train)].reshape(-1, 1)

    if len(train) < 2:
        return labels

    model = GaussianHMM(
        n_components=n_states,
        covariance_type="diag",
        n_iter=100,
        random_state=seed,
    )
    model.fit(train)

    # --- Map HMM states to L/M/H by ascending variance ---
    variances = model.covars_.flatten()
    order = np.argsort(variances)  # lowest var first
    state_map = {}
    regime_names = ["L", "M", "H"]
    for rank, state_idx in enumerate(order):
        state_map[state_idx] = regime_names[rank]

    # --- Causal prediction: expand window one step at a time ---
    for t in range(warmup_end, n):
        window = sigma_hat[warmup_end : t + 1]
        valid = window[~np.isnan(window)]
        if len(valid) == 0:
            continue
        obs = valid.reshape(-1, 1)
        hidden = model.predict(obs)
        labels[t] = state_map[hidden[-1]]

    return labels


# WP2 ana iş akışı: rejim üretimi → mid-price → RV hesaplama → eşik kalibrasyonu → tespit.
# Active downstream pipelines consume the returned dataframe, not the disk CSVs.
# Disk outputs are artifacts:
#   - data/processed/wp2_synth.csv is the latest convenience snapshot / backward-compatible alias.
#   - ctx.run_dir/wp2_synth.csv is the provenance artifact for a specific run.
#   - ctx.run_dir/wp2_synth_snapshot.csv is kept as the legacy per-run alias.
def run_wp2(
    cfg: dict,
    seed: int,
    ctx=None,
) -> tuple[pd.DataFrame, float, float]:
    rng = np.random.default_rng(seed)

    n_steps = int(cfg["episode"]["n_steps"])
    rv_window = int(cfg["regime"]["rv_window"])
    warmup_steps = int(cfg["regime"]["warmup_steps"])

    regime_true_int = generate_regime_series(n_steps, seed, cfg=cfg, rng=rng)
    mid, ret = generate_mid_series(regime_true_int, cfg, rng)

    # mid has n_steps+1 elements; compute rv on full mid array
    rv, sigma_hat = compute_rolling_rv(mid, rv_window, cfg["market"]["tick_size"])

    # calibrate on warmup portion (indices 0..warmup_steps-1)
    thresh_LM, thresh_MH = calibrate_thresholds(sigma_hat, warmup_steps)

    # assign detected regime
    regime_hat = assign_regime_hat(sigma_hat, thresh_LM, thresh_MH, warmup_steps)

    # convert true regime to string labels
    regime_true_str = [REGIME_LABELS[r] for r in regime_true_int]

    # Build DataFrame — align to n_steps+1 length (mid array)
    # regime_true_int has n_steps elements, ret has n_steps elements
    # pad regime_true with initial state for t=0
    regime_true_full = ["M"] + regime_true_str  # length n_steps+1
    ret_full = np.concatenate([[0.0], ret])  # length n_steps+1

    df = pd.DataFrame({
        "t": np.arange(len(mid)),
        "mid": mid,
        "ret": ret_full,
        "rv": rv,
        "sigma_hat": sigma_hat,
        "regime_true": regime_true_full,
        "regime_hat": regime_hat,
    })

    out_dir = Path("data/processed")
    out_dir.mkdir(parents=True, exist_ok=True)
    df.to_csv(out_dir / "wp2_synth.csv", index=False)

    if ctx is not None and hasattr(ctx, "run_dir"):
        run_dir = Path(ctx.run_dir)
        df.to_csv(run_dir / "wp2_synth.csv", index=False)
        df.to_csv(run_dir / "wp2_synth_snapshot.csv", index=False)

    return df, thresh_LM, thresh_MH

# ============================================================
# FILE: src/wp2/job_w2_synth.py
# PURPOSE: WP2 run.py job entry for synthetic regime generation, metrics, and diagnostic plots.
# STATUS: active
# ============================================================

from __future__ import annotations

import json
from pathlib import Path

import numpy as np
import matplotlib.pyplot as plt

from src.wp2.synth_regime import run_wp2, REGIME_LABELS


REGIME_COLORS = {"L": "#4183c4", "M": "#e8b730", "H": "#d9534f"}
REGIME_ORDER = ["L", "M", "H"]


def _plot_mid_series(df, plots_dir: Path) -> None:
    fig, ax = plt.subplots(figsize=(14, 4))
    mid = df["mid"].values
    t = df["t"].values
    regime = df["regime_true"].values

    ax.plot(t, mid, linewidth=0.4, color="black", zorder=2)

    prev = 0
    for i in range(1, len(regime)):
        if regime[i] != regime[prev] or i == len(regime) - 1:
            end = i if regime[i] != regime[prev] else i + 1
            color = REGIME_COLORS.get(regime[prev], "gray")
            ax.axvspan(t[prev], t[min(end, len(t) - 1)], alpha=0.15, color=color, linewidth=0)
            prev = i

    ax.set_title("Mid-price series with true regime background")
    ax.set_xlabel("step")
    ax.set_ylabel("mid")
    ax.legend(
        [plt.Rectangle((0, 0), 1, 1, fc=REGIME_COLORS[r], alpha=0.3) for r in REGIME_ORDER],
        REGIME_ORDER,
        loc="upper right",
    )
    fig.tight_layout()
    fig.savefig(plots_dir / "mid_series.png", dpi=150)
    plt.close(fig)


def _plot_sigma_hat(df, thresh_LM: float, thresh_MH: float, plots_dir: Path) -> None:
    fig, ax = plt.subplots(figsize=(14, 4))
    t = df["t"].values
    sigma_hat = df["sigma_hat"].values

    ax.plot(t, sigma_hat, linewidth=0.5, label="sigma_hat")
    ax.axhline(thresh_LM, color="blue", linestyle="--", linewidth=1, label=f"thresh_LM={thresh_LM:.4f}")
    ax.axhline(thresh_MH, color="red", linestyle="--", linewidth=1, label=f"thresh_MH={thresh_MH:.4f}")
    ax.set_title("Rolling realized volatility (sigma_hat)")
    ax.set_xlabel("step")
    ax.set_ylabel("sigma_hat (ticks)")
    ax.legend()
    fig.tight_layout()
    fig.savefig(plots_dir / "sigma_hat.png", dpi=150)
    plt.close(fig)


def _plot_regime_comparison(df, warmup_end: int, plots_dir: Path) -> None:
    post = df[df["t"] >= warmup_end].copy()
    t = post["t"].values

    regime_map = {"L": 0, "M": 1, "H": 2}
    true_num = np.array([regime_map.get(r, -1) for r in post["regime_true"]])
    hat_num = np.array([regime_map.get(r, -1) for r in post["regime_hat"]])

    fig, axes = plt.subplots(2, 1, figsize=(14, 5), sharex=True)
    axes[0].step(t, true_num, where="post", linewidth=0.6, color="black")
    axes[0].set_yticks([0, 1, 2])
    axes[0].set_yticklabels(["L", "M", "H"])
    axes[0].set_title("True regime")
    axes[0].set_ylabel("regime")

    axes[1].step(t, hat_num, where="post", linewidth=0.6, color="blue")
    axes[1].set_yticks([0, 1, 2])
    axes[1].set_yticklabels(["L", "M", "H"])
    axes[1].set_title("Detected regime")
    axes[1].set_xlabel("step")
    axes[1].set_ylabel("regime")

    fig.tight_layout()
    fig.savefig(plots_dir / "regime_comparison.png", dpi=150)
    plt.close(fig)


def _plot_confusion_matrix(df, warmup_end: int, plots_dir: Path) -> None:
    post = df[df["t"] >= warmup_end].copy()
    true_vals = post["regime_true"].values
    hat_vals = post["regime_hat"].values

    labels = REGIME_ORDER
    n_labels = len(labels)
    cm = np.zeros((n_labels, n_labels), dtype=int)

    label_idx = {l: i for i, l in enumerate(labels)}
    for tr, pr in zip(true_vals, hat_vals):
        if tr in label_idx and pr in label_idx:
            cm[label_idx[tr], label_idx[pr]] += 1

    fig, ax = plt.subplots(figsize=(5, 4))
    im = ax.imshow(cm, cmap="Blues")

    ax.set_xticks(range(n_labels))
    ax.set_yticks(range(n_labels))
    ax.set_xticklabels(labels)
    ax.set_yticklabels(labels)
    ax.set_xlabel("Predicted")
    ax.set_ylabel("True")
    ax.set_title("Confusion Matrix (post-warmup)")

    for i in range(n_labels):
        for j in range(n_labels):
            color = "white" if cm[i, j] > cm.max() / 2 else "black"
            ax.text(j, i, str(cm[i, j]), ha="center", va="center", color=color, fontsize=12)

    fig.colorbar(im, ax=ax, shrink=0.8)
    fig.tight_layout()
    fig.savefig(plots_dir / "confusion_matrix.png", dpi=150)
    plt.close(fig)


def job_entry(cfg: dict, ctx) -> None:
    seed = int(cfg["seed"])
    warmup_end = int(cfg["regime"]["warmup_steps"])

    df, thresh_LM, thresh_MH = run_wp2(cfg, seed, ctx=ctx)

    # --- Detection accuracy (post-warmup) ---
    post = df[df["t"] >= warmup_end].copy()
    correct = (post["regime_true"] == post["regime_hat"]).sum()
    total = len(post)
    accuracy = correct / total if total > 0 else 0.0
    ctx.logger.info(f"Detection accuracy (post-warmup): {accuracy:.4f} ({correct}/{total})")

    # --- Per-regime counts (based on regime_true, post-warmup) ---
    regime_counts = {}
    for r in REGIME_ORDER:
        cnt = int((post["regime_true"] == r).sum())
        pct = cnt / total if total > 0 else 0.0
        regime_counts[r] = {"count": cnt, "pct": round(pct, 4)}
        ctx.logger.info(f"Regime {r}: {cnt} steps ({pct:.2%})")

    # --- Empirical transition matrix (from regime_true over full series) ---
    true_labels = df["regime_true"].values
    label_idx = {l: i for i, l in enumerate(REGIME_ORDER)}
    trans_counts = np.zeros((3, 3), dtype=int)
    for i in range(1, len(true_labels)):
        fr = true_labels[i - 1]
        to = true_labels[i]
        if fr in label_idx and to in label_idx:
            trans_counts[label_idx[fr], label_idx[to]] += 1

    row_sums = trans_counts.sum(axis=1, keepdims=True)
    row_sums[row_sums == 0] = 1
    trans_probs = trans_counts / row_sums

    ctx.logger.info("Empirical transition matrix:")
    for i, r in enumerate(REGIME_ORDER):
        ctx.logger.info(f"  {r}: {trans_probs[i].round(4).tolist()}")

    # --- Expected duration per regime: 1 / (1 - p_ii) ---
    expected_duration = {}
    for i, r in enumerate(REGIME_ORDER):
        p_ii = trans_probs[i, i]
        dur = 1.0 / (1.0 - p_ii) if p_ii < 1.0 else float("inf")
        expected_duration[r] = round(dur, 2)
        ctx.logger.info(f"Expected duration {r}: {dur:.2f} steps")

    # --- Log metrics ---
    ctx.metrics.log({
        "accuracy": round(accuracy, 4),
        "thresh_LM": round(thresh_LM, 4),
        "thresh_MH": round(thresh_MH, 4),
        **{f"count_{r}": regime_counts[r]["count"] for r in REGIME_ORDER},
        **{f"pct_{r}": regime_counts[r]["pct"] for r in REGIME_ORDER},
        **{f"expected_dur_{r}": expected_duration[r] for r in REGIME_ORDER},
    })

    # --- summary.json ---
    summary = {
        "accuracy": round(accuracy, 4),
        "thresh_LM": round(thresh_LM, 4),
        "thresh_MH": round(thresh_MH, 4),
        "per_regime": regime_counts,
        "expected_duration": expected_duration,
        "empirical_transition_matrix": {
            r: trans_probs[i].round(4).tolist() for i, r in enumerate(REGIME_ORDER)
        },
    }
    summary_path = Path(ctx.run_dir) / "summary.json"
    summary_path.write_text(json.dumps(summary, indent=2), encoding="utf-8")
    ctx.logger.info(f"Summary written to {summary_path}")

    # --- Plots ---
    plots_dir = Path(ctx.plots_dir)
    _plot_mid_series(df, plots_dir)
    ctx.logger.info("Plot saved: mid_series.png")

    _plot_sigma_hat(df, thresh_LM, thresh_MH, plots_dir)
    ctx.logger.info("Plot saved: sigma_hat.png")

    _plot_regime_comparison(df, warmup_end, plots_dir)
    ctx.logger.info("Plot saved: regime_comparison.png")

    _plot_confusion_matrix(df, warmup_end, plots_dir)
    ctx.logger.info("Plot saved: confusion_matrix.png")

    ctx.logger.info("WP2 synth regime job completed.")

# ============================================================
# FILE: src/wp2/compare_detectors.py
# PURPOSE: Standalone WP2 detector comparison script for rv_baseline, rv_dwell, and HMM.
# STATUS: support
# ============================================================

"""Compare regime detectors on the same synthetic data."""
# Dedektör Karşılaştırması (WP2)
# -------------------------------
# rv_baseline, rv_dwell ve HMM dedektörlerinin tespit doğruluklarını
# gerçek rejim etiketlerine karşı ölçer ve raporlar.

from __future__ import annotations

import json
from pathlib import Path

from src.wp2.synth_regime import (
    assign_regime_hat_dwell,
    assign_regime_hat_hmm,
    run_wp2,
)

SEED = 123
CONFIG_PATH = Path("config/w2_synth.json")


def _accuracy(true: list[str], pred: list[str], warmup_end: int) -> float:
    """Post-warmup accuracy (fraction of matching labels)."""
    t = true[warmup_end:]
    p = pred[warmup_end:]
    return sum(a == b for a, b in zip(t, p)) / len(t)


def main() -> None:
    with open(CONFIG_PATH) as f:
        cfg = json.load(f)

    warmup = int(cfg["regime"]["warmup_steps"])

    # --- generate data (also produces the RV baseline) ---
    df, thresh_LM, thresh_MH = run_wp2(cfg, seed=SEED)

    regime_true = df["regime_true"].tolist()
    sigma_hat = df["sigma_hat"].to_numpy()

    # --- Detector 1: rolling RV baseline (already in df) ---
    det_rv = df["regime_hat"].tolist()

    # --- Detector 2: RV + dwell filter ---
    det_dwell = assign_regime_hat_dwell(
        sigma_hat, thresh_LM, thresh_MH, warmup, min_dwell=5,
    )

    # --- Detector 3: HMM ---
    det_hmm = assign_regime_hat_hmm(sigma_hat, warmup_end=warmup, seed=SEED)

    # --- compute accuracies ---
    results = {
        "rv_baseline": _accuracy(regime_true, det_rv, warmup),
        "rv_dwell": _accuracy(regime_true, det_dwell, warmup),
        "hmm": _accuracy(regime_true, det_hmm, warmup),
    }

    # --- print summary ---
    print(f"{'Detector':<20} {'Accuracy':>8}")
    print("-" * 30)
    for name, acc in results.items():
        print(f"{name:<20} {acc:>8.4f}")

    # --- save per-step results ---
    df["det_rv"] = det_rv
    df["det_dwell"] = det_dwell
    df["det_hmm"] = det_hmm

    out_dir = Path("data/processed")
    out_dir.mkdir(parents=True, exist_ok=True)
    df.to_csv(out_dir / "detector_comparison.csv", index=False)
    print(f"\nSaved to {out_dir / 'detector_comparison.csv'}")


if __name__ == "__main__":
    main()

# ============================================================
# FILE: src/wp3/__init__.py
# PURPOSE: Package marker for WP3 Gymnasium environment code.
# STATUS: support
# ============================================================



# ============================================================
# FILE: src/wp3/env.py
# PURPOSE: Gymnasium market-making environment defining observations, actions, rewards, exogenous series, and misspecification hooks.
# STATUS: active
# ============================================================

"""Gymnasium environment for market-making with MMSimulator backend."""
# Bu modül, PPO ajanının eğitim ve değerlendirme için kullandığı
# OpenAI Gymnasium uyumlu piyasa yapıcılığı ortamını tanımlar.
# Ajan her adımda half-spread (h) ve skew (m) seçer.

from __future__ import annotations

import gymnasium
import numpy as np
from gymnasium import spaces

from src.wp1.sim import ExecParams, MarketParams, MMSimulator, MMState


class MMEnv(gymnasium.Env):
    """Market-making env: agent chooses half-spread h and skew m each tick."""

    metadata = {"render_modes": []}

    def __init__(self, cfg: dict):
        """Initialise spaces and simulation parameters from *cfg*."""
        super().__init__()

        # Piyasa ve emir çalıştırma parametrelerini config'den yükle
        # Gözlem uzayı 6 boyutlu: [q_norm, sigma_hat, tau, r_L, r_M, r_H]
        # Eylem uzayı: h_idx(0-4) → h=1..5 tick, m_idx(0-4) → m=-2..+2 tick
        self.market = MarketParams(**cfg["market"])
        self.execp = ExecParams(**cfg["exec"])
        self.n_steps = int(cfg["episode"]["n_steps"])
        self.inv_max_clip = int(cfg["episode"].get("inv_max_clip", 50))

        wp3 = cfg.get("wp3", {})
        self.eta = float(wp3.get("eta", 0.01))
        # Regime-conditional eta: {"L": ..., "M": ..., "H": ...}
        # Eğer config'de yoksa None — fallback olarak self.eta kullanılır
        eta_regime_cfg = wp3.get("eta_regime", None)
        if eta_regime_cfg is not None:
            self.eta_regime = {
                "L": float(eta_regime_cfg["L"]),
                "M": float(eta_regime_cfg["M"]),
                "H": float(eta_regime_cfg["H"]),
            }
        else:
            self.eta_regime = None
        self.skew_penalty_c = float(wp3.get("skew_penalty_c", 0.0))
        self.use_sigma = bool(wp3.get("use_sigma", True))
        regime_source = str(wp3.get("regime_source", "hat"))
        # backward compat: eski use_regime flag'ini dönüştür
        if "regime_source" not in wp3 and "use_regime" in wp3:
            regime_source = "hat" if wp3["use_regime"] else "none"
        if regime_source not in ("none", "hat", "true"):
            raise ValueError(f"regime_source must be 'none', 'hat' or 'true', got: {regime_source!r}")
        self.regime_source = regime_source
        self.seed_val = int(cfg["seed"])
        self.cfg = cfg

        # Model misspecification: regime-dependent fill parameters
        misspec = cfg.get("misspec", {})
        self.misspec_enabled = bool(misspec.get("enabled", False))
        self.misspec_params = misspec.get("params", {})
        # e.g. {"L": {"A": 4.0, "k": 1.8}, "M": {"A": 5.0, "k": 1.5}, "H": {"A": 6.0, "k": 1.2}}

        # obs: [q_norm, sigma_hat, tau, regime_L, regime_M, regime_H]
        self.observation_space = spaces.Box(
            low=-np.inf, high=np.inf, shape=(6,), dtype=np.float32,
        )
        # action: [h_idx(0..4), m_idx(0..4)]
        self.action_space = spaces.MultiDiscrete([5, 5])

        self._sim: MMSimulator | None = None
        self._state: MMState | None = None
        self._t: int = 0
        self._exog = None

    # ------------------------------------------------------------------
    def _get_obs(self) -> np.ndarray:
        """Build 6-dim observation vector."""
        q = np.clip(self._state.inv, -self.inv_max_clip, self.inv_max_clip)
        q_norm = q / self.inv_max_clip

        tau = (self.n_steps - self._t) / self.n_steps

        # sigma_hat
        if self._exog is not None and self.use_sigma:
            idx = min(self._t, len(self._exog) - 1)
            sh = float(self._exog["sigma_hat"].iloc[idx])
            if np.isnan(sh):
                sh = 0.0
        else:
            sh = 0.0

        # regime one-hot
        r_l = r_m = r_h = 0.0
        if self.regime_source != "none" and self._exog is not None:
            idx = min(self._t, len(self._exog) - 1)
            if self.regime_source == "true":
                label = str(self._exog["regime_true"].iloc[idx])
            else:
                label = str(self._exog["regime_hat"].iloc[idx])
            # warmup veya geçersiz label → zero one-hot (M'ye map etme)
            if label == "L":
                r_l = 1.0
            elif label == "M":
                r_m = 1.0
            elif label == "H":
                r_h = 1.0

        return np.array([q_norm, sh, tau, r_l, r_m, r_h], dtype=np.float32)

    # ------------------------------------------------------------------
    def reset(self, *, seed=None, options=None):
        """Reset env; optionally inject exogenous series via *options*."""
        super().reset(seed=seed)

        # Derive a sim seed: use explicit seed if given, else sample from
        # the Gymnasium-managed np_random (seeded by super().reset).
        if seed is not None:
            sim_seed = seed
        else:
            sim_seed = int(self.np_random.integers(0, 2**31 - 1))

        self._sim = MMSimulator(self.market, self.execp, seed=sim_seed)
        self._state = self._sim.reset()
        self._t = 0

        # Update exog only when explicitly provided; otherwise keep previous.
        if options and "exog" in options:
            self._exog = options["exog"]

        # If exog series is available, override initial mid from it.
        if self._exog is not None:
            self._state = MMState(
                t=0,
                mid=float(self._exog["mid"].iloc[0]),
                cash=0.0,
                inv=0,
            )
            self._sim._mid_hist = [self._state.mid]

        obs = self._get_obs()
        info = {"equity": self._state.equity, "inv": self._state.inv}
        return obs, info

    # ------------------------------------------------------------------
    def step(self, action):
        """Execute one tick: decode action, run sim, compute reward."""
        h_idx, m_idx = int(action[0]), int(action[1])
        # Eylem decode: h_idx → half-spread (1-5 tick), m_idx → skew (-2 ile +2 arası)
        # delta_bid = h + m (en az 1 tick), delta_ask = h - m (en az 1 tick)
        # Ödül = ΔEquity − η × inv² (PnL artışından envanter cezası düş)
        h = h_idx + 1        # ticks 1..5
        m = m_idx - 2        # skew  -2..+2

        delta_bid = max(1, h + m)
        delta_ask = max(1, h - m)

        equity_before = self._state.equity

        # Apply misspec: override simulator exec params based on regime_true
        if self.misspec_enabled and self._exog is not None and "regime_true" in self._exog.columns:
            idx = min(self._t, len(self._exog) - 1)
            regime_true = str(self._exog["regime_true"].iloc[idx])
            if regime_true in self.misspec_params:
                p = self.misspec_params[regime_true]
                from src.wp1.sim import ExecParams
                self._sim.e = ExecParams(
                    A=float(p.get("A", self.execp.A)),
                    k=float(p.get("k", self.execp.k)),
                    fee_bps=self.execp.fee_bps,
                    latency_steps=self.execp.latency_steps,
                )

        # when exogenous mid is available, override simulator's BM
        if self._exog is not None and (self._t + 1) < len(self._exog):
            exog_mid = float(self._exog["mid"].iloc[self._t + 1])
            orig_evolve = self._sim._evolve_mid
            self._sim._evolve_mid = lambda _mid, _m=exog_mid: _m
            try:
                self._state, info = self._sim.step(
                    self._state, delta_bid, delta_ask,
                )
            finally:
                self._sim._evolve_mid = orig_evolve
        else:
            self._state, info = self._sim.step(
                self._state, delta_bid, delta_ask,
            )

        self._t += 1

        equity_after = self._state.equity
        inv_after = self._state.inv

        # Regime-conditional eta seçimi
        if self.eta_regime is not None and self._exog is not None:
            idx = min(self._t - 1, len(self._exog) - 1)
            regime_now = str(self._exog["regime_hat"].iloc[idx])
            if regime_now not in ("L", "M", "H"):
                regime_now = "M"
            eta_now = self.eta_regime[regime_now]
        else:
            eta_now = self.eta
        # reward: PnL increment minus inventory penalty
        reward = (equity_after - equity_before) - eta_now * (inv_after ** 2) - self.skew_penalty_c * abs(m)

        terminated = self._t >= self.n_steps
        truncated = False

        obs = self._get_obs()
        info["equity"] = equity_after
        info["inv"] = inv_after

        return obs, float(reward), terminated, truncated, info

# ============================================================
# FILE: src/wp3/w3_sanity.py
# PURPOSE: WP3 sanity-check job comparing naive, AS, and random policies through the environment.
# STATUS: active
# ============================================================

"""WP3 sanity check: compare naive, AS and random policies through the Gym env."""

from __future__ import annotations

import copy
import json
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd

from src.wp1.w1_as_baseline import as_deltas_ticks, compute_metrics
from src.wp1.sim import ExecParams, MarketParams, MMSimulator
from src.wp2.synth_regime import run_wp2
from src.wp3.env import MMEnv


def _run_episode(env, n, action_fn):
    """Run one episode and return equity, inventory, fills arrays."""
    equity = np.zeros(n + 1)
    inv = np.zeros(n + 1, dtype=int)
    fills = np.zeros(n, dtype=int)

    obs, info = equity[0], None  # placeholder; caller must reset env first
    # re-read from env after caller's reset
    equity[0] = env._state.equity
    inv[0] = env._state.inv

    for t in range(n):
        action = action_fn(t)
        obs, _r, _term, _trunc, info = env.step(action)
        equity[t + 1] = info["equity"]
        inv[t + 1] = info["inv"]
        fills[t] = info["fills"]

    return equity, inv, fills


def run(cfg: dict, ctx=None) -> None:
    """Run sanity check: naive sweep, AS, random — all through MMEnv."""
    if ctx is not None and hasattr(ctx, "run_dir"):
        out_dir = Path(ctx.run_dir)
    else:
        out_dir = Path("results/runs/manual_w3_sanity")
        out_dir.mkdir(parents=True, exist_ok=True)

    (out_dir / "plots").mkdir(exist_ok=True)
    (out_dir / "config_snapshot.json").write_text(
        json.dumps(cfg, indent=2), encoding="utf-8",
    )

    market = MarketParams(**cfg["market"])
    execp = ExecParams(**cfg["exec"])
    n = int(cfg["episode"]["n_steps"])
    seed = int(cfg["seed"])

    # generate exogenous WP2 series (mid, sigma_hat, regime_hat)
    df_exog, _, _ = run_wp2(cfg, seed)

    env = MMEnv(cfg)
    reset_opts = {"exog": df_exog}
    rows: list[dict] = []

    # ---- Policy 1: Naive (h=1..5, m=0) ----
    for h in cfg["sweep"]["half_spreads_ticks"]:
        env.reset(seed=seed, options=reset_opts)
        eq, iv, fl = _run_episode(
            env, n, action_fn=lambda _t, _h=int(h): np.array([_h - 1, 2]),
        )
        m = compute_metrics(eq, iv, fl, dt=market.dt)
        rows.append({"strategy": f"naive_h{h}", **m})

    # ---- Policy 2: AS (deltas clamped to action grid) ----
    env.reset(seed=seed, options=reset_opts)
    eq = np.zeros(n + 1)
    iv = np.zeros(n + 1, dtype=int)
    fl = np.zeros(n, dtype=int)
    eq[0] = env._state.equity
    iv[0] = env._state.inv

    for t in range(n):
        db, da = as_deltas_ticks(
            env._state.mid, env._state.inv, t, cfg, market, execp,
        )
        h_as = int(np.clip((db + da) // 2, 1, 5))
        m_as = int(np.clip((db - da) // 2, -2, 2))
        _obs, _r, _term, _trunc, info = env.step(
            np.array([h_as - 1, m_as + 2]),
        )
        eq[t + 1] = info["equity"]
        iv[t + 1] = info["inv"]
        fl[t] = info["fills"]

    m = compute_metrics(eq, iv, fl, dt=market.dt)
    rows.append({"strategy": "AS", **m})

    # ---- Policy 3: Random ----
    env.reset(seed=seed, options=reset_opts)
    env.action_space.seed(seed)
    eq = np.zeros(n + 1)
    iv = np.zeros(n + 1, dtype=int)
    fl = np.zeros(n, dtype=int)
    eq[0] = env._state.equity
    iv[0] = env._state.inv

    for t in range(n):
        action = env.action_space.sample()
        _obs, _r, _term, _trunc, info = env.step(action)
        eq[t + 1] = info["equity"]
        iv[t + 1] = info["inv"]
        fl[t] = info["fills"]

    m = compute_metrics(eq, iv, fl, dt=market.dt)
    rows.append({"strategy": "random", **m})

    # ---- Results ----
    df = pd.DataFrame(rows)
    df.to_csv(out_dir / "metrics_w3_sanity.csv", index=False)

    # bar plot
    plt.figure(figsize=(10, 5))
    plt.bar(df["strategy"], df["final_equity"])
    plt.xticks(rotation=45, ha="right")
    plt.title("WP3 Sanity: Final Equity Comparison")
    plt.ylabel("Final Equity")
    plt.tight_layout()
    plt.savefig(out_dir / "plots" / "compare_final_equity.png", dpi=150)
    plt.close()

    # ---- WP1 baseline comparison (direct sim, naive h=2) ----
    sim = MMSimulator(market, execp, seed=seed)
    s = sim.reset()
    eq_wp1 = np.zeros(n + 1)
    eq_wp1[0] = s.equity
    for t in range(n):
        s, _ = sim.step(s, delta_bid_ticks=2, delta_ask_ticks=2)
        eq_wp1[t + 1] = s.equity

    wp1_final = float(eq_wp1[-1])
    gym_h2 = float(df.loc[df["strategy"] == "naive_h2", "final_equity"].iloc[0])
    diff = gym_h2 - wp1_final

    log = ctx.logger if ctx else None
    if log:
        log.info(f"WP1 naive h=2 final_equity: {wp1_final:.4f}")
        log.info(f"WP3 gym naive_h2 final_equity: {gym_h2:.4f}")
        log.info(f"Diff (gym - WP1): {diff:.6f}")

    print(
        df[["strategy", "final_equity", "fill_rate",
            "inv_p99", "max_drawdown", "sharpe_like"]],
    )
    print("\nWP1 baseline comparison (naive h=2):")
    print(f"  WP1 direct:  {wp1_final:.4f}")
    print(f"  WP3 gym:     {gym_h2:.4f}")
    print(f"  Diff:        {diff:.6f}")


def job_entry(cfg: dict, ctx) -> None:
    """Entry point: run regime-aware then regime-blind, compare AS equity."""
    out_dir = Path(ctx.run_dir)

    # --- regime-aware ---
    run(cfg, ctx)
    df_aware = pd.read_csv(out_dir / "metrics_w3_sanity.csv")
    (out_dir / "metrics_w3_sanity.csv").rename(
        out_dir / "metrics_w3_sanity_aware.csv",
    )
    p_aware = out_dir / "plots" / "compare_final_equity.png"
    if p_aware.exists():
        p_aware.rename(out_dir / "plots" / "compare_final_equity_aware.png")

    # --- regime-blind ---
    cfg_blind = copy.deepcopy(cfg)
    cfg_blind["wp3"]["use_regime"] = False
    run(cfg_blind, ctx)
    df_blind = pd.read_csv(out_dir / "metrics_w3_sanity.csv")
    (out_dir / "metrics_w3_sanity.csv").rename(
        out_dir / "metrics_w3_sanity_blind.csv",
    )
    p_blind = out_dir / "plots" / "compare_final_equity.png"
    if p_blind.exists():
        p_blind.rename(out_dir / "plots" / "compare_final_equity_blind.png")

    # --- compare ---
    as_aware = float(
        df_aware.loc[df_aware["strategy"] == "AS", "final_equity"].iloc[0],
    )
    as_blind = float(
        df_blind.loc[df_blind["strategy"] == "AS", "final_equity"].iloc[0],
    )
    diff = as_aware - as_blind

    ctx.logger.info(f"regime-aware AS final_equity: {as_aware:.4f}")
    ctx.logger.info(f"regime-blind AS final_equity: {as_blind:.4f}")
    ctx.logger.info(f"diff: {diff:.6f}")

# ============================================================
# FILE: src/wp4/__init__.py
# PURPOSE: Package marker for WP4 PPO training code.
# STATUS: support
# ============================================================



# ============================================================
# FILE: src/wp4/job_w4_ppo.py
# PURPOSE: WP4 PPO training infrastructure and pilot/in-sample evaluation boundary.
# STATUS: active
# ============================================================

"""WP4 — PPO Training Infrastructure Pilot.

PILOT/IN-SAMPLE only. This job trains and evaluates a PPO agent on
the same full synthetic exogenous series with no temporal split. It
is used for smoke-testing the PPO training stack (Stable-Baselines3
integration, observation/action shapes, reward stability) and NOT
for reporting out-of-sample performance.

All thesis numerical claims about PPO performance come from WP5
(src/wp5/job_w5_eval.py), which implements a 70/30 chronological
train/test split with n_train=5600 and n_test=2400 over the
exogenous series.

Do not cite metrics produced by this job as out-of-sample results.
"""

from __future__ import annotations

import copy
import json
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from stable_baselines3 import PPO
from stable_baselines3.common.monitor import Monitor
from stable_baselines3.common.vec_env import DummyVecEnv

from src.wp1.w1_as_baseline import compute_metrics
from src.wp2.synth_regime import run_wp2
from src.wp3.env import MMEnv


def _eval_model(model, cfg, df_exog, seed, stage, out_dir, ctx):
    """Deterministic rollout → metrics CSV + equity/inventory plots."""
    env = MMEnv(cfg)
    obs, _ = env.reset(seed=seed, options={"exog": df_exog})

    n = int(cfg["episode"]["n_steps"])
    equity = np.zeros(n + 1)
    inv = np.zeros(n + 1, dtype=int)
    fills = np.zeros(n, dtype=int)

    equity[0] = env._state.equity
    inv[0] = env._state.inv

    for t in range(n):
        action, _ = model.predict(obs, deterministic=True)
        obs, _r, _term, _trunc, info = env.step(action)
        equity[t + 1] = info["equity"]
        inv[t + 1] = info["inv"]
        fills[t] = info["fills"]

    m = compute_metrics(equity, inv, fills, dt=cfg["market"]["dt"])

    # CSV
    pd.DataFrame([m]).to_csv(
        out_dir / f"metrics_w4_eval_{stage}.csv", index=False,
    )

    # Plots
    plots_dir = out_dir / "plots"
    ts = np.arange(n + 1)

    plt.figure()
    plt.plot(ts, equity)
    plt.title(f"Equity ({stage})")
    plt.xlabel("step")
    plt.ylabel("equity")
    plt.tight_layout()
    plt.savefig(plots_dir / f"equity_{stage}.png", dpi=150)
    plt.close()

    plt.figure()
    plt.plot(ts, inv)
    plt.title(f"Inventory ({stage})")
    plt.xlabel("step")
    plt.ylabel("inventory")
    plt.tight_layout()
    plt.savefig(plots_dir / f"inventory_{stage}.png", dpi=150)
    plt.close()

    # Log to run metrics CSV
    ctx.metrics.log({
        "stage": stage,
        "total_timesteps": int(cfg["wp4"]["total_timesteps"]),
        "final_equity": m["final_equity"],
        "inv_p99": m["inv_p99"],
        "max_drawdown": m["max_drawdown"],
        "sharpe_like": m["sharpe_like"],
        "fill_rate": m["fill_rate"],
        "turnover": m["turnover"],
    })

    return m


def job_entry(cfg: dict, ctx) -> None:
    """Train PPO (regime-aware + regime-blind) and evaluate both."""
    out_dir = Path(ctx.run_dir)
    models_dir = out_dir / "models"
    models_dir.mkdir(exist_ok=True)
    (out_dir / "plots").mkdir(exist_ok=True)

    seed = int(cfg["seed"])
    wp4 = cfg["wp4"]

    # Generate exogenous WP2 series
    df_exog, _, _ = run_wp2(cfg, seed, ctx=ctx)
    ctx.logger.info(f"WP2 exog generated: {len(df_exog)} rows")

    results = {}

    for stage, use_regime in [("aware", True), ("blind", False)]:
        ctx.logger.info(f"--- Training PPO ({stage}) ---")

        cfg_local = copy.deepcopy(cfg)
        cfg_local["wp3"]["use_regime"] = use_regime

        # Build env → Monitor → DummyVecEnv
        env = MMEnv(cfg_local)
        env.reset(seed=seed, options={"exog": df_exog})
        monitor = Monitor(env)
        vec_env = DummyVecEnv([lambda _m=monitor: _m])

        device = cfg.get("wp4", {}).get("device", "cpu")
        model = PPO(
            "MlpPolicy",
            vec_env,
            seed=seed,
            learning_rate=float(wp4["learning_rate"]),
            n_steps=int(wp4["n_steps"]),
            batch_size=int(wp4["batch_size"]),
            n_epochs=int(wp4["n_epochs"]),
            gamma=float(wp4["gamma"]),
            gae_lambda=float(wp4["gae_lambda"]),
            clip_range=float(wp4["clip_range"]),
            ent_coef=float(wp4["ent_coef"]),
            verbose=1,
            device=device,
        )

        model.learn(total_timesteps=int(wp4["total_timesteps"]))
        model.save(str(models_dir / f"ppo_{stage}"))
        ctx.logger.info(f"Model saved: models/ppo_{stage}.zip")

        vec_env.close()

        # Deterministic evaluation
        m = _eval_model(model, cfg_local, df_exog, seed, stage, out_dir, ctx)
        results[stage] = m
        ctx.logger.info(
            f"Eval {stage}: equity={m['final_equity']:.4f}  "
            f"sharpe={m['sharpe_like']:.4f}  inv_p99={m['inv_p99']:.1f}",
        )

    # Compare bar plot
    stages = list(results.keys())
    equities = [results[s]["final_equity"] for s in stages]

    plt.figure(figsize=(6, 4))
    plt.bar(stages, equities, color=["steelblue", "salmon"])
    plt.title("PPO Final Equity: Aware vs Blind")
    plt.ylabel("Final Equity")
    plt.tight_layout()
    plt.savefig(
        out_dir / "plots" / "compare_final_equity_aware_vs_blind.png", dpi=150,
    )
    plt.close()

    # Summary JSON
    summary = {
        "wp4_hyperparams": wp4,
        "aware": results["aware"],
        "blind": results["blind"],
    }
    (out_dir / "summary_w4.json").write_text(
        json.dumps(summary, indent=2), encoding="utf-8",
    )
    ctx.logger.info("summary_w4.json written")

# ============================================================
# FILE: src/wp5/__init__.py
# PURPOSE: Package marker for WP5 OOS evaluation, robustness, statistics, and figure code.
# STATUS: support
# ============================================================



# ============================================================
# FILE: src/wp5/job_w5_eval.py
# PURPOSE: WP5 main out-of-sample evaluation for baseline and PPO strategies.
# STATUS: active
# ============================================================

"""WP5 — Out-of-sample evaluation: naive, AS, PPO variants (5 ablation configs)."""
# WP5 Ana Değerlendirme
# ----------------------
# 7 stratejiyi (naive, AS + 5 PPO varyantı) 20 bağımsız seed üzerinde
# out-of-sample değerlendirir. PPO varyantları: sigma_only, regime_only,
# combined, oracle_pure, oracle_full.
# Sonuçlar: metrics_wp5_oos.csv ve metrics_wp5_oos_by_regime.csv

from __future__ import annotations

import copy
import math
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from stable_baselines3 import PPO
from stable_baselines3.common.monitor import Monitor
from stable_baselines3.common.vec_env import DummyVecEnv

from src.wp1.w1_as_baseline import as_deltas_ticks, compute_metrics
from src.wp1.sim import ExecParams, MarketParams
from src.wp2.synth_regime import run_wp2
from src.wp3.env import MMEnv


# ------------------------------------------------------------------
# Helpers
# ------------------------------------------------------------------

def _run_wp2_safe(cfg, seed, ctx):
    """Call run_wp2, handling old signatures that don't accept ctx."""
    try:
        return run_wp2(cfg, seed, ctx=ctx)
    except TypeError:
        df, thresh_LM, thresh_MH = run_wp2(cfg, seed)
        return df, thresh_LM, thresh_MH


def _compute_regime_metrics(equity, inv, fills, fees, regime_labels, dt):
    """Per-regime performance breakdown."""
    n = len(fills)
    results = []
    for regime in sorted(set(regime_labels)):
        # indices where this regime is active (use first n+1 for equity/inv)
        mask = np.array([regime_labels[t] == regime for t in range(n)])
        steps_count = int(mask.sum())
        if steps_count < 10:
            continue
        step_pnl = np.diff(equity)[mask]
        mean_step_pnl = float(step_pnl.mean()) if len(step_pnl) else 0.0
        std_pnl = float(step_pnl.std(ddof=1)) if len(step_pnl) > 1 else 0.0
        sharpe_like = (mean_step_pnl / std_pnl * np.sqrt(1.0 / dt)) if std_pnl > 0 else 0.0
        fill_rate = float(fills[mask].sum() / steps_count) if steps_count > 0 else 0.0
        inv_vals = np.abs(inv[:-1][mask])
        inv_p99 = float(np.quantile(inv_vals, 0.99)) if len(inv_vals) else 0.0
        results.append({
            "regime": regime,
            "mean_step_pnl": mean_step_pnl,
            "sharpe_like": sharpe_like,
            "fill_rate": fill_rate,
            "inv_p99": inv_p99,
            "steps_count": steps_count,
        })
    return results


# ------------------------------------------------------------------
# Job entry
# ------------------------------------------------------------------

def job_entry(cfg: dict, ctx) -> None:
    out_dir = Path(ctx.run_dir)
    (out_dir / "models").mkdir(exist_ok=True)
    (out_dir / "plots").mkdir(exist_ok=True)
    (out_dir / "curves").mkdir(exist_ok=True)

    market = MarketParams(**cfg["market"])
    execp = ExecParams(**cfg["exec"])
    wp5 = cfg["wp5"]
    seeds = wp5["seeds"]
    train_frac = float(wp5["train_frac"])
    naive_h = int(wp5["naive"]["h"])
    naive_m = int(wp5["naive"]["m"])
    n_full = int(cfg["episode"]["n_steps"])
    n_train = math.floor(train_frac * n_full)
    n_test = n_full - n_train

    VARIANTS = {
        "sigma_only":  {"use_sigma": True,  "regime_source": "none"},
        "regime_only": {"use_sigma": False, "regime_source": "hat"},
        "combined":    {"use_sigma": True,  "regime_source": "hat"},
        "oracle_pure": {"use_sigma": False, "regime_source": "true"},
        "oracle_full": {"use_sigma": True,  "regime_source": "true"},
    }

    rows_oos = []
    rows_regime = []

    for seed in seeds:
        ctx.logger.info(f"=== Seed {seed} ===")

        # 1) Exog series
        df_exog, _, _ = _run_wp2_safe(cfg, seed, ctx)

        # 2) Split — +1 row trick
        exog_train = df_exog.iloc[: n_train + 1].reset_index(drop=True)
        exog_test = df_exog.iloc[n_train : n_train + n_test + 1].reset_index(drop=True)

        # 3) Train PPO variants
        for stage, vcfg in VARIANTS.items():
            ctx.logger.info(f"Training PPO-{stage} seed={seed}")
            cfg_tr = copy.deepcopy(cfg)
            cfg_tr["wp3"]["use_sigma"] = vcfg["use_sigma"]
            cfg_tr["wp3"]["regime_source"] = vcfg["regime_source"]
            cfg_tr["episode"] = {**cfg_tr["episode"], "n_steps": n_train}
            cfg_tr["as"]["horizon_steps"] = n_train

            env_tr = MMEnv(cfg_tr)
            env_tr.reset(seed=seed, options={"exog": exog_train})
            vec_env = DummyVecEnv([lambda _e=Monitor(env_tr): _e])

            wp4 = cfg["wp4"]
            device = cfg.get("wp4", {}).get("device", "cpu")
            model = PPO(
                "MlpPolicy", vec_env, seed=seed,
                learning_rate=float(wp4["learning_rate"]),
                n_steps=int(wp4["n_steps"]),
                batch_size=int(wp4["batch_size"]),
                n_epochs=int(wp4["n_epochs"]),
                gamma=float(wp4["gamma"]),
                gae_lambda=float(wp4["gae_lambda"]),
                clip_range=float(wp4["clip_range"]),
                ent_coef=float(wp4["ent_coef"]),
                verbose=0,
                device=device,
            )
            model.learn(total_timesteps=int(wp4["total_timesteps"]))

            seed_dir = out_dir / "models" / f"seed{seed}"
            seed_dir.mkdir(parents=True, exist_ok=True)
            model.save(str(seed_dir / f"ppo_{stage}"))
            ctx.logger.info(f"Saved: models/seed{seed}/ppo_{stage}.zip")
            vec_env.close()

        # 4) OOS Evaluation
        models = {
            name: PPO.load(str(out_dir / "models" / f"seed{seed}" / f"ppo_{name}"), device="cpu")
            for name in VARIANTS
        }

        def _base_eval_cfg():
            c = copy.deepcopy(cfg)
            c["episode"] = {**c["episode"], "n_steps": n_test}
            c["wp3"]["use_sigma"] = True
            c["wp3"]["regime_source"] = "hat"
            c["as"]["horizon_steps"] = n_test
            return c

        strategies = {
            "naive": (_base_eval_cfg(), None),
            "AS":    (_base_eval_cfg(), None),
        }
        for vname, vcfg in VARIANTS.items():
            cfg_ev = copy.deepcopy(cfg)
            cfg_ev["episode"] = {**cfg_ev["episode"], "n_steps": n_test}
            cfg_ev["wp3"]["use_sigma"] = vcfg["use_sigma"]
            cfg_ev["wp3"]["regime_source"] = vcfg["regime_source"]
            cfg_ev["as"]["horizon_steps"] = n_test
            strategies[f"ppo_{vname}"] = (cfg_ev, models[vname])

        for strat_name, (cfg_ev, model_ev) in strategies.items():
            env_ev = MMEnv(cfg_ev)
            obs, _ = env_ev.reset(seed=seed, options={"exog": exog_test})

            eq = np.zeros(n_test + 1)
            iv = np.zeros(n_test + 1, dtype=int)
            fl = np.zeros(n_test, dtype=int)
            fe = np.zeros(n_test)
            eq[0] = env_ev._state.equity
            iv[0] = env_ev._state.inv
            h_arr = np.full(n_test + 1, np.nan)
            m_arr = np.full(n_test + 1, np.nan)
            rh_arr = [""] * (n_test + 1)
            rh_arr[0] = str(exog_test["regime_hat"].iloc[0]) if "regime_hat" in exog_test.columns else ""

            for t in range(n_test):
                if strat_name == "naive":
                    action = np.array([naive_h - 1, naive_m + 2])
                elif strat_name == "AS":
                    db, da = as_deltas_ticks(
                        env_ev._state.mid, env_ev._state.inv, t, cfg_ev, market, execp,
                    )
                    h_as = int(np.clip((db + da) // 2, 1, 5))
                    m_as = int(np.clip((db - da) // 2, -2, 2))
                    action = np.array([h_as - 1, m_as + 2])
                else:
                    action, _ = model_ev.predict(obs, deterministic=True)

                h_val = int(action[0]) + 1
                m_val = int(action[1]) - 2
                h_arr[t] = h_val
                m_arr[t] = m_val

                obs, _r, _term, _trunc, info = env_ev.step(action)
                eq[t + 1] = info["equity"]
                iv[t + 1] = info["inv"]
                fl[t] = info["fills"]
                fe[t] = float(info.get("fee_total", 0.0))
                idx = min(t + 1, len(exog_test) - 1)
                rh_arr[t + 1] = str(exog_test["regime_hat"].iloc[idx]) if "regime_hat" in exog_test.columns else ""

            # General metrics
            m = compute_metrics(eq, iv, fl, dt=market.dt)
            total_fees = float(fe.sum())
            n_trades = int(fl.sum())
            fee_per_trade = total_fees / n_trades if n_trades > 0 else 0.0

            row = {
                "seed": seed, "strategy": strat_name, "split": "test",
                **m, "total_fees": total_fees, "fee_per_trade": fee_per_trade,
            }
            rows_oos.append(row)
            ctx.metrics.log({
                "seed": seed, "strategy": strat_name,
                "final_equity": m["final_equity"],
                "inv_p99": m["inv_p99"],
                "sharpe_like": m["sharpe_like"],
            })

            # Save curve
            # regime_true her zaman mevcut (sentetik veri)
            rt_arr = list(exog_test["regime_true"].values[:n_test + 1]) if "regime_true" in exog_test.columns else [""] * (n_test + 1)
            # obs_regime_source: bu strateji hangi kaynağı kullandı
            ev_cfg_this = strategies[strat_name][0]
            obs_regime_source = ev_cfg_this["wp3"].get("regime_source", "hat") if ev_cfg_this is not None else "hat"

            pd.DataFrame({
                "t": np.arange(n_test + 1),
                "equity": eq,
                "inv": iv,
                "h": h_arr,
                "m": m_arr,
                "regime_hat": rh_arr,
                "regime_true": rt_arr,
                "obs_regime_source": obs_regime_source,
            }).to_csv(
                out_dir / "curves" / f"seed{seed}_{strat_name}_test.csv", index=False,
            )

            # Regime-wise metrics
            if "regime_true" in exog_test.columns:
                regime_labels = list(exog_test["regime_true"].values[:n_test + 1])
            else:
                regime_labels = ["M"] * (n_test + 1)

            rw = _compute_regime_metrics(eq, iv, fl, fe, regime_labels, market.dt)
            for r in rw:
                rows_regime.append({"seed": seed, "strategy": strat_name, **r})

            ctx.logger.info(
                f"seed={seed} {strat_name}: equity={m['final_equity']:.4f} "
                f"sharpe={m['sharpe_like']:.4f}"
            )

    # 5) Save CSVs
    df_oos = pd.DataFrame(rows_oos)
    df_oos.to_csv(out_dir / "metrics_wp5_oos.csv", index=False)

    df_regime = pd.DataFrame(rows_regime)
    df_regime.to_csv(out_dir / "metrics_wp5_oos_by_regime.csv", index=False)

    # 6) Plots
    # Plot 1: Final equity by seed (grouped bar)
    strat_names = df_oos["strategy"].unique().tolist()
    seed_vals = sorted(df_oos["seed"].unique().tolist())
    x = np.arange(len(seed_vals))
    width = 0.2
    fig, ax = plt.subplots(figsize=(10, 5))
    for i, s in enumerate(strat_names):
        vals = [
            float(df_oos[(df_oos["seed"] == sd) & (df_oos["strategy"] == s)]["final_equity"].iloc[0])
            for sd in seed_vals
        ]
        ax.bar(x + i * width, vals, width, label=s)
    ax.set_xticks(x + width * 1.5)
    ax.set_xticklabels([f"seed={s}" for s in seed_vals])
    ax.set_title("WP5 Ablation OOS Final Equity by Seed")
    ax.set_ylabel("Final Equity")
    ax.legend()
    fig.tight_layout()
    fig.savefig(out_dir / "plots" / "wp5_final_equity_by_seed.png", dpi=150)
    plt.close(fig)

    # Plot 2: Mean +/- std across seeds
    agg = df_oos.groupby("strategy")["final_equity"].agg(["mean", "std"]).reset_index()
    fig, ax = plt.subplots(figsize=(7, 4))
    ax.bar(agg["strategy"], agg["mean"], yerr=agg["std"].fillna(0), capsize=5)
    ax.set_title("WP5 Ablation: Mean +/- Std Across Seeds")
    ax.set_ylabel("Final Equity")
    fig.tight_layout()
    fig.savefig(out_dir / "plots" / "wp5_final_equity_mean_std.png", dpi=150)
    plt.close(fig)

    ctx.logger.info("WP5 complete.")
    print(df_oos[["seed", "strategy", "final_equity", "sharpe_like", "inv_p99", "max_drawdown"]])

# ============================================================
# FILE: src/wp5/job_w5_ablation_eta.py
# PURPOSE: WP5 inventory-penalty eta ablation job.
# STATUS: active
# ============================================================

"""WP5.1 — Eta ablation sweep (PPO-aware only, OOS test)."""
# Eta (η) Ablasyon Deneyi (WP5)
# ------------------------------
# Envanter ceza katsayısının (η) farklı değerleri için PPO performansını ölçer.
# η = [0.0001, 0.0005, 0.001, 0.005, 0.01] gibi değerler test edilir.
# Optimal η = 0.001 bu deney ile belirlendi.

from __future__ import annotations

import copy
import math
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from stable_baselines3 import PPO
from stable_baselines3.common.monitor import Monitor
from stable_baselines3.common.vec_env import DummyVecEnv

from src.wp1.w1_as_baseline import compute_metrics
from src.wp2.synth_regime import run_wp2
from src.wp3.env import MMEnv


# ------------------------------------------------------------------
# Helpers
# ------------------------------------------------------------------

def _run_wp2_safe(cfg, seed, ctx):
    try:
        return run_wp2(cfg, seed, ctx=ctx)
    except TypeError:
        df, thresh_LM, thresh_MH = run_wp2(cfg, seed)
        return df, thresh_LM, thresh_MH


# ------------------------------------------------------------------
# Job entry
# ------------------------------------------------------------------

def job_entry(cfg: dict, ctx) -> None:
    out_dir = Path(ctx.run_dir)
    (out_dir / "models").mkdir(exist_ok=True)
    (out_dir / "plots").mkdir(exist_ok=True)

    wp5 = cfg["wp5"]
    seeds = wp5["seeds"]
    eta_values = wp5["eta_values"]
    train_frac = float(wp5.get("train_frac", 0.7))
    n_full = int(cfg["episode"]["n_steps"])
    n_train = math.floor(train_frac * n_full)
    n_test = n_full - n_train
    dt = float(cfg["market"]["dt"])

    rows = []

    for eta in eta_values:
        for seed in seeds:
            ctx.logger.info(f"eta={eta} seed={seed}")

            cfg_local = copy.deepcopy(cfg)
            cfg_local["seed"] = seed
            cfg_local["wp3"]["eta"] = eta
            cfg_local["wp3"]["use_regime"] = True

            # (A) Exog series
            df_exog, _, _ = _run_wp2_safe(cfg_local, seed, ctx)
            exog_train = df_exog.iloc[: n_train + 1].reset_index(drop=True)
            exog_test = df_exog.iloc[n_train : n_train + n_test + 1].reset_index(drop=True)

            # (B) Train
            cfg_train = copy.deepcopy(cfg_local)
            cfg_train["episode"]["n_steps"] = n_train

            env_tr = MMEnv(cfg_train)
            env_tr.reset(seed=seed, options={"exog": exog_train})
            vec_env = DummyVecEnv([lambda _e=Monitor(env_tr): _e])

            wp4 = cfg["wp4"]
            device = cfg.get("wp4", {}).get("device", "cpu")
            model = PPO(
                "MlpPolicy", vec_env, seed=seed,
                learning_rate=float(wp4["learning_rate"]),
                n_steps=int(wp4["n_steps"]),
                batch_size=int(wp4["batch_size"]),
                n_epochs=int(wp4["n_epochs"]),
                gamma=float(wp4["gamma"]),
                gae_lambda=float(wp4["gae_lambda"]),
                clip_range=float(wp4["clip_range"]),
                ent_coef=float(wp4["ent_coef"]),
                verbose=0,
                device=device,
            )
            model.learn(total_timesteps=int(wp4["total_timesteps"]))

            seed_dir = out_dir / "models" / f"eta{eta}" / f"seed{seed}"
            seed_dir.mkdir(parents=True, exist_ok=True)
            model.save(str(seed_dir / "ppo_aware"))
            ctx.logger.info(f"Saved: models/eta{eta}/seed{seed}/ppo_aware.zip")
            vec_env.close()

            # (C) OOS test
            cfg_test = copy.deepcopy(cfg_local)
            cfg_test["episode"]["n_steps"] = n_test

            env_test = MMEnv(cfg_test)
            obs, _ = env_test.reset(seed=seed, options={"exog": exog_test})

            eq = np.zeros(n_test + 1)
            iv = np.zeros(n_test + 1, dtype=int)
            fl = np.zeros(n_test, dtype=int)
            eq[0] = env_test._state.equity
            iv[0] = env_test._state.inv

            for t in range(n_test):
                action, _ = model.predict(obs, deterministic=True)
                obs, _r, _term, _trunc, info = env_test.step(action)
                eq[t + 1] = info["equity"]
                iv[t + 1] = info["inv"]
                fl[t] = info["fills"]

            m = compute_metrics(eq, iv, fl, dt=dt)
            row = {"eta": eta, "seed": seed, **m}
            rows.append(row)

            ctx.metrics.log({
                "eta": eta, "seed": seed,
                "final_equity": m["final_equity"],
                "inv_p99": m["inv_p99"],
                "sharpe_like": m["sharpe_like"],
                "fill_rate": m["fill_rate"],
            })
            ctx.logger.info(
                f"eta={eta} seed={seed}: equity={m['final_equity']:.4f} "
                f"sharpe={m['sharpe_like']:.4f} fill_rate={m['fill_rate']:.4f}"
            )

    # Save CSV
    df = pd.DataFrame(rows)
    df.to_csv(out_dir / "metrics_wp5_ablation_eta.csv", index=False)

    # Plots
    eta_vals_sorted = sorted(df["eta"].unique())
    seed_vals = sorted(df["seed"].unique())

    for metric, ylabel, fname in [
        ("fill_rate", "Fill Rate", "wp5_ablation_eta_fill_rate.png"),
        ("final_equity", "Final Equity", "wp5_ablation_eta_final_equity.png"),
    ]:
        fig, ax = plt.subplots(figsize=(7, 4))
        for s in seed_vals:
            sub = df[df["seed"] == s].sort_values("eta")
            ax.plot(sub["eta"], sub[metric], marker="o", label=f"seed={s}")
        ax.set_xscale("log")
        ax.set_xlabel("eta")
        ax.set_ylabel(ylabel)
        ax.set_title(f"WP5 Eta Ablation: {ylabel}")
        ax.legend()
        fig.tight_layout()
        fig.savefig(out_dir / "plots" / fname, dpi=150)
        plt.close(fig)

    ctx.logger.info("WP5.1 eta ablation complete.")
    print(df[["eta", "seed", "final_equity", "sharpe_like", "inv_p99", "fill_rate"]])

# ============================================================
# FILE: src/wp5/job_w5_ablation_skew.py
# PURPOSE: WP5 skew-penalty ablation job and action-distribution diagnostics.
# STATUS: active
# ============================================================

"""WP5 — Skew penalty ablation (PPO-aware only, OOS test with action histograms)."""
# Skew Penalty Ablasyon Deneyi (WP5)
# ------------------------------------
# Skew ceza katsayısının (c) etkisini ölçer: R_t = ΔEquity - η*inv² - c*|m|
# c=0 (kontrol grubu) ile c>0 karşılaştırılır.
# ppo_aware'in H rejimindeki bimodal skew dağılımını düzeltir.

from __future__ import annotations

import copy
import math
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from stable_baselines3 import PPO
from stable_baselines3.common.monitor import Monitor
from stable_baselines3.common.vec_env import DummyVecEnv

from src.wp1.w1_as_baseline import compute_metrics
from src.wp2.synth_regime import run_wp2
from src.wp3.env import MMEnv


# ------------------------------------------------------------------
# Helpers
# ------------------------------------------------------------------

def _run_wp2_safe(cfg, seed, ctx):
    try:
        return run_wp2(cfg, seed, ctx=ctx)
    except TypeError:
        df, thresh_LM, thresh_MH = run_wp2(cfg, seed)
        return df, thresh_LM, thresh_MH


REGIMES = ["L", "M", "H"]


def _hist_plot(curves_for_c: pd.DataFrame, col: str, title: str, outpath: Path):
    """Plot action distribution per regime for a single c value."""
    fig, axes = plt.subplots(1, 3, figsize=(12, 3.5), sharey=False)
    for j, rg in enumerate(REGIMES):
        ax = axes[j]
        x = curves_for_c[curves_for_c["regime_hat"] == rg][col].dropna().values
        if len(x) == 0:
            ax.set_title(f"{rg} (n=0)")
            continue
        vals, counts = np.unique(x, return_counts=True)
        probs = counts / counts.sum()
        ax.bar(vals, probs)
        ax.set_title(f"{rg} (n={len(x)})")
        ax.set_ylim(0, max(0.30, probs.max() * 1.2))
        ax.grid(True, axis="y", alpha=0.3)
        ax.set_xlabel(col)
    fig.suptitle(title)
    fig.tight_layout(rect=[0, 0.02, 1, 0.95])
    fig.savefig(outpath, dpi=150)
    plt.close(fig)


# ------------------------------------------------------------------
# Job entry
# ------------------------------------------------------------------

def job_entry(cfg: dict, ctx) -> None:
    out_dir = Path(ctx.run_dir)
    (out_dir / "models").mkdir(exist_ok=True)
    (out_dir / "plots").mkdir(exist_ok=True)
    (out_dir / "curves").mkdir(exist_ok=True)

    wp5 = cfg["wp5"]
    seeds = wp5["seeds"]
    skew_c_values = wp5["skew_c_values"]
    train_frac = float(wp5.get("train_frac", 0.7))
    n_full = int(cfg["episode"]["n_steps"])
    n_train = math.floor(train_frac * n_full)
    n_test = n_full - n_train
    dt = float(cfg["market"]["dt"])

    rows = []

    for c_val in skew_c_values:
        all_curves_for_c = []

        for seed in seeds:
            ctx.logger.info(f"skew_c={c_val} seed={seed}")

            cfg_local = copy.deepcopy(cfg)
            cfg_local["seed"] = seed
            cfg_local["wp3"]["eta"] = float(cfg["wp3"].get("eta", 0.001))
            cfg_local["wp3"]["use_regime"] = True
            cfg_local["wp3"]["skew_penalty_c"] = c_val

            # (A) Exog series
            df_exog, _, _ = _run_wp2_safe(cfg_local, seed, ctx)
            exog_train = df_exog.iloc[: n_train + 1].reset_index(drop=True)
            exog_test = df_exog.iloc[n_train : n_train + n_test + 1].reset_index(drop=True)

            # (B) Train PPO-aware with skew penalty
            cfg_train = copy.deepcopy(cfg_local)
            cfg_train["episode"]["n_steps"] = n_train
            cfg_train["as"]["horizon_steps"] = n_train

            env_tr = MMEnv(cfg_train)
            env_tr.reset(seed=seed, options={"exog": exog_train})
            vec_env = DummyVecEnv([lambda _e=Monitor(env_tr): _e])

            wp4 = cfg["wp4"]
            device = cfg.get("wp4", {}).get("device", "cpu")
            model = PPO(
                "MlpPolicy", vec_env, seed=seed,
                learning_rate=float(wp4["learning_rate"]),
                n_steps=int(wp4["n_steps"]),
                batch_size=int(wp4["batch_size"]),
                n_epochs=int(wp4["n_epochs"]),
                gamma=float(wp4["gamma"]),
                gae_lambda=float(wp4["gae_lambda"]),
                clip_range=float(wp4["clip_range"]),
                ent_coef=float(wp4["ent_coef"]),
                verbose=0,
                device=device,
            )
            model.learn(total_timesteps=int(wp4["total_timesteps"]))

            seed_dir = out_dir / "models" / f"c{c_val}" / f"seed{seed}"
            seed_dir.mkdir(parents=True, exist_ok=True)
            model.save(str(seed_dir / "ppo_aware"))
            ctx.logger.info(f"Saved: models/c{c_val}/seed{seed}/ppo_aware.zip")
            vec_env.close()

            # (C) OOS test — no skew penalty in eval env (penalty only shapes training)
            cfg_test = copy.deepcopy(cfg_local)
            cfg_test["episode"]["n_steps"] = n_test
            cfg_test["as"]["horizon_steps"] = n_test
            cfg_test["wp3"]["skew_penalty_c"] = 0.0  # eval uses raw PnL

            env_test = MMEnv(cfg_test)
            obs, _ = env_test.reset(seed=seed, options={"exog": exog_test})

            eq = np.zeros(n_test + 1)
            iv = np.zeros(n_test + 1, dtype=int)
            fl = np.zeros(n_test, dtype=int)
            h_arr = np.full(n_test + 1, np.nan)
            m_arr = np.full(n_test + 1, np.nan)
            rh_arr = [""] * (n_test + 1)
            eq[0] = env_test._state.equity
            iv[0] = env_test._state.inv
            rh_arr[0] = str(exog_test["regime_hat"].iloc[0]) if "regime_hat" in exog_test.columns else ""

            for t in range(n_test):
                action, _ = model.predict(obs, deterministic=True)
                h_val = int(action[0]) + 1
                m_val = int(action[1]) - 2
                h_arr[t] = h_val
                m_arr[t] = m_val

                obs, _r, _term, _trunc, info = env_test.step(action)
                eq[t + 1] = info["equity"]
                iv[t + 1] = info["inv"]
                fl[t] = info["fills"]
                idx = min(t + 1, len(exog_test) - 1)
                rh_arr[t + 1] = str(exog_test["regime_hat"].iloc[idx]) if "regime_hat" in exog_test.columns else ""

            # Metrics
            m = compute_metrics(eq, iv, fl, dt=dt)

            # Per-regime step PnL in H
            mask_h = np.array([rh_arr[t] == "H" for t in range(n_test)])
            step_pnl = np.diff(eq)
            mean_step_pnl_H = float(step_pnl[mask_h].mean()) if mask_h.sum() > 0 else 0.0

            row = {
                "skew_c": c_val, "seed": seed, **m,
                "mean_step_pnl_H": mean_step_pnl_H,
            }
            rows.append(row)

            ctx.metrics.log({
                "skew_c": c_val, "seed": seed,
                "final_equity": m["final_equity"],
                "inv_p99": m["inv_p99"],
                "sharpe_like": m["sharpe_like"],
            })
            ctx.logger.info(
                f"skew_c={c_val} seed={seed}: equity={m['final_equity']:.4f} "
                f"sharpe={m['sharpe_like']:.4f} inv_p99={m['inv_p99']:.0f}"
            )

            # Save curve CSV
            curve_df = pd.DataFrame({
                "t": np.arange(n_test + 1),
                "equity": eq, "inv": iv,
                "h": h_arr, "m": m_arr, "regime_hat": rh_arr,
            })
            curve_df.to_csv(
                out_dir / "curves" / f"c{c_val}_seed{seed}_ppo_aware_test.csv",
                index=False,
            )
            all_curves_for_c.append(curve_df)

        # Action histograms for this c value
        combined = pd.concat(all_curves_for_c, ignore_index=True)
        combined = combined[combined["regime_hat"].isin(REGIMES)]
        _hist_plot(
            combined, "h",
            f"h distribution (c={c_val})",
            out_dir / "plots" / f"hist_h_c{c_val}.png",
        )
        _hist_plot(
            combined, "m",
            f"m distribution (c={c_val})",
            out_dir / "plots" / f"hist_m_c{c_val}.png",
        )

    # Save summary CSV
    df = pd.DataFrame(rows)
    df.to_csv(out_dir / "metrics_wp5_ablation_skew.csv", index=False)

    # Summary plot: mean equity and inv_p99 vs c
    agg = df.groupby("skew_c").agg(
        mean_equity=("final_equity", "mean"),
        std_equity=("final_equity", "std"),
        mean_inv_p99=("inv_p99", "mean"),
        mean_sharpe=("sharpe_like", "mean"),
        mean_step_pnl_H=("mean_step_pnl_H", "mean"),
    ).reset_index()

    fig, axes = plt.subplots(1, 3, figsize=(14, 4))
    axes[0].errorbar(agg["skew_c"], agg["mean_equity"], yerr=agg["std_equity"], marker="o", capsize=4)
    axes[0].set_xlabel("skew_c")
    axes[0].set_ylabel("Final Equity")
    axes[0].set_title("Equity vs skew penalty c")
    axes[0].grid(alpha=0.3)

    axes[1].plot(agg["skew_c"], agg["mean_inv_p99"], marker="s", color="tab:orange")
    axes[1].set_xlabel("skew_c")
    axes[1].set_ylabel("inv_p99")
    axes[1].set_title("inv_p99 vs skew penalty c")
    axes[1].grid(alpha=0.3)

    axes[2].plot(agg["skew_c"], agg["mean_step_pnl_H"], marker="^", color="tab:red")
    axes[2].set_xlabel("skew_c")
    axes[2].set_ylabel("Mean step PnL (H regime)")
    axes[2].set_title("H-regime step PnL vs c")
    axes[2].grid(alpha=0.3)

    fig.tight_layout()
    fig.savefig(out_dir / "plots" / "skew_ablation_summary.png", dpi=150)
    plt.close(fig)

    ctx.logger.info("WP5 skew ablation complete.")
    print(df[["skew_c", "seed", "final_equity", "sharpe_like", "inv_p99", "mean_step_pnl_H"]])

# ============================================================
# FILE: src/wp5/job_w5_detector_compare.py
# PURPOSE: WP5 detector-robustness experiment across rv_baseline, rv_dwell, and HMM.
# STATUS: active
# ============================================================

"""WP5 — Detector comparison: rv_baseline vs rv_dwell vs hmm (PPO-aware & PPO-blind)."""
# Detector Robustness Deneyi (WP5)
# ----------------------------------
# 3 dedektör × 20 seed × 2 strateji = 120 model eğitir ve değerlendirir.
# Dedektörler: rv_baseline (%60.7), rv_dwell (%60.4), HMM (%81.8)
# Null sonucun dedektör kalitesinden bağımsız olduğunu kanıtlar.

from __future__ import annotations

import copy
import math
from pathlib import Path

import numpy as np
import pandas as pd
from stable_baselines3 import PPO
from stable_baselines3.common.monitor import Monitor
from stable_baselines3.common.vec_env import DummyVecEnv

from src.wp1.w1_as_baseline import compute_metrics
from src.wp2.synth_regime import (
    assign_regime_hat_dwell,
    assign_regime_hat_hmm,
    run_wp2,
)
from src.wp3.env import MMEnv


# ------------------------------------------------------------------
# Helpers
# ------------------------------------------------------------------

DETECTOR_TYPES = ["rv_baseline", "rv_dwell", "hmm"]


def _run_wp2_safe(cfg, seed, ctx):
    try:
        return run_wp2(cfg, seed, ctx=ctx)
    except TypeError:
        return run_wp2(cfg, seed)


def _apply_detector(detector: str, df_exog: pd.DataFrame,
                    thresh_LM: float, thresh_MH: float,
                    warmup_end: int) -> pd.DataFrame:
    """Return a copy of df_exog with regime_hat replaced by the chosen detector."""
    df = df_exog.copy()
    sigma_hat = df["sigma_hat"].to_numpy()

    if detector == "rv_baseline":
        pass  # already correct from run_wp2
    elif detector == "rv_dwell":
        df["regime_hat"] = assign_regime_hat_dwell(
            sigma_hat, thresh_LM, thresh_MH, warmup_end, min_dwell=5,
        )
    elif detector == "hmm":
        df["regime_hat"] = assign_regime_hat_hmm(
            sigma_hat, warmup_end,
        )
    else:
        raise ValueError(f"Unknown detector: {detector}")

    return df


# ------------------------------------------------------------------
# Job entry
# ------------------------------------------------------------------

def job_entry(cfg: dict, ctx) -> None:
    out_dir = Path(ctx.run_dir)
    (out_dir / "models").mkdir(exist_ok=True)

    wp5 = cfg["wp5"]
    seeds = wp5["seeds"]
    train_frac = float(wp5.get("train_frac", 0.7))
    n_full = int(cfg["episode"]["n_steps"])
    n_train = math.floor(train_frac * n_full)
    n_test = n_full - n_train
    dt = float(cfg["market"]["dt"])
    warmup = int(cfg["regime"]["warmup_steps"])

    rows = []

    for detector in DETECTOR_TYPES:
        for seed in seeds:
            ctx.logger.info(f"detector={detector} seed={seed}")

            # 1) Generate exog series
            df_exog, thresh_LM, thresh_MH = _run_wp2_safe(cfg, seed, ctx)

            # 2) Replace regime_hat with chosen detector
            df_exog = _apply_detector(
                detector, df_exog, thresh_LM, thresh_MH, warmup,
            )

            # 3) Train-test split
            exog_train = df_exog.iloc[: n_train + 1].reset_index(drop=True)
            exog_test = df_exog.iloc[n_train : n_train + n_test + 1].reset_index(drop=True)

            # 4) Train PPO-aware and PPO-blind
            for stage, use_regime in [("ppo_aware", True), ("ppo_blind", False)]:
                ctx.logger.info(f"  Training {stage}")
                cfg_tr = copy.deepcopy(cfg)
                cfg_tr["wp3"]["use_regime"] = use_regime
                cfg_tr["episode"] = {**cfg_tr["episode"], "n_steps": n_train}
                cfg_tr["as"]["horizon_steps"] = n_train

                env_tr = MMEnv(cfg_tr)
                env_tr.reset(seed=seed, options={"exog": exog_train})
                vec_env = DummyVecEnv([lambda _e=Monitor(env_tr): _e])

                wp4 = cfg["wp4"]
                device = cfg.get("wp4", {}).get("device", "cpu")
                model = PPO(
                    "MlpPolicy", vec_env, seed=seed,
                    learning_rate=float(wp4["learning_rate"]),
                    n_steps=int(wp4["n_steps"]),
                    batch_size=int(wp4["batch_size"]),
                    n_epochs=int(wp4["n_epochs"]),
                    gamma=float(wp4["gamma"]),
                    gae_lambda=float(wp4["gae_lambda"]),
                    clip_range=float(wp4["clip_range"]),
                    ent_coef=float(wp4["ent_coef"]),
                    verbose=0,
                    device=device,
                )
                model.learn(total_timesteps=int(wp4["total_timesteps"]))

                model_dir = out_dir / "models" / detector / f"seed{seed}"
                model_dir.mkdir(parents=True, exist_ok=True)
                model.save(str(model_dir / stage))
                vec_env.close()

                # 5) OOS evaluation
                cfg_ev = copy.deepcopy(cfg)
                cfg_ev["episode"] = {**cfg_ev["episode"], "n_steps": n_test}
                cfg_ev["wp3"]["use_regime"] = use_regime
                cfg_ev["as"]["horizon_steps"] = n_test

                env_ev = MMEnv(cfg_ev)
                obs, _ = env_ev.reset(seed=seed, options={"exog": exog_test})

                eq = np.zeros(n_test + 1)
                iv = np.zeros(n_test + 1, dtype=int)
                fl = np.zeros(n_test, dtype=int)
                eq[0] = env_ev._state.equity
                iv[0] = env_ev._state.inv

                for t in range(n_test):
                    action, _ = model.predict(obs, deterministic=True)
                    obs, _r, _term, _trunc, info = env_ev.step(action)
                    eq[t + 1] = info["equity"]
                    iv[t + 1] = info["inv"]
                    fl[t] = info["fills"]

                m = compute_metrics(eq, iv, fl, dt=dt)
                rows.append({
                    "detector": detector,
                    "seed": seed,
                    "strategy": stage,
                    "sharpe_like": m["sharpe_like"],
                    "final_equity": m["final_equity"],
                })

                ctx.logger.info(
                    f"  {stage}: equity={m['final_equity']:.4f} "
                    f"sharpe={m['sharpe_like']:.4f}"
                )

    # 6) Save results
    df_out = pd.DataFrame(rows)
    df_out.to_csv(out_dir / "metrics_detector_pilot.csv", index=False)

    # 7) Summary table: mean sharpe_like by detector x strategy
    summary = (
        df_out.groupby(["detector", "strategy"])["sharpe_like"]
        .mean()
        .reset_index()
        .rename(columns={"sharpe_like": "mean_sharpe_like"})
    )
    print("\n=== Detector Pilot Summary (mean sharpe_like) ===")
    print(summary.to_string(index=False))

    ctx.logger.info("WP5 detector comparison complete.")

# ============================================================
# FILE: src/wp5/analyze_actions.py
# PURPOSE: WP5 post-hoc action analysis and action-plot generation.
# STATUS: support
# ============================================================

"""WP5 action analysis: h/m distributions by regime across strategies."""
# Eylem Dağılımı Analizi (WP5)
# ------------------------------
# PPO ajanlarının rejim bazında half-spread (h) ve skew (m) dağılımlarını analiz eder.
# Cross-seed standart sapma kullanır (seed-level aggregation).
# Tez Şekil 3'ü (Action Distribution by Regime) üretir.

from pathlib import Path
import re
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt


def _find_latest_eval_run() -> Path:
    runs = list(Path("results/runs").glob("*wp5-eval*"))
    if not runs:
        raise FileNotFoundError("No wp5-eval run found under results/runs/")
    return max(runs, key=lambda p: p.stat().st_mtime)


def _load_curves(run_dir: Path) -> pd.DataFrame:
    """Load all *_test.csv from curves/, parse strategy & seed from filename."""
    curves_dir = run_dir / "curves"
    pattern = re.compile(r"seed(\d+)_(.+)_test\.csv$")
    frames = []
    for csv_path in sorted(curves_dir.glob("*_test.csv")):
        m = pattern.search(csv_path.name)
        if m is None:
            continue
        seed, strategy = int(m.group(1)), m.group(2)
        df = pd.read_csv(csv_path)
        df["seed"] = seed
        df["strategy"] = strategy
        frames.append(df)
    if not frames:
        raise FileNotFoundError(f"No *_test.csv files in {curves_dir}")
    return pd.concat(frames, ignore_index=True)


def _clean(df: pd.DataFrame) -> pd.DataFrame:
    """Drop warmup rows and rows with NaN h/m."""
    df = df[df["regime_hat"].isin(["L", "M", "H"])].copy()
    df = df.dropna(subset=["h", "m"])
    return df


def _seed_level_stats(df: pd.DataFrame) -> pd.DataFrame:
    g = df.groupby(["seed", "strategy", "regime_hat"], as_index=False)
    out = g.agg(
        mean_h=("h", "mean"),
        mean_m=("m", "mean"),
        ph5=("h", lambda x: float((x == 5).mean())),
    )
    return out


# ── plots ────────────────────────────────────────────────────────────────

REGIMES = ["L", "M", "H"]
REGIME_COLORS = {"L": "#4CAF50", "M": "#FF9800", "H": "#F44336"}


STRATEGIES = ["AS", "naive", "ppo_aware", "ppo_blind"]


def _grouped_bar(ax, agg: pd.DataFrame, val_col: str, err_col: str, title: str):
    strategies = STRATEGIES
    x = np.arange(len(strategies))
    width = 0.25
    for i, reg in enumerate(REGIMES):
        sub = agg[agg["regime_hat"] == reg]
        # align on strategy order
        vals = [sub.loc[sub["strategy"] == s, val_col].values[0] for s in strategies]
        errs = [sub.loc[sub["strategy"] == s, err_col].values[0] for s in strategies]
        ax.bar(x + i * width, vals, width, yerr=errs, label=reg,
               color=REGIME_COLORS[reg], capsize=3)
    ax.set_xticks(x + width)
    ax.set_xticklabels(strategies)
    ax.set_title(title)
    ax.legend(title="Regime")
    ax.grid(axis="y", alpha=0.3)


def plot_h_by_regime(df: pd.DataFrame, out: Path):
    df_seed = _seed_level_stats(df)
    agg = (df_seed.groupby(["strategy", "regime_hat"])["mean_h"]
             .agg(mean_h="mean", std_h="std").reset_index())
    fig, ax = plt.subplots(figsize=(7, 4))
    _grouped_bar(ax, agg, "mean_h", "std_h", "Mean Half-Spread by Regime")
    ax.set_ylabel("h (ticks)")
    fig.tight_layout()
    fig.savefig(out, dpi=150)
    plt.close(fig)
    print(f"  saved {out}")


def plot_m_by_regime(df: pd.DataFrame, out: Path):
    df_seed = _seed_level_stats(df)
    agg = (df_seed.groupby(["strategy", "regime_hat"])["mean_m"]
             .agg(mean_m="mean", std_m="std").reset_index())
    fig, ax = plt.subplots(figsize=(7, 4))
    _grouped_bar(ax, agg, "mean_m", "std_m", "Mean Skew by Regime")
    ax.set_ylabel("m (ticks)")
    fig.tight_layout()
    fig.savefig(out, dpi=150)
    plt.close(fig)
    print(f"  saved {out}")


def plot_ph5_by_regime(df: pd.DataFrame, out: Path):
    df_seed = _seed_level_stats(df)
    agg = (df_seed.groupby(["strategy", "regime_hat"])["ph5"]
             .agg(ph5="mean", std_ph5="std").reset_index())
    x = np.arange(len(STRATEGIES))
    width = 0.25
    fig, ax = plt.subplots(figsize=(7, 4))
    for i, reg in enumerate(REGIMES):
        sub = agg[agg["regime_hat"] == reg]
        vals = [sub.loc[sub["strategy"] == s, "ph5"].values[0] for s in STRATEGIES]
        errs = [sub.loc[sub["strategy"] == s, "std_ph5"].values[0] for s in STRATEGIES]
        ax.bar(x + i * width, vals, width, yerr=errs, label=reg,
               color=REGIME_COLORS[reg], capsize=3)
    ax.set_xticks(x + width)
    ax.set_xticklabels(STRATEGIES)
    ax.set_ylabel("P(h=5)")
    ax.set_title("P(h=5 | Regime) — Undertrading Indicator")
    ax.legend(title="Regime")
    ax.grid(axis="y", alpha=0.3)
    fig.tight_layout()
    fig.savefig(out, dpi=150)
    plt.close(fig)
    print(f"  saved {out}")


def print_table(df: pd.DataFrame):
    df_seed = _seed_level_stats(df)
    rows = []
    for (strat, reg), g in df_seed.groupby(["strategy", "regime_hat"]):
        rows.append({
            "strategy": strat,
            "regime": reg,
            "mean_h": g["mean_h"].mean(),
            "std_h": g["mean_h"].std(),
            "mean_m": g["mean_m"].mean(),
            "std_m": g["mean_m"].std(),
            "P(h=5)": g["ph5"].mean(),
            "std_ph5": g["ph5"].std(),
        })
    tbl = pd.DataFrame(rows)
    fmt = {"mean_h": "{:.2f}", "std_h": "{:.2f}",
           "mean_m": "{:.2f}", "std_m": "{:.2f}", "P(h=5)": "{:.3f}",
           "std_ph5": "{:.3f}"}
    print("\n" + tbl.to_string(index=False, formatters={
        k: v.format for k, v in fmt.items()
    }))


# ── main ─────────────────────────────────────────────────────────────────

def main():
    run_dir = _find_latest_eval_run()
    print(f"Run dir: {run_dir}")

    df = _load_curves(run_dir)
    print(f"Loaded {len(df)} rows, strategies: {sorted(df['strategy'].unique())}")

    df = _clean(df)
    print(f"After clean: {len(df)} rows")

    plots_dir = run_dir / "plots"
    plots_dir.mkdir(exist_ok=True)

    plot_h_by_regime(df, plots_dir / "action_h_by_regime.png")
    plot_m_by_regime(df, plots_dir / "action_m_by_regime.png")
    plot_ph5_by_regime(df, plots_dir / "ph5_by_regime.png")
    print_table(df)


if __name__ == "__main__":
    main()

# ============================================================
# FILE: src/wp5/stats_detector_robustness.py
# PURPOSE: WP5 statistical summary script for detector-robustness tests and ANOVA framing.
# STATUS: audit
# ============================================================

"""Statistical analysis of detector robustness experiment results."""
# İstatistiksel Testler — Detector Robustness (WP5)
# ---------------------------------------------------
# 120 modelin (3 dedektör × 20 seed × 2 strateji) istatistiksel analizini yapar.
# - Paired t-test: ppo_aware vs ppo_blind (her dedektör için ayrı)
# - One-way ANOVA: dedektör seçiminin ppo_aware'e etkisi (destekleyici
#   robustness check; primary inferential machinery per-detector
#   paired t-test'lerdir)
# - Sonuçlar: stats_detector_robustness.txt
#
# Design note (Lane C audit follow-up):
# The ANOVA below is one-way (between-subjects). The actual design is
# repeated-measures: the same 20 seeds are shared across all three
# detectors, so each seed contributes a matched triple of observations
# (hmm, rv_baseline, rv_dwell) drawn from the same underlying market
# path. The RM-appropriate non-parametric analog is the Friedman test
# (scipy.stats.friedmanchisquare); the parametric analog is
# repeated-measures ANOVA.
#
# The one-way ANOVA is retained here because the observed F-statistic
# (F = 0.003, p = 0.997) is so far from any rejection threshold that
# the qualitative conclusion (detector choice does not meaningfully
# affect ppo_aware sharpe) is robust to test choice — an RM analysis
# would yield a different p-value but the same null conclusion. The
# per-detector paired t-tests (which ARE correctly matched on seed)
# carry the primary inferential weight in §4.6.
#
# Lane D follow-up: implement Friedman / RM-ANOVA as a supplementary
# check and report alongside the one-way result. Not implemented here
# to stay within Lane C scope.

from __future__ import annotations

import sys
from pathlib import Path

import numpy as np
import pandas as pd
from scipy import stats


def find_results_csv() -> Path:
    """Find the most recent detector-full run's metrics CSV."""
    runs_dir = Path("results/runs")
    candidates = sorted(runs_dir.glob("*detector-full*"), key=lambda p: p.name)
    if not candidates:
        print("ERROR: No detector-full run found under results/runs/")
        sys.exit(1)
    run_dir = candidates[-1]
    csv_path = run_dir / "metrics_detector_pilot.csv"
    if not csv_path.exists():
        csv_path = run_dir / "metrics_detector_compare.csv"
    if not csv_path.exists():
        print(f"ERROR: No metrics CSV in {run_dir}")
        sys.exit(1)
    return csv_path


def paired_tests(df: pd.DataFrame) -> pd.DataFrame:
    """Paired t-test: ppo_aware vs ppo_blind for each detector x metric."""
    rows = []
    for detector in sorted(df["detector"].unique()):
        sub = df[df["detector"] == detector]
        aware = sub[sub["strategy"] == "ppo_aware"].sort_values("seed")
        blind = sub[sub["strategy"] == "ppo_blind"].sort_values("seed")

        for metric in ["sharpe_like", "final_equity"]:
            a = aware[metric].values
            b = blind[metric].values
            diff = a - b
            t_stat, p_val = stats.ttest_rel(a, b)
            rows.append({
                "detector": detector,
                "metric": metric,
                "mean_aware": float(np.mean(a)),
                "mean_blind": float(np.mean(b)),
                "mean_diff": float(np.mean(diff)),
                "t_stat": float(t_stat),
                "p_value": float(p_val),
                "significant": "Yes" if p_val < 0.05 else "No",
            })
    return pd.DataFrame(rows)


def anova_across_detectors(df: pd.DataFrame) -> dict:
    """One-way ANOVA: does ppo_aware sharpe differ across detectors?"""
    aware = df[df["strategy"] == "ppo_aware"]
    groups = [
        g["sharpe_like"].values
        for _, g in aware.groupby("detector")
    ]
    f_stat, p_val = stats.f_oneway(*groups)
    return {
        "test": "one-way ANOVA (ppo_aware sharpe_like across detectors)",
        "F_stat": float(f_stat),
        "p_value": float(p_val),
        "significant": "Yes" if p_val < 0.05 else "No",
    }


def main():
    csv_path = find_results_csv()
    print(f"Reading: {csv_path}\n")

    df = pd.read_csv(csv_path)
    print(f"Rows: {len(df)}, Detectors: {sorted(df['detector'].unique())}")
    print(f"Seeds: {sorted(df['seed'].unique())}\n")

    # --- Paired t-tests ---
    results = paired_tests(df)

    fmt = {
        "mean_aware": "{:.4f}",
        "mean_blind": "{:.4f}",
        "mean_diff": "{:.4f}",
        "t_stat": "{:.4f}",
        "p_value": "{:.4f}",
    }
    table_str = results.to_string(
        index=False,
        formatters={k: v.format for k, v in fmt.items()},
    )

    print("=" * 90)
    print("PAIRED T-TEST: ppo_aware vs ppo_blind (per detector, per metric)")
    print("=" * 90)
    print(table_str)

    # --- ANOVA ---
    anova = anova_across_detectors(df)
    anova_str = (
        f"\n{'=' * 90}\n"
        f"ONE-WAY ANOVA: ppo_aware sharpe_like across detectors\n"
        f"{'=' * 90}\n"
        f"F-stat: {anova['F_stat']:.4f}, p-value: {anova['p_value']:.4f}, "
        f"significant: {anova['significant']}\n"
    )
    print(anova_str)

    # --- Summary ---
    summary_lines = [
        "\nSUMMARY",
        "-" * 40,
    ]
    sig_count = int((results["significant"] == "Yes").sum())
    summary_lines.append(
        f"Significant tests (p<0.05): {sig_count} / {len(results)}"
    )
    if sig_count == 0:
        summary_lines.append(
            "Null result: no detector shows significant aware vs blind difference."
        )
    else:
        sig_rows = results[results["significant"] == "Yes"]
        for _, r in sig_rows.iterrows():
            direction = "aware > blind" if r["mean_diff"] > 0 else "blind > aware"
            summary_lines.append(
                f"  {r['detector']} / {r['metric']}: p={r['p_value']:.4f} ({direction})"
            )
    summary_str = "\n".join(summary_lines)
    print(summary_str)

    # --- Save to file ---
    out_path = Path("results/stats_detector_robustness.txt")
    out_path.parent.mkdir(parents=True, exist_ok=True)
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(f"Source: {csv_path}\n\n")
        f.write("PAIRED T-TEST: ppo_aware vs ppo_blind (per detector, per metric)\n")
        f.write(table_str + "\n")
        f.write(anova_str + "\n")
        f.write(summary_str + "\n")
    print(f"\nSaved to {out_path}")


if __name__ == "__main__":
    main()

# ============================================================
# FILE: src/wp5/figure_thesis.py
# PURPOSE: Active thesis figure generator for the original WP5 figure set.
# STATUS: support
# ============================================================

"""Generate thesis figures from WP5 OOS results."""
# Thesis figure script — main figure suite (Fig 1–5).
# ---------------------------------------------------
# Ownership scope: this script produces and owns the following
# figures embedded in thesis_28 (sha256-verified):
#   - fig1_sharpe_inv.png           → §4.1
#   - fig2_paired_seed.png          → §4.1
#   - fig3_regime_sharpe.png        → §4.5
#   - fig4_detector_robustness.png  → §4.6
#   - fig5_action_analysis.png      → §4.4
# Output directory: results/plots/thesis/
#
# Companion script: src/wp5/figure_thesis_23.py owns the
# supplementary thesis figure suite (Fig 6–9: pure ablation,
# oracle paired seed, eta-regime, model misspecification). The
# "_23" suffix in that filename is preserved for appendix/
# file-index stability across gen_thesis_{23..28}.py — renaming
# would invalidate the appendix references.
#
# Out of scope: Chapter 5 (WP6) figures are produced by the WP6
# sweep toolchain in docs/internal/wp6_sweep_full/plots/, NOT
# by this script.

import argparse
from pathlib import Path
import re
import numpy as np
import pandas as pd
import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt

# ── paths ────────────────────────────────────────────────────────────────

DEFAULT_MAIN_RUN = "results/runs/20260228-093733_seed1_wp5-eval-main_3e8eacc"
DEFAULT_DETECTOR_CSV = (
    "results/runs/20260316-223842_seed1_wp5-detector-full_a67e381"
    "/metrics_detector_pilot.csv"
)
DEFAULT_OUT_DIR = "results/plots/thesis"

MAIN_RUN = Path(DEFAULT_MAIN_RUN)
DETECTOR_CSV = Path(DEFAULT_DETECTOR_CSV)
OOS_CSV = MAIN_RUN / "metrics_wp5_oos.csv"
REGIME_CSV = MAIN_RUN / "metrics_wp5_oos_by_regime.csv"
CURVES_DIR = MAIN_RUN / "curves"
OUT_DIR = Path(DEFAULT_OUT_DIR)

# ── style ────────────────────────────────────────────────────────────────

COLORS = {
    "AS": "#888888",
    "naive": "#5B9BD5",
    "ppo_blind": "#ED7D31",
    "ppo_aware": "#2E75B6",
}
STRAT_ORDER = ["AS", "naive", "ppo_blind", "ppo_aware"]
STRAT_LABELS = {"AS": "AS", "naive": "Naive", "ppo_blind": "PPO-blind", "ppo_aware": "PPO-aware"}
REGIMES = ["L", "M", "H"]

plt.rcParams.update({
    "font.size": 11,
    "figure.dpi": 150,
    "figure.autolayout": True,
})


def _parse_args(argv=None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--main-run",
        type=Path,
        default=DEFAULT_MAIN_RUN,
        help="WP5 main evaluation run directory.",
    )
    parser.add_argument(
        "--detector-csv",
        type=Path,
        default=DEFAULT_DETECTOR_CSV,
        help="Detector robustness metrics CSV.",
    )
    parser.add_argument(
        "--out-dir",
        type=Path,
        default=DEFAULT_OUT_DIR,
        help="Directory for generated thesis figures.",
    )
    return parser.parse_args(argv)


def _configure_paths(args: argparse.Namespace) -> None:
    global MAIN_RUN, DETECTOR_CSV, OOS_CSV, REGIME_CSV, CURVES_DIR, OUT_DIR

    MAIN_RUN = Path(args.main_run)
    DETECTOR_CSV = Path(args.detector_csv)
    OOS_CSV = MAIN_RUN / "metrics_wp5_oos.csv"
    REGIME_CSV = MAIN_RUN / "metrics_wp5_oos_by_regime.csv"
    CURVES_DIR = MAIN_RUN / "curves"
    OUT_DIR = Path(args.out_dir)


# ── helpers ──────────────────────────────────────────────────────────────

def _bar_labels(ax, bars, fmt=".2f"):
    for bar in bars:
        h = bar.get_height()
        ax.text(bar.get_x() + bar.get_width() / 2, h,
                f"{h:{fmt}}", ha="center", va="bottom", fontsize=9)


def _load_curves() -> pd.DataFrame:
    pattern = re.compile(r"seed(\d+)_(.+)_test\.csv$")
    frames = []
    for csv_path in sorted(CURVES_DIR.glob("*_test.csv")):
        m = pattern.search(csv_path.name)
        if m is None:
            continue
        seed, strategy = int(m.group(1)), m.group(2)
        df = pd.read_csv(csv_path)
        df["seed"] = seed
        df["strategy"] = strategy
        frames.append(df)
    return pd.concat(frames, ignore_index=True)


# ── figure 1 ─────────────────────────────────────────────────────────────

def fig1_sharpe_inv(oos: pd.DataFrame):
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 4))

    labels = [STRAT_LABELS[s] for s in STRAT_ORDER]

    for ax, col, ylabel, panel_label in [
        (ax1, "sharpe_like", "Sharpe", "Sharpe"),
        (ax2, "inv_p99", "inv_p99 (lots)", "inv_p99"),
    ]:
        means, stds = [], []
        for s in STRAT_ORDER:
            sub = oos[oos["strategy"] == s][col]
            means.append(sub.mean())
            stds.append(sub.std())
        x = np.arange(len(STRAT_ORDER))
        bars = ax.bar(x, means, yerr=stds, capsize=4,
                      color=[COLORS[s] for s in STRAT_ORDER])
        ax.set_xticks(x)
        ax.set_xticklabels(labels, fontsize=9)
        ax.set_ylim(bottom=0)
        ax.set_ylabel(ylabel)
        ax.set_title(panel_label, fontsize=11, fontweight="bold")
        ax.grid(axis="y", alpha=0.3)
        _bar_labels(ax, bars)

    out = OUT_DIR / "fig1_sharpe_inv.png"
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)
    print(f"  saved {out}")


# ── figure 2 ─────────────────────────────────────────────────────────────

def fig2_paired_seed(oos: pd.DataFrame):
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 4))

    aware = oos[oos["strategy"] == "ppo_aware"].set_index("seed")
    blind = oos[oos["strategy"] == "ppo_blind"].set_index("seed")
    seeds = sorted(set(aware.index) & set(blind.index))

    for ax, col, panel_label, pval in [
        (ax1, "sharpe_like",
         "Sharpe", 0.261),
        (ax2, "final_equity",
         "Final equity", 0.023),
    ]:
        x_vals = blind.loc[seeds, col].values
        y_vals = aware.loc[seeds, col].values
        ax.scatter(x_vals, y_vals, c=COLORS["ppo_aware"], edgecolors="k",
                   linewidths=0.5, s=40, zorder=3)
        lo = min(x_vals.min(), y_vals.min()) * 0.95
        hi = max(x_vals.max(), y_vals.max()) * 1.05
        ax.plot([lo, hi], [lo, hi], "--", color="grey", linewidth=1, zorder=1)
        ax.set_xlim(lo, hi)
        ax.set_ylim(lo, hi)
        ax.set_aspect("equal", adjustable="box")
        xlabel = "PPO-blind Sharpe" if col == "sharpe_like" else "PPO-blind Equity"
        ylabel = "PPO-aware Sharpe" if col == "sharpe_like" else "PPO-aware Equity"
        ax.set_xlabel(xlabel)
        ax.set_ylabel(ylabel)
        ax.set_title(panel_label, fontsize=11, fontweight="bold")
        ax.text(0.05, 0.05, f"p = {pval}", transform=ax.transAxes,
                fontsize=10, verticalalignment="bottom")
        ax.grid(alpha=0.3)

    out = OUT_DIR / "fig2_paired_seed.png"
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)
    print(f"  saved {out}")


# ── figure 3 ─────────────────────────────────────────────────────────────

def fig3_regime_sharpe(regime: pd.DataFrame):
    fig, ax = plt.subplots(figsize=(10, 4))
    x = np.arange(len(REGIMES))
    n = len(STRAT_ORDER)
    width = 0.18

    for i, s in enumerate(STRAT_ORDER):
        means, stds = [], []
        for r in REGIMES:
            sub = regime[(regime["strategy"] == s) & (regime["regime"] == r)]
            means.append(sub["sharpe_like"].mean())
            stds.append(sub["sharpe_like"].std())
        ax.bar(x + i * width, means, width, yerr=stds, capsize=3,
               color=COLORS[s], label=STRAT_LABELS[s])

    ax.set_xticks(x + width * (n - 1) / 2)
    ax.set_xticklabels(REGIMES)
    ax.set_xlabel("Volatility Regime")
    ax.set_ylabel("Mean Sharpe")
    ax.set_ylim(bottom=0)
    ax.legend(loc="upper right")
    ax.grid(axis="y", alpha=0.3)

    out = OUT_DIR / "fig3_regime_sharpe.png"
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)
    print(f"  saved {out}")


# ── figure 4 ─────────────────────────────────────────────────────────────

def fig4_detector_robustness(det: pd.DataFrame):
    detectors = ["rv_baseline", "rv_dwell", "hmm"]
    fig, ax = plt.subplots(figsize=(8, 5))

    x = np.arange(len(detectors))
    for i, d in enumerate(detectors):
        sub = det[det["detector"] == d]
        aware_vals = sub[sub["strategy"] == "ppo_aware"]["sharpe_like"]
        blind_vals = sub[sub["strategy"] == "ppo_blind"]["sharpe_like"]
        aware_mean, aware_std = aware_vals.mean(), aware_vals.std()
        blind_mean, blind_std = blind_vals.mean(), blind_vals.std()

        # vertical line
        ax.plot([i, i], [blind_mean, aware_mean], color="grey", linewidth=1.5,
                zorder=1)
        # aware dot with error bar
        ax.errorbar(i, aware_mean, yerr=aware_std, fmt="o",
                    color=COLORS["ppo_aware"], markersize=10, capsize=4,
                    zorder=3, label="PPO-aware" if i == 0 else "")
        ax.text(i + 0.10, aware_mean, f"{aware_mean:.3f}", va="center", fontsize=9)
        # blind dot with error bar
        ax.errorbar(i, blind_mean, yerr=blind_std, fmt="s",
                    color=COLORS["ppo_blind"], markersize=9, capsize=4,
                    zorder=3, label="PPO-blind" if i == 0 else "")
        ax.text(i + 0.10, blind_mean, f"{blind_mean:.3f}", va="center", fontsize=9)

    ax.set_xticks(x)
    ax.set_xticklabels(["rv_baseline", "rv_dwell", "HMM"])
    ax.set_ylim(0.60, 0.80)
    ax.set_ylabel("Mean Sharpe")
    ax.legend(loc="upper right")
    ax.grid(axis="y", alpha=0.3)

    out = OUT_DIR / "fig4_detector_robustness.png"
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)
    print(f"  saved {out}")


# ── figure 5 ─────────────────────────────────────────────────────────────

def fig5_action_analysis(curves: pd.DataFrame):
    curves = curves[curves["regime_hat"].isin(REGIMES)].dropna(subset=["h", "m"])

    # per-seed mean h/m by strategy+regime
    seed_agg = (curves.groupby(["seed", "strategy", "regime_hat"])
                .agg(mean_h=("h", "mean"), mean_m=("m", "mean"))
                .reset_index())

    fig, axes = plt.subplots(2, 2, figsize=(10, 6))

    strat_colors = {"ppo_aware": "#2E75B6", "ppo_blind": "#ED7D31"}

    configs = [
        (axes[0, 0], "ppo_aware", "mean_h", "PPO-aware", "Mean h", False),
        (axes[0, 1], "ppo_blind", "mean_h", "PPO-blind", "Mean h", False),
        (axes[1, 0], "ppo_aware", "mean_m", "PPO-aware", "Mean m", True),
        (axes[1, 1], "ppo_blind", "mean_m", "PPO-blind", "Mean m", True),
    ]

    for ax, strat, col, panel_label, ylabel, show_zero in configs:
        sub = seed_agg[seed_agg["strategy"] == strat]
        x = np.arange(len(REGIMES))
        means, stds = [], []
        for r in REGIMES:
            vals = sub[sub["regime_hat"] == r][col]
            means.append(vals.mean())
            stds.append(vals.std())
        bars = ax.bar(x, means, yerr=stds, capsize=4,
                      color=strat_colors[strat])
        ax.set_xticks(x)
        ax.set_xticklabels(REGIMES)
        ax.set_title(panel_label, fontsize=11, fontweight="bold")
        ax.set_ylabel(ylabel)
        ax.grid(axis="y", alpha=0.3)
        if show_zero:
            ax.axhline(0, color="black", linewidth=0.8, linestyle="-")

    # sync y-axis limits: h subplots 0–3.0, m subplots -0.3–0.3
    for ax in axes[0]:
        ax.set_ylim(0, 3.0)
    for ax in axes[1]:
        ax.set_ylim(-0.3, 0.3)

    out = OUT_DIR / "fig5_action_analysis.png"
    fig.savefig(out)
    plt.close(fig)
    print(f"  saved {out}")


# ── main ─────────────────────────────────────────────────────────────────

def main(argv=None):
    args = _parse_args(argv)
    _configure_paths(args)

    OUT_DIR.mkdir(parents=True, exist_ok=True)

    oos = pd.read_csv(OOS_CSV)
    oos = oos[oos["split"] == "test"]
    regime = pd.read_csv(REGIME_CSV)
    det = pd.read_csv(DETECTOR_CSV)
    curves = _load_curves()

    print(f"OOS rows: {len(oos)}, Regime rows: {len(regime)}, "
          f"Detector rows: {len(det)}, Curve rows: {len(curves)}")

    fig1_sharpe_inv(oos)
    fig2_paired_seed(oos)
    fig3_regime_sharpe(regime)
    fig4_detector_robustness(det)
    fig5_action_analysis(curves)

    print("\nAll figures saved to", OUT_DIR)


if __name__ == "__main__":
    main()

# ============================================================
# FILE: src/wp5/figure_thesis_23.py
# PURPOSE: Active thesis figure generator for five-variant, eta-regime, and misspecification figures.
# STATUS: support
# ============================================================

"""Generate thesis_23 extension figures from existing WP5 result files."""
# Thesis figure script — supplementary thesis figure suite (Fig 6–9).
# -------------------------------------------------------------------
# Ownership scope: this script produces and owns the following
# figures embedded in thesis_28 (sha256-verified):
#   - fig6_ablation_summary.png     → §4.2 (pure ablation / signal redundancy)
#   - fig7_oracle_paired_seed.png   → §4.3 (oracle vs PPO_aware paired-seed)
#   - fig8_eta_regime_summary.png   → §4.7 (regime-conditional eta)
#   - fig9_misspec_summary.png      → §4.8 (model misspecification)
# Output directory: results/plots/thesis_23/
#
# Filename rationale: the "_23" suffix is preserved for
# appendix/file-index stability across gen_thesis_{23..28}.py —
# renaming would invalidate appendix references chained through
# those generator scripts.
#
# Companion script: src/wp5/figure_thesis.py owns the main
# thesis figure suite (Fig 1–5).
#
# Out of scope: Chapter 5 (WP6) figures are produced by the WP6
# sweep toolchain in docs/internal/wp6_sweep_full/plots/, NOT
# by this script.

import argparse
from pathlib import Path

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from scipy import stats


DEFAULT_ABLATION_CSV = (
    "results/runs/20260327-171914_seed1_wp5-ablation_e1545a5"
    "/metrics_wp5_oos_combined.csv"
)
DEFAULT_ETA_CSV = (
    "results/runs/20260330-155235_seed42_w5-eta-regime_af82a9f"
    "/metrics_wp5_oos.csv"
)
DEFAULT_MISSPEC_CSV = (
    "results/runs/20260408-160248_seed1_w5-misspec-mild_5d9dc23"
    "/metrics_wp5_oos.csv"
)
DEFAULT_OUT_DIR = "results/plots/thesis_23"

ABLATION_CSV = Path(DEFAULT_ABLATION_CSV)
ETA_CSV = Path(DEFAULT_ETA_CSV)
MISSPEC_CSV = Path(DEFAULT_MISSPEC_CSV)
OUT_DIR = Path(DEFAULT_OUT_DIR)

VARIANT_ORDER = [
    "ppo_sigma_only",
    "ppo_oracle_full",
    "ppo_regime_only",
    "ppo_combined",
    "ppo_oracle_pure",
]
SHORT_LABELS = {
    "ppo_sigma_only": "sigma_only",
    "ppo_oracle_full": "oracle_full",
    "ppo_regime_only": "regime_only",
    "ppo_combined": "combined",
    "ppo_oracle_pure": "oracle_pure",
}
COLORS = {
    "ppo_sigma_only": "#2E75B6",
    "ppo_oracle_full": "#70AD47",
    "ppo_regime_only": "#A5A5A5",
    "ppo_combined": "#ED7D31",
    "ppo_oracle_pure": "#8064A2",
}

plt.rcParams.update({
    "font.size": 10,
    "figure.dpi": 150,
    "figure.autolayout": True,
})


def _parse_args(argv=None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--ablation-csv",
        type=Path,
        default=DEFAULT_ABLATION_CSV,
        help="WP5 ablation combined OOS metrics CSV.",
    )
    parser.add_argument(
        "--eta-csv",
        type=Path,
        default=DEFAULT_ETA_CSV,
        help="WP5 eta-regime OOS metrics CSV.",
    )
    parser.add_argument(
        "--misspec-csv",
        type=Path,
        default=DEFAULT_MISSPEC_CSV,
        help="WP5 mild misspecification OOS metrics CSV.",
    )
    parser.add_argument(
        "--out-dir",
        type=Path,
        default=DEFAULT_OUT_DIR,
        help="Directory for generated thesis_23 extension figures.",
    )
    return parser.parse_args(argv)


def _configure_paths(args: argparse.Namespace) -> None:
    global ABLATION_CSV, ETA_CSV, MISSPEC_CSV, OUT_DIR

    ABLATION_CSV = Path(args.ablation_csv)
    ETA_CSV = Path(args.eta_csv)
    MISSPEC_CSV = Path(args.misspec_csv)
    OUT_DIR = Path(args.out_dir)


def _load_test(path: Path) -> pd.DataFrame:
    df = pd.read_csv(path)
    if "split" in df.columns:
        df = df[df["split"] == "test"].copy()
    return df


def _paired_t(df: pd.DataFrame, a: str, b: str, metric: str) -> float:
    aa = df[df["strategy"] == a].set_index("seed")
    bb = df[df["strategy"] == b].set_index("seed")
    seeds = sorted(set(aa.index) & set(bb.index))
    return stats.ttest_rel(aa.loc[seeds, metric], bb.loc[seeds, metric]).pvalue


def _bar_panel(ax, df: pd.DataFrame, metric: str, ylabel: str):
    means = []
    stds = []
    for strategy in VARIANT_ORDER:
        vals = df[df["strategy"] == strategy][metric]
        means.append(vals.mean())
        stds.append(vals.std())
    x = np.arange(len(VARIANT_ORDER))
    bars = ax.bar(
        x,
        means,
        yerr=stds,
        capsize=4,
        color=[COLORS[s] for s in VARIANT_ORDER],
    )
    ax.set_xticks(x)
    ax.set_xticklabels([SHORT_LABELS[s] for s in VARIANT_ORDER], rotation=25, ha="right")
    ax.set_ylabel(ylabel)
    ax.set_ylim(bottom=0)
    ax.grid(axis="y", alpha=0.3)
    for bar in bars:
        h = bar.get_height()
        ax.text(bar.get_x() + bar.get_width() / 2, h, f"{h:.3f}",
                ha="center", va="bottom", fontsize=8)


def fig6_ablation_summary(ablation: pd.DataFrame):
    fig, axes = plt.subplots(1, 3, figsize=(12, 3.8))
    for ax, metric, ylabel in [
        (axes[0], "sharpe_like", "Sharpe"),
        (axes[1], "final_equity", "Final equity"),
        (axes[2], "inv_p99", "inv_p99"),
    ]:
        _bar_panel(ax, ablation, metric, ylabel)
    out = OUT_DIR / "fig6_ablation_summary.png"
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)
    print(f"  saved {out}")


def _paired_scatter(ax, df: pd.DataFrame, x_strategy: str, y_strategy: str, p_value: float):
    x_df = df[df["strategy"] == x_strategy].set_index("seed")
    y_df = df[df["strategy"] == y_strategy].set_index("seed")
    seeds = sorted(set(x_df.index) & set(y_df.index))
    x_vals = x_df.loc[seeds, "sharpe_like"].values
    y_vals = y_df.loc[seeds, "sharpe_like"].values
    lo = min(x_vals.min(), y_vals.min()) * 0.96
    hi = max(x_vals.max(), y_vals.max()) * 1.04
    ax.scatter(x_vals, y_vals, s=42, c="#2E75B6", edgecolors="black", linewidths=0.5)
    ax.plot([lo, hi], [lo, hi], "--", color="grey", linewidth=1)
    ax.set_xlim(lo, hi)
    ax.set_ylim(lo, hi)
    ax.set_aspect("equal", adjustable="box")
    ax.set_xlabel(f"{SHORT_LABELS[x_strategy]} Sharpe")
    ax.set_ylabel(f"{SHORT_LABELS[y_strategy]} Sharpe")
    ax.text(0.05, 0.05, f"p = {p_value:.3f}", transform=ax.transAxes,
            fontsize=10, va="bottom")
    ax.grid(alpha=0.3)


def fig7_oracle_paired_seed(ablation: pd.DataFrame):
    p_sigma_oracle = _paired_t(
        ablation, "ppo_sigma_only", "ppo_oracle_full", "sharpe_like"
    )
    p_oracle_combined = _paired_t(
        ablation, "ppo_oracle_full", "ppo_combined", "sharpe_like"
    )
    fig, axes = plt.subplots(1, 2, figsize=(9, 4.2))
    _paired_scatter(
        axes[0], ablation, "ppo_sigma_only", "ppo_oracle_full", p_sigma_oracle
    )
    _paired_scatter(
        axes[1], ablation, "ppo_combined", "ppo_oracle_full", p_oracle_combined
    )
    out = OUT_DIR / "fig7_oracle_paired_seed.png"
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)
    print(f"  saved {out}")


def fig8_eta_regime_summary(eta: pd.DataFrame):
    p_sharpe = _paired_t(eta, "ppo_combined", "ppo_sigma_only", "sharpe_like")
    p_equity = _paired_t(eta, "ppo_combined", "ppo_sigma_only", "final_equity")
    fig, axes = plt.subplots(1, 2, figsize=(9, 3.8))
    _bar_panel(axes[0], eta, "sharpe_like", "Sharpe")
    _bar_panel(axes[1], eta, "final_equity", "Final equity")
    axes[0].text(0.03, 0.95, f"combined vs sigma_only: p = {p_sharpe:.4f}",
                 transform=axes[0].transAxes, va="top", fontsize=9)
    axes[1].text(0.03, 0.95, f"combined vs sigma_only: p = {p_equity:.3f}",
                 transform=axes[1].transAxes, va="top", fontsize=9)
    out = OUT_DIR / "fig8_eta_regime_summary.png"
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)
    print(f"  saved {out}")


def fig9_misspec_summary(misspec: pd.DataFrame):
    p_value = _paired_t(misspec, "ppo_sigma_only", "ppo_oracle_full", "sharpe_like")
    fig, ax = plt.subplots(figsize=(7, 4))
    _bar_panel(ax, misspec, "sharpe_like", "Sharpe")
    ax.text(0.03, 0.95, f"sigma_only vs oracle_full: p = {p_value:.3f}",
            transform=ax.transAxes, va="top", fontsize=9)
    out = OUT_DIR / "fig9_misspec_summary.png"
    fig.savefig(out, bbox_inches="tight")
    plt.close(fig)
    print(f"  saved {out}")


def main(argv=None):
    args = _parse_args(argv)
    _configure_paths(args)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ablation = _load_test(ABLATION_CSV)
    eta = _load_test(ETA_CSV)
    misspec = _load_test(MISSPEC_CSV)

    fig6_ablation_summary(ablation)
    fig7_oracle_paired_seed(ablation)
    fig8_eta_regime_summary(eta)
    fig9_misspec_summary(misspec)
    print("\nAll thesis_23 extension figures saved to", OUT_DIR)


if __name__ == "__main__":
    main()

# ============================================================
# FILE: src/wp5_5/__init__.py
# PURPOSE: Package marker for WP5.5 signal-degradation audit and calibration code.
# STATUS: support
# ============================================================



# ============================================================
# FILE: src/wp5_5/signal_degradation.py
# PURPOSE: Shared signal-degradation transforms used by WP5.5 calibration and WP6 sweeps.
# STATUS: active
# ============================================================

"""Signal degradation transformations for WP5.5 audit (pure numpy).

Each transform takes a clean sigma_hat array (1-D, may contain NaN during
warmup) and returns a degraded array of identical shape and dtype.
"""

from __future__ import annotations

import numpy as np


def apply_noise(
    sigma_hat: np.ndarray,
    noise_std: float,
    rng: np.random.Generator,
) -> np.ndarray:
    """Add i.i.d. Gaussian noise (std=noise_std) to non-NaN positions."""
    out = sigma_hat.astype(np.float64, copy=True)
    mask = ~np.isnan(out)
    out[mask] = out[mask] + rng.normal(0.0, noise_std, size=int(mask.sum()))
    return out.astype(sigma_hat.dtype, copy=False)


def apply_lag(sigma_hat: np.ndarray, k_steps: int) -> np.ndarray:
    """Shift array right by k_steps; leading k_steps become NaN."""
    n = len(sigma_hat)
    out = np.empty(n, dtype=np.float64)
    k = max(0, int(k_steps))
    k = min(k, n)
    out[:k] = np.nan
    if n - k > 0:
        out[k:] = sigma_hat[: n - k].astype(np.float64, copy=False)
    return out.astype(sigma_hat.dtype, copy=False)


def apply_coarsen(sigma_hat: np.ndarray, cutpoints: np.ndarray) -> np.ndarray:
    """Discretize sigma_hat into bins defined by cutpoints; return bin midpoints.

    cutpoints: 1-D array of length (n_bins - 1) with interior boundaries.
    Bin midpoints are inferred assuming uniform cutpoint spacing (as produced
    by compute_clean_cutpoints).
    """
    cuts = np.asarray(cutpoints, dtype=np.float64).ravel()
    n_cuts = len(cuts)
    n_bins = n_cuts + 1

    if n_cuts >= 2:
        spacing = float(cuts[1] - cuts[0])
    else:
        # degenerate case — fall back to unit spacing
        spacing = 1.0

    mids = np.empty(n_bins, dtype=np.float64)
    if n_bins == 1:
        mids[0] = float(cuts[0]) if n_cuts else 0.0
    else:
        mids[0] = cuts[0] - spacing / 2.0
        for i in range(1, n_bins - 1):
            mids[i] = cuts[i - 1] + spacing / 2.0
        mids[-1] = cuts[-1] + spacing / 2.0

    out = np.full(sigma_hat.shape, np.nan, dtype=np.float64)
    valid = ~np.isnan(sigma_hat)
    if valid.any():
        idx = np.searchsorted(cuts, sigma_hat[valid].astype(np.float64), side="right")
        idx = np.clip(idx, 0, n_bins - 1)
        out[valid] = mids[idx]
    return out.astype(sigma_hat.dtype, copy=False)


def apply_remove(sigma_hat: np.ndarray, fill_value: float = 0.0) -> np.ndarray:
    """Replace all values (including warmup NaNs) with fill_value."""
    return np.full(sigma_hat.shape, fill_value, dtype=sigma_hat.dtype)


def compute_clean_cutpoints(
    sigma_hat_clean: np.ndarray,
    warmup_end: int,
    n_bins: int = 5,
) -> np.ndarray:
    """Fixed bin cutpoints from the clean warmup distribution.

    Uses np.linspace between min and max of warmup-region non-NaN values to
    build (n_bins + 1) edges, and returns the (n_bins - 1) interior
    boundaries. These cutpoints are reused identically across all conditions.
    """
    warm = sigma_hat_clean[:warmup_end]
    warm = warm[~np.isnan(warm)]
    if len(warm) == 0:
        raise ValueError("compute_clean_cutpoints: no finite warmup values")
    lo = float(np.min(warm))
    hi = float(np.max(warm))
    edges = np.linspace(lo, hi, int(n_bins) + 1)
    return edges[1:-1].astype(np.float64, copy=False)

# ============================================================
# FILE: src/wp5_5/signal_audit.py
# PURPOSE: Signal audit metrics for correlation, classification accuracy, separability, and threshold overlap.
# STATUS: audit
# ============================================================

"""Audit metrics for WP5.5 signal degradation ladder.

All functions operate on the post-warmup region and ignore NaNs / warmup labels.
"""

from __future__ import annotations

import numpy as np
from scipy import stats

from src.wp2.synth_regime import assign_regime_hat


def _drop_invalid(a: np.ndarray, b: np.ndarray) -> tuple[np.ndarray, np.ndarray]:
    a = np.asarray(a, dtype=np.float64)
    b = np.asarray(b, dtype=np.float64)
    keep = ~(np.isnan(a) | np.isnan(b))
    return a[keep], b[keep]


def spearman_correlation(
    clean: np.ndarray,
    degraded: np.ndarray,
    mask: np.ndarray,
) -> float:
    c = np.asarray(clean)[mask]
    d = np.asarray(degraded)[mask]
    c, d = _drop_invalid(c, d)
    if len(c) < 3:
        return 0.0
    if np.all(d == d[0]) or np.all(c == c[0]):
        return 0.0
    r, _ = stats.spearmanr(c, d)
    return 0.0 if np.isnan(r) else float(r)


def pearson_correlation(
    clean: np.ndarray,
    degraded: np.ndarray,
    mask: np.ndarray,
) -> float:
    c = np.asarray(clean)[mask]
    d = np.asarray(degraded)[mask]
    c, d = _drop_invalid(c, d)
    if len(c) < 3:
        return 0.0
    if np.all(d == d[0]) or np.all(c == c[0]):
        return 0.0
    r, _ = stats.pearsonr(c, d)
    return 0.0 if np.isnan(r) else float(r)


def regime_classification_accuracy(
    degraded_sigma: np.ndarray,
    regime_true: list,
    thresh_LM: float,
    thresh_MH: float,
    warmup_end: int,
) -> float:
    """Apply rv_baseline detector to degraded signal with clean thresholds."""
    # assign_regime_hat handles NaN via comparisons (NaN < x is False in numpy),
    # which would route NaNs into the "H" bucket. Replace NaN post-warmup with
    # a sentinel that lands in a deterministic bucket — 0.0 (forces "L"), but
    # the accuracy comparison still counts that as a misclassification when
    # regime_true differs, which is the desired behaviour.
    safe = np.asarray(degraded_sigma, dtype=np.float64).copy()
    nan_mask = np.isnan(safe)
    if nan_mask.any():
        safe[nan_mask] = 0.0
    pred = assign_regime_hat(safe, thresh_LM, thresh_MH, warmup_end)
    total = 0
    correct = 0
    for t in range(warmup_end, len(safe)):
        rt = str(regime_true[t]) if t < len(regime_true) else ""
        rh = pred[t]
        if rt not in ("L", "M", "H"):
            continue
        if rh == "warmup":
            continue
        total += 1
        if rh == rt:
            correct += 1
    return 0.0 if total == 0 else correct / total


def class_separability(
    degraded_sigma: np.ndarray,
    regime_true: list,
    warmup_end: int,
) -> float:
    """Kruskal-Wallis H statistic for degraded sigma grouped by true regime."""
    n = len(degraded_sigma)
    groups: dict[str, list[float]] = {"L": [], "M": [], "H": []}
    for t in range(warmup_end, n):
        val = float(degraded_sigma[t])
        if np.isnan(val):
            continue
        lab = str(regime_true[t]) if t < len(regime_true) else ""
        if lab in groups:
            groups[lab].append(val)

    arrays = [np.asarray(g) for g in groups.values() if len(g) > 0]
    if len(arrays) < 2:
        return 0.0
    flat = np.concatenate(arrays)
    if np.all(flat == flat[0]):
        return 0.0
    try:
        H, _ = stats.kruskal(*arrays)
    except ValueError:
        return 0.0
    return 0.0 if np.isnan(H) else float(H)


def threshold_overlap_rate(
    degraded_sigma: np.ndarray,
    thresh_LM: float,
    thresh_MH: float,
    warmup_end: int,
    band_pct: float = 0.05,
) -> float:
    """Fraction of post-warmup samples within ±band_pct of either threshold."""
    post = np.asarray(degraded_sigma, dtype=np.float64)[warmup_end:]
    post = post[~np.isnan(post)]
    if len(post) == 0:
        return 0.0
    lo_lm = thresh_LM * (1.0 - band_pct)
    hi_lm = thresh_LM * (1.0 + band_pct)
    lo_mh = thresh_MH * (1.0 - band_pct)
    hi_mh = thresh_MH * (1.0 + band_pct)
    in_lm = (post >= lo_lm) & (post <= hi_lm)
    in_mh = (post >= lo_mh) & (post <= hi_mh)
    return float((in_lm | in_mh).mean())


def nrmse(
    clean: np.ndarray,
    degraded: np.ndarray,
    mask: np.ndarray,
) -> float:
    """Normalised RMSE: sqrt(mean((c-d)^2)) / std(c), restricted to mask, NaN-safe."""
    c = np.asarray(clean)[mask]
    d = np.asarray(degraded)[mask]
    c, d = _drop_invalid(c, d)
    if len(c) < 2:
        return 0.0
    sd = float(np.std(c, ddof=0))
    if sd == 0.0:
        return 0.0
    rmse = float(np.sqrt(np.mean((c - d) ** 2)))
    return rmse / sd


def regime_crossing_rate(
    degraded_sigma: np.ndarray,
    thresh_LM: float,
    thresh_MH: float,
    warmup_end: int,
) -> float:
    """Fraction of consecutive post-warmup steps where rv_baseline regime label flips."""
    safe = np.asarray(degraded_sigma, dtype=np.float64).copy()
    nan_mask = np.isnan(safe)
    if nan_mask.any():
        safe[nan_mask] = 0.0
    pred = assign_regime_hat(safe, thresh_LM, thresh_MH, warmup_end)
    n = len(safe)
    transitions = 0
    counted = 0
    prev = None
    for t in range(warmup_end, n):
        cur = pred[t]
        if cur == "warmup":
            prev = None
            continue
        if prev is not None:
            counted += 1
            if cur != prev:
                transitions += 1
        prev = cur
    return 0.0 if counted == 0 else transitions / counted

# ============================================================
# FILE: src/wp5_5/job_w55_audit.py
# PURPOSE: WP5.5 offline signal-degradation audit job; no PPO training.
# STATUS: audit
# ============================================================

"""WP5.5 signal audit job.

Generates a clean synthetic volatility signal, builds four degraded variants
(noisy, lagged, coarsened, none) using cutpoints fixed from the clean warmup
distribution, computes a small audit metric ladder, and writes a summary.

No PPO training here — the audit is purely offline and must finish in <10 min.
"""

from __future__ import annotations

import copy
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd

from src.wp2.synth_regime import (
    REGIME_LABELS,
    calibrate_thresholds,
    compute_rolling_rv,
    generate_mid_series,
    generate_regime_series,
)
from src.wp5_5.signal_audit import (
    class_separability,
    pearson_correlation,
    regime_classification_accuracy,
    spearman_correlation,
    threshold_overlap_rate,
)
from src.wp5_5.signal_degradation import (
    apply_coarsen,
    apply_lag,
    apply_noise,
    apply_remove,
    compute_clean_cutpoints,
)


CONDITION_ORDER = ["clean", "noisy", "lagged", "coarsened", "none"]
METRIC_ORDER = [
    "spearman",
    "pearson",
    "classification_accuracy",
    "separability",
    "threshold_overlap",
]


def _build_clean_signal(cfg: dict, seed: int):
    """Generate clean (mid, sigma_hat, regime_true, thresholds, warmup_end)."""
    rng = np.random.default_rng(seed)
    n_steps = int(cfg["episode"]["n_steps"])
    rv_window = int(cfg["regime"]["rv_window"])
    warmup_end = int(cfg["regime"]["warmup_steps"])

    regime_cfg = copy.deepcopy(cfg.get("regime", {}))
    regime_cfg.setdefault("sigma_mid_ticks_base", float(cfg["market"].get("sigma_mid_ticks", 0.5)))
    regime_cfg.setdefault("sigma_mult", [0.6, 1.0, 1.8])
    cfg_local = {**cfg, "regime": regime_cfg}

    regime_true_int = generate_regime_series(n_steps, seed, cfg=cfg_local, rng=rng)
    mid, _ = generate_mid_series(regime_true_int, cfg_local, rng)
    _, sigma_hat = compute_rolling_rv(mid, rv_window, float(cfg["market"]["tick_size"]))
    thresh_LM, thresh_MH = calibrate_thresholds(sigma_hat, warmup_end)

    regime_true_str = ["M"] + [REGIME_LABELS[int(r)] for r in regime_true_int]
    return sigma_hat, regime_true_str, warmup_end, thresh_LM, thresh_MH


def _plot_ladder(results: dict, plots_dir: Path) -> Path:
    fig, axes = plt.subplots(2, 3, figsize=(14, 8))
    flat = axes.flatten()
    for i, metric in enumerate(METRIC_ORDER):
        ax = flat[i]
        vals = [results[c][metric] for c in CONDITION_ORDER]
        colors = ["seagreen" if c == "clean" else "steelblue" for c in CONDITION_ORDER]
        ax.bar(CONDITION_ORDER, vals, color=colors)
        ax.set_title(metric)
        ax.tick_params(axis="x", rotation=25)
    flat[-1].axis("off")
    fig.suptitle("WP5.5 Signal Audit Ladder", y=1.02)
    plt.tight_layout()
    out = plots_dir / "signal_audit_ladder.png"
    plt.savefig(out, dpi=150, bbox_inches="tight")
    plt.close(fig)
    return out


NONE_UNGATED_METRICS = ("classification_accuracy", "threshold_overlap")


def _evaluate_flags(results: dict) -> tuple[dict, str]:
    clean = results["clean"]
    none = results["none"]

    # (a) Direction: expected monotone change from clean → none.
    # Under the constant-input 'none' condition, classification_accuracy and
    # threshold_overlap are structurally undefined (they reflect fill-value /
    # threshold geometry, not signal information), so they are reported but
    # not gated.
    direction_checks = {
        "spearman": none["spearman"] < clean["spearman"],
        "pearson": none["pearson"] < clean["pearson"],
        "classification_accuracy": (
            none["classification_accuracy"] < clean["classification_accuracy"]
        ),
        "separability": none["separability"] < clean["separability"],
        "threshold_overlap": none["threshold_overlap"] > clean["threshold_overlap"],
    }
    gated_direction_keys = [k for k in direction_checks if k not in NONE_UNGATED_METRICS]
    violations = sum(1 for k in gated_direction_keys if not direction_checks[k])
    flag_direction = violations <= 1

    # (b) Separation: no two conditions identical on every metric (tol 0.01).
    flag_separation = True
    for i in range(len(CONDITION_ORDER)):
        for j in range(i + 1, len(CONDITION_ORDER)):
            a, b = CONDITION_ORDER[i], CONDITION_ORDER[j]
            max_diff = max(abs(results[a][m] - results[b][m]) for m in METRIC_ORDER)
            if max_diff <= 0.01:
                flag_separation = False

    # (c) Coarsened must not "cheat" by beating clean classification.
    flag_coarsen_safety = (
        results["coarsened"]["classification_accuracy"]
        <= clean["classification_accuracy"] + 0.05
    )

    # (d) Monotonicity across the full ladder, for ≥3 of 5 metrics.
    # For the ungated metrics on the 'none' condition (classification_accuracy,
    # threshold_overlap) the 'none' endpoint is excluded from the monotonicity
    # check: those values are structurally undefined under constant input.
    mono_count = 0
    mono_detail = {}
    for metric in METRIC_ORDER:
        if metric in NONE_UNGATED_METRICS:
            vals = [results[c][metric] for c in CONDITION_ORDER if c != "none"]
        else:
            vals = [results[c][metric] for c in CONDITION_ORDER]
        if metric == "threshold_overlap":
            is_mono = all(vals[i] <= vals[i + 1] for i in range(len(vals) - 1))
        else:
            is_mono = all(vals[i] >= vals[i + 1] for i in range(len(vals) - 1))
        mono_detail[metric] = is_mono
        if is_mono:
            mono_count += 1
    flag_monotonicity = mono_count >= 3

    flags = {
        "direction": flag_direction,
        "separation": flag_separation,
        "coarsen_safety": flag_coarsen_safety,
        "monotonicity": flag_monotonicity,
    }

    n_fail = sum(1 for v in flags.values() if not v)
    if n_fail == 0:
        recommendation = "PROCEED"
    elif n_fail == 1:
        recommendation = "REVIEW"
    else:
        recommendation = "REDESIGN"

    flags["_direction_detail"] = direction_checks
    flags["_monotonicity_detail"] = mono_detail
    return flags, recommendation


def _write_summary(
    ctx,
    results: dict,
    flags: dict,
    recommendation: str,
    cutpoints: np.ndarray,
    clean_sigma_std: float,
    thresh_LM: float,
    thresh_MH: float,
    noise_std: float,
    lag_k: int,
    n_bins: int,
    fill_value: float,
    band_pct: float,
) -> Path:
    md = []
    md.append("# WP5.5 Signal Audit Summary")
    md.append("")
    md.append(f"Run ID: `{ctx.run_id}`")
    md.append("")
    md.append("## Parameters")
    md.append("")
    ratio = noise_std / clean_sigma_std if clean_sigma_std > 0 else float("nan")
    md.append(f"- clean_sigma_std (post-warmup): {clean_sigma_std:.6f}")
    md.append(f"- noise_std: {noise_std}  (noise_std / clean_sigma_std = {ratio:.4f})")
    md.append(f"- lag_k_steps: {lag_k}")
    md.append(f"- n_bins: {n_bins}")
    md.append(f"- fill_value: {fill_value}")
    md.append(f"- threshold_band_pct: {band_pct}")
    md.append(f"- thresh_LM: {thresh_LM:.6f}")
    md.append(f"- thresh_MH: {thresh_MH:.6f}")
    md.append(f"- cutpoints (interior, len={len(cutpoints)}): {cutpoints.tolist()}")
    md.append("")
    md.append("## Metric Table")
    md.append("")
    header = "| condition | " + " | ".join(METRIC_ORDER) + " |"
    sep = "|" + "---|" * (len(METRIC_ORDER) + 1)
    md.append(header)
    md.append(sep)
    for c in CONDITION_ORDER:
        row = " | ".join(f"{results[c][m]:.4f}" for m in METRIC_ORDER)
        md.append(f"| {c} | {row} |")
    md.append("")
    md.append("## Policy: ungated metrics under the none condition")
    md.append("")
    md.append(
        "Under constant-input conditions, regime_classification_accuracy and "
        "threshold_overlap_rate become structurally undefined; they reflect the "
        "relationship between the chosen fill value and calibrated thresholds "
        "rather than signal information content. These metrics are therefore "
        "reported but not used in PASS/FAIL evaluation for the none condition."
    )
    md.append("")
    md.append("## Flags")
    md.append("")
    for key in ("direction", "separation", "coarsen_safety", "monotonicity"):
        status = "PASS" if flags[key] else "FAIL"
        md.append(f"- {key}: **{status}**")
    md.append("")
    md.append("### Direction detail (clean → none)")
    for k, v in flags["_direction_detail"].items():
        ungated_tag = " [REPORTED, NOT GATED]" if k in NONE_UNGATED_METRICS else ""
        md.append(f"  - {k}: {'OK' if v else 'VIOLATION'}{ungated_tag}")
    md.append("")
    md.append("### Monotonicity detail (clean → noisy → lagged → coarsened → none)")
    for k, v in flags["_monotonicity_detail"].items():
        ungated_tag = (
            " [none endpoint excluded; REPORTED, NOT GATED]"
            if k in NONE_UNGATED_METRICS
            else ""
        )
        md.append(f"  - {k}: {'monotone' if v else 'non-monotone'}{ungated_tag}")
    md.append("")
    md.append(f"## Overall Recommendation: **{recommendation}**")
    md.append("")

    out = Path(ctx.run_dir) / "audit_summary.md"
    out.write_text("\n".join(md), encoding="utf-8")
    return out


def run(cfg: dict, ctx) -> None:
    seed = int(cfg["seed"])
    audit_cfg = cfg["audit"]
    noise_std_raw = audit_cfg["noise_std"]
    lag_k = int(audit_cfg["lag_k_steps"])
    n_bins = int(audit_cfg["n_bins"])
    fill_value = float(audit_cfg["fill_value"])
    band_pct = float(audit_cfg.get("threshold_band_pct", 0.05))

    # 1. Clean signal
    sigma_hat_clean, regime_true_str, warmup_end, thresh_LM, thresh_MH = (
        _build_clean_signal(cfg, seed)
    )
    n = len(sigma_hat_clean)
    clean_sigma_std = float(np.nanstd(sigma_hat_clean[warmup_end:]))
    ctx.logger.info(
        f"Clean signal: n={n}, warmup_end={warmup_end}, "
        f"clean_sigma_std(post)={clean_sigma_std:.6f}, "
        f"thresh_LM={thresh_LM:.6f}, thresh_MH={thresh_MH:.6f}"
    )

    # Resolve noise_std: 'auto' -> 0.5 * clean_sigma_std_post_warmup.
    if isinstance(noise_std_raw, str) and noise_std_raw.strip().lower() == "auto":
        noise_std = 0.5 * clean_sigma_std
        ctx.logger.info(
            f"noise_std='auto' resolved to {noise_std:.6f} "
            f"(= 0.5 * clean_sigma_std_post_warmup = 0.5 * {clean_sigma_std:.6f})"
        )
    else:
        noise_std = float(noise_std_raw)
        ctx.logger.info(f"noise_std={noise_std:.6f} (fixed from config)")

    # 2. Fixed cutpoints — computed exactly once from the clean warmup region.
    cutpoints = compute_clean_cutpoints(sigma_hat_clean, warmup_end, n_bins=n_bins)
    ctx.logger.info(f"Clean cutpoints (n_bins={n_bins}): {cutpoints.tolist()}")

    # 3. Build four degraded versions.
    deg_rng = np.random.default_rng(seed)
    sigma_noisy = apply_noise(sigma_hat_clean, noise_std, deg_rng)
    sigma_lagged = apply_lag(sigma_hat_clean, lag_k)
    sigma_coarse = apply_coarsen(sigma_hat_clean, cutpoints)
    sigma_none = apply_remove(sigma_hat_clean, fill_value=fill_value)

    # 4. Shape/dtype assertions.
    for name, arr in [
        ("noisy", sigma_noisy),
        ("lagged", sigma_lagged),
        ("coarsened", sigma_coarse),
        ("none", sigma_none),
    ]:
        assert arr.shape == sigma_hat_clean.shape, f"{name}: shape mismatch"
        assert arr.dtype == sigma_hat_clean.dtype, f"{name}: dtype mismatch"

    # 5. None condition post-warmup must be exactly fill_value.
    assert np.all(sigma_none[warmup_end:] == fill_value), (
        "none condition post-warmup contains non-fill values"
    )

    # 6. Compute metrics for each condition.
    conditions = {
        "clean": sigma_hat_clean,
        "noisy": sigma_noisy,
        "lagged": sigma_lagged,
        "coarsened": sigma_coarse,
        "none": sigma_none,
    }
    post_mask = np.zeros(n, dtype=bool)
    post_mask[warmup_end:] = True

    rows = []
    results: dict[str, dict[str, float]] = {}
    for cond in CONDITION_ORDER:
        deg = conditions[cond]
        m = {
            "spearman": spearman_correlation(sigma_hat_clean, deg, post_mask),
            "pearson": pearson_correlation(sigma_hat_clean, deg, post_mask),
            "classification_accuracy": regime_classification_accuracy(
                deg, regime_true_str, thresh_LM, thresh_MH, warmup_end
            ),
            "separability": class_separability(deg, regime_true_str, warmup_end),
            "threshold_overlap": threshold_overlap_rate(
                deg, thresh_LM, thresh_MH, warmup_end, band_pct=band_pct
            ),
        }
        results[cond] = m
        for metric, value in m.items():
            rows.append({"condition": cond, "metric": metric, "value": float(value)})
        ctx.logger.info(
            f"[{cond}] spearman={m['spearman']:.4f} pearson={m['pearson']:.4f} "
            f"acc={m['classification_accuracy']:.4f} sep={m['separability']:.2f} "
            f"overlap={m['threshold_overlap']:.4f}"
        )

    df = pd.DataFrame(rows, columns=["condition", "metric", "value"])
    csv_path = Path(ctx.run_dir) / "metrics_signal_audit.csv"
    df.to_csv(csv_path, index=False)
    ctx.logger.info(f"Wrote {csv_path.as_posix()}")

    # 7. Ladder figure.
    fig_path = _plot_ladder(results, Path(ctx.plots_dir))
    ctx.logger.info(f"Wrote {fig_path.as_posix()}")

    # 8. Flags + markdown summary.
    flags, recommendation = _evaluate_flags(results)
    md_path = _write_summary(
        ctx,
        results=results,
        flags=flags,
        recommendation=recommendation,
        cutpoints=cutpoints,
        clean_sigma_std=clean_sigma_std,
        thresh_LM=thresh_LM,
        thresh_MH=thresh_MH,
        noise_std=noise_std,
        lag_k=lag_k,
        n_bins=n_bins,
        fill_value=fill_value,
        band_pct=band_pct,
    )
    ctx.logger.info(f"Wrote {md_path.as_posix()}")
    ctx.logger.info(
        f"Flags: direction={flags['direction']} separation={flags['separation']} "
        f"coarsen_safety={flags['coarsen_safety']} monotonicity={flags['monotonicity']}"
    )
    ctx.logger.info(f"Recommendation: {recommendation}")

# ============================================================
# FILE: src/wp5_5/job_w55_runtime.py
# PURPOSE: WP5.5 runtime-profiling job for candidate sweep scale decisions.
# STATUS: audit
# ============================================================

"""WP5.5 runtime benchmark: PPO on CPU vs GPU (2 seeds each).

Purely a device-selection helper for downstream phases; training is short
(default 100k steps) and the resulting policy is discarded.
"""

from __future__ import annotations

import copy
import time
from pathlib import Path

import numpy as np
import pandas as pd
from stable_baselines3 import PPO
from stable_baselines3.common.monitor import Monitor
from stable_baselines3.common.vec_env import DummyVecEnv

from src.wp2.synth_regime import (
    REGIME_LABELS,
    assign_regime_hat,
    calibrate_thresholds,
    compute_rolling_rv,
    generate_mid_series,
    generate_regime_series,
)
from src.wp3.env import MMEnv


def _build_exog_df(cfg: dict, seed: int) -> pd.DataFrame:
    rng = np.random.default_rng(seed)
    n_steps = int(cfg["episode"]["n_steps"])
    rv_window = int(cfg["regime"]["rv_window"])
    warmup = int(cfg["regime"]["warmup_steps"])

    regime_cfg = copy.deepcopy(cfg.get("regime", {}))
    regime_cfg.setdefault("sigma_mid_ticks_base", float(cfg["market"].get("sigma_mid_ticks", 0.5)))
    regime_cfg.setdefault("sigma_mult", [0.6, 1.0, 1.8])
    cfg_local = {**cfg, "regime": regime_cfg}

    regime_true_int = generate_regime_series(n_steps, seed, cfg=cfg_local, rng=rng)
    mid, _ = generate_mid_series(regime_true_int, cfg_local, rng)
    _, sigma_hat = compute_rolling_rv(mid, rv_window, float(cfg["market"]["tick_size"]))
    thresh_LM, thresh_MH = calibrate_thresholds(sigma_hat, warmup)
    regime_hat = assign_regime_hat(sigma_hat, thresh_LM, thresh_MH, warmup)
    regime_true_str = ["M"] + [REGIME_LABELS[int(r)] for r in regime_true_int]

    return pd.DataFrame({
        "t": np.arange(len(mid)),
        "mid": mid,
        "sigma_hat": sigma_hat,
        "regime_hat": regime_hat,
        "regime_true": regime_true_str,
    })


def _build_env_cfg(cfg: dict, seed: int) -> dict:
    env_cfg = copy.deepcopy(cfg)
    env_cfg["seed"] = int(seed)
    env_block = cfg.get("env", {})
    env_cfg["wp3"] = {
        "eta": float(env_block.get("eta", 1e-3)),
        "use_regime": bool(env_block.get("use_regime", False)),
        "use_sigma": True,
    }
    env_cfg.setdefault("episode", {})
    env_cfg["episode"].setdefault("inv_max_clip", 50)
    return env_cfg


def _train_once(cfg: dict, exog: pd.DataFrame, device: str, seed: int, timesteps: int) -> float:
    env_cfg = _build_env_cfg(cfg, seed)
    env = MMEnv(env_cfg)
    env.reset(seed=seed, options={"exog": exog})
    monitor = Monitor(env)
    vec_env = DummyVecEnv([lambda _m=monitor: _m])

    ppo_cfg = cfg["ppo"]
    model = PPO(
        ppo_cfg.get("policy", "MlpPolicy"),
        vec_env,
        seed=seed,
        learning_rate=float(ppo_cfg["learning_rate"]),
        n_steps=int(ppo_cfg["n_steps"]),
        batch_size=int(ppo_cfg["batch_size"]),
        n_epochs=int(ppo_cfg["n_epochs"]),
        gamma=float(ppo_cfg["gamma"]),
        verbose=0,
        device=device,
    )
    t0 = time.perf_counter()
    model.learn(total_timesteps=int(timesteps))
    wall = time.perf_counter() - t0
    vec_env.close()
    return wall


def _write_decision(
    ctx,
    df: pd.DataFrame,
    cpu_times: pd.Series,
    gpu_times: pd.Series,
) -> tuple[str, str, float]:
    lines = ["# WP5.5 Runtime Benchmark Decision", ""]
    lines.append("## Raw wall times")
    lines.append("")
    for _, r in df.iterrows():
        wt = r["wall_time_seconds"]
        wt_s = "NaN" if pd.isna(wt) else f"{float(wt):.2f}s"
        lines.append(
            f"- device={r['device']}, seed={int(r['seed'])}, "
            f"time={wt_s}, status={r['status']}"
        )
    lines.append("")

    if len(gpu_times) == 0:
        recommendation = "USE_CPU"
        reason = "no_cuda"
        ratio = float("nan")
        cpu_mean = float(cpu_times.mean()) if len(cpu_times) else float("nan")
        lines.append("## Means")
        lines.append("")
        lines.append(f"- CPU mean: {cpu_mean:.2f}s")
        lines.append("- GPU mean: N/A (skipped — CUDA unavailable)")
        lines.append("")
        lines.append("## Ratio: N/A")
        lines.append("")
        lines.append("## Applied rule: GPU runs skipped → USE_CPU")
    else:
        cpu_mean = float(cpu_times.mean())
        gpu_mean = float(gpu_times.mean())
        ratio = gpu_mean / cpu_mean if cpu_mean > 0 else float("nan")
        lines.append("## Means")
        lines.append("")
        lines.append(f"- CPU mean: {cpu_mean:.2f}s")
        lines.append(f"- GPU mean: {gpu_mean:.2f}s")
        lines.append("")
        lines.append(f"## Ratio (GPU/CPU): {ratio:.3f}")
        lines.append("")
        if ratio <= 0.67:
            recommendation = "USE_GPU"
            reason = f"GPU >=1.5x faster (ratio={ratio:.3f})"
            lines.append("## Applied rule: ratio <= 0.67 → USE_GPU")
        elif ratio >= 1.5:
            recommendation = "USE_CPU"
            reason = f"GPU slower than CPU (ratio={ratio:.3f})"
            lines.append("## Applied rule: ratio >= 1.5 → USE_CPU (GPU slower)")
        else:
            recommendation = "USE_CPU"
            reason = f"GPU benefit marginal (ratio={ratio:.3f})"
            lines.append("## Applied rule: 0.67 < ratio < 1.5 → USE_CPU (small/no benefit)")

    lines.append("")
    lines.append(f"## Recommendation: **{recommendation}**  ({reason})")
    lines.append("")

    out = Path(ctx.run_dir) / "runtime_decision.md"
    out.write_text("\n".join(lines), encoding="utf-8")
    return recommendation, reason, ratio


def run(cfg: dict, ctx) -> None:
    import torch

    cuda_ok = bool(torch.cuda.is_available())
    if not cuda_ok:
        ctx.logger.warning("CUDA not available — GPU runs will be skipped.")

    bench = cfg["benchmark"]
    devices = list(bench["devices"])
    seeds = list(bench["seeds"])
    timesteps = int(cfg["ppo"].get("total_timesteps", 100_000))

    exog_per_seed = {int(s): _build_exog_df(cfg, int(s)) for s in seeds}

    rows: list[dict] = []
    for device in devices:
        if device == "cuda" and not cuda_ok:
            for s in seeds:
                rows.append({
                    "device": "cuda",
                    "seed": int(s),
                    "wall_time_seconds": float("nan"),
                    "total_timesteps": timesteps,
                    "status": "skipped_no_cuda",
                })
            continue
        for s in seeds:
            s_int = int(s)
            ctx.logger.info(
                f"[benchmark] device={device} seed={s_int} timesteps={timesteps}"
            )
            wall = _train_once(cfg, exog_per_seed[s_int], device, s_int, timesteps)
            ctx.logger.info(f"[benchmark]   wall_time={wall:.2f}s")
            rows.append({
                "device": device,
                "seed": s_int,
                "wall_time_seconds": float(wall),
                "total_timesteps": timesteps,
                "status": "ok",
            })

    df = pd.DataFrame(rows, columns=[
        "device", "seed", "wall_time_seconds", "total_timesteps", "status",
    ])
    csv_path = Path(ctx.run_dir) / "runtime_summary.csv"
    df.to_csv(csv_path, index=False)
    ctx.logger.info(f"Wrote {csv_path.as_posix()}")

    cpu_times = df[(df["device"] == "cpu") & (df["status"] == "ok")]["wall_time_seconds"]
    gpu_times = df[(df["device"] == "cuda") & (df["status"] == "ok")]["wall_time_seconds"]

    recommendation, reason, ratio = _write_decision(ctx, df, cpu_times, gpu_times)
    ctx.logger.info(
        f"Runtime decision: {recommendation} (reason={reason}, ratio={ratio})"
    )

# ============================================================
# FILE: src/wp5_5/job_w55_calibration.py
# PURPOSE: WP5.5 calibration job for noisy/lagged/coarsened signal settings.
# STATUS: audit
# ============================================================

"""WP5.5 signal calibration sweep.

Two separate one-dimensional sweeps over the noise alpha and lag k
parameters of the volatility-signal degradation transforms. NO PPO,
NO training — pure offline numpy diagnostic to inform parameter
choices for the upcoming Signal Informativeness Sweep.

For each seed, the same clean signal is generated once, then:
  - alpha sweep (k = 0): degraded = apply_noise(clean, alpha * sigma_std)
  - k     sweep (alpha = 0): degraded = apply_lag(clean, k)

Metrics per (parameter, value, seed):
  pearson, nrmse, classification_accuracy, accuracy_drop, regime_crossing_rate.
"""

from __future__ import annotations

from pathlib import Path

import numpy as np
import pandas as pd

from src.wp5_5.job_w55_audit import _build_clean_signal
from src.wp5_5.signal_audit import (
    nrmse,
    pearson_correlation,
    regime_classification_accuracy,
    regime_crossing_rate,
)
from src.wp5_5.signal_degradation import apply_lag, apply_noise


METRIC_COLS = [
    "pearson",
    "nrmse",
    "classification_accuracy",
    "accuracy_drop",
    "regime_crossing_rate",
]


def _compute_metrics(
    clean: np.ndarray,
    degraded: np.ndarray,
    regime_true: list,
    thresh_LM: float,
    thresh_MH: float,
    warmup_end: int,
    clean_acc: float,
) -> dict:
    n = len(clean)
    post_mask = np.zeros(n, dtype=bool)
    post_mask[warmup_end:] = True
    deg_acc = regime_classification_accuracy(
        degraded, regime_true, thresh_LM, thresh_MH, warmup_end
    )
    return {
        "pearson": pearson_correlation(clean, degraded, post_mask),
        "nrmse": nrmse(clean, degraded, post_mask),
        "classification_accuracy": deg_acc,
        "accuracy_drop": clean_acc - deg_acc,
        "regime_crossing_rate": regime_crossing_rate(
            degraded, thresh_LM, thresh_MH, warmup_end
        ),
    }


def run(cfg: dict, ctx) -> None:
    cal = cfg["calibration"]
    alpha_values = [float(a) for a in cal["alpha_values"]]
    k_values = [int(k) for k in cal["k_values"]]
    n_seeds = int(cal["n_seeds"])
    base_seed = int(cfg["seed"])

    rows: list[dict] = []
    for seed_idx in range(n_seeds):
        seed = base_seed + seed_idx
        sigma_clean, regime_true, warmup_end, thresh_LM, thresh_MH = (
            _build_clean_signal(cfg, seed)
        )
        n = len(sigma_clean)
        sigma_std_post = float(np.nanstd(sigma_clean[warmup_end:]))
        clean_acc = regime_classification_accuracy(
            sigma_clean, regime_true, thresh_LM, thresh_MH, warmup_end
        )
        ctx.logger.info(
            f"seed={seed} n={n} warmup_end={warmup_end} "
            f"sigma_std_post={sigma_std_post:.6f} clean_acc={clean_acc:.4f} "
            f"thresh_LM={thresh_LM:.6f} thresh_MH={thresh_MH:.6f}"
        )

        rng = np.random.default_rng(seed)
        for alpha in alpha_values:
            noise_std = alpha * sigma_std_post
            degraded = apply_noise(sigma_clean, noise_std, rng)
            m = _compute_metrics(
                sigma_clean, degraded, regime_true,
                thresh_LM, thresh_MH, warmup_end, clean_acc,
            )
            rows.append({"parameter": "alpha", "value": float(alpha), "seed": seed, **m})

        for k in k_values:
            degraded = apply_lag(sigma_clean, k)
            m = _compute_metrics(
                sigma_clean, degraded, regime_true,
                thresh_LM, thresh_MH, warmup_end, clean_acc,
            )
            rows.append({"parameter": "k", "value": float(k), "seed": seed, **m})

    df_per_seed = pd.DataFrame(rows)
    per_seed_path = Path(ctx.run_dir) / "metrics_calibration_per_seed.csv"
    df_per_seed.to_csv(per_seed_path, index=False)
    ctx.logger.info(f"Wrote {per_seed_path.as_posix()}")

    agg = (
        df_per_seed
        .groupby(["parameter", "value"], sort=False)[METRIC_COLS]
        .agg(["mean", "std"])
    )
    agg.columns = [f"{m}_{stat}" for m, stat in agg.columns]
    agg = agg.reset_index()
    agg_path = Path(ctx.run_dir) / "metrics_calibration_aggregated.csv"
    agg.to_csv(agg_path, index=False)
    ctx.logger.info(f"Wrote {agg_path.as_posix()}")

    print()
    print(
        "| parameter | value | pearson_mean | nrmse_mean | accuracy_drop_mean | crossing_rate_mean |"
    )
    print("|---|---:|---:|---:|---:|---:|")
    for _, r in agg.iterrows():
        print(
            f"| {r['parameter']} | {r['value']:g} | "
            f"{r['pearson_mean']:.4f} | {r['nrmse_mean']:.4f} | "
            f"{r['accuracy_drop_mean']:.4f} | {r['regime_crossing_rate_mean']:.4f} |"
        )
    print()

# ============================================================
# FILE: src/wp6/__init__.py
# PURPOSE: Package marker for WP6 signal-informativeness sweep code.
# STATUS: support
# ============================================================



# ============================================================
# FILE: src/wp6/_resume.py
# PURPOSE: WP6 resume/checkpoint consistency helpers for long-running sweep jobs.
# STATUS: active
# ============================================================

"""Shared helpers for --resume aware sweep jobs (wp6).

A cell is considered complete only if BOTH:
  * the expected model .zip exists on disk, AND
  * a (seed, condition, variant) row exists in the metrics CSV.

If exactly one of those exists, the run is in an inconsistent state.
The job must stop and report — silently skipping such an orphan would
risk data loss (the metric row may have been written for a partially
trained model, or the model may exist for a cell whose metrics never
landed). Re-evaluation from a saved model is not implemented; until
it is, fail loudly so the operator can decide.
"""

from __future__ import annotations

from pathlib import Path

import pandas as pd


def load_completed_set(metrics_path: Path):
    """Return (rows, completed_set) from an existing metrics CSV.

    rows           — list[dict], original CSV rows preserved as-is for
                     re-emission at the end of the run.
    completed_set  — set[tuple[int, str, str]] keyed by
                     (seed, condition, variant).

    If the file does not exist, returns ([], set()).
    """
    if not metrics_path.exists():
        return [], set()
    df = pd.read_csv(metrics_path)
    rows = df.to_dict("records")
    completed: set[tuple[int, str, str]] = set()
    for r in rows:
        completed.add((int(r["seed"]), str(r["condition"]), str(r["variant"])))
    return rows, completed


def check_cell_consistency(
    model_exists: bool,
    metric_exists: bool,
    *,
    seed: int,
    condition: str,
    variant: str,
) -> str:
    """Classify a cell during resume and raise on inconsistency.

    Returns:
        "skip"  — both model and metric row already present; skip the cell.
        "train" — neither present; train normally.

    Raises:
        RuntimeError — exactly one of (model, metric) is present. The
            error message names the orphan tuple and which artifact is
            missing.
    """
    if model_exists and metric_exists:
        return "skip"
    if not model_exists and not metric_exists:
        return "train"

    cell = (seed, condition, variant)
    if model_exists and not metric_exists:
        missing = "metrics row"
        present = "model file"
    else:
        missing = "model file"
        present = "metrics row"
    raise RuntimeError(
        f"Inconsistent resume state for cell {cell}: {present} exists but "
        f"{missing} is missing. Re-evaluation from saved models is not "
        f"implemented. To proceed: either delete the orphan {present.lower()} "
        f"so the cell re-trains from scratch, or implement re-evaluation. "
        f"Stopping to avoid silent data loss."
    )

# ============================================================
# FILE: src/wp6/job_w6_sweep_pilot.py
# PURPOSE: WP6 pilot sweep implementation and design rationale for degradation conditions and variants.
# STATUS: active
# ============================================================

"""WP6 Signal Informativeness Sweep — PILOT.

Do not interpret pilot outputs as final statistical evidence; pilot is
for pipeline correctness, condition/variant wiring, and directional
sanity only. The full WP6 sweep is planned at 20 seeds × 24 cells × 1M
timesteps and will be run separately after pilot review.

Grid: conditions × variants minus omit_cells. Per cell, train one PPO
model on a degraded sigma_hat series (or the clean one for the `full`
condition; `none` flips use_sigma=False on the env). Coarse-layer
checkpoint: a model file already on disk causes the cell to be SKIPPED
without appending a metrics row, so resumption after interruption does
not double-count.

For the `none` condition, use_sigma is forced to False regardless of
variant flags. The (none, sigma_only) cell is omitted entirely;
remaining 4 variants under `none` train under use_sigma=False with
their respective regime_source. By design, regime_only and oracle_pure
are condition-invariant: the regime detector operates on the clean
sigma upstream of the observation manipulation, so the regime label
quality is constant across conditions. This is intended scope —
the experiment isolates the marginal value of the explicit regime
label given a fixed-quality regime estimate.
"""

from __future__ import annotations

import copy
import math
import time
from pathlib import Path

import numpy as np
import pandas as pd
from stable_baselines3 import PPO
from stable_baselines3.common.monitor import Monitor
from stable_baselines3.common.vec_env import DummyVecEnv

from src.wp1.w1_as_baseline import compute_metrics
from src.wp2.synth_regime import run_wp2
from src.wp3.env import MMEnv
from src.wp5_5.signal_degradation import (
    apply_coarsen,
    apply_lag,
    apply_noise,
    compute_clean_cutpoints,
)
from src.wp6._resume import check_cell_consistency, load_completed_set


VARIANT_FLAGS = {
    "sigma_only":  {"use_sigma": True,  "regime_source": "none"},
    "regime_only": {"use_sigma": False, "regime_source": "hat"},
    "combined":    {"use_sigma": True,  "regime_source": "hat"},
    "oracle_pure": {"use_sigma": False, "regime_source": "true"},
    "oracle_full": {"use_sigma": True,  "regime_source": "true"},
}

PILOT_WARNING = (
    "Do not interpret pilot outputs as final statistical evidence; pilot is\n"
    "for pipeline correctness, condition/variant wiring, and directional\n"
    "sanity only. The full WP6 sweep is planned at 20 seeds × 24 cells × 1M\n"
    "timesteps and will be run separately after pilot review."
)


def _build_degraded(condition, sigma_clean, sigma_std_post,
                    alpha, k_steps, cutpoints, rng):
    if condition == "full":
        return sigma_clean.copy()
    if condition == "noisy":
        return apply_noise(sigma_clean, alpha * sigma_std_post, rng)
    if condition == "lagged":
        return apply_lag(sigma_clean, k_steps)
    if condition == "coarsened":
        return apply_coarsen(sigma_clean, cutpoints)
    if condition == "none":
        # Series is unused: use_sigma_eff is forced False at the variant level
        # for the none condition, so MMEnv zero-fills the sigma slot.
        return sigma_clean.copy()
    raise ValueError(f"Unknown condition: {condition}")


def _eval_model(model, cfg_eval, df_exog_eval, seed):
    env = MMEnv(cfg_eval)
    obs, _ = env.reset(seed=seed, options={"exog": df_exog_eval})
    n = int(cfg_eval["episode"]["n_steps"])
    equity = np.zeros(n + 1)
    inv = np.zeros(n + 1, dtype=int)
    fills = np.zeros(n, dtype=int)
    equity[0] = env._state.equity
    inv[0] = env._state.inv
    for t in range(n):
        action, _ = model.predict(obs, deterministic=True)
        obs, _r, _term, _trunc, info = env.step(action)
        equity[t + 1] = info["equity"]
        inv[t + 1] = info["inv"]
        fills[t] = info["fills"]
    return compute_metrics(equity, inv, fills, dt=cfg_eval["market"]["dt"])


def run(cfg: dict, ctx) -> None:
    sweep = cfg["sweep"]
    conditions = list(sweep["conditions"])
    variants = list(sweep["variants"])
    omit_cells = [tuple(c) for c in sweep.get("omit_cells", [])]
    seeds = list(sweep["seeds"])
    alpha = float(sweep["noisy_alpha"])
    k_steps = int(sweep["lagged_k"])
    n_bins = int(sweep["coarsened_n_bins"])

    # Generic grid validation
    total_cells = len(conditions) * len(variants) - len(omit_cells)
    total_trainings = total_cells * len(seeds)
    assert total_cells > 0, f"total_cells must be > 0, got {total_cells}"
    assert total_trainings > 0, f"total_trainings must be > 0, got {total_trainings}"

    # Pilot-specific assertion
    is_real_pilot = (
        cfg.get("run_tag") == "wp6-sweep-pilot"
        and cfg.get("job") == "w6_sweep_pilot"
        and len(seeds) > 1
    )
    if is_real_pilot:
        assert total_cells == 24, f"Real pilot must have 24 cells, got {total_cells}"
        assert len(seeds) == 3, f"Real pilot must have 3 seeds, got {len(seeds)}"
        assert total_trainings == 72, f"Real pilot must have 72 trainings, got {total_trainings}"

    # Banner
    est_min = total_trainings * 17
    print()
    print("=" * 72)
    print("WP6 Signal Informativeness Sweep — PILOT")
    print(f"Cells: {total_cells}  Trainings: {total_trainings}  Seeds: {seeds}")
    print(f"Est. wall time @ 17 min/cell: {est_min} min (~{est_min/60:.1f} h)")
    print("=" * 72)
    print(PILOT_WARNING)
    print("=" * 72)
    print(flush=True)

    out_dir = Path(ctx.run_dir)
    models_root = out_dir / "models"
    models_root.mkdir(exist_ok=True)

    n_full = int(cfg["episode"]["n_steps"])
    train_frac = float(sweep.get("train_frac", 0.8))
    n_train = math.floor(train_frac * n_full)
    n_test = n_full - n_train
    warmup_end = int(cfg["regime"]["warmup_steps"])
    total_timesteps = int(cfg["wp4"]["total_timesteps"])

    metrics_path = out_dir / "metrics_sweep_pilot.csv"
    is_resume = getattr(ctx, "resume_run_id", None) is not None
    if is_resume:
        existing_rows, completed_set = load_completed_set(metrics_path)
        ctx.logger.info(
            f"Resume mode: {len(existing_rows)} existing metric rows, "
            f"{len(completed_set)} completed (seed, condition, variant) cells"
        )
    else:
        existing_rows, completed_set = [], set()

    rows = []
    trained = 0
    skipped = 0
    omit_set = set(omit_cells)

    idx_train_end = n_train + 1
    idx_test_start = n_train
    idx_test_end = n_train + n_test + 1

    for seed in seeds:
        ctx.logger.info(f"=== Seed {seed} ===")
        df_exog, _, _ = run_wp2(cfg, seed, ctx=ctx)
        sigma_clean = df_exog["sigma_hat"].to_numpy(dtype=np.float64, copy=True)
        # NOTE (calibration scope): sigma_std_post is computed from the full
        # post-warmup span (training + OOS test). The synthetic series is
        # design-stationary (sticky Markov chain on a stationary process), so
        # the std differs between train-only and full-span calibration by an
        # estimated <10%. The directional Chapter 5 findings (Cohen's dz
        # between -0.57 and -0.91 for combined < sigma_only) are robust to
        # this scale. See thesis §5.1 calibration-range caveat. Future work
        # may recalibrate using training-only data for stricter OOS hygiene.
        sigma_std_post = float(np.nanstd(sigma_clean[warmup_end:]))
        cutpoints = compute_clean_cutpoints(sigma_clean, warmup_end, n_bins=n_bins)
        deg_rng = np.random.default_rng(seed)

        seed_dir = models_root / f"seed{seed}"
        seed_dir.mkdir(exist_ok=True)

        for condition in conditions:
            sigma_deg = _build_degraded(
                condition, sigma_clean, sigma_std_post,
                alpha, k_steps, cutpoints, deg_rng,
            )
            df_full = df_exog.copy()
            df_full["sigma_hat"] = sigma_deg

            for variant in variants:
                if (condition, variant) in omit_set:
                    continue

                model_path = seed_dir / f"{condition}__{variant}.zip"
                if is_resume:
                    decision = check_cell_consistency(
                        model_exists=model_path.exists(),
                        metric_exists=(seed, condition, variant) in completed_set,
                        seed=seed, condition=condition, variant=variant,
                    )
                    if decision == "skip":
                        msg = f"RESUME-SKIP {condition}/{variant}"
                        print(msg, flush=True)
                        ctx.logger.info(msg)
                        skipped += 1
                        continue
                else:
                    if model_path.exists():
                        msg = f"SKIP {condition}/{variant} (already trained)"
                        print(msg, flush=True)
                        ctx.logger.info(msg)
                        skipped += 1
                        continue

                vflags = VARIANT_FLAGS[variant]
                use_sigma_eff = vflags["use_sigma"] and condition != "none"
                cfg_tr = copy.deepcopy(cfg)
                cfg_tr["wp3"] = {
                    **cfg_tr.get("wp3", {}),
                    "use_sigma": use_sigma_eff,
                    "regime_source": vflags["regime_source"],
                }
                cfg_tr["episode"] = {**cfg_tr["episode"], "n_steps": n_train}

                exog_train = df_full.iloc[:idx_train_end].reset_index(drop=True)
                exog_test = df_full.iloc[idx_test_start:idx_test_end].reset_index(drop=True)

                env_tr = MMEnv(cfg_tr)
                env_tr.reset(seed=seed, options={"exog": exog_train})
                vec_env = DummyVecEnv([lambda _e=Monitor(env_tr): _e])

                wp4 = cfg["wp4"]
                t_train_start = time.time()
                model = PPO(
                    "MlpPolicy", vec_env, seed=seed,
                    learning_rate=float(wp4["learning_rate"]),
                    n_steps=int(wp4["n_steps"]),
                    batch_size=int(wp4["batch_size"]),
                    n_epochs=int(wp4["n_epochs"]),
                    gamma=float(wp4["gamma"]),
                    gae_lambda=float(wp4["gae_lambda"]),
                    clip_range=float(wp4["clip_range"]),
                    ent_coef=float(wp4["ent_coef"]),
                    verbose=0,
                    device=cfg.get("wp4", {}).get("device", "cpu"),
                )
                model.learn(total_timesteps=total_timesteps)
                model.save(str(model_path))
                vec_env.close()
                train_seconds = time.time() - t_train_start

                cfg_ev = copy.deepcopy(cfg)
                cfg_ev["wp3"] = {
                    **cfg_ev.get("wp3", {}),
                    "use_sigma": use_sigma_eff,
                    "regime_source": vflags["regime_source"],
                }
                cfg_ev["episode"] = {**cfg_ev["episode"], "n_steps": n_test}

                t_eval_start = time.time()
                m = _eval_model(model, cfg_ev, exog_test, seed)
                eval_seconds = time.time() - t_eval_start

                rows.append({
                    "seed": seed,
                    "condition": condition,
                    "variant": variant,
                    "sharpe_like": m["sharpe_like"],
                    "final_equity": m["final_equity"],
                    "fill_rate": m["fill_rate"],
                    "inv_p99": m["inv_p99"],
                    "train_seconds": train_seconds,
                    "eval_seconds": eval_seconds,
                })
                trained += 1
                ctx.logger.info(
                    f"[{seed}] {condition}/{variant}: trained {train_seconds:.1f}s "
                    f"eval {eval_seconds:.1f}s sharpe={m['sharpe_like']:.4f} "
                    f"inv_p99={m['inv_p99']:.0f}"
                )

    all_rows = existing_rows + rows
    df_metrics = pd.DataFrame(all_rows)
    df_metrics.to_csv(metrics_path, index=False)
    ctx.logger.info(
        f"Wrote {metrics_path.as_posix()} "
        f"(existing={len(existing_rows)}, new={len(rows)}, total={len(all_rows)})"
    )

    md = []
    md.append("# WP6 Signal Informativeness Sweep — Pilot Summary")
    md.append("")
    md.append(f"Run ID: `{ctx.run_id}`")
    md.append("")
    md.append(f"- trained = {trained}")
    md.append(f"- skipped = {skipped}")
    md.append("")
    if len(df_metrics) > 0:
        ts = df_metrics["train_seconds"]
        md.append("## Cell timing distribution (train_seconds)")
        md.append("")
        md.append(f"- min: {ts.min():.1f}")
        md.append(f"- median: {ts.median():.1f}")
        md.append(f"- max: {ts.max():.1f}")
        md.append(f"- p95: {ts.quantile(0.95):.1f}")
        md.append("")
        md.append("## Per-cell mean across seeds")
        md.append("")
        agg = (df_metrics
               .groupby(["condition", "variant"], sort=False)
               [["sharpe_like", "inv_p99"]]
               .mean()
               .reset_index()
               .sort_values(["condition", "variant"]))
        md.append("| condition | variant | sharpe_like_mean | inv_p99_mean |")
        md.append("|---|---|---:|---:|")
        for _, r in agg.iterrows():
            md.append(
                f"| {r['condition']} | {r['variant']} | "
                f"{r['sharpe_like']:.4f} | {r['inv_p99']:.1f} |"
            )
        md.append("")
    md.append("## Reminder")
    md.append("")
    md.append(PILOT_WARNING)
    md.append("")
    md.append("## Next step")
    md.append("")
    md.append(
        "Review pilot_summary.md, then run the full WP6 sweep at "
        "20 seeds × 1M timesteps."
    )
    summary_path = out_dir / "pilot_summary.md"
    summary_path.write_text("\n".join(md), encoding="utf-8")
    ctx.logger.info(f"Wrote {summary_path.as_posix()}")
    print(f"Trained: {trained}   Skipped: {skipped}", flush=True)

# ============================================================
# FILE: src/wp6/job_w6_sweep_full.py
# PURPOSE: WP6 full 20-seed signal-informativeness sweep used for Chapter 5 evidence.
# STATUS: active
# ============================================================

"""Full WP6 Signal Informativeness Sweep — 20 seeds × 24 cells × 1M timesteps.

This is the thesis's central empirical experiment. Outputs feed Chapter 5
directly. Do not interpret partial results as final until all 480 cells
have completed. Do not interrupt unless necessary; if interrupted, resume
with `python run.py --config config/w6_sweep_full.json --resume <run_id>`
where <run_id> is the original run dir name.

Uses the same pipeline as the pilot (verified bug-free in commit 3144051).
The `none` condition forces use_sigma=False at the variant level — see
src/wp6/job_w6_sweep_pilot.py docstring for the full design rationale.
By design, regime_only and oracle_pure are condition-invariant: the
regime detector operates on the clean upstream sigma, so the regime
label quality is constant across conditions; the experiment isolates the
marginal value of the explicit regime label given a fixed-quality regime
estimate.

Coarse-layer checkpointing: a model file already on disk causes the cell
to be skipped. In --resume mode, both the model .zip and the
corresponding metrics row must exist; orphan state raises a RuntimeError
(see src/wp6/_resume.py).
"""

from __future__ import annotations

import copy
import math
import time
from pathlib import Path

import numpy as np
import pandas as pd
from stable_baselines3 import PPO
from stable_baselines3.common.monitor import Monitor
from stable_baselines3.common.vec_env import DummyVecEnv

from src.wp1.w1_as_baseline import compute_metrics
from src.wp2.synth_regime import run_wp2
from src.wp3.env import MMEnv
from src.wp5_5.signal_degradation import (
    apply_coarsen,
    apply_lag,
    apply_noise,
    compute_clean_cutpoints,
)
from src.wp6._resume import check_cell_consistency, load_completed_set


VARIANT_FLAGS = {
    "sigma_only":  {"use_sigma": True,  "regime_source": "none"},
    "regime_only": {"use_sigma": False, "regime_source": "hat"},
    "combined":    {"use_sigma": True,  "regime_source": "hat"},
    "oracle_pure": {"use_sigma": False, "regime_source": "true"},
    "oracle_full": {"use_sigma": True,  "regime_source": "true"},
}

FULL_BANNER_NOTE = (
    "FULL SWEEP — thesis Chapter 5 input. 20 seeds × 24 cells × 1M timesteps.\n"
    "If interrupted, resume with --resume <run_id>."
)


def _build_degraded(condition, sigma_clean, sigma_std_post,
                    alpha, k_steps, cutpoints, rng):
    if condition == "full":
        return sigma_clean.copy()
    if condition == "noisy":
        return apply_noise(sigma_clean, alpha * sigma_std_post, rng)
    if condition == "lagged":
        return apply_lag(sigma_clean, k_steps)
    if condition == "coarsened":
        return apply_coarsen(sigma_clean, cutpoints)
    if condition == "none":
        # Series is unused: use_sigma_eff is forced False at the variant level
        # for the none condition, so MMEnv zero-fills the sigma slot.
        return sigma_clean.copy()
    raise ValueError(f"Unknown condition: {condition}")


def _eval_model(model, cfg_eval, df_exog_eval, seed):
    env = MMEnv(cfg_eval)
    obs, _ = env.reset(seed=seed, options={"exog": df_exog_eval})
    n = int(cfg_eval["episode"]["n_steps"])
    equity = np.zeros(n + 1)
    inv = np.zeros(n + 1, dtype=int)
    fills = np.zeros(n, dtype=int)
    equity[0] = env._state.equity
    inv[0] = env._state.inv
    for t in range(n):
        action, _ = model.predict(obs, deterministic=True)
        obs, _r, _term, _trunc, info = env.step(action)
        equity[t + 1] = info["equity"]
        inv[t + 1] = info["inv"]
        fills[t] = info["fills"]
    return compute_metrics(equity, inv, fills, dt=cfg_eval["market"]["dt"])


def run(cfg: dict, ctx) -> None:
    sweep = cfg["sweep"]
    conditions = list(sweep["conditions"])
    variants = list(sweep["variants"])
    omit_cells = [tuple(c) for c in sweep.get("omit_cells", [])]
    seeds = list(sweep["seeds"])
    alpha = float(sweep["noisy_alpha"])
    k_steps = int(sweep["lagged_k"])
    n_bins = int(sweep["coarsened_n_bins"])

    # Generic grid validation
    total_cells = len(conditions) * len(variants) - len(omit_cells)
    total_trainings = total_cells * len(seeds)
    assert total_cells > 0, f"total_cells must be > 0, got {total_cells}"
    assert total_trainings > 0, f"total_trainings must be > 0, got {total_trainings}"

    # Full-sweep-specific assertion
    is_real_full = (
        cfg.get("run_tag") == "wp6-sweep-full"
        and cfg.get("job") == "w6_sweep_full"
        and len(seeds) > 3
    )
    if is_real_full:
        assert total_cells == 24, f"Real full sweep must have 24 cells, got {total_cells}"
        assert len(seeds) == 20, f"Real full sweep must have 20 seeds, got {len(seeds)}"
        assert total_trainings == 480, f"Real full sweep must have 480 trainings, got {total_trainings}"

    est_min = total_trainings * 17
    print()
    print("=" * 72)
    print("WP6 Signal Informativeness Sweep — FULL SWEEP")
    print(f"Cells: {total_cells}  Trainings: {total_trainings}  Seeds: {len(seeds)}")
    print(f"Est. wall time @ 17 min/cell: {est_min} min "
          f"(~{est_min/60:.1f} h, ~{est_min/60/24:.1f} days)")
    print("=" * 72)
    print(FULL_BANNER_NOTE)
    print("=" * 72)
    print(flush=True)

    out_dir = Path(ctx.run_dir)
    models_root = out_dir / "models"
    models_root.mkdir(exist_ok=True)

    n_full = int(cfg["episode"]["n_steps"])
    train_frac = float(sweep.get("train_frac", 0.8))
    n_train = math.floor(train_frac * n_full)
    n_test = n_full - n_train
    warmup_end = int(cfg["regime"]["warmup_steps"])
    total_timesteps = int(cfg["wp4"]["total_timesteps"])

    metrics_path = out_dir / "metrics_sweep_full.csv"
    is_resume = getattr(ctx, "resume_run_id", None) is not None
    if is_resume:
        existing_rows, completed_set = load_completed_set(metrics_path)
        ctx.logger.info(
            f"Resume mode: {len(existing_rows)} existing metric rows, "
            f"{len(completed_set)} completed (seed, condition, variant) cells"
        )
    else:
        existing_rows, completed_set = [], set()

    rows = []
    trained = 0
    skipped = 0
    omit_set = set(omit_cells)

    idx_train_end = n_train + 1
    idx_test_start = n_train
    idx_test_end = n_train + n_test + 1

    for seed in seeds:
        ctx.logger.info(f"=== Seed {seed} ===")
        df_exog, _, _ = run_wp2(cfg, seed, ctx=ctx)
        sigma_clean = df_exog["sigma_hat"].to_numpy(dtype=np.float64, copy=True)
        # NOTE (calibration scope): sigma_std_post is computed from the full
        # post-warmup span (training + OOS test). The synthetic series is
        # design-stationary (sticky Markov chain on a stationary process), so
        # the std differs between train-only and full-span calibration by an
        # estimated <10%. The directional Chapter 5 findings (Cohen's dz
        # between -0.57 and -0.91 for combined < sigma_only) are robust to
        # this scale. See thesis §5.1 calibration-range caveat. Future work
        # may recalibrate using training-only data for stricter OOS hygiene.
        sigma_std_post = float(np.nanstd(sigma_clean[warmup_end:]))
        cutpoints = compute_clean_cutpoints(sigma_clean, warmup_end, n_bins=n_bins)
        deg_rng = np.random.default_rng(seed)

        seed_dir = models_root / f"seed{seed}"
        seed_dir.mkdir(exist_ok=True)

        for condition in conditions:
            sigma_deg = _build_degraded(
                condition, sigma_clean, sigma_std_post,
                alpha, k_steps, cutpoints, deg_rng,
            )
            df_full = df_exog.copy()
            df_full["sigma_hat"] = sigma_deg

            for variant in variants:
                if (condition, variant) in omit_set:
                    continue

                model_path = seed_dir / f"{condition}__{variant}.zip"
                if is_resume:
                    decision = check_cell_consistency(
                        model_exists=model_path.exists(),
                        metric_exists=(seed, condition, variant) in completed_set,
                        seed=seed, condition=condition, variant=variant,
                    )
                    if decision == "skip":
                        msg = f"RESUME-SKIP {condition}/{variant}"
                        print(msg, flush=True)
                        ctx.logger.info(msg)
                        skipped += 1
                        continue
                else:
                    if model_path.exists():
                        msg = f"SKIP {condition}/{variant} (already trained)"
                        print(msg, flush=True)
                        ctx.logger.info(msg)
                        skipped += 1
                        continue

                vflags = VARIANT_FLAGS[variant]
                use_sigma_eff = vflags["use_sigma"] and condition != "none"
                cfg_tr = copy.deepcopy(cfg)
                cfg_tr["wp3"] = {
                    **cfg_tr.get("wp3", {}),
                    "use_sigma": use_sigma_eff,
                    "regime_source": vflags["regime_source"],
                }
                cfg_tr["episode"] = {**cfg_tr["episode"], "n_steps": n_train}

                exog_train = df_full.iloc[:idx_train_end].reset_index(drop=True)
                exog_test = df_full.iloc[idx_test_start:idx_test_end].reset_index(drop=True)

                env_tr = MMEnv(cfg_tr)
                env_tr.reset(seed=seed, options={"exog": exog_train})
                vec_env = DummyVecEnv([lambda _e=Monitor(env_tr): _e])

                wp4 = cfg["wp4"]
                t_train_start = time.time()
                model = PPO(
                    "MlpPolicy", vec_env, seed=seed,
                    learning_rate=float(wp4["learning_rate"]),
                    n_steps=int(wp4["n_steps"]),
                    batch_size=int(wp4["batch_size"]),
                    n_epochs=int(wp4["n_epochs"]),
                    gamma=float(wp4["gamma"]),
                    gae_lambda=float(wp4["gae_lambda"]),
                    clip_range=float(wp4["clip_range"]),
                    ent_coef=float(wp4["ent_coef"]),
                    verbose=0,
                    device=cfg.get("wp4", {}).get("device", "cpu"),
                )
                model.learn(total_timesteps=total_timesteps)
                model.save(str(model_path))
                vec_env.close()
                train_seconds = time.time() - t_train_start

                cfg_ev = copy.deepcopy(cfg)
                cfg_ev["wp3"] = {
                    **cfg_ev.get("wp3", {}),
                    "use_sigma": use_sigma_eff,
                    "regime_source": vflags["regime_source"],
                }
                cfg_ev["episode"] = {**cfg_ev["episode"], "n_steps": n_test}

                t_eval_start = time.time()
                m = _eval_model(model, cfg_ev, exog_test, seed)
                eval_seconds = time.time() - t_eval_start

                new_row = {
                    "seed": seed,
                    "condition": condition,
                    "variant": variant,
                    "sharpe_like": m["sharpe_like"],
                    "final_equity": m["final_equity"],
                    "fill_rate": m["fill_rate"],
                    "inv_p99": m["inv_p99"],
                    "train_seconds": train_seconds,
                    "eval_seconds": eval_seconds,
                }
                rows.append(new_row)
                # Persist metric rows incrementally so a kill mid-cell never leaves
                # an orphan model.zip without its metrics row in the next resume.
                pd.DataFrame(existing_rows + rows).to_csv(metrics_path, index=False)
                trained += 1
                ctx.logger.info(
                    f"[{seed}] {condition}/{variant}: trained {train_seconds:.1f}s "
                    f"eval {eval_seconds:.1f}s sharpe={m['sharpe_like']:.4f} "
                    f"inv_p99={m['inv_p99']:.0f}"
                )

    all_rows = existing_rows + rows
    df_metrics = pd.DataFrame(all_rows)
    df_metrics.to_csv(metrics_path, index=False)
    ctx.logger.info(
        f"Wrote {metrics_path.as_posix()} "
        f"(existing={len(existing_rows)}, new={len(rows)}, total={len(all_rows)})"
    )

    md = []
    md.append("# WP6 Signal Informativeness Sweep — Full Summary")
    md.append("")
    md.append(f"Run ID: `{ctx.run_id}`")
    md.append("")
    md.append(f"- trained = {trained}")
    md.append(f"- skipped = {skipped}")
    md.append("")
    if len(df_metrics) > 0:
        ts = df_metrics["train_seconds"]
        md.append("## Cell timing distribution (train_seconds)")
        md.append("")
        md.append(f"- min: {ts.min():.1f}")
        md.append(f"- median: {ts.median():.1f}")
        md.append(f"- max: {ts.max():.1f}")
        md.append(f"- p95: {ts.quantile(0.95):.1f}")
        md.append("")
        md.append("## Per-cell mean across seeds")
        md.append("")
        agg = (df_metrics
               .groupby(["condition", "variant"], sort=False)
               [["sharpe_like", "inv_p99"]]
               .mean()
               .reset_index()
               .sort_values(["condition", "variant"]))
        md.append("| condition | variant | sharpe_like_mean | inv_p99_mean |")
        md.append("|---|---|---:|---:|")
        for _, r in agg.iterrows():
            md.append(
                f"| {r['condition']} | {r['variant']} | "
                f"{r['sharpe_like']:.4f} | {r['inv_p99']:.1f} |"
            )
        md.append("")
    md.append("## Reminder")
    md.append("")
    md.append(FULL_BANNER_NOTE)
    md.append("")
    md.append("## Next step")
    md.append("")
    md.append(
        "Review full_summary.md, then proceed with Chapter 5 analysis "
        "(per-condition × per-variant Sharpe distribution, paired-seed "
        "comparison, sensitivity check at α=0.20 / k=10)."
    )
    summary_path = out_dir / "full_summary.md"
    summary_path.write_text("\n".join(md), encoding="utf-8")
    ctx.logger.info(f"Wrote {summary_path.as_posix()}")
    print(f"Trained: {trained}   Skipped: {skipped}", flush=True)

# ============================================================
# FILE: scripts/eval_only_seed1to7.py
# PURPOSE: Active repair/evaluation script for seed 1-7 metrics with parameterized run paths.
# STATUS: support
# ============================================================

"""Eval-only script: seed 1-7 modellerini crash run'dan yükleyip OOS eval çalıştırır."""

from __future__ import annotations

import argparse
import copy
import json
import math
import sys
from pathlib import Path

import numpy as np
import pandas as pd
from stable_baselines3 import PPO

# Proje kökünü path'e ekle
ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from src.wp1.w1_as_baseline import as_deltas_ticks, compute_metrics
from src.wp1.sim import ExecParams, MarketParams
from src.wp2.synth_regime import run_wp2
from src.wp3.env import MMEnv

# ---------- Dizinler ----------
DEFAULT_MODEL_RUN = "results/runs/20260327-030624_seed1_wp5-ablation_e1545a5"
DEFAULT_OUTPUT_RUN = "results/runs/20260327-171914_seed1_wp5-ablation_e1545a5"
DEFAULT_CFG_PATH = "config/w5_main.json"

MODEL_RUN = ROOT / DEFAULT_MODEL_RUN
OUTPUT_RUN = ROOT / DEFAULT_OUTPUT_RUN
CFG_PATH = ROOT / DEFAULT_CFG_PATH

SEEDS = list(range(1, 8))  # seed 1-7

VARIANTS = {
    "sigma_only":  {"use_sigma": True,  "regime_source": "none"},
    "regime_only": {"use_sigma": False, "regime_source": "hat"},
    "combined":    {"use_sigma": True,  "regime_source": "hat"},
    "oracle_pure": {"use_sigma": False, "regime_source": "true"},
    "oracle_full": {"use_sigma": True,  "regime_source": "true"},
}


def _root_relative_path(path_value) -> Path:
    path = Path(path_value)
    return path if path.is_absolute() else ROOT / path


def _parse_args(argv=None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--model-run",
        default=DEFAULT_MODEL_RUN,
        help="Run directory containing the seed 1-7 trained models.",
    )
    parser.add_argument(
        "--output-run",
        default=DEFAULT_OUTPUT_RUN,
        help="Run directory containing/writing the seed 1-7 and combined metrics.",
    )
    parser.add_argument(
        "--cfg-path",
        default=DEFAULT_CFG_PATH,
        help="WP5 config path, relative to repo root unless absolute.",
    )
    return parser.parse_args(argv)


def _compute_regime_metrics(equity, inv, fills, fees, regime_labels, dt):
    n = len(fills)
    results = []
    for regime in sorted(set(regime_labels)):
        mask = np.array([regime_labels[t] == regime for t in range(n)])
        steps_count = int(mask.sum())
        if steps_count < 10:
            continue
        step_pnl = np.diff(equity)[mask]
        mean_step_pnl = float(step_pnl.mean()) if len(step_pnl) else 0.0
        std_pnl = float(step_pnl.std(ddof=1)) if len(step_pnl) > 1 else 0.0
        sharpe_like = (mean_step_pnl / std_pnl * np.sqrt(1.0 / dt)) if std_pnl > 0 else 0.0
        fill_rate = float(fills[mask].sum() / steps_count) if steps_count > 0 else 0.0
        inv_vals = np.abs(inv[:-1][mask])
        inv_p99 = float(np.quantile(inv_vals, 0.99)) if len(inv_vals) else 0.0
        results.append({
            "regime": regime,
            "mean_step_pnl": mean_step_pnl,
            "sharpe_like": sharpe_like,
            "fill_rate": fill_rate,
            "inv_p99": inv_p99,
            "steps_count": steps_count,
        })
    return results


def main(argv=None):
    args = _parse_args(argv)
    model_run = _root_relative_path(args.model_run)
    output_run = _root_relative_path(args.output_run)
    cfg_path = _root_relative_path(args.cfg_path)

    with open(cfg_path) as f:
        cfg = json.load(f)

    # Override seeds to [1..20] for correct exog generation params
    cfg["wp5"]["seeds"] = list(range(1, 21))

    market = MarketParams(**cfg["market"])
    execp = ExecParams(**cfg["exec"])
    wp5 = cfg["wp5"]
    train_frac = float(wp5["train_frac"])
    naive_h = int(wp5["naive"]["h"])
    naive_m = int(wp5["naive"]["m"])
    n_full = int(cfg["episode"]["n_steps"])
    n_train = math.floor(train_frac * n_full)
    n_test = n_full - n_train

    rows_oos = []
    rows_regime = []

    for seed in SEEDS:
        print(f"=== Eval seed {seed} ===")

        # 1) Exog series (aynı seed -> aynı veri)
        df_exog, _, _ = run_wp2(cfg, seed)

        # 2) Split
        exog_test = df_exog.iloc[n_train: n_train + n_test + 1].reset_index(drop=True)

        # 3) Load pre-trained models
        models = {}
        for name in VARIANTS:
            model_path = model_run / "models" / f"seed{seed}" / f"ppo_{name}"
            models[name] = PPO.load(str(model_path), device="cpu")
            print(f"  Loaded {model_path.name}")

        # 4) Build strategies
        def _base_eval_cfg():
            c = copy.deepcopy(cfg)
            c["episode"] = {**c["episode"], "n_steps": n_test}
            c["wp3"]["use_sigma"] = True
            c["wp3"]["regime_source"] = "hat"
            c["as"]["horizon_steps"] = n_test
            return c

        strategies = {
            "naive": (_base_eval_cfg(), None),
            "AS":    (_base_eval_cfg(), None),
        }
        for vname, vcfg in VARIANTS.items():
            cfg_ev = copy.deepcopy(cfg)
            cfg_ev["episode"] = {**cfg_ev["episode"], "n_steps": n_test}
            cfg_ev["wp3"]["use_sigma"] = vcfg["use_sigma"]
            cfg_ev["wp3"]["regime_source"] = vcfg["regime_source"]
            cfg_ev["as"]["horizon_steps"] = n_test
            strategies[f"ppo_{vname}"] = (cfg_ev, models[vname])

        # 5) Eval loop
        for strat_name, (cfg_ev, model_ev) in strategies.items():
            env_ev = MMEnv(cfg_ev)
            obs, _ = env_ev.reset(seed=seed, options={"exog": exog_test})

            eq = np.zeros(n_test + 1)
            iv = np.zeros(n_test + 1, dtype=int)
            fl = np.zeros(n_test, dtype=int)
            fe = np.zeros(n_test)
            eq[0] = env_ev._state.equity
            iv[0] = env_ev._state.inv

            for t in range(n_test):
                if strat_name == "naive":
                    action = np.array([naive_h - 1, naive_m + 2])
                elif strat_name == "AS":
                    db, da = as_deltas_ticks(
                        env_ev._state.mid, env_ev._state.inv, t, cfg_ev, market, execp,
                    )
                    h_as = int(np.clip((db + da) // 2, 1, 5))
                    m_as = int(np.clip((db - da) // 2, -2, 2))
                    action = np.array([h_as - 1, m_as + 2])
                else:
                    action, _ = model_ev.predict(obs, deterministic=True)

                obs, _r, _term, _trunc, info = env_ev.step(action)
                eq[t + 1] = info["equity"]
                iv[t + 1] = info["inv"]
                fl[t] = info["fills"]
                fe[t] = float(info.get("fee_total", 0.0))

            # Metrics
            m = compute_metrics(eq, iv, fl, dt=market.dt)
            total_fees = float(fe.sum())
            n_trades = int(fl.sum())
            fee_per_trade = total_fees / n_trades if n_trades > 0 else 0.0

            row = {
                "seed": seed, "strategy": strat_name, "split": "test",
                **m, "total_fees": total_fees, "fee_per_trade": fee_per_trade,
            }
            rows_oos.append(row)

            # Regime metrics
            if "regime_true" in exog_test.columns:
                regime_labels = list(exog_test["regime_true"].values[:n_test + 1])
            else:
                regime_labels = ["M"] * (n_test + 1)

            rw = _compute_regime_metrics(eq, iv, fl, fe, regime_labels, market.dt)
            for r in rw:
                rows_regime.append({"seed": seed, "strategy": strat_name, **r})

            print(f"  {strat_name}: equity={m['final_equity']:.4f} sharpe={m['sharpe_like']:.4f}")

    # 6) Save seed 1-7 CSVs
    df_oos = pd.DataFrame(rows_oos)
    df_regime = pd.DataFrame(rows_regime)

    oos_path = output_run / "metrics_wp5_oos_seed1to7.csv"
    regime_path = output_run / "metrics_wp5_oos_by_regime_seed1to7.csv"
    df_oos.to_csv(oos_path, index=False)
    df_regime.to_csv(regime_path, index=False)
    print(f"\nSaved: {oos_path}  ({len(df_oos)} rows)")
    print(f"Saved: {regime_path}  ({len(df_regime)} rows)")

    # 7) Merge with seed 8-20
    df_oos_8to20 = pd.read_csv(output_run / "metrics_wp5_oos.csv")
    df_regime_8to20 = pd.read_csv(output_run / "metrics_wp5_oos_by_regime.csv")

    df_combined = pd.concat([df_oos, df_oos_8to20], ignore_index=True).sort_values(
        ["seed", "strategy"]
    ).reset_index(drop=True)

    df_regime_combined = pd.concat([df_regime, df_regime_8to20], ignore_index=True).sort_values(
        ["seed", "strategy", "regime"]
    ).reset_index(drop=True)

    combined_path = output_run / "metrics_wp5_oos_combined.csv"
    regime_combined_path = output_run / "metrics_wp5_oos_by_regime_combined.csv"
    df_combined.to_csv(combined_path, index=False)
    df_regime_combined.to_csv(regime_combined_path, index=False)

    print(f"\n=== Combined ===")
    print(f"metrics_wp5_oos_combined.csv: {len(df_combined)} rows (expected 140)")
    print(f"metrics_wp5_oos_by_regime_combined.csv: {len(df_regime_combined)} rows (expected ~420)")

    # Doğrulama
    seeds_in = sorted(df_combined["seed"].unique())
    strats_in = sorted(df_combined["strategy"].unique())
    print(f"Seeds: {seeds_in}")
    print(f"Strategies: {strats_in}")
    print(f"Seeds x Strategies = {len(seeds_in)} x {len(strats_in)} = {len(seeds_in) * len(strats_in)}")


if __name__ == "__main__":
    main()

# ============================================================
# FILE: scripts/gen_thesis_28.py
# PURPOSE: Current thesis_28 DOCX generator.
# STATUS: support
# ============================================================

"""Generate thesis_28.docx.
Based on thesis_27 with Lane-A C1/C2/C3 caveat additions:
  (1) §3.3 footnote on rv_dwell offline non-causality;
  (2) §3.8 methodology clarification — all reported PPO numbers come
      from WP5 (70/30 split); WP4 is in-sample pilot only;
  (3) §5.1 calibration-range caveat extension on noisy α full-span vs
      training-only calibration (<10% scale shift, directional findings
      robust).
Lane-B B1-B3 fixes (Appendix Z removal, 5x5x20 arithmetic correction,
α direction in Future Work) are inherited from gen_thesis_27.py edits
already applied.
Cover-page marker updated to Sürüm 28.
Decisions log NOT updated in this round — Decision #49 will be logged
once Lane A and Lane B are both fully complete."""

import copy, sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

dst = Document()

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_normal(doc, text):
    p = doc.add_paragraph(text)
    return p

def add_table_2col(doc, title, rows):
    """2-column table with header row."""
    t = doc.add_table(rows=1 + len(rows), cols=2, style='Table Grid')
    t.rows[0].cells[0].text = rows[0][0] if isinstance(title, list) else title.split("|")[0]
    t.rows[0].cells[1].text = rows[0][1] if isinstance(title, list) else title.split("|")[1]
    # header bold
    for cell in t.rows[0].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for i, (c1, c2) in enumerate(rows[1:] if isinstance(title, list) else rows, start=1):
        t.rows[i].cells[0].text = c1
        t.rows[i].cells[1].text = c2
    return t

def add_table(doc, headers, rows):
    """General n-column table."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Table Grid')
    for j, h in enumerate(headers):
        t.rows[0].cells[j].text = h
        for p in t.rows[0].cells[j].paragraphs:
            for r in p.runs:
                r.bold = True
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            t.rows[i + 1].cells[j].text = str(val)
    return t

# ================================================================
# COVER PAGE  — CHANGE A: "Nisan 2026 — Surum 18"
# ================================================================
p = dst.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Karlsruhe Institute of Technology").font.size = Pt(14)
p = dst.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Financial Engineering MSc Programı")
dst.add_paragraph()
p = dst.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Yüksek Lisans Tez Taslağı").font.size = Pt(12)
dst.add_paragraph()
p = dst.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Farklı Volatilite Rejimleri Altında Pekiştirmeli Öğrenme ile Yüksek Frekanslı Piyasa Yapıcılığı")
r.bold = True; r.font.size = Pt(14)
dst.add_paragraph()
p = dst.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Onur Emiroğlu")
p = dst.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Financial Engineering MSc Programı")
p = dst.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Nisan 2026 — Sürüm 28")

# ================================================================
# 1. OZET — CHANGE B: son paragraf degistirildi, CHANGE C: anahtar kelimeler
# ================================================================
add_heading(dst, "1. ÖZET", level=1)

add_normal(dst, (
    "Bu çalışmada, volatilite rejim bilgisinin (Düşük/Orta/Yüksek) pekiştirmeli öğrenme tabanlı piyasa "
    "yapıcı ajanların performansına katkısı araştırılmaktadır. Proximal Policy Optimization (PPO) "
    "algoritması kullanılarak beş ajan varyantı eğitilmiştir: yalnızca sigma (ppo_sigma_only), yalnızca "
    "rejim (ppo_regime_only), sigma + tahmini rejim (ppo_combined), yalnızca gerçek rejim (ppo_oracle_pure) "
    "ve sigma + gerçek rejim (ppo_oracle_full). Ajanlar, Poisson varış süreçli yoğunluk tabanlı dolum modeli "
    "ve Markov zinciri rejim geçişleri içeren tamamen sentetik bir piyasa ortamında değerlendirilmiştir."
))

add_normal(dst, (
    "Yürütülen out-of-sample deney (20 bağımsız tohum, 1 milyon zaman adımı eğitim, %70/%30 "
    "kronolojik eğitim/test bölünmesi), tüm PPO varyantlarının Sharpe oranı açısından klasik referans "
    "stratejilerini (Avellaneda-Stoikov ve naif sabit-spread) yaklaşık 6-7 kat geçtiğini ortaya koymaktadır. "
    "Buna karşın, rejim bilgisinin PPO performansına net bir katkı sağladığı hipotezi yalnızca istatistiksel "
    "olarak desteklenmemekle kalmamış; ppo_sigma_only ile ppo_oracle_full'un Sharpe performansları "
    "TOST eşdeğerlik testi ile pratik olarak eşdeğer bulunmuştur (TOST p<0.05, ±0.10 sınır)."
))

add_normal(dst, (
    "Ana Bulgu (Ablasyon): oracle_full ajanı (sigma + regime_true), mükemmel rejim etiketlerine sahip olmasına "
    "karşın ppo_sigma_only'yi istatistiksel olarak anlamlı biçimde geçememiştir (Sharpe: 0.753 vs 0.722, "
    "p = 0.115). Bu sonuç, rejim etiketinin — gerçek ya da tahmini — sigma_hat'in zaten taşıdığı bilginin "
    "üzerine anlamlı katkı sağlamadığı hipoteziyle güçlü biçimde tutarlıdır."
))

# CHANGE B: two new paragraphs replacing the old "Detector robustness deneyi" paragraph
add_normal(dst, (
    "Detector robustness deneyi — rv_baseline (%60.7), rv_dwell (%60.4) ve HMM (%81.8) dedektörleri ile "
    "20 tohumda yinelenen analiz — null sonucun dedektör seçiminden bağımsız olduğunu doğrulamaktadır "
    "(tüm dedektörler için Sharpe bazlı p > 0.08). Oracle deneyi ise %100 doğrulukla eşdeğer bir test "
    "sunmaktadır: ppo_oracle_full mükemmel rejim bilgisiyle dahi sigma_only'yi geçememiştir (p=0.115)."
))

add_normal(dst, (
    "Model misspecification robustness deneyi kapsamında fill model parametreleri A ve k rejime bağlı "
    "kılınmış (L: A=4.0/k=1.8, M: A=5.0/k=1.5, H: A=6.0/k=1.2) ve 5-varyant değerlendirmesi bu "
    "zorlaştırılmış ortamda yinelenmiştir. ppo_sigma_only yine en yüksek Sharpe değerini elde etmiş "
    "(0.686); oracle_full ile aradaki fark istatistiksel olarak anlamsız kalmıştır (p=0.881). Bu bulgu, "
    "signal redundancy argümanının çevresel koşullara karşı da sağlam olduğunu göstermektedir."
))

# thesis_27 addition: WP6 one-sentence mention in abstract
add_normal(dst, (
    "Sinyal informativeness sweep'i (WP6), kategorik regime label'larının "
    "sigma_hat sürekli sinyali zayıfladığında dahi performansı iyileştirmediğini, "
    "hatta birlikte kullanıldığında yönsel olarak düşürdüğünü göstermiştir."
))

# CHANGE C: updated keywords
add_normal(dst, (
    "Anahtar kelimeler: piyasa yapıcılığı, pekiştirmeli öğrenme, PPO, volatilite rejimleri, "
    "Avellaneda-Stoikov, ablasyon, oracle agent, sinyal yedekliliği, model yanlış belirleme"
))

# ================================================================
# SEMBOLLER VE KISALTMALAR
# ================================================================
add_heading(dst, "SEMBOLLER VE KISALTMALAR", level=1)

add_normal(dst, (
    "Bu bölümde çalışma boyunca kullanılan matematiksel semboller, Yunan harfleri ve kısaltmalar "
    "tanımlanmaktadır."
))

# Matematiksel Semboller
add_heading(dst, "Matematiksel Semboller", level=2)
math_symbols = [
    ["Sembol", "Açıklama"],
    ["M_t", "t anındaki mid-price (orta fiyat) [para birimi]"],
    ["σ", "Volatilite parametresi — mid-price'ın birim zamandaki standart sapması [tick/√saniye]"],
    ["σ_base", "Temel volatilite parametresi; Orta (M) rejim için referans değer (0.8 tick/√saniye)"],
    ["σ_L, σ_M, σ_H", "Düşük, Orta ve Yüksek rejimlere karşılık gelen volatilite değerleri"],
    ["Δt", "Simülasyon zaman adımı büyüklüğü (0.2 saniye)"],
    ["z", "Standart normal rassal değişken; z ~ N(0,1)"],
    ["λ(δ)", "δ tick uzaklığındaki kotasyon için Poisson dolum yoğunluğu [dolum/saniye]"],
    ["A", "Sıfır uzaklıkta (δ=0) temel dolum yoğunluğu (5.0 dolum/saniye)"],
    ["k", "Dolum yoğunluğunun uzaklıkla üstel azalma hızı (1.5 tick başına)"],
    ["δ", "Kotasyon uzaklığı — bid veya ask fiyatının mid-price'tan tick cinsinden sapması"],
    ["δ_bid", "Bid (alış) kotasyonunun mid-price'tan uzaklığı [tick]"],
    ["δ_ask", "Ask (satış) kotasyonunun mid-price'tan uzaklığı [tick]"],
    ["P_fill", "Bir zaman adımında en az bir dolumun gerçekleşme olasılığı"],
    ["h", "Half-spread — kotasyon yarı-aralığı [tick]; h ∈ {1, 2, 3, 4, 5}"],
    ["m", "Skew — asimetrik kotasyon kayması [tick]; m ∈ {-2, -1, 0, +1, +2}"],
    ["q", "Envanter (inventory) — anda tutulan net pozisyon miktarı [lot]"],
    ["q_norm", "Normalize edilmiş envanter; q_norm = q / inv_max_clip ∈ [-1, +1]"],
    ["inv_max_clip", "Envanter kırpma sınırı (50 lot)"],
    ["τ", "Kalan zaman fraksiyonu; τ = (T - t) / T ∈ [0, 1]"],
    ["T", "Episode toplam adım sayısı (8000 adım)"],
    ["R_t", "t adımındaki ödül (reward)"],
    ["η", "Envanter ceza katsayısı; η = 0.001"],
    ["γ", "Avellaneda-Stoikov risk aversion parametresi; γ = 0.01"],
    ["r", "Avellaneda-Stoikov rezervasyon fiyatı"],
    ["d", "Avellaneda-Stoikov optimal yarı-spread"],
    ["k_price", "Fiyat cinsinden ifade edilen yoğunluk azalma parametresi; k_price = k / tick_size"],
    ["P", "Markov zinciri geçiş matrisi (3×3)"],
    ["Sharpe", "Sharpe oranı — risk-düzeltmeli getiri ölçütü; Sharpe = (μ/σ) × √(1/Δt)"],
    ["inv_p99", "Mutlak envanter değerinin 99. persantili [lot]"],
    ["equity", "Toplam varlık değeri; equity = cash + q × M_t"],
]
add_table(dst, math_symbols[0], math_symbols[1:])

# Yunan Harfleri
add_heading(dst, "Yunan Harfleri", level=2)
greek = [
    ["Sembol", "Açıklama"],
    ["γ (gamma)", "Risk aversion parametresi (Avellaneda-Stoikov modelinde)"],
    ["η (eta)", "Envanter ceza katsayısı (ödül fonksiyonunda)"],
    ["σ (sigma)", "Volatilite — fiyat değişiminin standart sapması"],
    ["τ (tau)", "Kalan zaman fraksiyonu"],
    ["λ (lambda)", "Poisson dolum yoğunluğu"],
    ["δ (delta)", "Kotasyon uzaklığı (mid-price'tan tick cinsinden sapma)"],
    ["μ (mu)", "Ortalama adım başına PnL (Sharpe hesabında)"],
]
add_table(dst, greek[0], greek[1:])

# Kısaltmalar
add_heading(dst, "Kısaltmalar", level=2)
abbrevs = [
    ["Kısaltma", "Açıklama"],
    ["PPO", "Proximal Policy Optimization — yakınsak politika optimizasyonu algoritması"],
    ["RL", "Reinforcement Learning — pekiştirmeli öğrenme"],
    ["AS", "Avellaneda-Stoikov — klasik stokastik kontrol tabanlı piyasa yapıcılığı modeli"],
    ["HFT", "High-Frequency Trading — yüksek frekanslı alım-satım"],
    ["OOS", "Out-of-Sample — modelin eğitilmediği veri üzerinde değerlendirme"],
    ["WP", "Work Package — iş paketi (proje aşaması)"],
    ["ABM", "Arithmetic Brownian Motion — aritmetik Brownian hareketi"],
    ["RV", "Realized Volatility — gerçekleşmiş volatilite (kayan pencere ile hesaplanan)"],
    ["PnL", "Profit and Loss — kar ve zarar"],
    ["CUDA", "Compute Unified Device Architecture — NVIDIA GPU hesaplama mimarisi"],
    ["SB3", "Stable-Baselines3 — PPO implementasyonu için kullanılan Python kütüphanesi"],
    ["MlpPolicy", "Multi-Layer Perceptron Policy — çok katmanlı algılayıcı politika ağı"],
    ["GAE", "Generalized Advantage Estimation — genelleştirilmiş avantaj tahmini (PPO bileşeni)"],
    ["L / M / H", "Low / Medium / High — Düşük / Orta / Yüksek volatilite rejimleri"],
]
add_table(dst, abbrevs[0], abbrevs[1:])

# ================================================================
# TERİMLER SÖZLÜĞÜ
# ================================================================
add_heading(dst, "TERİMLER SÖZLÜĞÜ", level=1)
add_normal(dst, (
    "Bu sözlük, çalışmada kullanılan teknik terimleri konuya aşina olmayan okuyucular için "
    "tanımlamaktadır. Terimler alfabetik sıraya göre düzenlenmiştir."
))

glossary = [
    ["Terim", "Tanım"],
    ["Ablation (Ablasyon)", "Bir modelin belirli bir bileşeninin çıkarılarak ya da devre dışı bırakılarak etkisinin ölçülmesi amacıyla yapılan kontrollü deney."],
    ["Action Space (Eylem Uzayı)", "Bir pekiştirmeli öğrenme ajanının her adımda seçebileceği tüm olası eylemlerin kümesi. Bu çalışmada half-spread (h) ve skew (m) değerlerinden oluşan ayrık bir yapıdadır."],
    ["Adverse Selection (Ters Seçim)", "Piyasa yapıcının bilgili yatırımcılara karşı dezavantajlı konuma düşmesi durumu."],
    ["Agent (Ajan)", "Pekiştirmeli öğrenmede, bir ortamda gözlem yaparak karar veren ve ödül sinyaline göre davranışını güncelleyen öğrenen sistem."],
    ["Ask (Satış Fiyatı)", "Piyasa yapıcının satmaya hazır olduğu fiyat. Her zaman bid fiyatının üzerindedir."],
    ["Bid (Alış Fiyatı)", "Piyasa yapıcının almaya hazır olduğu fiyat. Her zaman ask fiyatının altındadır."],
    ["Bid-Ask Spread", "Alış ve satış fiyatı arasındaki fark. Piyasa yapıcının temel gelir kaynağıdır."],
    ["Brownian Motion", "Rastlantısal yürüyüş modellerinin sürekli-zaman limiti. ABM'de fiyat değişimleri normal dağılımdan çekilmektedir."],
    ["Clip (Kırpma)", "Bir değerin belirli bir aralıkla sınırlandırılması işlemi. Bu çalışmada envanter ±50 lot ile kırpılmaktadır."],
    ["Convergence (Yakınsama)", "Pekiştirmeli öğrenme eğitiminde politikanın optimal ya da kararlı bir çözüme ulaşması süreci."],
    ["Entropy Coefficient", "PPO'da keşif-sömürü dengesini ayarlayan hiperparametre. Bu çalışmada ent_coef=0.01 kullanılmıştır."],
    ["Episode", "Pekiştirmeli öğrenmede başlangıçtan bitiş koşuluna kadar süren bir tam simülasyon koşusu. Bu çalışmada 8000 zaman adımından oluşmaktadır."],
    ["Exogenous Series (Dışsal Seri)", "Modelin dışında önceden üretilmiş ve tüm stratejilere aynı biçimde sunulan fiyat/volatilite serisi."],
    ["Fill (Dolum)", "Piyasa yapıcının kotasyon verdiği fiyattan gerçekleşen emir eşleşmesi."],
    ["Fill Rate (Dolum Oranı)", "Toplam adım sayısına bölünen dolum sayısı."],
    ["Gymnasium", "OpenAI tarafından geliştirilen, pekiştirmeli öğrenme ortamları için standart Python arayüzü."],
    ["Half-Spread (Yarı-Aralık)", "Bid veya ask fiyatının mid-price'tan uzaklığı. Bu çalışmada h sembolüyle gösterilir."],
    ["Inventory (Envanter)", "Piyasa yapıcının anda elinde bulundurduğu net pozisyon miktarı."],
    ["Inventory Risk (Envanter Riski)", "Piyasa yapıcının büyük envanter pozisyonu tutarken fiyat hareketleri nedeniyle maruz kaldığı zarar riski."],
    ["Latency (Gecikme)", "Piyasa yapıcının kotasyon güncellemesi ile piyasanın bu güncellemeyi işlemesi arasındaki zaman farkı."],
    ["Limit Order Book", "Bekleme sırasındaki tüm alış ve satış emirlerinin fiyat-miktar listesi."],
    ["Markov Chain", "Gelecekteki durumun yalnızca mevcut duruma bağlı olduğu olasılıksal süreç."],
    ["Market Making (Piyasa Yapıcılığı)", "Finansal piyasalarda sürekli alış ve satış kotasyonu vererek likidite sağlama faaliyeti."],
    ["Mid-Price", "Anlık alış (bid) ve satış (ask) fiyatlarının ortalaması."],
    ["Model Misspecification (Model Yanlış Belirleme)", "Modelin varsayımlarının gerçek veri üretim sürecinden sapması durumu. Bu çalışmada fill parametrelerinin rejime bağlı kılınmasıyla test edilmiştir."],
    ["Null Result", "Araştırma hipotezinin verilerce desteklenmediği bulgu."],
    ["Observation Space (Gözlem Uzayı)", "Ajanın her adımda çevreden aldığı bilginin vektör temsili. Bu çalışmada 6 boyutludur: [q_norm, σ̂, τ, r_L, r_M, r_H]."],
    ["One-Hot Encoding", "Kategorik bir değişkeni ikili vektörle temsil etme yöntemi."],
    ["Out-of-Sample (OOS)", "Modelin eğitilmediği, bağımsız bir veri kümesi üzerinde yapılan değerlendirme."],
    ["Percentile (Persantil)", "Veri setinin belirli bir yüzdesinin altında kaldığı değer."],
    ["Policy (Politika)", "Pekiştirmeli öğrenmede ajanın gözleme göre eylem seçme stratejisi."],
    ["Poisson Process", "Olayların bağımsız ve sabit ortalama hızda rastlantısal gerçekleştiği olasılıksal süreç."],
    ["PPO", "John Schulman ve ark. (2017) tarafından geliştirilen politika gradyanı tabanlı RL algoritması."],
    ["Regime (Rejim)", "Piyasanın belirli bir volatilite seviyesinde bulunduğu dönem. Bu çalışmada L/M/H."],
    ["Risk-Adjusted Return", "Getirinin, üstlenilen risk miktarına bölünmesiyle elde edilen performans ölçütü."],
    ["Seed (Tohum)", "Rastlantısal sayı üretecinin başlangıç değeri. Tekrarlanabilirlik için kullanılır."],
    ["Sharpe Ratio", "Birim risk başına elde edilen fazla getiriyi ölçen performans metriği."],
    ["Skew (Asimetrik Kotasyon)", "Bid ve ask kotasyonlarının mid-price'a göre eşit olmayan biçimde kaydırılması."],
    ["Stable-Baselines3 (SB3)", "PPO dahil çeşitli RL algoritmalarının güvenilir PyTorch implementasyonlarını içeren Python kütüphanesi."],
    ["Sticky Transition Matrix", "Köşegen elemanları yüksek Markov geçiş matrisi. Rejimlerin uzun süre devam etmesini sağlar."],
    ["Tick", "Finansal piyasalarda fiyatın alabileceği en küçük artış birimi. Bu çalışmada tick_size = 0.01."],
    ["Timestep (Zaman Adımı)", "Simülasyonun bir sonraki duruma geçtiği en küçük zaman birimi. Δt = 0.2 saniye."],
    ["Chronological Train/Test Split", "Zaman serisi verisinin kronolojik sıraya göre eğitim (%70) ve test (%30) kümelerine bölünmesi."],
]
add_table(dst, glossary[0], glossary[1:])

# ================================================================
# 2. GİRİŞ
# ================================================================
add_heading(dst, "2. GİRİŞ", level=1)

add_normal(dst, (
    "Piyasa yapıcılığı (market making), finansal piyasalarda likidite sağlayan ve sürekli olarak hem alış (bid) "
    "hem de satış (ask) kotasyonu veren ajanların stratejik davranışını inceleyen bir alandır. Piyasa yapıcılar, "
    "spread geliri elde etmeye çalışırken olumsuz fiyat hareketlerine maruz kalan envanter riskini de "
    "yönetmek zorundadır. Bu denge, özellikle volatilitenin değişken olduğu yüksek frekanslı piyasalarda "
    "kritik önem taşımaktadır."
))

add_normal(dst, (
    "Klasik yaklaşımlar arasında Avellaneda ve Stoikov (2008) tarafından geliştirilen stokastik kontrol "
    "çerçevesi öne çıkmaktadır. Bu model, envanter riskini rezervasyon fiyatı kavramıyla içselleştirerek "
    "optimal bid-ask spread'ini analitik olarak türetmektedir. Ancak bu yaklaşım, sabit bir volatilite "
    "varsayımına dayanmakta ve piyasa rejimlerindeki değişimlere uyum sağlayamamaktadır."
))

add_normal(dst, (
    "Bu çalışmanın temel araştırma sorusu şudur: Volatilite rejim bilgisine erişimi olan bir PPO ajanı, "
    "bu bilgiden yoksun olan eşdeğerine kıyasla daha iyi performans sergiler mi? Bu soruyu yanıtlamak için "
    "aşağıdaki katkılar sunulmaktadır:"
))

p = dst.add_paragraph("(1) Markov zinciri rejim geçişleri ve Poisson tabanlı dolum modeli içeren kontrollü sentetik bir piyasa ortamı tasarlanmıştır.", style='List Paragraph')
p = dst.add_paragraph("(2) Rejim farkındalığı ablasyonu gerçekleştirilmiş; 5 PPO varyantı 20 bağımsız tohum ile OOS protokolüyle karşılaştırılmıştır.", style='List Paragraph')
p = dst.add_paragraph("(3) Oracle agent deneyi ile mükemmel rejim bilgisinin üst sınır değeri ölçülmüştür.", style='List Paragraph')
p = dst.add_paragraph("(4) Detector robustness analizi gerçekleştirilmiş; null sonucun dedektör seçiminden bağımsızlığı doğrulanmıştır.", style='List Paragraph')

# ================================================================
# 3. TEORİ VE METODOLOJİ
# ================================================================
add_heading(dst, "3. TEORİ VE METODOLOJİ", level=1)

# 3.1 Piyasa Modeli
add_heading(dst, "3.1 Piyasa Modeli", level=2)
add_normal(dst, (
    "Piyasa fiyatı, ayrık zamanlı aritmetik Brownian hareketi (ABM) ile modellenmektedir:"
))
add_normal(dst, "M_{t+1} = M_t + σ_r · √(Δt) · z,    z ~ N(0,1)")
add_normal(dst, (
    "Burada M_t mid-price, σ_r rejime bağlı volatilite parametresi (ticks/√saniye cinsinden), "
    "Δt = 0.2 saniye zaman adımı büyüklüğüdür. Başlangıç fiyatı M_0 = 100.0, tick büyüklüğü 0.01'dir."
))
add_normal(dst, (
    "Temel volatilite parametresi σ_base = 0.8 tick olarak belirlenmiş; üç rejim için çarpanlar "
    "[0.6, 1.0, 1.8] olarak tanımlanmıştır. Buna göre σ_L = 0.48, σ_M = 0.80, σ_H = 1.44 tick/√saniye "
    "değerleri elde edilmektedir."
))

# 3.2 Dolum Modeli
add_heading(dst, "3.2 Dolum Modeli", level=2)
add_normal(dst, (
    "Emir dolumları, yoğunluk tabanlı Poisson süreci ile modellenmektedir. Delta tick cinsinden ifade "
    "edildiğinde, dolum yoğunluğu:"
))
add_normal(dst, "λ(δ) = A · exp(-k · δ)")
add_normal(dst, (
    "formülüyle hesaplanmaktadır. Burada A = 5.0 (delta=0'daki temel yoğunluk, dolum/saniye) ve k = 1.5 "
    "(tick başına üstel azalma) parametreleridir. Bir zaman adımında en az bir dolumun gerçekleşme olasılığı:"
))
add_normal(dst, "P_fill = 1 - exp(-λ · Δt)")
add_normal(dst, (
    "olarak verilmektedir. Hem bid hem ask tarafı için bağımsız Bernoulli denemeleri gerçekleştirilmektedir."
))
add_normal(dst, (
    "Gecikme modeli (latency_steps = 1), ajanın kotasyon hesaplamasında bir önceki adımın mid-price "
    "değerini kullanmasına yol açmaktadır; bu durum ters seçim (adverse selection) riskini "
    "içselleştirmektedir. Komisyon ücreti her işlem için 0.2 baz puan olarak uygulanmaktadır."
))

# 3.3 Rejim Modeli
add_heading(dst, "3.3 Rejim Modeli", level=2)
add_normal(dst, (
    "Volatilite rejimleri, üç durumlu (L: Düşük, M: Orta, H: Yüksek) birinci derece Markov zinciri ile "
    "üretilmektedir. Geçiş matrisi yapışkan (sticky) olacak şekilde tasarlanmıştır:"
))
add_normal(dst, "P = [[0.9967, 0.0023, 0.0010], [0.0042, 0.9917, 0.0041], [0.0010, 0.0030, 0.9960]]")
add_normal(dst, (
    "Beklenen rejim süreleri sırasıyla L: ~300, M: ~120, H: ~250 zaman adımıdır. Bu yapışkanlık, kayan "
    "gerçekleşmiş volatilite (rolling realized volatility — RV) ile rejim tespitinin mümkün olmasını "
    "sağlamaktadır."
))
add_normal(dst, (
    "Rejim tespiti için RV pencere boyutu 50 adım, ısınma süresi (warmup) 1000 adım olarak belirlenmiştir. "
    "Eşik değerleri ısınma döneminin 33. ve 66. persentillerinden kalibre edilmektedir. Gerçek zamanlı "
    "tespitte look-ahead bias engellenmiştir; regime_hat yalnızca geçmiş veriden hesaplanmaktadır. Elde "
    "edilen tespit doğruluğu %60.7'dir (rastgele sınır: %33.3)."
))
add_normal(dst, (
    "Detector robustness analizi kapsamında iki ek tespit yöntemi karşılaştırılmıştır: "
    "(i) dwell filtreli RV (rv_dwell): kısa süreli (< 5 adım) rejim geçişlerini düzelten filtre uygulanmış, "
    "tespit doğruluğu %60.4 olarak elde edilmiştir; "
    "(ii) Gaussian HMM (hmm): sigma_hat serisi üzerinde eğitilen üç durumlu Gaussian gizli Markov modeli, "
    "%81.8 doğruluk ile en yüksek tespit kalitesini sağlamıştır."
))
add_normal(dst, (
    "Not: rv_dwell offline bir dwell smoothing kullanır ve etiket kararı verilirken "
    "sınırlı bir ileri tarama içerir (en fazla min_dwell adım); bu nedenle rv_dwell "
    "doğruluk değeri kausal bir üst sınır olarak okunmalıdır. Bu sınırlama yalnızca §3.3 "
    "ve §4.6'da raporlanan rv_dwell metriklerini etkiler; tezde raporlanan tüm ana boru "
    "hatları — WP4 pilot, WP5 OOS değerlendirmesi ve WP6 sweep — kausal rv_baseline "
    "detector'ı kullandığından, Bulgu 1-7 bu sınırlamadan etkilenmemiştir."
))

# 3.4 Gymnasium Ortamı
add_heading(dst, "3.4 Gymnasium Ortamı", level=2)
add_normal(dst, (
    "OpenAI Gymnasium uyumlu MMEnv ortamı oluşturulmuştur. Gözlem uzayı 6 boyutludur:"
))
add_normal(dst, "obs = [q_norm, σ̂_t, τ, r_L, r_M, r_H]")
add_normal(dst, (
    "Burada q_norm = clip(inv, -50, 50) / inv_max_clip envanter normalizasyonu, σ̂_t kayan gerçekleşmiş "
    "volatilite, τ = (T-t)/T kalan zaman fraksiyonu, r_L/r_M/r_H ise rejim one-hot kodlamasıdır "
    "(ppo_blind için hepsi sıfır)."
))
add_normal(dst, "Eylem uzayı MultiDiscrete([5, 5]) olarak tanımlanmıştır:")
add_normal(dst, "h_idx ∈ {0,1,2,3,4} → h = h_idx + 1 tick (half-spread)")
add_normal(dst, "m_idx ∈ {0,1,2,3,4} → m = m_idx - 2 tick (skew)")
add_normal(dst, "δ_bid = max(1, h + m)")
add_normal(dst, "δ_ask = max(1, h - m)")

# 3.5 Referans Stratejiler
add_heading(dst, "3.5 Referans Stratejiler", level=2)

add_heading(dst, "3.5.1 Naif Sabit-Spread Stratejisi", level=3)
add_normal(dst, (
    "Her zaman adımında simetrik sabit half-spread uygulanmaktadır: δ_bid = δ_ask = h = 2 tick. "
    "Envanter farkındalığı veya skew mekanizması içermemektedir."
))

add_heading(dst, "3.5.2 Avellaneda-Stoikov Stratejisi", level=3)
add_normal(dst, "Rezervasyon fiyatı ve yarı-spread şu formüllerle hesaplanmaktadır:")
add_normal(dst, "r = mid - q · γ · σ² · τ")
add_normal(dst, "d = ½ · γ · σ² · τ + (1/γ) · ln(1 + γ/k_price)")
add_normal(dst, (
    "Burada γ = 0.01 risk aversion parametresi, k_price = k_ticks / tick_size fiyat cinsinden yoğunluk "
    "azalma parametresidir. Delta değerleri [1, 25] tick aralığında kırpılmaktadır."
))

# 3.6 Ödül Fonksiyonu
add_heading(dst, "3.6 Ödül Fonksiyonu", level=2)
add_normal(dst, "R_t = (equity_t - equity_{t-1}) - η · inv_t²")
add_normal(dst, (
    "Burada η = 0.001 envanter ceza katsayısıdır (η ablasyonu ile optimize edilmiştir). "
    "Kodda ek bir skew cezası terimi (c·|m|) bulunmakla birlikte, ana deneylerde c = 0 olarak sabitlenmiş "
    "ve bu terim etkisiz bırakılmıştır."
))

add_normal(dst, "")
add_normal(dst, (
    "Ek olarak, η'nın rejime bağlı kılındığı bir varyant denenmiştir: "
    "ηL = 0.0005, ηM = 0.001, ηH = 0.0025 (ηH = 5×ηL). Bu konfigürasyonda "
    "ajan, yüksek volatilite rejimlerinde daha ağır envanter cezasıyla karşılaşmakta; "
    "böylece rejime özgü davranış ödül kanalı aracılığıyla doğrudan teşvik edilmektedir. "
    "Sonuçlar Section 4.7'de sunulmaktadır."
))

# 3.7 Ablasyon Tasarımı
add_heading(dst, "3.7 Gözlem Uzayı ve Ablasyon Tasarımı", level=2)
add_normal(dst, (
    "Gözlem uzayı 6 boyutlu ve sabit mimariyle tasarlanmıştır: obs = [q_norm, σ̂, τ, r_L, r_M, r_H]. "
    "Beş PPO varyantı, use_sigma ve regime_source parametreleriyle kontrol edilmektedir:"
))
variant_headers = ["Varyant", "use_sigma", "regime_source", "Gözlemlediği Bilgi"]
variant_rows = [
    ["ppo_sigma_only", "True", "none", "σ̂ + sıfır rejim"],
    ["ppo_regime_only", "False", "hat", "rejim_hat + sıfır σ̂"],
    ["ppo_combined", "True", "hat", "σ̂ + rejim_hat"],
    ["ppo_oracle_pure", "False", "true", "rejim_true + sıfır σ̂"],
    ["ppo_oracle_full", "True", "true", "σ̂ + rejim_true"],
]
add_table(dst, variant_headers, variant_rows)
add_normal(dst, "Tüm varyantlar 20 seed, 1M timestep ile eğitilmiş ve OOS test seti üzerinde değerlendirilmiştir.")

# 3.8 PPO Eğitimi
add_heading(dst, "3.8 PPO Eğitimi ve Değerlendirme Protokolü", level=2)
add_normal(dst, "Stable-Baselines3 kütüphanesi kullanılarak PPO eğitimi gerçekleştirilmiştir. Hiperparametreler:")

hp_headers = ["Parametre", "Değer"]
hp_rows = [
    ["total_timesteps", "1.000.000"],
    ["learning_rate", "0.0003"],
    ["n_steps", "2048"],
    ["batch_size", "256"],
    ["n_epochs", "10"],
    ["gamma", "0.999"],
    ["gae_lambda", "0.95"],
    ["clip_range", "0.2"],
    ["ent_coef", "0.01"],
]
add_table(dst, hp_headers, hp_rows)
add_normal(dst, (
    "Her seed için episode uzunluğu 8000 adım olup kronolojik eğitim/test bölünmesi %70/%30 "
    "(train: 5600, test: 2400 adım) şeklinde uygulanmıştır. Değerlendirme yalnızca test bölümünden "
    "gerçekleştirilmektedir. Toplam 20 bağımsız tohum kullanılmıştır (seeds: 1-20). "
    "Tezde raporlanan tüm PPO performans değerleri, kronolojik 70/30 eğitim/test bölünmesi "
    "(n_train=5600, n_test=2400) ile çalışan WP5 OOS değerlendirme boru hattından "
    "(src/wp5/job_w5_eval.py) elde edilmiştir. WP4 yalnızca PPO eğitim altyapısı için bir "
    "pilot olarak hizmet etmiştir ve raporlanan sonuçların kaynağı değildir."
))

# ================================================================
# 4. SONUÇLAR VE TARTIŞMA
# ================================================================
add_heading(dst, "4. SONUÇLAR VE TARTIŞMA", level=1)

# 4.1 Ana Deney Sonuçları
add_heading(dst, "4.1 Ana Deney Sonuçları (20 Tohum, OOS)", level=2)
add_normal(dst, "Tablo 1, orijinal WP5 deneyinin (ppo_aware vs ppo_blind) out-of-sample sonuçlarını özetlemektedir.")

tbl1_h = ["Strateji", "Ort. Equity", "Std", "Ort. Sharpe", "inv_p99", "Fill Rate"]
tbl1_r = [
    ["AS", "5.05", "4.72", "0.105", "29.95", "0.444"],
    ["Naif", "4.49", "3.49", "0.126", "21.20", "0.119"],
    ["ppo_aware", "4.10", "0.78", "0.715", "2.00", "0.236"],
    ["ppo_blind", "4.42", "0.71", "0.740", "2.05", "0.232"],
]
add_table(dst, tbl1_h, tbl1_r)

add_normal(dst, (
    "Birinci bulgu — PPO'nun Sharpe Üstünlüğü: Her iki PPO varyantı da Sharpe oranı açısından klasik referans "
    "stratejilerini belirgin biçimde aşmaktadır (~6-7×). Bu üstünlüğün temel mekanizması envanter kontroldür: "
    "PPO varyantları inv_p99 ≈ 2 lot düzeyinde çalışırken, AS ve naif stratejiler sırasıyla 29.95 ve 21.20 "
    "değerlerine ulaşmaktadır. Mutlak equity açısından AS (5.05) ve naif (4.49) stratejilerin PPO varyantlarından "
    "(4.10-4.42) yüksek görünmesi, bu stratejilerin geniş spreadlerden yüksek brüt gelir elde etmesinden "
    "kaynaklanmaktadır; ancak söz konusu gelir yüksek envanter riski pahasına elde edilmektedir."
))

add_normal(dst, (
    "İkinci bulgu — Null Sonuç: ppo_aware, Sharpe bazında 20 tohumun 11'inde ppo_blind'ı geride bırakmıştır. "
    "Sharpe bazında istatistiksel olarak anlamlı fark gözlemlenmemiştir (paired t-test p = 0.261); ancak "
    "equity metriğinde ppo_blind lehine anlamlı fark mevcuttur (p = 0.023). Bu çalışmada risk-adjusted "
    "performans ölçütü olarak Sharpe oranı esas alınmaktadır."
))

add_normal(dst, (
    "Üçüncü bulgu — PPO Varyans Avantajı: PPO varyantlarının equity std değerleri (0.71-0.78) AS (4.72) ve "
    "naif (3.49) stratejilerine kıyasla belirgin biçimde düşüktür."
))

# Şekil 1
dst.add_picture(r"C:\Users\onure\Desktop\THESIS\results\plots\thesis\fig1_sharpe_inv.png", width=Inches(5.5))
add_normal(dst, (
    "Şekil 1. Out-of-sample performans özeti (20 tohum, ortalama ± std). Sol: Sharpe oranı. Sağ: envanter riski "
    "(inv_p99). Şekil, PPO'nun üstünlüğünün kaynağını görsel olarak ortaya koymaktadır. Sol panelde her iki PPO "
    "varyantının Sharpe oranı AS ve Naive stratejileri belirgin biçimde aşmaktadır. Sağ panelde PPO ajanları "
    "inv_p99 ≈ 2 lot düzeyinde çalışırken klasik stratejiler 20-30 lot envanter taşımaktadır."
))

# Üç hipotez
add_normal(dst, "Bu null sonuç için üç açıklayıcı hipotez öne sürülmektedir:")
dst.add_paragraph(
    "(a) Gözlem uzayında σ̂_t değeri hali hazırda mevcut olduğundan, one-hot rejim kodlaması "
    "σ̂_t'den türetilmiş bilgiyi yedekli biçimde sunmakta ve marjinal katkı sağlamamaktadır.",
    style='List Paragraph'
)
dst.add_paragraph(
    "(b) Ödül fonksiyonu rejime özgü davranışı doğrudan teşvik etmediğinden, rejim bilgisini "
    "performansa dönüştürmek için sinyal yetersiz kalmaktadır.",
    style='List Paragraph'
)
dst.add_paragraph(
    "(c) Detector robustness deneyi bu hipotezi geçersiz kılmaktadır: %81.8 doğruluklu HMM "
    "dedektörü ile dahi null sonuç korunmaktadır (p = 0.082). Dolayısıyla tespit kalitesi "
    "belirleyici değildir; açıklayıcı güç (a) hipotezinde yoğunlaşmaktadır.",
    style='List Paragraph'
)

# Şekil 2
dst.add_picture(r"C:\Users\onure\Desktop\THESIS\results\plots\thesis\fig2_paired_seed.png", width=Inches(5.5))
add_normal(dst, (
    "Şekil 2. Seed bazında PPO-aware ve PPO-blind karşılaştırması (her nokta bir seed, n=20). Kesikli çizgi y=x "
    "referansıdır. Sol: Sharpe (paired t-test p=0.261). Sağ: Final equity (paired t-test p=0.023). "
    "Sol panelde noktaların y=x çizgisi etrafında simetrik dağılımı Sharpe bazında istatistiksel anlamlılık "
    "eşiğine ulaşılmadığını göstermektedir. Sağ panelde ise noktaların çoğunluğunun y=x çizgisinin altında "
    "kalması, equity metriğinde PPO-blind'ın sistematik avantajına işaret etmektedir."
))

# 4.2 Pure Ablation
add_heading(dst, "4.2 Pure Ablation Deneyi (Sinyal Yedekliliği Testi)", level=2)
add_normal(dst, (
    "Motivasyon: Danışmanın temel eleştirisi, σ̂'nin her iki ajanda da gözlem uzayında mevcut olması "
    "nedeniyle aware-blind karşılaştırmasının confounded olduğuydu. Bu deneyle, σ̂'nin çıkarıldığı ve "
    "rejim etiketinin tek başına test edildiği pure ablation gerçekleştirilmiştir."
))

add_normal(dst, "Tablo 2, 5 varyantın 20 seed ortalamasını göstermektedir.")

tbl2_h = ["Varyant", "Sharpe (ort. ± std)", "Final Equity (ort. ± std)", "inv_p99 (ort. ± std)"]
tbl2_r = [
    ["ppo_sigma_only", "0.753 ± 0.111", "4.55 ± 0.65", "1.9 ± 0.7"],
    ["ppo_oracle_full", "0.722 ± 0.130", "4.07 ± 0.61", "2.0 ± 0.5"],
    ["ppo_regime_only", "0.698 ± 0.145", "4.01 ± 0.68", "2.1 ± 1.5"],
    ["ppo_combined", "0.696 ± 0.134", "3.91 ± 0.72", "1.8 ± 0.5"],
    ["ppo_oracle_pure", "0.684 ± 0.111", "3.93 ± 0.60", "1.9 ± 0.6"],
    ["naive", "0.127 ± 0.092", "4.49 ± 3.49", "21.2 ± 12.2"],
    ["AS", "0.105 ± 0.082", "5.05 ± 4.72", "29.9 ± 12.6"],
]
add_table(dst, tbl2_h, tbl2_r)

# thesis_23 extension: 5-variant visual summary
dst.add_page_break()
dst.add_picture(r"C:\Users\onure\Desktop\THESIS\results\plots\thesis_23\fig6_ablation_summary.png", width=Inches(5.5))
add_normal(dst, (
    "Şekil 3. Beş PPO varyantı için pure ablation özeti (20 tohum, ortalama ± std). Sol panel Sharpe, "
    "orta panel final equity, sağ panel inv_p99 değerlerini göstermektedir. AS ve naive referansları "
    "bu görsele dahil edilmemiştir; amaç explicit rejim etiketi, oracle rejim etiketi ve sigma_hat "
    "sinyalinin marjinal katkısını PPO varyantları içinde doğrudan karşılaştırmaktır."
))

add_normal(dst, (
    "Bu görsel, Tablo 2'deki ana örüntüyü özetlemektedir: sigma_only en yüksek Sharpe "
    "ortalamasına sahiptir ve oracle_full ya da combined varyantlarının üzerinde anlamlı bir iyileşme "
    "gözlenmemektedir. Bu desen, explicit rejim etiketinin sigma_hat üzerinde ek risk-adjusted bilgi "
    "taşımadığı yönündeki signal redundancy yorumuyla uyumludur."
))

add_normal(dst, (
    "Kritik Bulgu: ppo_sigma_only en yüksek Sharpe değerini (0.753) elde etmiştir — oracle_full'ı "
    "(σ̂ + regime_true, Sharpe = 0.722) ve combined'ı (σ̂ + regime_hat, Sharpe = 0.696) geçmektedir. "
    "Mükemmel rejim bilgisi bile σ̂'nin sağladığı bilginin üzerine anlamlı katkı sağlamamaktadır."
))

# 4.3 İstatistiksel Testler
add_heading(dst, "4.3 İstatistiksel Testler (Ablasyon)", level=2)
add_normal(dst, (
    "Tablo 3, 12 karşılaştırmanın eşleştirilmiş t-testi sonuçlarını göstermektedir "
    "(n=20 seed, Bonferroni düzeltmesi: α = 0.01/12 = 0.00083)."
))

tbl3_h = ["Karşılaştırma", "Metrik", "t", "p-değeri", "Cohen's d", "Anlamlı?"]
tbl3_r = [
    ["sigma_only vs combined", "sharpe", "2.014", "5.83e-02", "0.450", "Hayır"],
    ["sigma_only vs combined", "equity", "3.334", "3.49e-03", "0.746", "Hayır"],
    ["sigma_only vs oracle_full", "sharpe", "1.651", "1.15e-01", "0.369", "Hayır"],
    ["sigma_only vs oracle_full", "equity", "3.789", "1.24e-03", "0.847", "Hayır"],
    ["oracle_full vs combined", "sharpe", "1.062", "3.01e-01", "0.238", "Hayır"],
    ["oracle_full vs combined", "equity", "0.939", "3.60e-01", "0.210", "Hayır"],
    ["sigma_only vs regime_only", "sharpe", "2.283", "3.41e-02", "0.510", "Hayır"],
    ["sigma_only vs regime_only", "equity", "4.510", "2.40e-04", "1.008", "EVET"],
    ["combined vs AS", "sharpe", "16.615", "8.99e-13", "3.715", "EVET"],
    ["combined vs AS", "final_equity", "-1.064", "3.01e-01", "-0.238", "Hayır"],
    ["combined vs naive", "sharpe", "17.878", "2.42e-13", "3.998", "EVET"],
    ["combined vs naive", "final_equity", "-0.728", "4.76e-01", "-0.163", "Hayır"],
]
add_table(dst, tbl3_h, tbl3_r)

add_normal(dst, (
    "Oracle Paradoksu Özeti: oracle_full vs combined farkı Sharpe'ta p = 0.301, equity'de p = 0.360 — "
    "istatistiksel olarak hiçbir anlamlılık yok. Mükemmel etiket, tahmini etiketten ayırt edilemez."
))

add_normal(dst, (
    "TOST equivalence test (bounds ±0.10) confirms practical equivalence between ppo_sigma_only and "
    "ppo_oracle_full (p=0.00067). The 90% confidence interval for the Sharpe difference is "
    "[−0.001, +0.063] — this is the CI that corresponds to the TOST procedure at α=0.05 and it lies "
    "entirely within the ±0.10 equivalence bound. The wider 95% CI [−0.008, +0.069] is also contained "
    "within the bound. This provides positive evidence that the performance difference is practically "
    "negligible, not merely statistically inconclusive."
))

dst.add_picture(r"C:\Users\onure\Desktop\THESIS\results\plots\thesis_23\fig7_oracle_paired_seed.png", width=Inches(5.5))
add_normal(dst, (
    "Şekil 4. Oracle paradoksunun seed bazında görselleştirilmesi. Sol panel doğrudan karşılaştırmayı "
    "gösterir: ppo_sigma_only ile ppo_oracle_full Sharpe değerleri aynı test seed'leri üzerinde "
    "eşleştirilmiştir (paired t-test p = 0.115). Sağ panel oracle_full ile combined varyantlarını "
    "karşılaştırır (p = 0.301). Kesikli çizgi y=x referansıdır."
))

add_normal(dst, (
    "Bu şeklin birincil yorumu sol paneldir: mükemmel rejim etiketlerine sahip oracle_full, sigma_only "
    "varyantını Sharpe bazında istatistiksel olarak anlamlı biçimde geçememektedir. Dolayısıyla oracle "
    "varyantı, dedektör gürültüsü eleştirisini sınırlayan güçlü bir üst-sınır testi sunar; ancak p=0.115 "
    "sonucu, kanıtı deterministik bir eşdeğerlik iddiası olarak değil, anlamlı üstünlük bulunmaması "
    "olarak yorumlanmalıdır."
))

# 4.4 Eylem Analizi
add_heading(dst, "4.4 Eylem Analizi (Davranışsal Farklılaşma)", level=2)

add_normal(dst, "Tablo 4. Rejim Bazında Eylem Dağılımı (Ortalama ± Std, 20 Tohum)")
tbl_action_h = ["Strateji", "Rejim", "Ort. h", "Std h", "Ort. m", "Std m", "P(h=5)"]
tbl_action_r = [
    ["AS", "L", "1.00", "0.00", "-0.01", "0.02", "0.000"],
    ["AS", "M", "1.00", "0.00", "-0.02", "0.03", "0.000"],
    ["AS", "H", "1.00", "0.00", "-0.02", "0.03", "0.000"],
    ["Naif", "L", "2.00", "0.00", "0.00", "0.00", "0.000"],
    ["Naif", "M", "2.00", "0.00", "0.00", "0.00", "0.000"],
    ["Naif", "H", "2.00", "0.00", "0.00", "0.00", "0.000"],
    ["ppo_aware", "L", "1.43", "0.63", "-0.05", "0.22", "0.000"],
    ["ppo_aware", "M", "1.68", "0.72", "-0.01", "0.25", "0.008"],
    ["ppo_aware", "H", "1.74", "0.71", "+0.00", "0.22", "0.000"],
    ["ppo_blind", "L", "1.39", "0.47", "-0.03", "0.24", "0.000"],
    ["ppo_blind", "M", "1.42", "0.41", "-0.00", "0.23", "0.000"],
    ["ppo_blind", "H", "1.60", "0.42", "-0.00", "0.17", "0.000"],
]
add_table(dst, tbl_action_h, tbl_action_r)

add_normal(dst, (
    "ppo_aware ajanı, rejimler arasında sınırlı davranışsal farklılaşma sergilemektedir. Tüm rejimlerde "
    "ortalama half-spread h ≈ 1.4-1.8 tick aralığında seyretmekte, skew değerleri ise m ≈ 0 civarında "
    "kalmaktadır. ppo_blind ile karşılaştırıldığında anlamlı bir politika ayrışması gözlemlenmemektedir."
))

# Şekil 5
dst.add_picture(r"C:\Users\onure\Desktop\THESIS\results\plots\thesis\fig5_action_analysis.png", width=Inches(5.5))
add_normal(dst, (
    "Şekil 5. Rejim bazında eylem dağılımı (20 tohum, ortalama ± std). Üst: half-spread h. Alt: skew m. "
    "Sol sütun: PPO-aware. Sağ sütun: PPO-blind. Half-spread h değerleri L'den H'ye hafif artış göstermekte; "
    "ancak yüksek varyans nedeniyle bu artış istatistiksel olarak güçlü değildir."
))

# 4.5 Rejim Bazlı Performans Analizi
add_heading(dst, "4.5 Rejim Bazlı Performans Analizi", level=2)

add_normal(dst, (
    "Tablo 5, test döneminde üç volatilite rejiminde (L, M, H) elde edilen strateji performanslarını "
    "özetlemektedir."
))

add_normal(dst, (
    "Not: Rejim bazlı kırılım, gerçek rejim etiketleri (regime_true) üzerinden hesaplanmıştır; ajan ise "
    "yalnızca tahmini rejimi (regime_hat) gözlemlemektedir. Bu tablo, ex-post rejim attribution'ı "
    "sunmaktadır; ajan karar anında yalnızca regime_hat'i gözlemlemekte olup gerçek rejim bilgisine "
    "erişimi bulunmamaktadır."
))

tbl_regime_h = ["Strateji", "Rejim", "Sharpe", "inv_p99", "Fill Rate"]
tbl_regime_r = [
    ["AS", "L", "0.330", "22.26", "0.404"],
    ["AS", "M", "0.197", "27.49", "0.438"],
    ["AS", "H", "0.038", "27.66", "0.497"],
    ["naive", "L", "0.269", "16.35", "0.105"],
    ["naive", "M", "0.150", "18.87", "0.115"],
    ["naive", "H", "0.143", "19.20", "0.143"],
    ["PPO-aware", "L", "0.927", "1.85", "0.226"],
    ["PPO-aware", "M", "0.799", "1.60", "0.225"],
    ["PPO-aware", "H", "0.464", "1.79", "0.250"],
    ["PPO-blind", "L", "0.924", "2.00", "0.220"],
    ["PPO-blind", "M", "0.814", "1.88", "0.229"],
    ["PPO-blind", "H", "0.526", "1.91", "0.252"],
]
add_table(dst, tbl_regime_h, tbl_regime_r)

add_normal(dst, (
    "PPO ajanlarının her üç rejimde de AS ve naive baseline'larına kıyasla çok daha düşük envanter riski "
    "(inv_p99 ≈ 2 vs. 20-28) sergilediği görülmektedir. Her iki PPO ajanı da düşük volatilite (L) "
    "rejiminde en yüksek Sharpe değerlerini kaydetmiştir."
))

# Şekil 6
dst.add_picture(r"C:\Users\onure\Desktop\THESIS\results\plots\thesis\fig3_regime_sharpe.png", width=Inches(5.5))
add_normal(dst, (
    "Şekil 6. Volatilite rejimine göre Sharpe oranı (20 tohum, ortalama ± std). PPO'nun üstünlüğünün "
    "tüm volatilite rejimlerinde geçerli olduğu görülmektedir. Yüksek volatilite (H) rejiminde tüm "
    "stratejilerin Sharpe değerleri düşmekte; ancak PPO varyantları bu rejimde dahi klasik stratejileri "
    "belirgin biçimde geride bırakmaktadır."
))

# 4.6 Detector Robustness
add_heading(dst, "4.6 Detector Robustness Analizi", level=2)

tbl_det_h = ["Dedektör", "Doğruluk", "ppo_aware Sharpe", "ppo_blind Sharpe", "Sharpe p-değeri"]
tbl_det_r = [
    ["rv_baseline", "%60.7", "0.718", "0.753", "0.114"],
    ["rv_dwell", "%60.4", "0.716", "0.753", "0.110"],
    ["HMM", "%81.8", "0.719", "0.753", "0.082"],
]
add_table(dst, tbl_det_h, tbl_det_r)

add_normal(dst, (
    "Not: ppo_blind bu tabloda 0.753 Sharpe değeriyle görünmektedir; bu değer, farklı dedektörlerle "
    "yeniden eğitilen modellerin (wp5-detector-full run'ı) ortalamasını yansıtmaktadır. Section 4.1'deki "
    "0.740 değeri ise orijinal wp5-eval-main run'ından elde edilmiştir."
))

add_normal(dst, (
    "ANOVA testi: F = 0.003, p = 0.997 — ppo_aware performansı dedektör seçimine göre anlamlı biçimde "
    "değişmemektedir. Null sonuç, zayıf tespit kalitesinden kaynaklanmamaktadır."
))

# Şekil 7
dst.add_picture(r"C:\Users\onure\Desktop\THESIS\results\plots\thesis\fig4_detector_robustness.png", width=Inches(5.5))
add_normal(dst, (
    "Şekil 7. Dedektör bazında ortalama Sharpe (20 tohum). PPO-blind Sharpe değeri tüm dedektörlerde "
    "0.753 olarak sabit kalmakta; PPO-aware ise dedektör seçiminden bağımsız olarak ~0.717-0.719 "
    "düzeyinde seyretmektedir."
))

# 4.7 Rejime Koşullu Envanter Cezası Deneyi (moved BEFORE Section 5)
add_heading(dst, "4.7 Rejime Koşullu Envanter Cezası Deneyi", level=2)
add_normal(dst, (
    "Danışmanın önerisi doğrultusunda, null sonucun ödül tasarımı boyutunda da sınanması "
    "amacıyla rejime koşullu envanter cezası deneyi gerçekleştirilmiştir. Sabit η = 0.001 "
    "yerine ηL = 0.0005, ηM = 0.001, ηH = 0.0025 konfigürasyonu kullanılmış; beş varyant "
    "20 tohum, 1M timestep ile eğitilmiştir."
))
add_normal(dst, "Tablo 6, sonuçları özetlemektedir.")
add_table(dst,
    ["Varyant", "Ort. Sharpe", "Std Sharpe", "Ort. Equity"],
    [
        ["ppo_sigma_only", "0.714", "0.119", "3.91"],
        ["ppo_combined",   "0.629", "0.147", "3.50"],
        ["ppo_oracle_full","0.638", "0.158", "3.38"],
        ["ppo_oracle_pure","0.578", "0.247", "3.45"],
        ["ppo_regime_only","0.513", "0.238", "3.07"],
    ]
)

dst.add_picture(r"C:\Users\onure\Desktop\THESIS\results\plots\thesis_23\fig8_eta_regime_summary.png", width=Inches(5.5))
add_normal(dst, (
    "Şekil 8. Rejime koşullu envanter cezası deneyinde PPO varyantlarının Sharpe ve final equity "
    "sonuçları (20 tohum, ortalama ± std). Görsel, sigma_only ile combined varyantları arasındaki "
    "eşleştirilmiş karşılaştırmayı özellikle vurgular: Sharpe için p = 0.0016, final equity için "
    "p = 0.008."
))

add_normal(dst, (
    "Bu görsel, rejime koşullu ödül tasarımının explicit rejim etiketini faydalı hale getirmediğini "
    "göstermektedir. combined varyantı hem sigma_hat hem tahmini rejim etiketini gözlemlerken, "
    "sigma_only Sharpe bazında istatistiksel olarak daha güçlü kalmaktadır. Bu sonuç ödül kanalı "
    "üzerinden yapılan rejim hassaslaştırmasının signal redundancy yorumunu ortadan kaldırmadığını "
    "destekler."
))

add_normal(dst, (
    "Eşleştirilmiş t-testi (ppo_combined vs ppo_sigma_only): Sharpe p = 0.0016, "
    "Equity p = 0.008 — her iki metrikte ppo_sigma_only lehine istatistiksel olarak "
    "anlamlı fark. Rejime koşullu η, ppo_combined'ın performansını artırmak yerine "
    "anlamlı biçimde düşürmüştür. ppo_combined gözlem uzayında σ̂ aracılığıyla zaten "
    "rejim bilgisine sahipken, rejime koşullu η bu örtük sinyali gereksiz biçimde "
    "pekiştirmiş ve ajan optimizasyonunu zorlaştırmıştır. Bu bulgu, signal redundancy "
    "argümanını ödül tasarımı boyutunda da desteklemektedir."
))

# ================================================================
# CHANGE D — Section 4.8: Model Misspecification Robustness Deneyi
# ================================================================
add_heading(dst, "4.8 Model Misspecification Robustness Deneyi", level=2)

add_normal(dst, (
    "Motivasyon: Mevcut sentetik ortamda fill model parametreleri A ve k tüm rejimler için sabit "
    "tutulmaktadır. Bir eleştiri olarak, gerçek piyasalarda yüksek volatilite rejimlerinde likidite "
    "dinamiklerinin farklılaşabileceği öne sürülebilir. Bu deneyle, fill parametrelerinin rejime bağlı "
    "olduğu daha zorlu bir ortamda null result'ın geçerliliği test edilmektedir."
))

add_normal(dst, (
    "Ortam uyarısı (retrospektif not): Bu misspec-mild deneyleri tez yazım sürecinin erken bir "
    "aşamasında yürütüldüğünden, rejim parametreleri §3.1'de tanımlanan kanonik ortamla birebir "
    "örtüşmemektedir. Misspec-mild koşuları sigma_mult = [0.5, 1.0, 2.0] ve warmup_steps = 400 ile "
    "üretilmiştir; kanonik değerler ise sigma_mult = [0.6, 1.0, 1.8] ve warmup_steps = 1000'dir. "
    "Bu sapma WP5.5 sinyal denetimi öncesinde yapılan sanity-check sırasında tespit edilmiş, "
    "hesaplama maliyeti nedeniyle kanonik parametrelerle yeniden üretim gerçekleştirilmemiştir "
    "(bkz. Decisions Log #42, #45). Okuyucunun aşağıdaki tablo ve istatistiklerin bu daha "
    "agresif rejim-oranı ortamına ait olduğunu göz önünde bulundurması gerekir; ana null bulgusu "
    "(sigma_only ≈ oracle_full) her iki parametre setinde de aynı yönde raporlanmaktadır."
))

add_normal(dst, (
    "Tasarım: MMSimulator generic kalmış; env.py içinde her adımda regime_true'ya göre ExecParams "
    "override edilmiştir. Mild misspecification parametreleri: L rejimi A=4.0, k=1.8; M rejimi A=5.0, "
    "k=1.5 (baz değer); H rejimi A=6.0, k=1.2. Bu parametre konfigürasyonu, yüksek volatilite "
    "rejiminde gerçekçi bir likidite daralmasını değil, PPO ajanlarının değişen dolum dinamiklerine "
    "karşı dayanıklılığını test etmeye yönelik adversarial bir parametre şokunu temsil etmektedir. "
    "Aynı 5-varyant (sigma_only, regime_only, combined, "
    "oracle_pure, oracle_full), 20 tohum, 1M timestep eğitim protokolü uygulanmıştır."
))

add_normal(dst, "Tablo 7. Model Misspecification (Mild) OOS Sonuçları (20 Tohum)")
add_table(dst,
    ["Varyant", "Ort. Sharpe"],
    [
        ["ppo_sigma_only",  "0.685"],
        ["ppo_oracle_full", "0.682"],
        ["ppo_combined",    "0.651"],
        ["ppo_regime_only", "0.634"],
        ["ppo_oracle_pure", "0.602"],
    ]
)

dst.add_picture(r"C:\Users\onure\Desktop\THESIS\results\plots\thesis_23\fig9_misspec_summary.png", width=Inches(5.5))
add_normal(dst, (
    "Şekil 9. Mild model misspecification ortamında PPO varyantlarının Sharpe karşılaştırması "
    "(20 tohum, ortalama ± std). Fill parametreleri A ve k rejime bağlı olacak şekilde değiştirilmiştir; "
    "sigma_only ile oracle_full arasındaki eşleştirilmiş Sharpe karşılaştırması p = 0.881'dir."
))

add_normal(dst, (
    "Bu görsel, model misspecification sonucunu özetlemektedir: rejime bağlı A/k altında "
    "sigma_only ve oracle_full neredeyse aynı Sharpe düzeyinde kalmaktadır. Bu bulgu yalnızca bu mild "
    "misspecification tasarımı için bir robustness sonucu olarak yorumlanmalıdır; tüm piyasa model "
    "yanlış belirlemeleri altında evrensel eşdeğerlik iddiası değildir."
))

add_normal(dst, (
    "Eşleştirilmiş t-testi sonuçları: ppo_sigma_only vs ppo_oracle_full: p=0.881; "
    "ppo_sigma_only vs ppo_combined: p=0.217; ppo_sigma_only vs ppo_oracle_pure: p=0.098. "
    "Hiçbir karşılaştırma istatistiksel anlamlılık eşiğine ulaşmamıştır."
))

add_normal(dst, (
    "TOST equivalence test (bounds ±0.05) further confirms practical equivalence under model "
    "misspecification (p=0.042). The 90% confidence interval [−0.040, +0.048] — the CI corresponding "
    "to the TOST procedure at α=0.05 — lies entirely within the ±0.05 equivalence bound. The wider "
    "95% CI [−0.049, +0.057] remains within the ±0.10 bound. Cohen's d = 0.034 indicates negligible "
    "effect size, strengthening the signal redundancy interpretation across both standard and "
    "adversarial environments."
))

add_normal(dst, (
    "Bu bulgular, signal redundancy argümanının çevresel model parametrelerindeki değişimlere karşı "
    "sağlam olduğunu ortaya koymaktadır. Fill dinamiklerinin rejime bağlı olduğu daha zorlu bir ortamda "
    "dahi ppo_sigma_only, mükemmel rejim bilgisine sahip oracle_full'dan istatistiksel olarak ayırt "
    "edilemez performans göstermiştir (p=0.881)."
))

# ================================================================
# 4.9 Manipülasyon Geçerlilik Denetimi (WP5.5)
# ================================================================
add_heading(dst, "4.9 Manipülasyon Geçerlilik Denetimi", level=2)

add_normal(dst, (
    "Motivasyon: PPO ajanlarının eğitiminde kullanılan gözlem uzayı sigma_hat ve "
    "rejim one-hot bileşenlerini içermektedir. Null sonucu savunulur kılmanın gerekli "
    "koşullarından biri, bu sinyallerin taşıdığı bilginin gerçekten ajanın kullanabileceği "
    "yapıda olduğunun — ve sinyal kalitesi düşürüldüğünde metriklerin beklendiği gibi "
    "bozulduğunun — gösterilmesidir. Aksi hâlde 'rejim etiketi faydasız' ifadesi, "
    "etiketin öznitelik olarak hiç bilgi taşımadığı trivial durumdan ayırt edilemez. "
    "Bu bölüm, rejim sinyali üzerine uygulanan kontrollü manipülasyonların regresyon/"
    "sınıflandırma metriklerinde beklenen etkiyi ürettiğini raporlar."
))

add_normal(dst, (
    "Yöntem: Kanonik operating point (dt = 0.2, sigma_mid_ticks = 0.8, sigma_mult = "
    "[0.6, 1.0, 1.8], warmup_steps = 1000, rv_window = 50, n_steps = 8000, tohum = 123) "
    "altında tek bir epizot üretilir. Rolling realized volatility serisi üzerinde beş "
    "koşul uygulanır: clean (manipülasyon yok), noisy (noise_std = 0.5 × clean_sigma_std, "
    "yani 0.09 civarı), lagged (5 adım gecikme), coarsened (5 kuantil kovasına indirgeme, "
    "kova ortalamasıyla değiştirme) ve none (sabit fill_value = 0.0). Her koşul için "
    "sigma_hat ile regime_true arasındaki Spearman/Pearson korelasyonu, rejim "
    "sınıflandırma doğruluğu, ayrışabilirlik (Fisher benzeri skor) ve threshold overlap "
    "oranı hesaplanır. Audit iş paketinin kaynak dosyaları: config/w55_audit.json, "
    "src/wp5_5/job_w55_audit.py, src/wp5_5/signal_degradation.py, "
    "src/wp5_5/signal_audit.py."
))

add_normal(dst, (
    "Yapısal metrik politikası: 'none' koşulu sabit bir fill değeriyle çalıştığı için "
    "classification_accuracy ve threshold_overlap metrikleri bu koşulda yapısal olarak "
    "tanımsızdır — değerler sinyal bilgi içeriğini değil, fill değeri ile kalibre edilmiş "
    "eşikler arasındaki ilişkiyi yansıtır. Bu iki metrik 'none' koşulunda raporlanmaya "
    "devam eder ancak PASS/FAIL değerlendirmesinden çıkarılmıştır (reported-but-not-gated)."
))

add_normal(dst, "Tablo 8. WP5.5 Sinyal Denetim Metrikleri (kanonik operating point, tohum = 123)")
add_table(dst,
    ["Koşul", "Spearman", "Pearson", "Clas. Acc.", "Ayrışab.", "Thr. Overlap"],
    [
        ["clean",     "1.000", "1.000", "0.400", "5093.3", "0.300"],
        ["noisy",     "0.820", "0.894", "0.518", "4274.3", "0.103"],
        ["lagged",    "0.977", "0.991", "0.398", "4856.8", "0.300"],
        ["coarsened", "0.940", "0.877", "0.391", "5114.1", "0.374"],
        ["none",      "0.000", "0.000", "0.548", "0.000",  "0.000"],
    ]
)

add_normal(dst, (
    "Bulgular: (i) Direction — clean sinyali 'none' baseline'ına karşı Spearman, Pearson ve "
    "ayrışabilirlik boyutlarında PASS vermektedir; yapısal olarak tanımsız iki metrik "
    "(classification_accuracy ve threshold_overlap) 'none' koşulunda gate dışında "
    "bırakılmıştır. (ii) Separation — noisy, lagged ve coarsened koşullarının tümü clean "
    "baseline'ından belirgin biçimde sapmakta, yani rejim etiketine uygulanan kontrollü "
    "bozulumlar ölçülebilir etki üretmektedir. (iii) Coarsen safety — 5 kovalı "
    "ayrıklaştırma, metriklerin üzerinde yıkıcı bir bilgi kaybı yaratmamaktadır. "
    "(iv) Monotonicity — 'clean → noisy → lagged → coarsened → none' dizisi boyunca "
    "ölçütler yekpare (monotone) bir şiddet eğrisine oturmamaktadır (non-monotone FAIL). "
    "Bu sonuç, tek bir sıralama ekseni boyunca artan bozulmadan kaynaklanan yapısal bir "
    "başarısızlık değildir; üç degradasyon (gürültü, gecikme, kuantizasyon) sinyal "
    "yapısına farklı boyutlarda saldırdığı için ortak bir şiddet ölçütüne indirgenmeleri "
    "ölçüt tasarımı düzeyinde tartışmaya açıktır."
))

add_normal(dst, (
    "Yorum: Denetim, rejim sinyaline dokunan kontrollü manipülasyonların beklenen "
    "yönde ölçülebilir etki ürettiğini göstermektedir (direction PASS, separation PASS, "
    "coarsen safety PASS). Bu, sinyalin faydasız değil, *taşıdığı bilgi sigma_hat "
    "tarafından zaten yansıtıldığı için* RL performansına marjinal katkı sağladığı "
    "şeklindeki null-result yorumuyla tutarlıdır. Monotonicity FAIL'i, ayrı tipteki "
    "bozulumların tek bir şiddet ekseninde karşılaştırılmasına dair bir ölçüt tasarımı "
    "uyarısı olarak okunmalıdır; WP5.5 run'ın genel önerisi bu nedenle REVIEW "
    "düzeyindedir. Audit çıktıları: "
    "results/runs/20260422-170037_seed123_wp55-signal-audit_66fc17e/."
))

add_normal(dst, (
    "Reproducibility notu: config dosyasında noise_std alanı 'auto' değerine ayarlıdır; "
    "audit iş paketi koşu sırasında gürültü standart sapmasını 0.5 × clean_sigma_std "
    "(post-warmup) olarak hesaplar. Bu sayede gürültü seviyesi her zaman kanonik sinyal "
    "ölçeğine göre sabit oranda kalır ve operating point'e dair güncellemelere otomatik "
    "olarak adapte olur. WP5.5 öncesinde yapılan iki drift'li koşu "
    "(20260419-214718, 20260419-221930) süpersed edilmiş ve "
    "'results/runs/wp55_audit_archive_note.md' altında arşivlenmiştir."
))

# ================================================================
# 5. SİNYAL BİLGİLENDİRİCİLİK SÜPÜRMESİ (WP6)
# ================================================================
# New chapter for thesis_27 (Sürüm 27): three-level structure per
# wp6_notes.md §5 and Decision #47.  Hedge discipline (defense-driven):
#   - "directionally shifted below sigma_only" (NOT "consistently"/"harmful")
#   - "categorical-channel degradation" (NOT "encoding-interference",
#     NOT "representation-level effect")
#   - mechanism explicitly NOT identified
#   - Plot 3 framed as "TOST equivalence not formally established at n=20
#     (underpowered, MDED ≈ ±0.07–0.08)" — not as combined ≈ regime_only.
# ================================================================
add_heading(dst, "5. SİNYAL BİLGİLENDİRİCİLİK SÜPÜRMESİ", level=1)

# --- Chapter intro (incl. methodological-symmetry note + historical
# transparency disclosure required by Decision #47).
add_normal(dst, (
    "Bu bölüm, sürekli volatilite sinyalinin (σ̂_t) artan biçimde bozulduğu beş koşulda "
    "(full → noisy → lagged → coarsened → none) PPO ajanının beş gözlem-uzayı varyantı "
    "(sigma_only, combined, regime_only, oracle_full, oracle_pure) altında nasıl davrandığını "
    "raporlar. Kalibrasyon parametreleri Karar #46'da belirlenmiştir (noisy α = 0.40, lagged "
    "k = 20, coarsened binning). Toplam 4 koşul × 5 varyant + 1 koşul × 4 varyant = 24 hücre "
    "× 20 tohum = 480 koşum eğitilmiş ve her koşumda deterministik OOS değerlendirmesi yapılmıştır."
))

add_normal(dst, (
    "Tarihsel şeffaflık (Karar #47): WP6 başlangıçta, sürekli volatilite sinyali σ̂_t "
    "bozuldukça explicit kategorik rejim etiketlerinin değer kazandığı bir rejimi tespit "
    "etmek — yani bilgilendiricilik-eşiği hipotezini test etmek (Karar #46) — amacıyla "
    "tasarlanmıştır. Deney bunun yerine, test edilen kalibrasyon aralığı içinde böyle bir "
    "rejimin var olmadığını ortaya koymuş ve combined varyantında beklenmeyen bir bozulma "
    "örüntüsünü gün yüzüne çıkarmıştır. Bu bölüm her iki bulguyu gözlemlendiği şekilde "
    "raporlar; deney tasarımı ve analiz planı sonuçlar görüldükten sonra değiştirilmemiştir."
))

add_normal(dst, (
    "Metodolojik simetri notu — TOST kullanımı: WP5, ppo_aware ile ppo_blind arasındaki "
    "pratik eşitliği kanıtlamak amacıyla TOST'u eşitlik yönünde kullanmıştır. WP6 ise TOST'u "
    "her iki yönde kullanır: combined ile sigma_only karşılaştırmasında non-eşitlik (Şekil 11) "
    "ve combined ile regime_only karşılaştırmasında eşitlik (Şekil 12; n = 20'nin ikinci yön "
    "için underpowered olduğu kayda alınarak). Eşitlik testinde max(p_lower, p_upper) < α, "
    "non-eşitlik testinde min(p_lower, p_upper) < α karar kuralı uygulanır."
))

# --- Tablo 9: per-condition × variant summary (mean ± 95% CI, n=20)
add_normal(dst, (
    "Tablo 9, beş koşul × beş varyant matrisinde 20 tohum üzerinden OOS Sharpe ortalamasını "
    "ve %95 güven aralığını özetlemektedir (df = 19; t-kritik = 2.093). regime_only ve "
    "oracle_pure varyantları yapısı gereği σ̂_t'yi tüketmediğinden koşul-değişmezdir; bu "
    "tohum-bazında doğrulanmıştır (max |fark| = 0). sigma_only varyantı condition = none "
    "altında tanımsızdır (σ̂_t sıfırlanır)."
))

tbl9_h = ["Koşul", "sigma_only", "combined", "regime_only", "oracle_full", "oracle_pure"]
tbl9_r = [
    ["full",      "0.763 ± 0.060", "0.690 ± 0.061", "0.699 ± 0.067", "0.681 ± 0.070", "0.663 ± 0.077"],
    ["noisy",     "0.756 ± 0.062", "0.713 ± 0.063", "0.699 ± 0.067", "0.674 ± 0.101", "0.663 ± 0.077"],
    ["lagged",    "0.782 ± 0.045", "0.681 ± 0.054", "0.699 ± 0.067", "0.726 ± 0.069", "0.663 ± 0.077"],
    ["coarsened", "0.783 ± 0.053", "0.691 ± 0.083", "0.699 ± 0.067", "0.745 ± 0.055", "0.663 ± 0.077"],
    ["none",      "tanımsız",       "0.699 ± 0.067", "0.699 ± 0.067", "0.663 ± 0.077", "0.663 ± 0.077"],
]
add_table(dst, tbl9_h, tbl9_r)

# --- Şekil 10: monotonic-gap plot
dst.add_picture(
    r"C:\Users\onure\Desktop\THESIS\docs\internal\wp6_sweep_full\plots\monotonic_gap.png",
    width=Inches(5.5),
)
add_normal(dst, (
    "Şekil 10. Beş koşul boyunca beş varyantın OOS Sharpe ortalaması (20 tohum, ortalama ± "
    "%95 GA). x-ekseni sinyal bozulması ekseni (full → noisy → lagged → coarsened → none); "
    "y-ekseni Sharpe-benzeri performans metriği. sigma_only çizgisi dört bilgilendirici "
    "koşulda (0.756–0.783) yatay seyretmektedir; regime_only ve oracle_pure çizgileri "
    "yapısal olarak sabittir."
))

# ----------------------------------------------------------------
# 5.1 Reddedilen Hipotez: Bilgilendiricilik-Eşiği
# ----------------------------------------------------------------
add_heading(dst, "5.1 Reddedilen Hipotez: Bilgilendiricilik-Eşiği", level=2)

add_normal(dst, (
    "Pre-registered beklenti: σ̂_t aşamalı olarak bozuldukça (full → noisy → lagged → "
    "coarsened → none), sigma_only ile combined arasındaki Sharpe farkı monoton biçimde "
    "daralmalı; sigma yeterince gürültülü hale geldiğinde kategorik rejim etiketi anlamlı "
    "bilgi katmaya başlamalıdır. Bu hipotez WP5'in null sonucunu — explicit rejim "
    "etiketlerinin σ̂_t'nin halihazırda taşıdığı bilgiyle yedekli kaldığı bulgusunu — "
    "yöntemsel olarak izleyen bir testtir (Karar #46)."
))

add_normal(dst, (
    "Gözlem (Şekil 10, Tablo 9): sigma_only varyantı dört bilgilendirici koşulda "
    "0.756–0.783 aralığında esasen sabit kalır; sigma_only ile combined arasındaki fark "
    "0.06–0.10 Sharpe-benzeri bandında durur ve monoton bir daralma sergilemez. Bu, "
    "pre-registered hipotezin reddedildiği anlamına gelir: test edilen kalibrasyon aralığı "
    "içinde, sigma'nın bozulması combined varyantının sigma_only'yi yakalamasına yol açmaz."
))

add_normal(dst, (
    "Kalibrasyon aralığı uyarısı: Test edilen kalibrasyon aralığı içinde (Karar #46: "
    "noisy α = 0.40, lagged k = 20, coarsened binning), hiçbir bozulma koşulu, explicit "
    "kategorik etiketlerin faydalı hale geldiği bir rejim üretmemiştir. Daha agresif "
    "bozulmalar (örn. α = 0.80 veya k = 50) bu süpürmede test edilmemiştir ve bunlar "
    "ekarte edilemez. Dolayısıyla 'sigma_only düz' bulgusu, seçilmiş kalibrasyon bandına "
    "ilişkin bir ifadedir; tüm sigma bozulmaları için evrensel bir iddia değildir. "
    "Noisy condition'un α değeri, tüm post-warmup standart sapması kullanılarak kalibre "
    "edilmiştir; yalnızca eğitim verisi kullanılarak yeniden kalibre edildiğinde α tahmini "
    "olarak <%10 oranında kayacaktır. Bu kayma, §5.2'de raporlanan yönsel bulguların "
    "(combined < sigma_only, Cohen's dz arası -0.57 ile -0.91) sağlam kaldığı aralık içindedir."
))

# ----------------------------------------------------------------
# 5.2 Beklenmeyen Bulgu: combined Varyantının Yönsel Düşüşü
# ----------------------------------------------------------------
add_heading(dst, "5.2 Beklenmeyen Bulgu: combined Varyantının Yönsel Düşüşü", level=2)

add_normal(dst, (
    "Süpürme planlanmamış bir bulguyu ortaya çıkarmıştır: combined varyantı dört "
    "bilgilendirici koşulun tümünde sigma_only'nin altında yönsel olarak konumlanmaktadır. "
    "Tohum-bazında eşleştirilmiş analiz (Şekil 11, summary_paired_combined_vs_sigma.csv): "
    "her koşulda 20 tohumun 15–18'i kimlik çizgisinin (y = x) altında yer alır; eşleştirilmiş "
    "t-testi p-değeri her dört koşulda 0.05'in altındadır; Cohen's dz orta-büyüklükten "
    "büyüğe (−0.57'den −0.91'e) uzanır."
))

# Şekil 11
dst.add_picture(
    r"C:\Users\onure\Desktop\THESIS\docs\internal\wp6_sweep_full\plots\paired_seed_combined_vs_sigma.png",
    width=Inches(5.5),
)
add_normal(dst, (
    "Şekil 11. Tohum-bazında combined vs sigma_only karşılaştırması (her panel bir koşul, "
    "her nokta bir tohum, n = 20). Kesikli çizgi y = x referansıdır. Noktaların büyük "
    "çoğunluğunun y = x'in altında kalması, combined varyantının seed-by-seed temelinde "
    "sigma_only'nin altında yönsel olarak yer aldığını gösterir."
))

add_normal(dst, "Tablo 10. combined vs sigma_only — Eşleştirilmiş Tohum İstatistikleri")
tbl10_h = ["Koşul", "n_below", "mean_diff", "paired t p", "Cohen's dz", "TOST p (non-eşitlik)", "δ=0.05'te red"]
tbl10_r = [
    ["full",      "15/20", "−0.073", "0.012",   "−0.621", "0.198", "Hayır"],
    ["noisy",     "15/20", "−0.044", "0.020",   "−0.568", "0.640", "Hayır"],
    ["lagged",    "18/20", "−0.101", "0.00064", "−0.912", "0.027", "Evet"],
    ["coarsened", "15/20", "−0.092", "0.0088",  "−0.653", "0.099", "Hayır"],
]
add_table(dst, tbl10_h, tbl10_r)

add_normal(dst, (
    "Hedge — Yönsel vs pratik karakterizasyon: Yön dört koşulda da sağlamdır (paired t < "
    "0.05, dz orta-büyük). Ancak büyüklük yalnızca lagged koşulunda pratik anlamlılık "
    "eşiğini (δ = 0.05 Sharpe-benzeri) aşar (mean_diff = −0.101, TOST non-eşitlik p = 0.027). "
    "Diğer üç koşulda farkın 90% güven aralığı eşitlik bandına temas eder. Bu nedenle "
    "bölüm boyunca 'yönsel olarak sigma_only altında' ifadesi tercih edilmiş, 'tutarlı "
    "biçimde' veya 'zararlı' gibi nitelemelerden kaçınılmıştır."
))

add_normal(dst, (
    "Bu örüntü pre-registered yönün tersidir: 'sigma bozuldukça rejim etiketleri faydalı "
    "hale gelir' değil, 'rejim etiketleri sigma ile birlikte sunulduğunda combined "
    "performansı yönsel olarak sigma_only'nin altına iter' biçiminde tezahür eder. Bu "
    "örüntüye bölüm boyunca **kategorik-kanal performans bozulması** (categorical-channel "
    "degradation) diye atıfta bulunulur. Bu, mekanizma yüklü değil, betimleyici bir "
    "etikettir."
))

# ----------------------------------------------------------------
# 5.3 Yorumun Hassasiyeti
# ----------------------------------------------------------------
add_heading(dst, "5.3 Yorumun Hassasiyeti: Anlam-Eşitliği vs Politika-Eşitliği", level=2)

add_normal(dst, (
    "İkinci bir tohum-eşleştirilmiş analiz, combined varyantının none-koşul davranışına "
    "(yani regime_only'ye) deterministik biçimde çöküp çökmediğini sınar. Cevap: hayır. "
    "Mean diff'ler dört koşulda da küçüktür (|mean_diff| ≤ 0.018, üç tanesinde işaret "
    "negatif), ancak varyans heterojendir."
))

# Şekil 12
dst.add_picture(
    r"C:\Users\onure\Desktop\THESIS\docs\internal\wp6_sweep_full\plots\paired_seed_combined_vs_regime.png",
    width=Inches(5.5),
)
add_normal(dst, (
    "Şekil 12. Tohum-bazında combined vs regime_only karşılaştırması (her panel bir koşul, "
    "her nokta bir tohum, n = 20). Şekil 11'in aksine noktalar y = x etrafında dağılır; "
    "ortalama farklar küçüktür ancak tohum-bazında eşleşme deterministik değildir — "
    "policy'ler aynı per-tohum fonksiyonu öğrenmemektedir."
))

add_normal(dst, "Tablo 11. combined vs regime_only — Eşleştirilmiş Tohum İstatistikleri ve Varyans Yan-Çıktısı")
tbl11_h = ["Koşul", "mean_diff", "paired t p", "TOST p (eşitlik)", "std_combined", "std_regime_only", "std_diff"]
tbl11_r = [
    ["full",      "−0.009", "0.791", "0.128", "0.130", "0.143", "0.155"],
    ["noisy",     "+0.013", "0.685", "0.137", "0.134", "0.143", "0.145"],
    ["lagged",    "−0.018", "0.535", "0.136", "0.114", "0.143", "0.127"],
    ["coarsened", "−0.008", "0.814", "0.123", "0.177", "0.143", "0.156"],
]
add_table(dst, tbl11_h, tbl11_r)

add_normal(dst, (
    "TOST eşitliği δ = 0.05 düzeyinde **n = 20 tohumda formel olarak kurulamamıştır** "
    "(tost_p ≈ 0.12–0.14, dört koşulda da α = 0.05'i geçmez). Gözlenen std_diff ≈ 0.13–0.16 "
    "ile bu tasarım mevcut tohum sayısında ancak yaklaşık ±0.07–0.08 büyüklüğündeki "
    "minimum tespit edilebilir eşitlik bandını saptayabilir; ±0.05 değil. Bu nedenle "
    "Şekil 12 sonucu 'underpowered, etkinin yokluğu değil' olarak okunmalıdır. "
    "Muhafazakâr ifade: 'combined ortalaması, regime_only ortalamasından yaklaşık ±0.07 "
    "bandında ayırt edilemezdir' — 'δ = 0.05'te eşittir' DEĞİL."
))

add_normal(dst, (
    "Varyans yan-çıktısı eşitlik tablosunu zenginleştirir: std_diff > std_regime_only dört "
    "koşulun üçünde (full, noisy, coarsened) gözlenir. Bu, combined ile regime_only "
    "policy'lerinin tohum-eşleşmiş düzeyde aynı fonksiyonu öğrenmediğine işaret eder; "
    "etki anlam-düzeyinde eşitliğe yakın olsa da policy-düzeyinde eşitliğe değildir. "
    "coarsened panelinde std_combined = 0.177 ile en yüksek değere ulaşır; bu, panel "
    "üzerinde görülen birkaç yüksek-Sharpe tohumdan kaynaklanmaktadır."
))

add_normal(dst, (
    "Mekanizma kimliklendirilmemiştir. Veri çeşitli olası mekanizmalarla (gradyan girişimi, "
    "kapasite tahsisi, optimizasyon gürültüsü, veya başka faktörler) tutarlıdır; ancak bu "
    "deneyden tek başına hangi mekanizmanın işlediği saptanamaz. Bu nedenle bölüm boyunca "
    "'kategorik-kanal performans bozulması' betimleyici etiketi kullanılmış; mekanizma "
    "yüklü terimlerden (örn. 'representation-level effect', 'encoding-interference') "
    "kaçınılmıştır. Mekanizma ayrımı, bu deneyin kapsamı dışında kalan ek deneylerle "
    "(örn. ablation study, gradient norm tracking, ya da Bayesian linear probing) ele "
    "alınmalıdır."
))

# --- Kapanış: WP5 → WP6 sürekli argüman bağlantısı
add_normal(dst, (
    "Kapanış: Explicit rejim etiketleri ancak rejim sinyali gözlem uzayında halihazırda "
    "örtük olarak mevcut değilse fayda sağlamaktadır — bu, WP5 ana null sonucunun "
    "ifadesidir (Bulgu 2). WP6, σ̂_t bozulduğunda dahi kategorik kanalın yarar kazanmadığını "
    "ve bu deneysel ortamda sürekli kanalın yanına eklendiğinde performansı yönsel olarak "
    "düşürdüğünü göstererek bu bulguyu genişletmektedir. 'Bu deneysel ortamda' kaydı "
    "bilinçli bir savunma kaydıdır: bulgu Karar #46'da belirlenen kalibrasyon bandına "
    "ilişkindir; kategorik kanalın faydalı olabileceği daha agresif bozulma rejimleri "
    "gelecek çalışmaya bırakılmıştır."
))

# ================================================================
# 6. SONUÇ — renumbered from 5 (new Chapter 5 = WP6 inserted above)
# ================================================================
add_heading(dst, "6. SONUÇ", level=1)

add_normal(dst, "Bu çalışma, volatilite rejimleri altında pekiştirmeli öğrenme tabanlı piyasa yapıcılığını araştırmaktadır. Temel bulgular:")

add_normal(dst, (
    "Bulgu 1 — PPO üstünlüğü: Tüm PPO varyantları Sharpe oranı açısından Avellaneda-Stoikov ve naif "
    "referans stratejilerini yaklaşık 6-7 kat geçmiştir. Bu üstünlüğün temel kaynağı envanter yönetimi "
    "verimliliğidir: PPO inv_p99 ≈ 2 lot, baseline'lar 20-30 lot."
))

add_normal(dst, (
    "Bulgu 2 — Null Sonuç: Rejim bilgisinin performansı artırdığı hipotezi istatistiksel olarak "
    "desteklenmemiştir. Sharpe bazında istatistiksel olarak anlamlı fark gözlemlenmemiştir "
    "(paired t-test p = 0.261); equity metriğinde ppo_blind lehine anlamlı fark mevcuttur "
    "(p = 0.023). Bu null sonuç, "
    "gözlem uzayında hali hazırda bulunan σ̂_t'nin rejim bilgisini örtük olarak içermesiyle açıklanmaktadır."
))

add_normal(dst, (
    "Bulgu 3 — Pure Ablation (Sinyal Yedekliliği Testi): ppo_sigma_only en yüksek performansı elde "
    "etmiştir — mükemmel rejim etiketlerine sahip oracle_full ajanı dahi sigma_only'yi Sharpe bazında "
    "istatistiksel olarak anlamlı biçimde geçememiştir (p = 0.115). Bonferroni-düzeltilmiş eşik altında "
    "equity farkı da istatistiksel olarak anlamlı kabul edilmemiştir."
))

add_normal(dst, (
    "Bulgu 4 — Dedektör Bağımsızlığı: Null sonuç, %60.4'ten %81.8'e kadar üç farklı dedektör "
    "performansında tutarlı biçimde korunmuştur (ANOVA: F = 0.003, p = 0.997)."
))

add_normal(dst, (
    "Bulgu 5 — Rejime Koşullu Envanter Cezası: Basit rejime koşullu envanter cezası "
    "(ηH = 5×ηL), explicit rejim bilgisini faydalı hale getirmemiştir; sigma_only "
    "varyantı istatistiksel olarak anlamlı biçimde güçlü kalmıştır (p = 0.0016)."
))

# CHANGE E: Bulgu 6
add_normal(dst, (
    "Bulgu 6 — Model Misspecification Robustness: Fill model parametrelerinin (A, k) rejime bağlı "
    "kılındığı zorlaştırılmış ortamda da null sonuç geçerliliğini korumuştur. ppo_sigma_only yine en "
    "yüksek performansı göstermiş; oracle_full ile aradaki fark istatistiksel olarak anlamsız kalmıştır "
    "(p=0.881). Bu bulgu, signal redundancy argümanının çevresel koşullara karşı sağlamlığını "
    "(robustness) doğrulamaktadır."
))

# thesis_27 addition: Bulgu 7 — WP6 categorical-channel finding
add_normal(dst, (
    "Bulgu 7: Sinyal informativeness sweep'i (WP6), sigma_hat'in bozulduğu rejimlerde explicit "
    "kategorik regime label'larının değer kazanacağı hipotezini reddetti. Test edilen kalibrasyon "
    "aralığında, kategorik kanal sigma_only baseline'ının altına yönsel olarak çekti (tüm dört "
    "informative condition'da Cohen's dz arası -0.57 ile -0.91; pratik anlamlılık δ=0.05 yalnızca "
    "lagged condition'da). Kategorik ve sürekli kanalların birlikte kullanılması, bu deneysel "
    "ortamda performansı iyileştirmedi."
))

# CHANGE F: Gelecek Çalışmalar — item (2) updated
add_heading(dst, "Gelecek Çalışmalar", level=2)
dst.add_paragraph(
    "(1) Basit rejime koşullu envanter cezası denendi ve fayda sağlamadı; "
    "daha zengin rejime koşullu ödül tasarımları gelecekte incelenebilir.",
    style='List Paragraph'
)
dst.add_paragraph(
    "(2) Model misspecification strong variant — daha sert parametre farklılaşması "
    "(L: A=3.0/k=2.0, H: A=8.0/k=1.0) ile robustness testinin genişletilmesi.",
    style='List Paragraph'
)
dst.add_paragraph("(3) Paper yayını: iyi nitelendirilmiş null result ile kontrollü RL vs. klasik piyasa yapıcılığı kıyaslaması.", style='List Paragraph')
dst.add_paragraph(
    "(4) Daha agresif sinyal bozulması: WP6 sweep'i kalibrasyon aralığında "
    "(α=0.40 noisy, k=20 lagged, coarsened binning) sigma_hat'in monotonic decay'ini "
    "göstermedi. Daha agresif bozulmalar (örn. α=0.80, k=50) bu çalışmanın kapsamı "
    "dışındadır; kategorik kanalın değer kazandığı bir rejim mevcutsa, bu aralıkta "
    "ortaya çıkabilir.",
    style='List Paragraph'
)

# ================================================================
# KAYNAKLAR
# ================================================================
add_heading(dst, "KAYNAKLAR", level=1)
add_normal(dst, "[1] Avellaneda, M., & Stoikov, S. (2008). High-frequency trading in a limit order book. Quantitative Finance, 8(3), 217-224.")
add_normal(dst, "[2] Schulman, J., Wolski, F., Dhariwal, P., Radford, A., & Klimov, O. (2017). Proximal policy optimization algorithms. arXiv preprint arXiv:1707.06347.")
add_normal(dst, "[3] Spooner, T., Fearnley, J., Savani, R., & Koukorinis, A. (2018). Market making via reinforcement learning. In Proceedings of AAMAS (pp. 434-442).")
add_normal(dst, "[4] Raffin, A., Hill, A., Gleave, A., Kanervisto, A., Ernestus, M., & Dormann, N. (2021). Stable-baselines3: Reliable reinforcement learning implementations. JMLR, 22(268), 1-8.")
add_normal(dst, "[5] Cont, R. (2001). Empirical characteristics of asset returns: Stylized facts. Quantitative Finance, 1(2), 223-236.")
add_normal(dst, "[6] Hamilton, J. D. (1989). A new approach to the economic analysis of nonstationary time series and the business cycle. Econometrica, 57(2), 357-384.")

# ================================================================
# APPENDIX B — Source Code File Index
# ================================================================
add_heading(dst, "Appendix B: Source Code File Index", level=1)
add_normal(dst, "This appendix lists all active source code files in the project repository. Files are organized by module.")

app_h = ["Dosya Yolu", "Açıklama"]
app_r = [
    ["run.py", "Ana iş paketi dispatcher'ı"],
    ["src/run_context.py", "RunContext, logger, seed yönetimi, config snapshot"],
    ["src/w0_smoke.py", "WP0 smoke test"],
    ["src/w1_as_baseline.py", "Avellaneda-Stoikov baseline stratejisi"],
    ["src/w1_compare.py", "WP1 strateji karşılaştırması"],
    ["src/w1_naive_sweep.py", "Naive sabit-spread sweep deneyi"],
    ["src/w3_sanity.py", "Gymnasium ortamı sanity check"],
    ["src/wp1/sim.py", "Piyasa simülatörü (ABM + Poisson dolum modeli)"],
    ["src/wp2/synth_regime.py", "Sentetik rejim üretimi (Markov zinciri)"],
    ["src/wp2/job_w2_synth.py", "WP2 job entry"],
    ["src/wp2/compare_detectors.py", "Rejim dedektörü karşılaştırması"],
    ["src/wp3/env.py", "Gymnasium MMEnv ortamı"],
    ["src/wp4/job_w4_ppo.py", "PPO eğitimi (Stable-Baselines3)"],
    ["src/wp5/job_w5_eval.py", "Out-of-sample değerlendirme (20 tohum)"],
    ["src/wp5/job_w5_ablation_eta.py", "η envanter ceza katsayısı ablasyon deneyi"],
    ["src/wp5/job_w5_ablation_skew.py", "Skew penalty ablasyon deneyi"],
    ["src/wp5/job_w5_detector_compare.py", "Detector robustness deneyi (120 model)"],
    ["src/wp5/analyze_actions.py", "Eylem dağılımı analizi"],
    ["src/wp5/figure_thesis.py", "Tez figürlerinin üretimi"],
    ["src/wp5/figure_thesis_23.py", "Model misspecification ve eta-regime ek figürlerinin üretimi (legacy dosya adı korunmuştur)"],
    ["src/wp5/stats_detector_robustness.py", "İstatistiksel testler (paired t-test ve ANOVA)"],
    ["scripts/gen_thesis_docx.py", "Tez DOCX üretim scripti"],
    ["scripts/gen_thesis_draft.js", "Tez taslak üretim scripti (Node.js)"],
]
add_table(dst, app_h, app_r)

# ================================================================
# SAVE
# ================================================================
dst.save("manuscript/thesis_28.docx")
print("thesis_28.docx saved successfully.")

# ============================================================
# FILE: scripts/gen_decisions_log_12.py
# PURPOSE: Current decisions_log_12 generator documenting audit-remediation decisions.
# STATUS: audit
# ============================================================

"""decisions_log_11.docx -> decisions_log_12.docx

Appends Decisions #48-#51 documenting the audit-remediation sequence:
Lane A audit response, Lane B engineering fixes, Lane C remediation, and
procedural lessons.

Updates summary heading from "En Kritik 27 Karar" -> "En Kritik 31 Karar".
Updates the title-page marker from "Sürüm 27" -> "Sürüm 28" for the new
decision log only. decisions_log_11 remains unchanged.

Decision tables use the same v2 sub-section structure as #47:
  Karar / Seçenekler / Neden / Etki / Not  (no İkilem; trailing Not row)
The 6-row table template is reused; row labels are overwritten.
"""

from __future__ import annotations

import copy
import io
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

SRC = Path("manuscript/decisions_log_11.docx")
DST = Path("manuscript/decisions_log_12.docx")

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _set_cell_text(cell, text, *, bold=None):
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].text = text
        for r in para.runs[1:]:
            r.text = ""
        if bold is not None:
            para.runs[0].bold = bold
    else:
        r = para.add_run(text)
        if bold is not None:
            r.bold = bold


def _build_karar_table_v2(
    doc,
    template_tbl,
    number,
    title,
    wp,
    karar,
    secenekler,
    neden,
    etki,
    note,
):
    """Build a Karar table with Karar / Seçenekler / Neden / Etki / Not rows."""
    new_xml = copy.deepcopy(template_tbl._tbl)
    from docx.table import Table

    new_tbl = Table(new_xml, doc.part)
    _set_cell_text(new_tbl.rows[0].cells[0], wp, bold=True)
    _set_cell_text(new_tbl.rows[0].cells[1], f"#{number}  {title}", bold=True)
    _set_cell_text(new_tbl.rows[1].cells[0], "Karar", bold=True)
    _set_cell_text(new_tbl.rows[1].cells[1], karar)
    _set_cell_text(new_tbl.rows[2].cells[0], "Seçenekler", bold=True)
    _set_cell_text(new_tbl.rows[2].cells[1], secenekler)
    _set_cell_text(new_tbl.rows[3].cells[0], "Neden", bold=True)
    _set_cell_text(new_tbl.rows[3].cells[1], neden)
    _set_cell_text(new_tbl.rows[4].cells[0], "Etki", bold=True)
    _set_cell_text(new_tbl.rows[4].cells[1], etki)
    _set_cell_text(new_tbl.rows[5].cells[0], "Not", bold=True)
    _set_cell_text(new_tbl.rows[5].cells[1], note)
    return new_xml


def _replace_in_paragraph(para, old_text, new_text):
    full = "".join(r.text for r in para.runs)
    if old_text not in full:
        return False
    new_full = full.replace(old_text, new_text)
    if para.runs:
        para.runs[0].text = new_full
        for r in para.runs[1:]:
            r.text = ""
    return True


def _find_ozet_heading(doc):
    for child in doc.element.body.iterchildren():
        if child.tag != W_NS + "p":
            continue
        p_style = child.find(W_NS + "pPr/" + W_NS + "pStyle")
        if p_style is None or p_style.get(W_NS + "val") != "Heading1":
            continue
        text = "".join(t.text or "" for t in child.iter(W_NS + "t"))
        if "Özet" in text:
            return child
    raise RuntimeError("Özet heading not found")


def _blank_p():
    return OxmlElement("w:p")


DECISIONS = [
    {
        "number": 48,
        "title": "Lane A audit response scope",
        "wp": "Audit",
        "karar": (
            "Lane A bulguları thesis_28 içinde metodolojik caveat olarak "
            "yanıtlandı: C1 dwell filter look-ahead ana sonuç boru hattını "
            "etkilemeyen yardımcı detector comparison kapsamına alındı; C2 WP4 "
            "same-series train/eval yalnızca pilot/infrastructure olarak "
            "çerçevelendi; C3 WP6 noisy sigma calibration küçük noisy-only "
            "ölçek sorunu olarak caveat edildi."
        ),
        "secenekler": (
            "(a) Tüm Lane A bulguları için deneyleri yeniden koşturmak  |  "
            "(b) Boru hattı kapsamını ayırıp ana numerik iddiaları WP5 70/30 "
            "OOS ve WP6 sweep çıktıları üzerinden korumak  |  "
            "(c) Bulguları görmezden gelmek"
        ),
        "neden": (
            "C1 için main WP4/WP5/WP6 pipeline'ları causal rv_baseline kullanıyor; "
            "dwell filter offline auxiliary robustness comparison ile sınırlı. "
            "C2 için raporlanan PPO performans sayıları WP4'ten değil WP5'in "
            "70/30 OOS split sonuçlarından geliyor; WP4 eğitim/infrastructure "
            "pilot statüsünde. C3 için WP6 noisy sigma calibration etkisi "
            "noisy-only küçük ölçek kayması olarak sınıflandı; directional "
            "bulgular değişmediği için rerun yapılmadı."
        ),
        "etki": (
            "Thesis_28 sayısal iddiaları korunurken metodolojik sınırlar açık "
            "hale getirildi. Ana sonuçlar yeniden üretilmedi; caveat yaklaşımı "
            "kanıt hiyerarşisini netleştirdi: auxiliary detector analizi ayrı, "
            "reported PPO OOS sonuçları ayrı, noisy calibration sınırlaması ayrı."
        ),
        "note": (
            "Prosedürel ders: audit bulgusunun hangi evidence layer'a ait olduğu "
            "belirlenmeden apply yapılmadı. Bu ayrım yanlış audit mapping ve "
            "gereksiz rerun riskini düşürdü."
        ),
    },
    {
        "number": 49,
        "title": "Lane B engineering guardrails",
        "wp": "Infrastructure",
        "karar": (
            "Lane B kapsamında üç mühendislik güvence düzeltmesi kabul edildi: "
            "CSVMetricLogger schema consistency enforcement, resume sırasında "
            "config snapshot validation, ve CSV/TXT audit stability için "
            ".gitattributes + LF materialization."
        ),
        "secenekler": (
            "(a) Logger ve resume davranışını olduğu gibi bırakmak  |  "
            "(b) Sadece dokümantasyonla riski not etmek  |  "
            "(c) Küçük, local guardrail düzeltmeleriyle sessiz veri/konfigürasyon "
            "drift riskini azaltmak"
        ),
        "neden": (
            "CSVMetricLogger daha önce sessiz kolon uyumsuzluğu üretebilecek "
            "bir append pattern'ine açıktı; 26beecd ile explicit fieldnames "
            "schema check eklendi. Resume path'i yanlış config ile devam etme "
            "riskini taşıyordu; a63e640 ile snapshot mismatch default olarak "
            "reddedildi. Cross-platform hash/audit stabilitesi için ffe8d90 ile "
            ".gitattributes eklendi ve CSV/TXT line-ending davranışı sabitlendi."
        ),
        "etki": (
            "Bu değişiklikler mevcut numerik evidence artifact'larını yeniden "
            "üretmedi; gelecekteki koşuların izlenebilirliğini ve audit "
            "tekrarlanabilirliğini güçlendirdi. Hata yüzeyi runtime veya I/O "
            "guardrail seviyesinde daraltıldı."
        ),
        "note": (
            "Commit zinciri: 26beecd (B8 CSVMetricLogger), a63e640 (B9 resume "
            "validation), ffe8d90 (.gitattributes/LF audit stability)."
        ),
    },
    {
        "number": 50,
        "title": "Lane C active-code remediation",
        "wp": "Codebase",
        "karar": (
            "Lane C kapsamındaki aktif kod ve provenance düzeltmeleri ayrı "
            "küçük commit'lere bölündü: ANOVA framing netleştirildi; thesis "
            "figure ownership header'ları ayrıştırıldı; tarihsel generator'lar "
            "legacy alanına taşındı; aktif scriptlerde hardcoded run path'ler "
            "argparse default'larıyla parametrize edildi; WP2 per-run provenance "
            "artifact adı standardize edildi."
        ),
        "secenekler": (
            "(a) Tek büyük cleanup commit'i yapmak  |  "
            "(b) Sadece doküman notu eklemek  |  "
            "(c) Davranış yüzeyini koruyan, audit edilebilir küçük remediation "
            "commit'leri yapmak"
        ),
        "neden": (
            "0ed125e ile stats_detector_robustness.py ANOVA design framing'i "
            "açıklandı. 43ad288 ile figure_thesis.py ve figure_thesis_23.py "
            "ownership scope'u Fig 1-5 / Fig 6-9 olarak ayrıldı. 5ac5bcb ile "
            "historical thesis ve decisions-log generators scripts/legacy altına "
            "README ile taşındı. dca37b6 ile src/wp5/figure_thesis.py, "
            "src/wp5/figure_thesis_23.py ve scripts/eval_only_seed1to7.py "
            "hardcoded run path'leri byte-identical argparse defaults ile "
            "parametrize edildi. e0c47c5 ile src/wp2/synth_regime.py run-dir "
            "provenance artifact'ı ctx.run_dir/wp2_synth.csv olarak "
            "standardize edildi."
        ),
        "etki": (
            "No-arg davranışlar korunarak bakım yapılabilirlik ve provenance "
            "netliği artırıldı. thesis_28, CSV evidence ve PNG result artifacts "
            "değiştirilmedi; aktif downstream sayısal iddialar etkilenmedi."
        ),
        "note": (
            "C3 parametrize edilen aktif scriptler: src/wp5/figure_thesis.py, "
            "src/wp5/figure_thesis_23.py, scripts/eval_only_seed1to7.py. "
            "C4 WP2 global data/processed/wp2_synth.csv'yi latest convenience "
            "snapshot olarak korudu; run_dir/wp2_synth.csv specific-run "
            "provenance artifact oldu."
        ),
    },
    {
        "number": 51,
        "title": "Audit procedure and protected evidence",
        "wp": "Process",
        "karar": (
            "Audit-remediation sürecinde discovery-before-apply protokolü "
            "benimsendi; deney rerun yapılmadı; canonical Lane C protected "
            "SHA check her commit sonrası 4/4 MATCH kaldı; protected CSV "
            "evidence artifact'ları değiştirilmedi."
        ),
        "secenekler": (
            "(a) Hızlı apply ile tahmini düzeltmeler yapmak  |  "
            "(b) Discovery raporu, scoped apply, static verification ve protected "
            "hash check sırasını izlemek  |  "
            "(c) Rerun ile bütün evidence setini yeniden üretmek"
        ),
        "neden": (
            "Discovery adımı, ilk yanlış protected-file mapping riskini görünür "
            "kıldı ve canonical Lane C Rule dosyalarının doğru setini sabitledi: "
            "results/metrics_detector_compare.csv, "
            "docs/internal/wp6_sweep_full/summary_condition_variant.csv, "
            "docs/internal/wp6_sweep_full/summary_paired_combined_vs_sigma.csv, "
            "docs/internal/wp6_sweep_full/summary_paired_combined_vs_regime.csv. "
            "Her apply adımı bu set üzerinde post-commit 4/4 MATCH doğrulaması "
            "ile kapatıldı."
        ),
        "etki": (
            "Düzeltmeler evidence-preserving olarak kaldı: protected CSV "
            "artifact'ları, PNG result artifact'ları ve thesis_28 dokümanı "
            "değişmedi. Rerun yapılmadığı için sayısal iddialarda yeni "
            "stochastic variance veya artifact drift oluşmadı."
        ),
        "note": (
            "Bu karar, audit sürecinin kendisini reproducibility artifact olarak "
            "kayda geçirir. Ana ders: önce kapsam ve consumer/writer ilişkisi "
            "kanıtlanır, sonra küçük scoped apply yapılır, ardından canonical "
            "hash check ile evidence dokunulmazlığı doğrulanır."
        ),
    },
]


SUMMARY_ROWS = [
    (
        28,
        "Lane A audit response — C1 dwell auxiliary-only, C2 WP4 pilot-only, "
        "C3 WP6 noisy sigma calibration caveat; no rerun",
        "Audit",
    ),
    (
        29,
        "Lane B engineering guardrails — CSVMetricLogger schema, resume config "
        "validation, .gitattributes/LF audit stability",
        "Infrastructure",
    ),
    (
        30,
        "Lane C remediation — ANOVA framing, figure ownership, legacy generator "
        "archive, argparse path defaults, WP2 provenance artifact naming",
        "Codebase",
    ),
    (
        31,
        "Audit procedure — discovery before apply, no reruns, canonical Lane C "
        "protected SHA stayed 4/4 MATCH",
        "Process",
    ),
]


def main():
    shutil.copy2(SRC, DST)
    doc = Document(DST)

    # Last individual Karar table in log_11 is #47 — second-to-last table overall.
    template = doc.tables[-2]

    karar_tables = [
        _build_karar_table_v2(doc, template, **decision)
        for decision in DECISIONS
    ]

    ozet_heading = _find_ozet_heading(doc)
    for tbl_xml in karar_tables:
        ozet_heading.addprevious(_blank_p())
        ozet_heading.addprevious(tbl_xml)

    body = doc.element.body
    for child in body.iterchildren():
        if child.tag != W_NS + "p":
            continue
        text = "".join(t.text or "" for t in child.iter(W_NS + "t"))
        if any(s in text for s in ["En Kritik 27 Karar", "Sürüm 27"]):
            from docx.text.paragraph import Paragraph

            p_obj = Paragraph(child, doc.part)
            _replace_in_paragraph(p_obj, "En Kritik 27 Karar", "En Kritik 31 Karar")
            _replace_in_paragraph(p_obj, "Sürüm 27", "Sürüm 28")

    for para in doc.paragraphs:
        if "ChatGPT ve Claude" in para.text and "27 karar" in para.text:
            _replace_in_paragraph(para, "27 karar", "31 karar")

    sum_tbl = doc.tables[-1]
    template_row = sum_tbl.rows[1]._tr

    def _add_summary_row(num, text, wp):
        new_row = copy.deepcopy(template_row)
        sum_tbl._tbl.append(new_row)
        from docx.table import _Row

        row_obj = _Row(new_row, sum_tbl)
        _set_cell_text(row_obj.cells[0], str(num))
        _set_cell_text(row_obj.cells[1], text)
        _set_cell_text(row_obj.cells[2], wp)

    for row in SUMMARY_ROWS:
        _add_summary_row(*row)

    doc.save(DST)
    print(f"decisions_log_12.docx saved ({DST.stat().st_size:,} bytes)")

    doc2 = Document(DST)
    print("\n=== Verification ===")
    print(f"Total tables: {len(doc2.tables)}")
    print("New karar table headers:")
    for idx in range(-5, -1):
        print(f"  {doc2.tables[idx].rows[0].cells[1].text[:90]}")
    sum2 = doc2.tables[-1]
    print(f"Summary rows: {len(sum2.rows)} (should be 32 = header + 31)")
    print(f"  Last row: {[c.text[:70] for c in sum2.rows[-1].cells]}")
    for p in doc2.paragraphs[:40]:
        if "Sürüm" in p.text or "En Kritik" in p.text:
            print(f"Marker: {p.text}")


if __name__ == "__main__":
    main()

# ============================================================
# FILE: scripts/wp6_plot1_monotonic_gap.py
# PURPOSE: WP6 Plot 1 and condition-variant summary generator.
# STATUS: support
# ============================================================

"""WP6 Chapter 5 — Plot 1: monotonic-gap plot.

Reads docs/internal/wp6_sweep_full/metrics_sweep_full.csv, aggregates mean
sharpe_like by (condition, variant) with 95% CI across 20 seeds (t-critical,
df=n-1), and produces the monotonic-gap plot showing how each variant degrades
across the signal-degradation axis: full -> noisy -> lagged -> coarsened -> none.

By design:
  - regime_only and oracle_pure are condition-invariant (do not consume sigma_hat).
  - sigma_only is structurally undefined in the 'none' condition (NaN, not plotted).
  - 'combined' in 'none' collapses to regime_only (use_sigma=False, regime_source=hat).
"""
from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from scipy import stats

ROOT = Path(__file__).resolve().parents[1]
SWEEP_DIR = ROOT / "docs" / "internal" / "wp6_sweep_full"
METRICS_CSV = SWEEP_DIR / "metrics_sweep_full.csv"
PLOTS_DIR = SWEEP_DIR / "plots"
SUMMARY_CSV = SWEEP_DIR / "summary_condition_variant.csv"

CONDITION_ORDER = ["full", "noisy", "lagged", "coarsened", "none"]
VARIANT_ORDER = [
    "sigma_only",
    "combined",
    "regime_only",
    "oracle_full",
    "oracle_pure",
]

# Visual styling: emphasize sigma_only and combined; anchors are dashed/lighter.
VARIANT_STYLE = {
    "sigma_only":  {"color": "#1f77b4", "linestyle": "-",  "linewidth": 2.4, "marker": "o", "alpha": 1.0, "label": "sigma_only (emphasized)"},
    "combined":    {"color": "#d62728", "linestyle": "-",  "linewidth": 2.4, "marker": "s", "alpha": 1.0, "label": "combined (emphasized)"},
    "regime_only": {"color": "#7f7f7f", "linestyle": "--", "linewidth": 1.4, "marker": "^", "alpha": 0.75, "label": "regime_only (anchor, condition-invariant)"},
    "oracle_full": {"color": "#2ca02c", "linestyle": "-",  "linewidth": 1.6, "marker": "D", "alpha": 0.85, "label": "oracle_full"},
    "oracle_pure": {"color": "#9467bd", "linestyle": "--", "linewidth": 1.4, "marker": "v", "alpha": 0.75, "label": "oracle_pure (anchor, condition-invariant)"},
}


def aggregate(df: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for cond in CONDITION_ORDER:
        for var in VARIANT_ORDER:
            sub = df[(df["condition"] == cond) & (df["variant"] == var)]
            n = len(sub)
            if n == 0:
                rows.append({"condition": cond, "variant": var, "n": 0,
                             "mean": np.nan, "std": np.nan, "sem": np.nan,
                             "ci95_lo": np.nan, "ci95_hi": np.nan})
                continue
            vals = sub["sharpe_like"].to_numpy()
            mean = float(vals.mean())
            std = float(vals.std(ddof=1)) if n > 1 else 0.0
            sem = std / np.sqrt(n) if n > 1 else 0.0
            tcrit = float(stats.t.ppf(0.975, df=n - 1)) if n > 1 else 0.0
            half = tcrit * sem
            rows.append({"condition": cond, "variant": var, "n": n,
                         "mean": mean, "std": std, "sem": sem,
                         "ci95_lo": mean - half, "ci95_hi": mean + half})
    return pd.DataFrame(rows)


def make_plot(summary: pd.DataFrame, out_png: Path, out_pdf: Path) -> None:
    fig, ax = plt.subplots(figsize=(9.0, 5.6))
    x_pos = np.arange(len(CONDITION_ORDER))

    for var in VARIANT_ORDER:
        sub = summary[summary["variant"] == var].set_index("condition").reindex(CONDITION_ORDER)
        means = sub["mean"].to_numpy(dtype=float)
        lo = sub["ci95_lo"].to_numpy(dtype=float)
        hi = sub["ci95_hi"].to_numpy(dtype=float)
        yerr_lo = means - lo
        yerr_hi = hi - means
        mask = ~np.isnan(means)
        style = VARIANT_STYLE[var]
        ax.errorbar(
            x_pos[mask], means[mask],
            yerr=[yerr_lo[mask], yerr_hi[mask]],
            color=style["color"],
            linestyle=style["linestyle"],
            linewidth=style["linewidth"],
            marker=style["marker"],
            markersize=6.5,
            alpha=style["alpha"],
            label=style["label"],
            capsize=3.5,
            elinewidth=1.0,
        )

    ax.set_xticks(x_pos)
    ax.set_xticklabels(CONDITION_ORDER)
    ax.set_xlabel("Signal degradation condition (left = informative, right = sigma removed)")
    ax.set_ylabel("Mean OOS Sharpe-like (20 seeds, 95% CI)")
    ax.set_title("WP6 Signal Informativeness Sweep — monotonic-gap plot")
    ax.grid(True, axis="y", alpha=0.3)
    ax.legend(loc="lower left", fontsize=8.5, framealpha=0.92)

    caption = ("regime_only and oracle_pure are condition-invariant by design "
               "(do not consume sigma_hat); sigma_only is structurally undefined "
               "in the 'none' condition (sigma_hat = 0).")
    fig.text(0.5, 0.01, caption, ha="center", va="bottom",
             fontsize=8.0, style="italic", color="#333333", wrap=True)

    fig.tight_layout(rect=(0, 0.045, 1, 1))
    out_png.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_png, dpi=180, bbox_inches="tight")
    fig.savefig(out_pdf, bbox_inches="tight")
    plt.close(fig)


def main() -> None:
    df = pd.read_csv(METRICS_CSV)
    summary = aggregate(df)
    summary.to_csv(SUMMARY_CSV, index=False)

    out_png = PLOTS_DIR / "monotonic_gap.png"
    out_pdf = PLOTS_DIR / "monotonic_gap.pdf"
    make_plot(summary, out_png, out_pdf)

    print("=== Mean +/- 95% CI by condition x variant ===")
    disp = summary.copy()
    disp["mean+/-CI95"] = disp.apply(
        lambda r: ("NaN" if pd.isna(r["mean"])
                   else f"{r['mean']:.4f} +/- {(r['ci95_hi'] - r['mean']):.4f}"),
        axis=1,
    )
    print(disp[["condition", "variant", "n", "mean+/-CI95"]].to_string(index=False))
    print()
    print(f"Saved summary CSV: {SUMMARY_CSV}")
    print(f"Saved plot (PNG):  {out_png}")
    print(f"Saved plot (PDF):  {out_pdf}")


if __name__ == "__main__":
    main()

# ============================================================
# FILE: scripts/wp6_plot2_paired_seed.py
# PURPOSE: WP6 Plot 2 and paired combined-vs-sigma summary generator.
# STATUS: support
# ============================================================

"""WP6 Chapter 5 -- Plot 2: paired-seed scatter, combined vs sigma_only.

Reads docs/internal/wp6_sweep_full/metrics_sweep_full.csv and produces a
2x2 paired-seed scatter (one panel per informative condition: full, noisy,
lagged, coarsened). Excludes 'none' (sigma_only structurally undefined).

Per panel:
  - x = sigma_only sharpe_like, y = combined sharpe_like, paired on seed
  - y=x identity line; points below identity = combined underperforms
  - paired t-test p, Cohen's dz, n_below/n_total annotated
  - TOST for NON-equivalence (delta=0.05):
      H0: |mean_diff| <= delta  (combined ~equivalent to sigma_only)
      H1: |mean_diff|  > delta  (combined meaningfully different)
    Implemented as union-intersection of two one-sided tests at alpha=0.05;
    reject H0 iff min(p_lower, p_upper) < 0.05. This is the OPPOSITE of WP5's
    TOST usage (where we proved equivalence aware ~ blind).
"""
from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from scipy import stats

ROOT = Path(__file__).resolve().parents[1]
SWEEP_DIR = ROOT / "docs" / "internal" / "wp6_sweep_full"
METRICS_CSV = SWEEP_DIR / "metrics_sweep_full.csv"
PLOTS_DIR = SWEEP_DIR / "plots"
SUMMARY_CSV = SWEEP_DIR / "summary_paired_combined_vs_sigma.csv"

CONDITIONS = ["full", "noisy", "lagged", "coarsened"]  # 'none' excluded by design
DELTA = 0.05
ALPHA = 0.05


def paired_stats(seeds_x: np.ndarray, x: np.ndarray,
                 seeds_y: np.ndarray, y: np.ndarray, condition: str) -> dict:
    """Compute paired stats for combined (y) vs sigma_only (x).

    diff = combined - sigma_only  (negative => combined worse)
    """
    if not np.array_equal(seeds_x, seeds_y):
        raise RuntimeError(f"[{condition}] seed sets differ between sigma_only and combined")

    diff = y - x
    n = len(diff)
    mean_diff = float(diff.mean())
    sd_diff = float(diff.std(ddof=1))
    se_diff = sd_diff / np.sqrt(n)
    t_stat, p_paired = stats.ttest_rel(y, x)
    cohens_dz = mean_diff / sd_diff if sd_diff > 0 else float("nan")
    n_below = int((diff < 0).sum())

    # TOST for non-equivalence (inverse of WP5).
    # Null (equivalence) interval: [-DELTA, +DELTA]. Reject H0 if min one-sided p < alpha.
    df_t = n - 1
    t_lower = (mean_diff + DELTA) / se_diff   # large negative => mean_diff << -DELTA
    p_lower = float(stats.t.cdf(t_lower, df=df_t))
    t_upper = (mean_diff - DELTA) / se_diff   # large positive => mean_diff >>  DELTA
    p_upper = float(1.0 - stats.t.cdf(t_upper, df=df_t))
    tost_p = min(p_lower, p_upper)
    tost_reject = tost_p < ALPHA

    return {
        "condition": condition,
        "n_below": n_below,
        "n_total": int(n),
        "mean_diff": mean_diff,
        "t_stat": float(t_stat),
        "p_paired": float(p_paired),
        "cohens_dz": float(cohens_dz),
        "tost_p": float(tost_p),
        "tost_reject_h0_at_alpha_0.05": bool(tost_reject),
    }


def collect_pairs(df: pd.DataFrame) -> dict:
    """Return {condition: (seeds, sigma_vals, combined_vals)} sorted by seed."""
    out = {}
    for cond in CONDITIONS:
        sigma_df = df[(df["condition"] == cond) & (df["variant"] == "sigma_only")].sort_values("seed")
        comb_df = df[(df["condition"] == cond) & (df["variant"] == "combined")].sort_values("seed")
        seeds_s = sigma_df["seed"].to_numpy()
        seeds_c = comb_df["seed"].to_numpy()
        if not np.array_equal(seeds_s, seeds_c):
            raise RuntimeError(f"[{cond}] seed mismatch: sigma={seeds_s.tolist()} comb={seeds_c.tolist()}")
        out[cond] = (seeds_s,
                     sigma_df["sharpe_like"].to_numpy(dtype=float),
                     comb_df["sharpe_like"].to_numpy(dtype=float))
    return out


def make_plot(pairs: dict, summary_rows: list[dict], out_png: Path, out_pdf: Path) -> None:
    # Global axis limits across all panels for visual comparability.
    all_vals = np.concatenate([np.concatenate([s, c]) for (_, s, c) in pairs.values()])
    pad = 0.03 * (all_vals.max() - all_vals.min())
    lo = float(all_vals.min() - pad)
    hi = float(all_vals.max() + pad)

    fig, axes = plt.subplots(2, 2, figsize=(10.0, 9.4), sharex=True, sharey=True)
    axes = axes.flatten()

    summary_by_cond = {r["condition"]: r for r in summary_rows}

    for ax, cond in zip(axes, CONDITIONS):
        _, x, y = pairs[cond]
        ax.plot([lo, hi], [lo, hi], color="#888888", linestyle="--", linewidth=1.2, zorder=1)
        ax.scatter(x, y, s=58, color="#1f77b4", edgecolor="black",
                   linewidth=0.7, alpha=0.7, zorder=3)
        ax.set_xlim(lo, hi)
        ax.set_ylim(lo, hi)
        ax.set_aspect("equal", adjustable="box")
        ax.grid(True, alpha=0.3)
        ax.set_title(cond, fontsize=12, fontweight="bold")
        ax.set_xlabel("sigma_only Sharpe-like")
        ax.set_ylabel("combined Sharpe-like")

        s = summary_by_cond[cond]
        # Place annotation in upper-left (combined<sigma_only puts points below diag).
        text = (f"{s['n_below']}/{s['n_total']} below identity\n"
                f"mean_diff = {s['mean_diff']:+.3f}\n"
                f"paired t p = {s['p_paired']:.2e}\n"
                f"Cohen's dz = {s['cohens_dz']:+.3f}\n"
                f"TOST non-equiv: p={s['tost_p']:.2e}, "
                f"reject H0: {'yes' if s['tost_reject_h0_at_alpha_0.05'] else 'no'}")
        ax.text(0.03, 0.97, text, transform=ax.transAxes,
                fontsize=8.5, va="top", ha="left",
                bbox=dict(boxstyle="round,pad=0.35",
                          facecolor="white", edgecolor="#bbbbbb", alpha=0.92))

    fig.suptitle("Paired-seed comparison: combined vs sigma_only across conditions",
                 fontsize=13, fontweight="bold", y=0.995)
    caption = ("Points below identity line indicate combined underperforms sigma_only "
               f"for that seed. Delta for TOST = {DELTA:.2f}.")
    fig.text(0.5, 0.005, caption, ha="center", va="bottom",
             fontsize=8.5, style="italic", color="#333333")

    fig.tight_layout(rect=(0, 0.025, 1, 0.97))
    out_png.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_png, dpi=180, bbox_inches="tight")
    fig.savefig(out_pdf, bbox_inches="tight")
    plt.close(fig)


def main() -> None:
    df = pd.read_csv(METRICS_CSV)
    pairs = collect_pairs(df)

    summary_rows = []
    for cond, (seeds, x, y) in pairs.items():
        summary_rows.append(paired_stats(seeds, x, seeds, y, cond))

    summary_df = pd.DataFrame(summary_rows, columns=[
        "condition", "n_below", "n_total", "mean_diff", "t_stat", "p_paired",
        "cohens_dz", "tost_p", "tost_reject_h0_at_alpha_0.05",
    ])
    SUMMARY_CSV.parent.mkdir(parents=True, exist_ok=True)
    summary_df.to_csv(SUMMARY_CSV, index=False)

    out_png = PLOTS_DIR / "paired_seed_combined_vs_sigma.png"
    out_pdf = PLOTS_DIR / "paired_seed_combined_vs_sigma.pdf"
    make_plot(pairs, summary_rows, out_png, out_pdf)

    # === Sanity checks ===
    print("=== Sanity checks ===")
    # 1. n_total == 20 every condition
    n_ok = all(r["n_total"] == 20 for r in summary_rows)
    print(f"  [1] n_total == 20 for every condition: {n_ok}")
    for r in summary_rows:
        print(f"        {r['condition']}: n_total = {r['n_total']}")
    # 2. mean_diff negative everywhere
    neg_ok = all(r["mean_diff"] < 0 for r in summary_rows)
    print(f"  [2] mean_diff < 0 in all 4 conditions: {neg_ok}")
    for r in summary_rows:
        print(f"        {r['condition']}: mean_diff = {r['mean_diff']:+.6f}")
    # 3. Seed sets paired
    print(f"  [3] Seed sets identical between sigma_only and combined per condition:")
    for cond, (seeds, _, _) in pairs.items():
        sigma_seeds = set(df[(df["condition"] == cond) & (df["variant"] == "sigma_only")]["seed"].tolist())
        comb_seeds = set(df[(df["condition"] == cond) & (df["variant"] == "combined")]["seed"].tolist())
        match = sigma_seeds == comb_seeds
        print(f"        {cond}: match = {match} (n_sigma={len(sigma_seeds)}, n_comb={len(comb_seeds)})")
    # 4. 'none' excluded
    excl_ok = "none" not in CONDITIONS
    print(f"  [4] 'none' excluded from analysis: {excl_ok} (CONDITIONS = {CONDITIONS})")
    # 5. Print delta
    print(f"  [5] DELTA used in TOST = {DELTA:.4f}")
    print()

    # === Summary CSV contents ===
    print("=== summary_paired_combined_vs_sigma.csv ===")
    with pd.option_context("display.float_format", "{:.6f}".format,
                            "display.width", 200,
                            "display.max_columns", None):
        print(summary_df.to_string(index=False))
    print()
    print(f"Saved summary CSV: {SUMMARY_CSV}")
    print(f"Saved plot (PNG): {out_png}")
    print(f"Saved plot (PDF): {out_pdf}")


if __name__ == "__main__":
    main()

# ============================================================
# FILE: scripts/wp6_plot3_paired_seed_vs_regime.py
# PURPOSE: WP6 Plot 3 and paired combined-vs-regime summary generator.
# STATUS: support
# ============================================================

"""WP6 Chapter 5 -- Plot 3: paired-seed scatter, combined vs regime_only.

Reads docs/internal/wp6_sweep_full/metrics_sweep_full.csv and produces a
2x2 paired-seed scatter (one panel per informative condition: full, noisy,
lagged, coarsened). Excludes 'none' (combined collapses to regime_only there
by construction).

Question: in the combined variant does the policy still extract value from
sigma_hat beyond what regime_only alone provides?
  (a) combined > regime_only seed-paired (sigma channel still contributes)
  (b) combined ~ regime_only seed-paired (sigma channel crowded out)

TOST direction in this script: EQUIVALENCE (H1: |diff| <= delta).
This is the OPPOSITE of Plot 2 (which tested non-equivalence). For
equivalence we require BOTH one-sided tests to reject, hence
  tost_p = max(p_lower, p_upper)
and we reject H0 of non-equivalence iff tost_p < alpha.
"""
from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from scipy import stats

ROOT = Path(__file__).resolve().parents[1]
SWEEP_DIR = ROOT / "docs" / "internal" / "wp6_sweep_full"
METRICS_CSV = SWEEP_DIR / "metrics_sweep_full.csv"
PLOTS_DIR = SWEEP_DIR / "plots"
SUMMARY_CSV = SWEEP_DIR / "summary_paired_combined_vs_regime.csv"

CONDITIONS = ["full", "noisy", "lagged", "coarsened"]  # 'none' excluded by design
DELTA = 0.05
ALPHA = 0.05


def iqr(arr: np.ndarray) -> float:
    q75, q25 = np.percentile(arr, [75, 25])
    return float(q75 - q25)


def paired_stats(seeds_x: np.ndarray, x: np.ndarray,
                 seeds_y: np.ndarray, y: np.ndarray, condition: str) -> dict:
    """Compute paired stats for combined (y) vs regime_only (x).

    diff = combined - regime_only  (positive => combined adds value)
    """
    if not np.array_equal(seeds_x, seeds_y):
        raise RuntimeError(f"[{condition}] seed sets differ between regime_only and combined")

    diff = y - x
    n = len(diff)
    mean_diff = float(diff.mean())
    sd_diff = float(diff.std(ddof=1))
    se_diff = sd_diff / np.sqrt(n)
    t_stat, p_paired = stats.ttest_rel(y, x)
    cohens_dz = mean_diff / sd_diff if sd_diff > 0 else float("nan")
    n_above = int((diff > 0).sum())

    # TOST for EQUIVALENCE (standard direction, matches WP5 usage).
    # H0: |mean_diff| >  DELTA  (meaningfully different)
    # H1: |mean_diff| <= DELTA  (practically equivalent)
    # Two one-sided tests against the boundary; reject H0 iff BOTH reject.
    df_t = n - 1
    # Lower bound test: H0: mean_diff <= -DELTA vs H1: mean_diff > -DELTA
    t_lower = (mean_diff + DELTA) / se_diff
    p_lower = float(1.0 - stats.t.cdf(t_lower, df=df_t))
    # Upper bound test: H0: mean_diff >=  DELTA vs H1: mean_diff <  DELTA
    t_upper = (mean_diff - DELTA) / se_diff
    p_upper = float(stats.t.cdf(t_upper, df=df_t))
    tost_p = max(p_lower, p_upper)
    tost_reject = tost_p < ALPHA

    return {
        "condition": condition,
        "n_above": n_above,
        "n_total": int(n),
        "mean_diff": mean_diff,
        "t_stat": float(t_stat),
        "p_paired": float(p_paired),
        "cohens_dz": float(cohens_dz),
        "tost_p": float(tost_p),
        "tost_reject_h0_at_alpha_0.05": bool(tost_reject),
        "std_combined": float(np.std(y, ddof=1)),
        "std_regime_only": float(np.std(x, ddof=1)),
        "std_diff": sd_diff,
        "iqr_combined": iqr(y),
        "iqr_regime_only": iqr(x),
        "iqr_diff": iqr(diff),
    }


def collect_pairs(df: pd.DataFrame) -> dict:
    """Return {condition: (seeds, regime_only_vals, combined_vals)} sorted by seed."""
    out = {}
    for cond in CONDITIONS:
        reg_df = df[(df["condition"] == cond) & (df["variant"] == "regime_only")].sort_values("seed")
        comb_df = df[(df["condition"] == cond) & (df["variant"] == "combined")].sort_values("seed")
        seeds_r = reg_df["seed"].to_numpy()
        seeds_c = comb_df["seed"].to_numpy()
        if not np.array_equal(seeds_r, seeds_c):
            raise RuntimeError(f"[{cond}] seed mismatch: regime={seeds_r.tolist()} comb={seeds_c.tolist()}")
        out[cond] = (seeds_r,
                     reg_df["sharpe_like"].to_numpy(dtype=float),
                     comb_df["sharpe_like"].to_numpy(dtype=float))
    return out


def make_plot(pairs: dict, summary_rows: list[dict], out_png: Path, out_pdf: Path) -> None:
    all_vals = np.concatenate([np.concatenate([r, c]) for (_, r, c) in pairs.values()])
    pad = 0.03 * (all_vals.max() - all_vals.min())
    lo = float(all_vals.min() - pad)
    hi = float(all_vals.max() + pad)

    fig, axes = plt.subplots(2, 2, figsize=(10.0, 9.4), sharex=True, sharey=True)
    axes = axes.flatten()
    summary_by_cond = {r["condition"]: r for r in summary_rows}

    for ax, cond in zip(axes, CONDITIONS):
        _, x, y = pairs[cond]
        ax.plot([lo, hi], [lo, hi], color="#888888", linestyle="--", linewidth=1.2, zorder=1)
        ax.scatter(x, y, s=58, color="#1f77b4", edgecolor="black",
                   linewidth=0.7, alpha=0.7, zorder=3)
        ax.set_xlim(lo, hi)
        ax.set_ylim(lo, hi)
        ax.set_aspect("equal", adjustable="box")
        ax.grid(True, alpha=0.3)
        ax.set_title(cond, fontsize=12, fontweight="bold")
        ax.set_xlabel("regime_only Sharpe-like")
        ax.set_ylabel("combined Sharpe-like")

        s = summary_by_cond[cond]
        text = (f"{s['n_above']}/{s['n_total']} above identity\n"
                f"mean_diff = {s['mean_diff']:+.3f}\n"
                f"paired t p = {s['p_paired']:.2e}\n"
                f"Cohen's dz = {s['cohens_dz']:+.3f}\n"
                f"TOST equiv: p={s['tost_p']:.2e}, "
                f"reject H0: {'yes' if s['tost_reject_h0_at_alpha_0.05'] else 'no'}")
        ax.text(0.03, 0.97, text, transform=ax.transAxes,
                fontsize=8.5, va="top", ha="left",
                bbox=dict(boxstyle="round,pad=0.35",
                          facecolor="white", edgecolor="#bbbbbb", alpha=0.92))

    fig.suptitle("Paired-seed comparison: combined vs regime_only across conditions",
                 fontsize=13, fontweight="bold", y=0.995)
    caption = ("Points above identity line indicate combined adds value over "
               f"regime_only for that seed. Delta for TOST = {DELTA:.2f}.")
    fig.text(0.5, 0.005, caption, ha="center", va="bottom",
             fontsize=8.5, style="italic", color="#333333")

    fig.tight_layout(rect=(0, 0.025, 1, 0.97))
    out_png.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_png, dpi=180, bbox_inches="tight")
    fig.savefig(out_pdf, bbox_inches="tight")
    plt.close(fig)


def label_outcome(row: dict) -> str:
    if row["mean_diff"] > 0 and row["p_paired"] < 0.05:
        return "OUTCOME (a): combined > regime_only"
    if row["tost_reject_h0_at_alpha_0.05"]:
        return "OUTCOME (b): combined ~ regime_only"
    return "OUTCOME (c): inconclusive"


def main() -> None:
    df = pd.read_csv(METRICS_CSV)
    pairs = collect_pairs(df)

    summary_rows = []
    for cond, (seeds, x, y) in pairs.items():
        summary_rows.append(paired_stats(seeds, x, seeds, y, cond))

    cols = [
        "condition", "n_above", "n_total", "mean_diff", "t_stat", "p_paired",
        "cohens_dz", "tost_p", "tost_reject_h0_at_alpha_0.05",
        "std_combined", "std_regime_only", "std_diff",
        "iqr_combined", "iqr_regime_only", "iqr_diff",
    ]
    summary_df = pd.DataFrame(summary_rows, columns=cols)
    SUMMARY_CSV.parent.mkdir(parents=True, exist_ok=True)
    summary_df.to_csv(SUMMARY_CSV, index=False)

    out_png = PLOTS_DIR / "paired_seed_combined_vs_regime.png"
    out_pdf = PLOTS_DIR / "paired_seed_combined_vs_regime.pdf"
    make_plot(pairs, summary_rows, out_png, out_pdf)

    # === Sanity checks ===
    print("=== Sanity checks ===")
    # 1. n_total == 20
    n_ok = all(r["n_total"] == 20 for r in summary_rows)
    print(f"  [1] n_total == 20 for every condition: {n_ok}")
    for r in summary_rows:
        print(f"        {r['condition']}: n_total = {r['n_total']}")
    # 2. regime_only is condition-invariant per seed
    print(f"  [2] regime_only is condition-invariant per seed (max |diff| should be 0):")
    ref_seeds, ref_vals, _ = pairs[CONDITIONS[0]]
    ref_map = dict(zip(ref_seeds.tolist(), ref_vals.tolist()))
    for cond in CONDITIONS[1:]:
        s, r, _ = pairs[cond]
        diffs = np.array([r[i] - ref_map[s[i]] for i in range(len(s))])
        max_abs = float(np.abs(diffs).max())
        print(f"        max |regime_only[{CONDITIONS[0]}] - regime_only[{cond}]| = {max_abs:.6e}")
    # 3. Seed sets paired
    print(f"  [3] Seed sets identical between regime_only and combined per condition:")
    for cond in CONDITIONS:
        rseeds = set(df[(df["condition"] == cond) & (df["variant"] == "regime_only")]["seed"].tolist())
        cseeds = set(df[(df["condition"] == cond) & (df["variant"] == "combined")]["seed"].tolist())
        match = rseeds == cseeds
        print(f"        {cond}: match = {match} (n_regime={len(rseeds)}, n_comb={len(cseeds)})")
    # 4. 'none' excluded
    excl_ok = "none" not in CONDITIONS
    print(f"  [4] 'none' excluded from analysis: {excl_ok} (CONDITIONS = {CONDITIONS})")
    # 5. delta
    print(f"  [5] DELTA used in TOST = {DELTA:.4f}")
    # 6. TOST direction
    print(f"  [6] TOST direction in this script: EQUIVALENCE (H1: |diff| <= delta)")
    print()

    # === Summary CSV contents ===
    print("=== summary_paired_combined_vs_regime.csv ===")
    with pd.option_context("display.float_format", "{:.6f}".format,
                            "display.width", 220,
                            "display.max_columns", None):
        print(summary_df.to_string(index=False))
    print()

    # === Outcome labels ===
    print("=== Per-condition outcome (auto-classified) ===")
    for r in summary_rows:
        print(f"  {r['condition']:<10s} -> {label_outcome(r)}")
    print()
    print(f"Saved summary CSV: {SUMMARY_CSV}")
    print(f"Saved plot (PNG): {out_png}")
    print(f"Saved plot (PDF): {out_pdf}")


if __name__ == "__main__":
    main()

# ============================================================
# FILE: tests/test_csv_metric_logger.py
# PURPOSE: Audit test for CSVMetricLogger schema consistency guardrail.
# STATUS: audit
# ============================================================

"""Unit tests for CSVMetricLogger schema enforcement (B8)."""
from pathlib import Path
import sys
import tempfile

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
from src.run_context import CSVMetricLogger


def test_inferred_schema_rejects_mismatched_row():
    with tempfile.TemporaryDirectory() as td:
        path = Path(td) / "metrics.csv"
        logger = CSVMetricLogger(path)
        logger.log({"a": 1, "b": 2})
        try:
            logger.log({"a": 3, "c": 4})
        except ValueError as exc:
            msg = str(exc)
            assert "schema mismatch" in msg, f"unexpected error: {msg}"
            assert "missing=['b']" in msg, f"missing key not reported: {msg}"
            assert "extra=['c']" in msg, f"extra key not reported: {msg}"
        else:
            raise AssertionError("expected ValueError on schema mismatch")


def test_explicit_schema_rejects_mismatched_row_at_first_log():
    with tempfile.TemporaryDirectory() as td:
        path = Path(td) / "metrics.csv"
        logger = CSVMetricLogger(path, fieldnames=["x", "y"])
        try:
            logger.log({"x": 1, "z": 2})
        except ValueError as exc:
            assert "schema mismatch" in str(exc), f"unexpected error: {exc}"
        else:
            raise AssertionError("expected ValueError on schema mismatch at first log")


def test_consistent_rows_succeed():
    with tempfile.TemporaryDirectory() as td:
        path = Path(td) / "metrics.csv"
        logger = CSVMetricLogger(path)
        logger.log({"a": 1, "b": 2})
        logger.log({"a": 3, "b": 4})
        content = path.read_text(encoding="utf-8")
        lines = [l for l in content.splitlines() if l]
        assert len(lines) == 3, f"expected 3 lines (header + 2 rows), got {len(lines)}: {content!r}"
        assert lines[0] == "a,b", f"header should be 'a,b', got {lines[0]!r}"


if __name__ == "__main__":
    test_inferred_schema_rejects_mismatched_row()
    test_explicit_schema_rejects_mismatched_row_at_first_log()
    test_consistent_rows_succeed()
    print("CSVMetricLogger tests passed (3/3)")

# ============================================================
# FILE: tests/test_resume_validation.py
# PURPOSE: Audit test for resume config-snapshot validation guardrail.
# STATUS: audit
# ============================================================

"""Unit tests for resume-mode config validation (B9)."""
from pathlib import Path
import json
import logging
import sys
import tempfile

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
from src.run_context import setup_run, _diff_config


def _write_json(path: Path, obj):
    path.write_text(json.dumps(obj, indent=2), encoding="utf-8")


def _close_thesis_logger():
    """Release log file handles so Windows tempdir cleanup can delete run.log."""
    lg = logging.getLogger("thesis")
    for h in list(lg.handlers):
        h.close()
        lg.removeHandler(h)


def _setup_resumable_run(td: Path, original_cfg: dict, run_id: str = "20260101-000000_seed1_test_abc1234"):
    """Create a fake completed run dir with config_snapshot.json."""
    results_root = td / "results"
    run_dir = results_root / run_id
    run_dir.mkdir(parents=True)
    (run_dir / "plots").mkdir()
    snapshot = {**original_cfg, "run_id": run_id}
    _write_json(run_dir / "config_snapshot.json", snapshot)
    return results_root, run_id


def test_resume_with_matching_config_succeeds():
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        cfg = {"seed": 1, "market": {"mid0": 100, "tick_size": 0.01}, "episode": {"n_steps": 8000}}
        results_root, run_id = _setup_resumable_run(td, cfg)
        cfg_path = td / "cfg.json"
        _write_json(cfg_path, cfg)
        try:
            ctx = setup_run(cfg_path, results_root=results_root, resume_run_id=run_id)
            assert ctx.run_id == run_id
            assert ctx.resume_run_id == run_id
        finally:
            _close_thesis_logger()


def test_resume_with_mismatched_config_refuses():
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        original = {"seed": 1, "market": {"mid0": 100, "tick_size": 0.01}, "episode": {"n_steps": 8000}}
        results_root, run_id = _setup_resumable_run(td, original)
        mismatched = {"seed": 1, "market": {"mid0": 200, "tick_size": 0.01}, "episode": {"n_steps": 8000}}
        cfg_path = td / "cfg.json"
        _write_json(cfg_path, mismatched)
        try:
            try:
                setup_run(cfg_path, results_root=results_root, resume_run_id=run_id)
            except ValueError as exc:
                msg = str(exc)
                assert "Resume config mismatch" in msg, f"unexpected error: {msg}"
                assert "market.mid0" in msg, f"differing key not reported in diff: {msg}"
                assert "100" in msg and "200" in msg, f"old/new values not reported: {msg}"
                assert "resume_force=True" in msg, f"override hint missing: {msg}"
            else:
                raise AssertionError("expected ValueError on mismatched resume config")
        finally:
            _close_thesis_logger()


def test_resume_with_mismatched_config_force_proceeds():
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        original = {"seed": 1, "market": {"mid0": 100, "tick_size": 0.01}, "episode": {"n_steps": 8000}}
        results_root, run_id = _setup_resumable_run(td, original)
        mismatched = {"seed": 1, "market": {"mid0": 200, "tick_size": 0.01}, "episode": {"n_steps": 8000}}
        cfg_path = td / "cfg.json"
        _write_json(cfg_path, mismatched)
        try:
            ctx = setup_run(
                cfg_path,
                results_root=results_root,
                resume_run_id=run_id,
                resume_force=True,
            )
            assert ctx.run_id == run_id
        finally:
            _close_thesis_logger()


def test_resume_with_missing_snapshot_proceeds():
    """Legacy runs predating snapshot saving must remain resumable."""
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        cfg = {"seed": 1, "market": {"mid0": 100, "tick_size": 0.01}, "episode": {"n_steps": 8000}}
        results_root = td / "results"
        run_id = "20260101-000000_seed1_legacy_xyz9876"
        run_dir = results_root / run_id
        run_dir.mkdir(parents=True)
        (run_dir / "plots").mkdir()
        # NOTE: no config_snapshot.json
        cfg_path = td / "cfg.json"
        _write_json(cfg_path, cfg)
        try:
            ctx = setup_run(cfg_path, results_root=results_root, resume_run_id=run_id)
            assert ctx.run_id == run_id
        finally:
            _close_thesis_logger()


def test_diff_config_finds_nested_keys():
    a = {"x": 1, "nested": {"y": 2, "z": 3}}
    b = {"x": 1, "nested": {"y": 2, "z": 4}}
    diffs = _diff_config(a, b)
    assert any("nested.z" in d for d in diffs), f"nested key not found: {diffs}"
    assert any("3" in d and "4" in d for d in diffs), f"values not in diff: {diffs}"


if __name__ == "__main__":
    test_resume_with_matching_config_succeeds()
    test_resume_with_mismatched_config_refuses()
    test_resume_with_mismatched_config_force_proceeds()
    test_resume_with_missing_snapshot_proceeds()
    test_diff_config_finds_nested_keys()
    print("Resume validation tests passed (5/5)")

# ============================================================
# FILE: config/base.json
# PURPOSE: Minimal default dispatcher config.
# STATUS: active config
# ============================================================

{
  "project": "thesis-hfmm",
  "seed": 123,
  "run_tag": "wp0-smoke",
  "n_steps": 200
}

# ============================================================
# FILE: config/w1_naive_sweep.json
# PURPOSE: Canonical WP1 naive-spread sweep config.
# STATUS: active config
# ============================================================

{
  "job": "w1_naive_sweep",
  "seed": 123,
  "market": {
    "mid0": 100.0,
    "tick_size": 0.01,
    "dt": 0.2,
    "sigma_mid_ticks": 0.8
  },
  "exec": {
    "A": 5.0,
    "k": 1.5,
    "fee_bps": 0.2,
    "latency_steps": 1
  },
  "episode": {
    "n_steps": 8000
  },
  "sweep": {
    "half_spreads_ticks": [1, 2, 3, 4, 5]
  }
}

# ============================================================
# FILE: config/w1_as_baseline.json
# PURPOSE: Canonical WP1 Avellaneda-Stoikov baseline config.
# STATUS: active config
# ============================================================

{
  "job": "w1_as_baseline",
  "seed": 123,
  "market": {
    "mid0": 100.0,
    "tick_size": 0.01,
    "dt": 0.2,
    "sigma_mid_ticks": 0.8
  },
  "exec": {
    "A": 5.0,
    "k": 1.5,
    "fee_bps": 0.2,
    "latency_steps": 1
  },
  "episode": {
    "n_steps": 8000
  },
  "as": {
    "gamma": 0.01,
    "horizon_steps": 8000,
    "min_delta_ticks": 1,
    "max_delta_ticks": 25
  }
}

# ============================================================
# FILE: config/w1_compare.json
# PURPOSE: Canonical WP1 naive-vs-AS comparison config.
# STATUS: active config
# ============================================================

{
  "job": "w1_compare",
  "seed": 123,
  "market": {
    "mid0": 100.0,
    "tick_size": 0.01,
    "dt": 0.2,
    "sigma_mid_ticks": 0.8
  },
  "exec": {
    "A": 5.0,
    "k": 1.5,
    "fee_bps": 0.2,
    "latency_steps": 1
  },
  "episode": {
    "n_steps": 8000
  },
  "sweep": {
    "half_spreads_ticks": [1, 2, 3, 4, 5]
  },
  "as": {
    "gamma": 0.01,
    "horizon_steps": 8000,
    "min_delta_ticks": 1,
    "max_delta_ticks": 25
  }
}

# ============================================================
# FILE: config/w2_synth.json
# PURPOSE: Canonical WP2 synthetic-regime generation config.
# STATUS: active config
# ============================================================

{
  "job": "w2_synth",
  "seed": 123,
  "run_tag": "wp2-synth",
  "market": {"mid0": 100.0, "tick_size": 0.01, "dt": 0.2},
  "regime": {
    "sigma_mid_ticks_base": 0.8,
    "sigma_mult": [0.5, 1.0, 2.0],
    "rv_window": 50,
    "warmup_steps": 1000,
    "trans_matrix": [
      [0.9967, 0.0023, 0.0010],
      [0.0042, 0.9917, 0.0041],
      [0.0010, 0.0030, 0.9960]
    ]
  },
  "episode": {"n_steps": 10000}
}

# ============================================================
# FILE: config/w3_sanity.json
# PURPOSE: Canonical WP3 single-mode sanity-check config.
# STATUS: active config
# ============================================================

{
  "job": "w3_sanity",
  "seed": 123,
  "run_tag": "wp3-sanity",
  "market": {"mid0": 100.0, "tick_size": 0.01, "dt": 0.2, "sigma_mid_ticks": 0.8},
  "exec": {"A": 5.0, "k": 1.5, "fee_bps": 0.2, "latency_steps": 1},
  "episode": {"n_steps": 8000, "inv_max_clip": 50},
  "wp3": {"eta": 0.01, "use_regime": true},
  "as": {"gamma": 0.01, "horizon_steps": 8000, "min_delta_ticks": 1, "max_delta_ticks": 25},
  "sweep": {"half_spreads_ticks": [1, 2, 3, 4, 5]},
  "regime": {
    "rv_window": 50,
    "warmup_steps": 1000,
    "sigma_mid_ticks_base": 0.8,
    "sigma_mult": [0.6, 1.0, 1.8],
    "trans_matrix": [
      [0.9967, 0.0023, 0.0010],
      [0.0042, 0.9917, 0.0041],
      [0.0010, 0.0030, 0.9960]
    ]
  }
}

# ============================================================
# FILE: config/w3_sanity_both.json
# PURPOSE: Canonical WP3 aware/blind sanity-check config.
# STATUS: active config
# ============================================================

{
  "job": "w3_sanity",
  "seed": 123,
  "run_tag": "wp3-sanity-both",
  "market": {"mid0": 100.0, "tick_size": 0.01, "dt": 0.2, "sigma_mid_ticks": 0.8},
  "exec": {"A": 5.0, "k": 1.5, "fee_bps": 0.2, "latency_steps": 1},
  "episode": {"n_steps": 8000, "inv_max_clip": 50},
  "wp3": {"eta": 0.01, "use_regime": false},
  "as": {"gamma": 0.01, "horizon_steps": 8000, "min_delta_ticks": 1, "max_delta_ticks": 25},
  "sweep": {"half_spreads_ticks": [1, 2, 3, 4, 5]},
  "regime": {
    "rv_window": 50,
    "warmup_steps": 1000,
    "sigma_mid_ticks_base": 0.8,
    "sigma_mult": [0.6, 1.0, 1.8],
    "trans_matrix": [
      [0.9967, 0.0023, 0.0010],
      [0.0042, 0.9917, 0.0041],
      [0.0010, 0.0030, 0.9960]
    ]
  }
}

# ============================================================
# FILE: config/w4_ppo.json
# PURPOSE: Canonical WP4 PPO pilot-training config.
# STATUS: active config
# ============================================================

{
  "job": "w4_ppo",
  "seed": 123,
  "run_tag": "wp4-ppo",
  "market": {"mid0": 100.0, "tick_size": 0.01, "dt": 0.2, "sigma_mid_ticks": 0.8},
  "exec": {"A": 5.0, "k": 1.5, "fee_bps": 0.2, "latency_steps": 1},
  "episode": {"n_steps": 8000, "inv_max_clip": 50},
  "wp3": {"eta": 0.01, "use_regime": true},
  "as": {"gamma": 0.01, "horizon_steps": 8000, "min_delta_ticks": 1, "max_delta_ticks": 25},
  "sweep": {"half_spreads_ticks": [1, 2, 3, 4, 5]},
  "regime": {
    "rv_window": 50,
    "warmup_steps": 1000,
    "sigma_mid_ticks_base": 0.8,
    "sigma_mult": [0.6, 1.0, 1.8],
    "trans_matrix": [
      [0.9967, 0.0023, 0.0010],
      [0.0042, 0.9917, 0.0041],
      [0.0010, 0.0030, 0.9960]
    ]
  },
  "wp4": {
    "total_timesteps": 200000,
    "learning_rate": 0.0003,
    "n_steps": 2048,
    "batch_size": 256,
    "n_epochs": 10,
    "gamma": 0.999,
    "gae_lambda": 0.95,
    "clip_range": 0.2,
    "ent_coef": 0.0
  }
}

# ============================================================
# FILE: config/w5_main.json
# PURPOSE: Canonical WP5 main OOS evaluation config.
# STATUS: active config
# ============================================================

{
  "job": "w5_eval",
  "seed": 1,
  "run_tag": "wp5-ablation",
  "market": {"mid0": 100.0, "tick_size": 0.01, "dt": 0.2, "sigma_mid_ticks": 0.8},
  "exec": {"A": 5.0, "k": 1.5, "fee_bps": 0.2, "latency_steps": 1},
  "episode": {"n_steps": 8000, "inv_max_clip": 50},
  "wp3": {"eta": 0.001, "use_sigma": true, "regime_source": "hat"},
  "as": {"gamma": 0.01, "horizon_steps": 8000, "min_delta_ticks": 1, "max_delta_ticks": 25},
  "sweep": {"half_spreads_ticks": [1, 2, 3, 4, 5]},
  "regime": {
    "rv_window": 50,
    "warmup_steps": 1000,
    "sigma_mid_ticks_base": 0.8,
    "sigma_mult": [0.6, 1.0, 1.8],
    "trans_matrix": [
      [0.9967, 0.0023, 0.0010],
      [0.0042, 0.9917, 0.0041],
      [0.0010, 0.0030, 0.9960]
    ]
  },
  "wp4": {
    "total_timesteps": 1000000,
    "learning_rate": 0.0003,
    "n_steps": 2048,
    "batch_size": 256,
    "n_epochs": 10,
    "gamma": 0.999,
    "gae_lambda": 0.95,
    "clip_range": 0.2,
    "ent_coef": 0.01,
    "device": "cpu"
  },
  "wp5": {
    "seeds": [1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 18, 19, 20],
    "train_frac": 0.7,
    "naive": {"h": 2, "m": 0}
  }
}

# ============================================================
# FILE: config/w5_eval.json
# PURPOSE: Smaller WP5 OOS evaluation config retained for active dispatcher coverage.
# STATUS: active config
# ============================================================

{
  "job": "w5_eval",
  "seed": 11,
  "run_tag": "wp5-eval-oos",
  "market": {"mid0": 100.0, "tick_size": 0.01, "dt": 0.2, "sigma_mid_ticks": 0.8},
  "exec": {"A": 5.0, "k": 1.5, "fee_bps": 0.2, "latency_steps": 1},
  "episode": {"n_steps": 8000, "inv_max_clip": 50},
  "wp3": {"eta": 0.001, "use_regime": true},
  "as": {"gamma": 0.01, "horizon_steps": 8000, "min_delta_ticks": 1, "max_delta_ticks": 25},
  "sweep": {"half_spreads_ticks": [1, 2, 3, 4, 5]},
  "regime": {
    "rv_window": 50,
    "warmup_steps": 1000,
    "sigma_mid_ticks_base": 0.8,
    "sigma_mult": [0.6, 1.0, 1.8],
    "trans_matrix": [
      [0.9967, 0.0023, 0.0010],
      [0.0042, 0.9917, 0.0041],
      [0.0010, 0.0030, 0.9960]
    ]
  },
  "wp4": {
    "total_timesteps": 200000,
    "learning_rate": 0.0003,
    "n_steps": 2048,
    "batch_size": 256,
    "n_epochs": 10,
    "gamma": 0.999,
    "gae_lambda": 0.95,
    "clip_range": 0.2,
    "ent_coef": 0.0
  },
  "wp5": {
    "seeds": [11, 22, 33],
    "train_frac": 0.7,
    "naive": {"h": 2, "m": 0}
  }
}

# ============================================================
# FILE: config/w5_ablation_eta.json
# PURPOSE: Canonical WP5 eta-ablation config.
# STATUS: active config
# ============================================================

{
  "job": "w5_ablation_eta",
  "seed": 11,
  "run_tag": "wp5-ablation-eta",
  "market": {"mid0": 100.0, "tick_size": 0.01, "dt": 0.2, "sigma_mid_ticks": 0.8},
  "exec": {"A": 5.0, "k": 1.5, "fee_bps": 0.2, "latency_steps": 1},
  "episode": {"n_steps": 8000, "inv_max_clip": 50},
  "wp3": {"eta": 0.01, "use_regime": true},
  "as": {"gamma": 0.01, "horizon_steps": 8000, "min_delta_ticks": 1, "max_delta_ticks": 25},
  "sweep": {"half_spreads_ticks": [1, 2, 3, 4, 5]},
  "regime": {
    "rv_window": 50,
    "warmup_steps": 1000,
    "sigma_mid_ticks_base": 0.8,
    "sigma_mult": [0.6, 1.0, 1.8],
    "trans_matrix": [
      [0.9967, 0.0023, 0.0010],
      [0.0042, 0.9917, 0.0041],
      [0.0010, 0.0030, 0.9960]
    ]
  },
  "wp4": {
    "total_timesteps": 200000,
    "learning_rate": 0.0003,
    "n_steps": 2048,
    "batch_size": 256,
    "n_epochs": 10,
    "gamma": 0.999,
    "gae_lambda": 0.95,
    "clip_range": 0.2,
    "ent_coef": 0.0
  },
  "wp5": {
    "seeds": [11, 22, 33],
    "train_frac": 0.7,
    "naive": {"h": 2, "m": 0},
    "eta_values": [0.0001, 0.001, 0.01]
  }
}

# ============================================================
# FILE: config/w5_ablation_skew.json
# PURPOSE: Canonical WP5 skew-ablation config.
# STATUS: active config
# ============================================================

{
  "job": "w5_ablation_skew",
  "seed": 1,
  "run_tag": "wp5-ablation-skew",
  "market": {"mid0": 100.0, "tick_size": 0.01, "dt": 0.2, "sigma_mid_ticks": 0.8},
  "exec": {"A": 5.0, "k": 1.5, "fee_bps": 0.2, "latency_steps": 1},
  "episode": {"n_steps": 8000, "inv_max_clip": 50},
  "wp3": {"eta": 0.001, "use_regime": true, "skew_penalty_c": 0.0},
  "as": {"gamma": 0.01, "horizon_steps": 8000, "min_delta_ticks": 1, "max_delta_ticks": 25},
  "sweep": {"half_spreads_ticks": [1, 2, 3, 4, 5]},
  "regime": {
    "rv_window": 50,
    "warmup_steps": 1000,
    "sigma_mid_ticks_base": 0.8,
    "sigma_mult": [0.6, 1.0, 1.8],
    "trans_matrix": [
      [0.9967, 0.0023, 0.0010],
      [0.0042, 0.9917, 0.0041],
      [0.0010, 0.0030, 0.9960]
    ]
  },
  "wp4": {
    "total_timesteps": 1000000,
    "learning_rate": 0.0003,
    "n_steps": 2048,
    "batch_size": 256,
    "n_epochs": 10,
    "gamma": 0.999,
    "gae_lambda": 0.95,
    "clip_range": 0.2,
    "ent_coef": 0.01
  },
  "wp5": {
    "seeds": [1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 18, 19, 20],
    "train_frac": 0.7,
    "naive": {"h": 2, "m": 0},
    "skew_c_values": [1e-4]
  }
}

# ============================================================
# FILE: config/w5_detector_pilot.json
# PURPOSE: WP5 detector-robustness pilot config.
# STATUS: active config
# ============================================================

{
  "job": "w5_detector_compare",
  "seed": 1,
  "run_tag": "wp5-detector-pilot",
  "market": {"mid0": 100.0, "tick_size": 0.01, "dt": 0.2, "sigma_mid_ticks": 0.8},
  "exec": {"A": 5.0, "k": 1.5, "fee_bps": 0.2, "latency_steps": 1},
  "episode": {"n_steps": 8000, "inv_max_clip": 50},
  "wp3": {"eta": 0.001, "use_regime": true},
  "as": {"gamma": 0.01, "horizon_steps": 8000, "min_delta_ticks": 1, "max_delta_ticks": 25},
  "sweep": {"half_spreads_ticks": [1, 2, 3, 4, 5]},
  "regime": {
    "rv_window": 50,
    "warmup_steps": 1000,
    "sigma_mid_ticks_base": 0.8,
    "sigma_mult": [0.6, 1.0, 1.8],
    "trans_matrix": [
      [0.9967, 0.0023, 0.0010],
      [0.0042, 0.9917, 0.0041],
      [0.0010, 0.0030, 0.9960]
    ]
  },
  "wp4": {
    "total_timesteps": 1000000,
    "learning_rate": 0.0003,
    "n_steps": 2048,
    "batch_size": 256,
    "n_epochs": 10,
    "gamma": 0.999,
    "gae_lambda": 0.95,
    "clip_range": 0.2,
    "ent_coef": 0.01
  },
  "wp5": {
    "seeds": [1, 2, 3],
    "train_frac": 0.7,
    "naive": {"h": 2, "m": 0}
  }
}

# ============================================================
# FILE: config/w5_detector_full.json
# PURPOSE: Canonical WP5 detector-robustness full-run config.
# STATUS: active config
# ============================================================

{
  "job": "w5_detector_compare",
  "seed": 1,
  "run_tag": "wp5-detector-full",
  "market": {"mid0": 100.0, "tick_size": 0.01, "dt": 0.2, "sigma_mid_ticks": 0.8},
  "exec": {"A": 5.0, "k": 1.5, "fee_bps": 0.2, "latency_steps": 1},
  "episode": {"n_steps": 8000, "inv_max_clip": 50},
  "wp3": {"eta": 0.001, "use_regime": true},
  "as": {"gamma": 0.01, "horizon_steps": 8000, "min_delta_ticks": 1, "max_delta_ticks": 25},
  "sweep": {"half_spreads_ticks": [1, 2, 3, 4, 5]},
  "regime": {
    "rv_window": 50,
    "warmup_steps": 1000,
    "sigma_mid_ticks_base": 0.8,
    "sigma_mult": [0.6, 1.0, 1.8],
    "trans_matrix": [
      [0.9967, 0.0023, 0.0010],
      [0.0042, 0.9917, 0.0041],
      [0.0010, 0.0030, 0.9960]
    ]
  },
  "wp4": {
    "total_timesteps": 1000000,
    "learning_rate": 0.0003,
    "n_steps": 2048,
    "batch_size": 256,
    "n_epochs": 10,
    "gamma": 0.999,
    "gae_lambda": 0.95,
    "clip_range": 0.2,
    "ent_coef": 0.01
  },
  "wp5": {
    "seeds": [1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 18, 19, 20],
    "train_frac": 0.7,
    "naive": {"h": 2, "m": 0}
  }
}

# ============================================================
# FILE: config/w5_eta_regime.json
# PURPOSE: Canonical WP5 regime-conditional eta config.
# STATUS: active config
# ============================================================

{
  "job": "w5_eval",
  "seed": 42,
  "run_tag": "w5-eta-regime",
  "market": {"mid0": 100.0, "tick_size": 0.01, "dt": 0.2, "sigma_mid_ticks": 0.8},
  "exec": {"A": 5.0, "k": 1.5, "fee_bps": 0.2, "latency_steps": 1},
  "episode": {"n_steps": 8000, "inv_max_clip": 50},
  "wp3": {
    "eta": 0.001,
    "eta_regime": {
      "L": 0.0005,
      "M": 0.001,
      "H": 0.0025
    },
    "use_sigma": true,
    "regime_source": "hat"
  },
  "as": {"gamma": 0.01, "horizon_steps": 8000, "min_delta_ticks": 1, "max_delta_ticks": 25},
  "sweep": {"half_spreads_ticks": [1, 2, 3, 4, 5]},
  "regime": {
    "rv_window": 50,
    "warmup_steps": 1000,
    "sigma_mid_ticks_base": 0.8,
    "sigma_mult": [0.6, 1.0, 1.8],
    "trans_matrix": [
      [0.9967, 0.0023, 0.0010],
      [0.0042, 0.9917, 0.0041],
      [0.0010, 0.0030, 0.9960]
    ]
  },
  "wp4": {
    "total_timesteps": 1000000,
    "learning_rate": 0.0003,
    "n_steps": 2048,
    "batch_size": 256,
    "n_epochs": 10,
    "gamma": 0.999,
    "gae_lambda": 0.95,
    "clip_range": 0.2,
    "ent_coef": 0.01,
    "device": "cpu"
  },
  "wp5": {
    "seeds": [1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 18, 19, 20],
    "train_frac": 0.7,
    "naive": {"h": 2, "m": 0}
  }
}

# ============================================================
# FILE: config/w5_misspec_mild.json
# PURPOSE: Canonical WP5 mild model-misspecification config.
# STATUS: active config
# ============================================================

{
  "job": "w5_eval",
  "run_tag": "w5-misspec-mild",
  "seed": 1,
  "market": {
    "mid0": 100.0,
    "tick_size": 0.01,
    "dt": 0.2,
    "sigma_mid_ticks": 0.8
  },
  "exec": {
    "A": 5.0,
    "k": 1.5,
    "fee_bps": 0.2,
    "latency_steps": 1
  },
  "misspec": {
    "enabled": true,
    "params": {
      "L": {"A": 4.0, "k": 1.8},
      "M": {"A": 5.0, "k": 1.5},
      "H": {"A": 6.0, "k": 1.2}
    }
  },
  "episode": {
    "n_steps": 8000,
    "inv_max_clip": 50
  },
  "wp3": {
    "eta": 0.001,
    "use_regime": true,
    "use_sigma": true,
    "regime_source": "hat"
  },
  "wp4": {
    "total_timesteps": 1000000,
    "learning_rate": 0.0003,
    "n_steps": 2048,
    "batch_size": 256,
    "n_epochs": 10,
    "gamma": 0.999,
    "gae_lambda": 0.95,
    "clip_range": 0.2,
    "ent_coef": 0.01,
    "device": "cuda"
  },
  "wp5": {
    "seeds": [1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 18, 19, 20],
    "train_frac": 0.7,
    "naive": {"h": 2, "m": 0},
    "variants": ["sigma_only", "regime_only", "combined", "oracle_pure", "oracle_full"]
  },
  "as": {
    "gamma": 0.01,
    "horizon_steps": 8000,
    "min_delta_ticks": 1,
    "max_delta_ticks": 25
  },
  "regime": {
    "rv_window": 50,
    "warmup_steps": 400,
    "sigma_mid_ticks_base": 0.8,
    "sigma_mult": [0.5, 1.0, 2.0],
    "trans_matrix": [[0.9967, 0.0023, 0.001], [0.0042, 0.9917, 0.0041], [0.001, 0.003, 0.996]]
  }
}

# ============================================================
# FILE: config/w55_audit.json
# PURPOSE: Canonical WP5.5 offline signal-degradation audit config.
# STATUS: active config
# ============================================================

{
  "project": "thesis-hfmm",
  "seed": 123,
  "run_tag": "wp55-signal-audit",
  "job": "w55_audit",
  "market": {"mid0": 100.0, "tick_size": 0.01, "dt": 0.2, "sigma_mid_ticks": 0.8},
  "exec": {"A": 5.0, "k": 1.5, "fee_bps": 0.2, "latency_steps": 1},
  "regime": {
    "rv_window": 50,
    "warmup_steps": 1000,
    "sigma_mid_ticks_base": 0.8,
    "sigma_mult": [0.6, 1.0, 1.8],
    "trans_matrix": [
      [0.9967, 0.0023, 0.0010],
      [0.0042, 0.9917, 0.0041],
      [0.0010, 0.0030, 0.9960]
    ]
  },
  "episode": {"n_steps": 8000},
  "audit": {
    "noise_std": "auto",
    "noise_std_rationale": "When 'auto', job computes noise_std = 0.5 * clean_sigma_std_post_warmup so the 'moderate estimation error' regime tracks the canonical signal scale.",
    "lag_k_steps": 5,
    "n_bins": 5,
    "fill_value": 0.0,
    "threshold_band_pct": 0.05
  }
}

# ============================================================
# FILE: config/w55_runtime.json
# PURPOSE: WP5.5 runtime profiling config.
# STATUS: active config
# ============================================================

{
  "project": "thesis-hfmm",
  "seed": 42,
  "run_tag": "wp55-runtime-benchmark",
  "job": "w55_runtime",
  "market": {"mid0": 100.0, "tick_size": 0.01, "dt": 1.0, "sigma_mid_ticks": 0.5},
  "regime": {"rv_window": 50, "warmup_steps": 1000},
  "episode": {"n_steps": 8000},
  "exec": {"A": 5.0, "k": 1.5, "fee_bps": 0.1, "latency_steps": 0},
  "ppo": {
    "total_timesteps": 100000,
    "n_steps": 2048,
    "batch_size": 64,
    "n_epochs": 10,
    "learning_rate": 3e-4,
    "gamma": 0.99,
    "policy": "MlpPolicy"
  },
  "env": {"use_regime": false, "eta": 1e-3},
  "benchmark": {"devices": ["cpu", "cuda"], "seeds": [42, 123]}
}

# ============================================================
# FILE: config/w55_calibration.json
# PURPOSE: WP5.5 signal-degradation calibration config.
# STATUS: active config
# ============================================================

{
  "project": "thesis-hfmm",
  "seed": 42,
  "run_tag": "wp55-calibration",
  "job": "w55_calibration",
  "market": {"mid0": 100.0, "tick_size": 0.01, "dt": 0.2, "sigma_mid_ticks": 0.8},
  "exec": {"A": 5.0, "k": 1.5, "fee_bps": 0.2, "latency_steps": 1},
  "regime": {
    "rv_window": 50,
    "warmup_steps": 1000,
    "sigma_mid_ticks_base": 0.8,
    "sigma_mult": [0.6, 1.0, 1.8],
    "trans_matrix": [
      [0.9967, 0.0023, 0.0010],
      [0.0042, 0.9917, 0.0041],
      [0.0010, 0.0030, 0.9960]
    ]
  },
  "episode": {"n_steps": 8000},
  "calibration": {
    "alpha_values": [0.0, 0.05, 0.10, 0.20, 0.40, 0.80, 1.60],
    "k_values": [0, 1, 2, 5, 10, 20, 50],
    "n_seeds": 5
  }
}

# ============================================================
# FILE: config/w6_sweep_pilot.json
# PURPOSE: WP6 signal-informativeness pilot sweep config.
# STATUS: active config
# ============================================================

{
  "project": "thesis-hfmm",
  "seed": 42,
  "run_tag": "wp6-sweep-pilot",
  "job": "w6_sweep_pilot",
  "market": {"mid0": 100.0, "tick_size": 0.01, "dt": 0.2, "sigma_mid_ticks": 0.8},
  "exec": {"A": 5.0, "k": 1.5, "fee_bps": 0.2, "latency_steps": 1},
  "episode": {"n_steps": 8000, "inv_max_clip": 50},
  "regime": {
    "rv_window": 50,
    "warmup_steps": 1000,
    "sigma_mid_ticks_base": 0.8,
    "sigma_mult": [0.6, 1.0, 1.8],
    "trans_matrix": [
      [0.9967, 0.0023, 0.0010],
      [0.0042, 0.9917, 0.0041],
      [0.0010, 0.0030, 0.9960]
    ]
  },
  "wp3": {"eta": 0.001, "use_sigma": true, "regime_source": "hat"},
  "wp4": {
    "total_timesteps": 200000,
    "learning_rate": 0.0003,
    "n_steps": 2048,
    "batch_size": 256,
    "n_epochs": 10,
    "gamma": 0.999,
    "gae_lambda": 0.95,
    "clip_range": 0.2,
    "ent_coef": 0.01,
    "device": "cpu"
  },
  "sweep": {
    "conditions": ["full", "noisy", "lagged", "coarsened", "none"],
    "variants": ["sigma_only", "regime_only", "combined", "oracle_pure", "oracle_full"],
    "omit_cells": [["none", "sigma_only"]],
    "seeds": [42, 43, 44],
    "timesteps": 200000,
    "noisy_alpha": 0.40,
    "lagged_k": 20,
    "coarsened_n_bins": 5
  }
}

# ============================================================
# FILE: config/w6_sweep_full.json
# PURPOSE: Canonical WP6 full signal-informativeness sweep config.
# STATUS: active config
# ============================================================

{
  "project": "thesis-hfmm",
  "seed": 42,
  "run_tag": "wp6-sweep-full",
  "job": "w6_sweep_full",
  "market": {"mid0": 100.0, "tick_size": 0.01, "dt": 0.2, "sigma_mid_ticks": 0.8},
  "exec": {"A": 5.0, "k": 1.5, "fee_bps": 0.2, "latency_steps": 1},
  "episode": {"n_steps": 8000, "inv_max_clip": 50},
  "regime": {
    "rv_window": 50,
    "warmup_steps": 1000,
    "sigma_mid_ticks_base": 0.8,
    "sigma_mult": [0.6, 1.0, 1.8],
    "trans_matrix": [
      [0.9967, 0.0023, 0.0010],
      [0.0042, 0.9917, 0.0041],
      [0.0010, 0.0030, 0.9960]
    ]
  },
  "wp3": {"eta": 0.001, "use_sigma": true, "regime_source": "hat"},
  "wp4": {
    "total_timesteps": 1000000,
    "learning_rate": 0.0003,
    "n_steps": 2048,
    "batch_size": 256,
    "n_epochs": 10,
    "gamma": 0.999,
    "gae_lambda": 0.95,
    "clip_range": 0.2,
    "ent_coef": 0.01,
    "device": "cpu"
  },
  "sweep": {
    "conditions": ["full", "noisy", "lagged", "coarsened", "none"],
    "variants": ["sigma_only", "regime_only", "combined", "oracle_pure", "oracle_full"],
    "omit_cells": [["none", "sigma_only"]],
    "seeds": [42, 43, 44, 45, 46, 47, 48, 49, 50, 51, 52, 53, 54, 55, 56, 57, 58, 59, 60, 61],
    "timesteps": 1000000,
    "noisy_alpha": 0.40,
    "lagged_k": 20,
    "coarsened_n_bins": 5
  }
}
