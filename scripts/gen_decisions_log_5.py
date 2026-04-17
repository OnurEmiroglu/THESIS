"""decisions_log_4.docx -> decisions_log_5.docx:
Karar #33 (misspec deney tasarimi), #34 (strong variant skip), #35 (oracle >=90%),
#36 (thesis_19 duzeltmeler), #37 (thesis_20 Tablo 3 Bonferroni),
#38 (thesis_21 Sec4.7 varyant, Tablo7 Std Sharpe, oracle paradox) ekle, ozet tabloyu guncelle."""

from __future__ import annotations
import sys, io, shutil
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from lxml import etree

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
SRC = Path("manuscript/decisions_log_4.docx")
DST = Path("manuscript/decisions_log_5.docx")


def _styled_run(paragraph, text, bold=False, size_pt=10, font_name="Arial", color_hex=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = font_name
    if color_hex:
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor.from_string(color_hex)
    return run


def _set_cell_props(cell, width_dxa, fill_color):
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = etree.SubElement(tc, qn("w:tcPr"))
        tc.insert(0, tcPr)
    for old in tcPr.findall(qn("w:tcW")):
        tcPr.remove(old)
    tcW = etree.SubElement(tcPr, qn("w:tcW"))
    tcW.set(qn("w:w"), width_dxa)
    tcW.set(qn("w:type"), "dxa")
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    borders = etree.SubElement(tcPr, qn("w:tcBorders"))
    for side in ("top", "left", "bottom", "right"):
        b = etree.SubElement(borders, qn(f"w:{side}"))
        b.set(qn("w:val"), "single")
        b.set(qn("w:color"), "BBCFE0")
        b.set(qn("w:sz"), "1")
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = etree.SubElement(tcPr, qn("w:shd"))
    shd.set(qn("w:fill"), fill_color)
    shd.set(qn("w:val"), "clear")
    for old in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(old)
    tcMar = etree.SubElement(tcPr, qn("w:tcMar"))
    for side, val in [("top", "80"), ("left", "160"), ("bottom", "80"), ("right", "160")]:
        m = etree.SubElement(tcMar, qn(f"w:{side}"))
        m.set(qn("w:type"), "dxa")
        m.set(qn("w:w"), val)


def _add_decision_table(doc, body, insert_before_elem, wp_tag, number, title, rows_data):
    """Insert a 2-column, 6-row decision table matching the style of #26-#32."""
    from docx.oxml import OxmlElement

    # Add a blank spacer paragraph before the table
    spacer = OxmlElement("w:p")
    body.insert(list(body).index(insert_before_elem), spacer)

    # Create table element
    tbl = doc.add_table(rows=6, cols=2, style='Normal Table')

    # Row 0: [WP tag, #N title]
    labels = [wp_tag, f"#{number}  {title}"]
    row_labels = ["", "İkilem", "Seçenekler", "Karar", "Neden", "Etki"]
    all_data = [labels] + [[row_labels[i], rows_data[i-1]] for i in range(1, 6)]

    for ri, (c0_text, c1_text) in enumerate(all_data):
        c0 = tbl.rows[ri].cells[0]
        c1 = tbl.rows[ri].cells[1]

        c0.paragraphs[0].clear()
        c1.paragraphs[0].clear()

        is_header = (ri == 0)
        fill_0 = "1F4E79" if is_header else "E8F4FD"
        fill_1 = "1F4E79" if is_header else "FFFFFF"
        color_0 = "FFFFFF" if is_header else "1F4E79"
        color_1 = "FFFFFF" if is_header else "2C2C2C"

        _styled_run(c0.paragraphs[0], c0_text,
                     bold=True, size_pt=10, color_hex=color_0)
        _styled_run(c1.paragraphs[0], c1_text,
                     bold=is_header, size_pt=10 if not is_header else 11,
                     color_hex=color_1)

        _set_cell_props(c0, "1400", fill_0)
        _set_cell_props(c1, "6960", fill_1)

    # Move table from end of document to before insert_before_elem
    tbl_elem = tbl._tbl
    body.remove(tbl_elem)
    body.insert(list(body).index(insert_before_elem), tbl_elem)

    return tbl


def _add_summary_row(sum_tbl, num_str, description, wp_tag):
    """Add a styled row to the summary table."""
    new_row = sum_tbl.add_row()
    cells = new_row.cells

    cells[0].paragraphs[0].clear()
    _styled_run(cells[0].paragraphs[0], num_str, bold=True, size_pt=10, color_hex="1F4E79")
    _set_cell_props(cells[0], "600", "E8F4FD")

    cells[1].paragraphs[0].clear()
    _styled_run(cells[1].paragraphs[0], description, bold=False, size_pt=10, color_hex="2C2C2C")
    _set_cell_props(cells[1], "6960", "E8F4FD")

    cells[2].paragraphs[0].clear()
    _styled_run(cells[2].paragraphs[0], wp_tag, bold=False, size_pt=10, color_hex="2C2C2C")
    _set_cell_props(cells[2], "1000", "E8F4FD")


def main():
    shutil.copy2(SRC, DST)
    doc = Document(DST)
    body = doc.element.body

    # ── Find the summary heading to insert decisions before it ──
    summary_heading = None
    for p in doc.paragraphs:
        if "Kritik 12 Karar" in p.text:
            summary_heading = p
            break
    assert summary_heading is not None, "Summary heading 'Kritik 12 Karar' not found"

    insert_elem = summary_heading._element

    # ── KARAR #33 ──
    _add_decision_table(
        doc, body,
        insert_before_elem=insert_elem,
        wp_tag="WP5",
        number=33,
        title="Model Misspecification Deney Tasarımı",
        rows_data=[
            (
                "Gerçek piyasalarda fill parametreleri rejime göre değişebilir. "
                "Sabit A, k varsayımı altındaki null result ne kadar sağlam?"
            ),
            (
                "A) Sabit A=5.0, k=1.5 ile devam et  |  "
                "B) Rejime bağlı A, k ile misspecification deneyi kur"
            ),
            (
                "B — Fill model parametreleri A ve k rejime bağlanmıştır. "
                "MMSimulator generic kalır; env.py step() içinde regime_true'ya göre "
                "ExecParams override edilir. Mild parametreler: L(A=4.0, k=1.8), "
                "M(A=5.0, k=1.5), H(A=6.0, k=1.2)."
            ),
            (
                "sim.py'a dokunmadan çevresel misspecification eklenmesi mimari "
                "temizliği korur. ExecParams dataclass frozen=True kısıtı kaldırılarak "
                "mutable hale getirildi (yalnızca ExecParams; MarketParams frozen kaldı)."
            ),
            (
                "20 tohum, 1M timestep, 5-varyant full run tamamlandı. "
                "ppo_sigma_only vs ppo_oracle_full: p=0.881. Null result korundu."
            ),
        ],
    )

    # ── KARAR #34 ──
    _add_decision_table(
        doc, body,
        insert_before_elem=insert_elem,
        wp_tag="WP5",
        number=34,
        title="Strong Variant Çalıştırılmaması",
        rows_data=[
            (
                "Mild misspecification null result verdi. Strong variant "
                "(L: A=3.0/k=2.0, H: A=8.0/k=1.0) çalıştırılmalı mı?"
            ),
            (
                "A) Strong variant çalıştır  |  "
                "B) Mild sonuçlarla yetinip danışman toplantısına hazırlan"
            ),
            (
                "B — Strong variant çalıştırılmamıştır."
            ),
            (
                "Mild misspecification sonuçları (ppo_sigma_only vs ppo_oracle_full: "
                "p=0.881; vs ppo_combined: p=0.217) null result'ın bu ortamda da "
                "geçerli olduğunu açıkça göstermiştir. p=0.88 gibi güçlü bir null "
                "result varlığında strong variant'ın anlamlı ek katkı sağlaması "
                "beklenmemekte; danışman toplantısı yaklaşmaktadır."
            ),
            (
                "Strong variant gelecek çalışmalar listesine eklendi. "
                "Mevcut bulgular danışman sunumu için yeterlidir."
            ),
        ],
    )

    # ── KARAR #35 ──
    _add_decision_table(
        doc, body,
        insert_before_elem=insert_elem,
        wp_tag="WP5",
        number=35,
        title="Oracle Deneyi >=90% Dedektör Hedefini Karşılıyor",
        rows_data=[
            (
                "Danışman >=90% doğruluklu dedektör hedefi belirtmişti. "
                "HMM %81.8'de kaldı. Yeterli mi?"
            ),
            (
                "A) Daha iyi dedektör geliştir  |  "
                "B) Oracle deneyi ile %100 doğruluğun bile fark yaratmadığını göster"
            ),
            (
                "B — Oracle deneyi (ppo_oracle_full, ppo_oracle_pure) %100 "
                "doğrulukla eşdeğerdir; agent gerçek rejim etiketini (regime_true) "
                "doğrudan gözlemlemektedir."
            ),
            (
                "oracle_full sigma_only'yi geçememiştir (p=0.115). Bu bulgu "
                ">=90% hedefini aşan, daha güçlü bir testtir — mükemmel rejim "
                "bilgisi bile katkı sağlamamaktadır."
            ),
            (
                "Dedektör kalitesi argümanı kapatılmıştır. Oracle deneyi, "
                "herhangi bir dedektör doğruluğunun yeterli olacağını kanıtlamıştır."
            ),
        ],
    )

    # ── KARAR #36 ──
    _add_decision_table(
        doc, body,
        insert_before_elem=insert_elem,
        wp_tag="WP5",
        number=36,
        title="thesis_19 — Birim Hatası ve p-Değeri Düzeltmesi",
        rows_data=[
            (
                "thesis_18'de iki faktüel hata tespit edildi: "
                "inv_p99 birimi 'tick' olarak yazılmış (doğrusu 'lot'), "
                "oracle paradox p-değeri sigma_only vs oracle_full için 0.301 olarak etiketlenmiş "
                "(doğrusu 0.115; 0.301 oracle_full vs combined'a aittir)."
            ),
            (
                "A) Errata notu ekle  |  "
                "B) Yeni sürüm (thesis_19) oluştur"
            ),
            (
                "B — thesis_19 oluşturuldu. inv_p99 birimi tick → lot olarak düzeltildi. "
                "sigma_only vs oracle_full p-değeri 0.301 → 0.115 olarak düzeltildi "
                "(stats_ablation.txt: ppo_sigma_only vs ppo_oracle_full sharpe_like p=1.15e-01)."
            ),
            (
                "inv_p99 envanter pozisyon büyüklüğünü ölçer — birimi lot/kontrat, tick değil. "
                "p=0.301, oracle_full vs combined karşılaştırmasına aittir; "
                "sigma_only vs oracle_full Sharpe p-değeri 0.115'tir."
            ),
            (
                "İki faktüel hata giderildi; tez argümanı değişmedi, "
                "yalnızca raporlama doğruluğu artırıldı."
            ),
        ],
    )

    # ── KARAR #37 ──
    _add_decision_table(
        doc, body,
        insert_before_elem=insert_elem,
        wp_tag="WP5",
        number=37,
        title="thesis_20 — Tablo 3 Bonferroni Düzeltmesi",
        rows_data=[
            (
                "thesis_19 Tablo 3'te 10 karşılaştırma rapor edilmiş ancak "
                "combined vs AS ve combined vs naive equity satırları eksikti. "
                "Bonferroni α = 0.01/10 = 0.001 olarak yazılmıştı."
            ),
            (
                "A) Sadece metin düzelt, tabloyu 10 satır bırak  |  "
                "B) Eksik 2 equity satırını ekle, Bonferroni'yi 12'ye güncelle"
            ),
            (
                "B — Tablo 3'e combined vs AS (final_equity, t=-1.064, p=3.01e-01) "
                "ve combined vs naive (final_equity, t=-0.728, p=4.76e-01) satırları eklendi. "
                "Bonferroni düzeltmesi 10→12 karşılaştırma, α = 0.01/12 = 0.00083 olarak güncellendi."
            ),
            (
                "stats_ablation.txt'te 12 karşılaştırma mevcuttur; tablo ile kaynak "
                "arasındaki tutarsızlık giderildi. Yeni equity satırları 'Hayır' (anlamsız) "
                "olduğundan sonuçlar değişmedi, yalnızca raporlama bütünlüğü sağlandı."
            ),
            (
                "Tablo 3 artık stats_ablation.txt ile birebir uyumlu; "
                "Bonferroni eşiği daha muhafazakâr hale geldi (0.001→0.00083)."
            ),
        ],
    )

    # ── KARAR #38 ──
    _add_decision_table(
        doc, body,
        insert_before_elem=insert_elem,
        wp_tag="WP5",
        number=38,
        title="thesis_21 — Sec 4.7 Varyant, Tablo 7 Sütun, Oracle Paradox Güncelleme",
        rows_data=[
            (
                "thesis_20'de üç sorun tespit edildi: (1) Section 4.7'de 'üç varyant' yazılmış, "
                "doğrusu beş varyant; (2) Tablo 7'de 'Std Sharpe' sütunu tüm değerleri '—' iken gereksiz; "
                "(3) Abstract ve Conclusion'da oracle paradox, oracle_full vs combined p=0.301 üzerinden "
                "ifade edilmiş — daha doğrudan kanıt sigma_only vs oracle_full p=0.115'tir."
            ),
            (
                "A) Küçük düzeltmeleri atla  |  "
                "B) Üçünü birden düzelt, thesis_21 oluştur"
            ),
            (
                "B — (1) 'üç varyant' → 'beş varyant'; (2) Tablo 7'den Std Sharpe sütunu kaldırıldı; "
                "(3) Abstract ve Conclusion'da oracle paradox cümlesi sigma_only vs oracle_full p=0.115 "
                "ifadesiyle güncellendi."
            ),
            (
                "Beş varyant fiilen eğitilmiştir (sigma_only, combined, oracle_full, oracle_pure, regime_only). "
                "Std Sharpe sütunu bilgi taşımıyordu. p=0.115 doğrudan sigma_only vs oracle_full karşılaştırmasını "
                "yansıtır ve oracle paradox argümanını daha güçlü destekler."
            ),
            (
                "Faktüel tutarlılık artırıldı; ana argüman değişmedi."
            ),
        ],
    )

    # ── Özet tablosu güncelle: 12 -> 18 ──
    for p in doc.paragraphs:
        if "Kritik 12 Karar" in p.text:
            for run in p.runs:
                if "12" in run.text:
                    run.text = run.text.replace("En Kritik 12 Karar", "En Kritik 18 Karar")

    # Add rows #13, #14, #15, #16, #17, #18 to summary table (last table)
    sum_tbl = doc.tables[-1]

    _add_summary_row(
        sum_tbl, "13",
        "Model misspecification deney tasarımı — rejime bağlı A, k ile null result sağlamlığını test etmek (p=0.881)",
        "WP5",
    )
    _add_summary_row(
        sum_tbl, "14",
        "Strong variant çalıştırılmaması — mild sonuçların yeterliliği ve zaman kısıtı",
        "WP5",
    )
    _add_summary_row(
        sum_tbl, "15",
        "Oracle deneyi >=90% dedektör hedefini karşılıyor — %100 doğruluk bile sigma_only'yi geçemedi (p=0.115)",
        "WP5",
    )
    _add_summary_row(
        sum_tbl, "16",
        "thesis_19 — inv_p99 birim hatası (tick→lot) ve oracle paradox p-değeri etiketi (0.301→0.115) düzeltildi",
        "WP5",
    )
    _add_summary_row(
        sum_tbl, "17",
        "thesis_20 — Tablo 3 Bonferroni 10→12 karşılaştırma (α=0.001→0.00083); combined vs AS/naive equity eklendi",
        "WP5",
    )
    _add_summary_row(
        sum_tbl, "18",
        "thesis_21 — Sec 4.7 üç→beş varyant, Tablo 7 Std Sharpe kaldırıldı, oracle paradox p=0.115 güncellendi",
        "WP5",
    )

    doc.save(DST)
    print(f"decisions_log_5.docx saved ({DST.stat().st_size:,} bytes)")

    # ── Verify ──
    doc2 = Document(DST)
    sum2 = doc2.tables[-1]
    print(f"\n=== SUMMARY TABLE ({len(sum2.rows)} rows) ===")
    for ri, row in enumerate(sum2.rows):
        cells = [c.text.strip()[:70] for c in row.cells]
        print(f"  row[{ri:2d}]: {cells}")

    print()
    for p in doc2.paragraphs:
        if "Kritik" in p.text:
            print(f"Heading: '{p.text}'")

    # Verify #33, #34, #35, #36, #37, #38 exist
    for ti, tbl in enumerate(doc2.tables):
        r0 = tbl.rows[0].cells[1].text.strip()[:60] if len(tbl.rows[0].cells) > 1 else ""
        for tag in ("#33", "#34", "#35", "#36", "#37", "#38"):
            if tag in r0:
                print(f"\n{tag} table found at table[{ti}]")
                for ri, row in enumerate(tbl.rows):
                    c0 = row.cells[0].text.strip()[:15]
                    c1 = row.cells[1].text.strip()[:70]
                    print(f"  row[{ri}]: [{c0}] {c1}")


if __name__ == "__main__":
    main()
