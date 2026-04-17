"""decisions_log_7.docx -> decisions_log_8.docx
Appends decision #41 (thesis_25 TOST recompute) and adds a "Sürüm 25" title-page
marker. Summary: 20 karar -> 21 karar."""

from __future__ import annotations
import sys, io, shutil, copy
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC = Path("manuscript/decisions_log_7.docx")
DST = Path("manuscript/decisions_log_8.docx")

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _set_cell_text(cell, text, *, bold=None):
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].text = text
        for r in para.runs[1:]:
            r.text = ""
        if bold is not None:
            para.runs[0].bold = bold
    else:
        r = para.add_run(text)
        if bold is not None:
            r.bold = bold


def _build_karar_table_from_template(doc, template_tbl, number, title, wp, ikilem, secenekler, karar, neden, etki):
    new_xml = copy.deepcopy(template_tbl._tbl)
    from docx.table import Table
    new_tbl = Table(new_xml, doc.part)
    _set_cell_text(new_tbl.rows[0].cells[0], wp, bold=True)
    _set_cell_text(new_tbl.rows[0].cells[1], f"#{number}  {title}", bold=True)
    _set_cell_text(new_tbl.rows[1].cells[0], "İkilem", bold=True)
    _set_cell_text(new_tbl.rows[1].cells[1], ikilem)
    _set_cell_text(new_tbl.rows[2].cells[0], "Seçenekler", bold=True)
    _set_cell_text(new_tbl.rows[2].cells[1], secenekler)
    _set_cell_text(new_tbl.rows[3].cells[0], "Karar", bold=True)
    _set_cell_text(new_tbl.rows[3].cells[1], karar)
    _set_cell_text(new_tbl.rows[4].cells[0], "Neden", bold=True)
    _set_cell_text(new_tbl.rows[4].cells[1], neden)
    _set_cell_text(new_tbl.rows[5].cells[0], "Etki", bold=True)
    _set_cell_text(new_tbl.rows[5].cells[1], etki)
    return new_xml


def _replace_in_paragraph(para, old_text, new_text):
    full = "".join(r.text for r in para.runs)
    if old_text not in full:
        return False
    new_full = full.replace(old_text, new_text)
    if para.runs:
        para.runs[0].text = new_full
        for r in para.runs[1:]:
            r.text = ""
    return True


def main():
    shutil.copy2(SRC, DST)
    doc = Document(DST)

    template = doc.tables[38]  # last Karar table (#40) in decisions_log_7

    k41 = _build_karar_table_from_template(
        doc, template,
        number=41,
        title="thesis_25 — TOST değerleri birincil CSV recompute'a göre güncellendi",
        wp="WP5",
        ikilem=(
            "Codex denetimi, thesis_24'teki TOST p-değerleri ve CI sınırlarının birincil CSV "
            "recompute'a birebir uymadığını tespit etti (normal p=0.0005 vs gerçek 0.00067; "
            "misspec p=0.039 vs gerçek 0.042). Ayrıca TOST α=0.05 konvansiyonunun 90% CI'a "
            "karşılık geldiği tezde açık değildi."
        ),
        secenekler=(
            "A) Eski değerleri bırak  |  B) Birincil CSV'den recompute değerleri ile güncelle "
            "ve 90% CI konvansiyonunu netleştir"
        ),
        karar=(
            "B — Her iki bölümde p ve 95% CI recompute değerlerine güncellendi; ek olarak 90% CI "
            "(TOST α=0.05'e karşılık gelen aralık) eklendi. Normal: 90% CI [−0.001, +0.063] ±0.10 "
            "içinde. Misspec: 90% CI [−0.040, +0.048] ±0.05 içinde."
        ),
        neden=(
            "Reproducibility (independently verified, 6-decimal match). TOST metodolojik "
            "şeffaflığı artırıldı."
        ),
        etki=(
            "Null result argümanı değişmedi, aksine güçlendi — misspec'te 90% CI ±0.05'in *içinde*, "
            "yani eşdeğerlik sıkı sınırda destekleniyor."
        ),
    )

    body = doc.element.body
    ozet_heading = None
    for child in body.iterchildren():
        if child.tag == W_NS + "p":
            pStyle = child.find(W_NS + "pPr/" + W_NS + "pStyle")
            if pStyle is not None and pStyle.get(W_NS + "val") == "Heading1":
                text = "".join(t.text or "" for t in child.iter(W_NS + "t"))
                if "Özet" in text:
                    ozet_heading = child
                    break
    if ozet_heading is None:
        raise RuntimeError("Özet heading not found")

    def _blank_p():
        return OxmlElement("w:p")

    ozet_heading.addprevious(_blank_p())
    ozet_heading.addprevious(k41)

    # Summary heading: 20 -> 21
    for child in body.iterchildren():
        if child.tag == W_NS + "p":
            text = "".join(t.text or "" for t in child.iter(W_NS + "t"))
            if "En Kritik 20 Karar" in text:
                from docx.text.paragraph import Paragraph
                p_obj = Paragraph(child, doc.part)
                _replace_in_paragraph(p_obj, "En Kritik 20 Karar", "En Kritik 21 Karar")

    for para in doc.paragraphs:
        if "ChatGPT ve Claude" in para.text and "20 karar" in para.text:
            _replace_in_paragraph(para, "20 karar", "21 karar")

    # Title-page Sürüm 25 marker after KIT paragraph
    kit_para = None
    for child in body.iterchildren():
        if child.tag == W_NS + "p":
            text = "".join(t.text or "" for t in child.iter(W_NS + "t"))
            if "KIT Financial Engineering MSc Thesis" in text:
                kit_para = child
                break
    if kit_para is not None:
        new_marker = copy.deepcopy(kit_para)
        for r in list(new_marker.findall(W_NS + "r")):
            new_marker.remove(r)
        run_el = OxmlElement("w:r")
        t_el = OxmlElement("w:t")
        t_el.text = "Nisan 2026 — Sürüm 25"
        run_el.append(t_el)
        new_marker.append(run_el)
        kit_para.addnext(new_marker)

    # Summary row #21
    sum_tbl = doc.tables[-1]
    template_row = sum_tbl.rows[1]._tr
    new_row = copy.deepcopy(template_row)
    sum_tbl._tbl.append(new_row)
    from docx.table import _Row
    row_obj = _Row(new_row, sum_tbl)
    _set_cell_text(row_obj.cells[0], "21")
    _set_cell_text(
        row_obj.cells[1],
        "thesis_25 TOST recompute — p/CI değerleri birincil CSV'ye göre düzeltildi, 90% CI eklendi"
    )
    _set_cell_text(row_obj.cells[2], "WP5")

    doc.save(DST)
    print(f"decisions_log_8.docx saved ({DST.stat().st_size:,} bytes)")

    doc2 = Document(DST)
    print(f"\n=== Verification ===")
    print(f"Total tables: {len(doc2.tables)}")
    last_karar = doc2.tables[-2]
    print(f"Last Karar table header: {last_karar.rows[0].cells[1].text[:80]}")
    sum2 = doc2.tables[-1]
    print(f"Summary rows: {len(sum2.rows)} (should be 22 = header + 21)")
    print(f"  Last row: {[c.text[:50] for c in sum2.rows[-1].cells]}")
    for p in doc2.paragraphs[:15]:
        if "Sürüm" in p.text:
            print(f"Version marker: {p.text}")


if __name__ == "__main__":
    main()
