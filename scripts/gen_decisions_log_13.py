"""Generate decisions_log_13.docx from decisions_log_12.docx.

Appends Decisions #52-#55 for the post-hoc signal redundancy diagnostics.
"""

from __future__ import annotations

import copy
import io
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

SRC = Path("manuscript/decisions_log_12.docx")
DST = Path("manuscript/decisions_log_13.docx")
W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _set_cell_text(cell, text: str, *, bold=None):
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].text = text
        for run in para.runs[1:]:
            run.text = ""
        if bold is not None:
            para.runs[0].bold = bold
    else:
        run = para.add_run(text)
        if bold is not None:
            run.bold = bold


def _replace_in_paragraph(para, old_text: str, new_text: str) -> bool:
    full = "".join(run.text for run in para.runs)
    if old_text not in full:
        return False
    new_full = full.replace(old_text, new_text)
    if para.runs:
        para.runs[0].text = new_full
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(new_full)
    return True


def _find_ozet_heading(doc: Document):
    for child in doc.element.body.iterchildren():
        if child.tag != W_NS + "p":
            continue
        p_style = child.find(W_NS + "pPr/" + W_NS + "pStyle")
        if p_style is None or p_style.get(W_NS + "val") != "Heading1":
            continue
        text = "".join(t.text or "" for t in child.iter(W_NS + "t"))
        if "Özet" in text:
            return child
    raise RuntimeError("Özet heading not found")


def _blank_p():
    return OxmlElement("w:p")


def _build_karar_table_v2(doc: Document, template_tbl, number, title, wp, karar, secenekler, neden, etki, note):
    new_xml = copy.deepcopy(template_tbl._tbl)
    from docx.table import Table

    new_tbl = Table(new_xml, doc.part)
    _set_cell_text(new_tbl.rows[0].cells[0], wp, bold=True)
    _set_cell_text(new_tbl.rows[0].cells[1], f"#{number}  {title}", bold=True)
    _set_cell_text(new_tbl.rows[1].cells[0], "Karar", bold=True)
    _set_cell_text(new_tbl.rows[1].cells[1], karar)
    _set_cell_text(new_tbl.rows[2].cells[0], "Seçenekler", bold=True)
    _set_cell_text(new_tbl.rows[2].cells[1], secenekler)
    _set_cell_text(new_tbl.rows[3].cells[0], "Neden", bold=True)
    _set_cell_text(new_tbl.rows[3].cells[1], neden)
    _set_cell_text(new_tbl.rows[4].cells[0], "Etki", bold=True)
    _set_cell_text(new_tbl.rows[4].cells[1], etki)
    _set_cell_text(new_tbl.rows[5].cells[0], "Not", bold=True)
    _set_cell_text(new_tbl.rows[5].cells[1], note)
    return new_xml


DECISIONS = [
    {
        "number": 52,
        "title": "Post-hoc diagnostics added as supporting evidence",
        "wp": "Post-hoc",
        "karar": (
            "Post-hoc signal redundancy diagnostics ana tez sonuçlarına destekleyici "
            "evidence olarak eklendi; appendix-only bırakılmadı. Diagnostics, WP6 "
            "sonrasında signal redundancy yorumunu güçlendiren interpretive support "
            "olarak konumlandırıldı."
        ),
        "secenekler": (
            "(a) Diagnostics'i hiç dahil etmemek  |  "
            "(b) Appendix-only provenance notu olarak bırakmak  |  "
            "(c) Ana metodoloji ve sonuçlar içinde kısa, defense-safe supporting "
            "evidence olarak raporlamak"
        ),
        "neden": (
            "WP5/WP6 sonuçları explicit categorical regime label'ın sigma_hat sonrası "
            "robust incremental value sağlamadığı yönündeydi. Post-hoc classifier, "
            "incremental prediction ve action-explanation diagnostics bu yoruma "
            "mekanik destek sağlar; ancak yeni bir primary discovery veya yeni tez "
            "yönü değildir."
        ),
        "etki": (
            "thesis_29 içinde 'Post-hoc Signal Redundancy Diagnostics' metodoloji ve "
            "sonuç alt bölümleri eklendi. Ana claim güçlendirildi fakat bounded kaldı: "
            "test edilen sentetik HFMM ortamında limited incremental value."
        ),
        "note": (
            "Kaynak paket: docs/internal/posthoc_signal_analysis/. Ana okuma: "
            "final_signal_redundancy_assessment.md."
        ),
    },
    {
        "number": 53,
        "title": "No new PPO training / frozen artifact policy preserved",
        "wp": "Audit",
        "karar": (
            "Post-hoc diagnostics yalnızca mevcut donmuş artifact'ları okudu. PPO "
            "yeniden eğitilmedi, WP5/WP6 yeniden koşturulmadı ve protected evidence "
            "artifact'ları değiştirilmedi."
        ),
        "secenekler": (
            "(a) Yeni PPO koşuları veya full sweep yapmak  |  "
            "(b) Existing frozen WP2/WP5/WP6 artifact'ları üzerinden post-hoc tanı "
            "analizi yapmak  |  "
            "(c) thesis_28 numerik claims'i değiştirmek"
        ),
        "neden": (
            "Amaç yeni performans iddiası üretmek değil, mevcut signal redundancy "
            "yorumunu mekanik olarak desteklemekti. Yeni PPO training veya WP5/WP6 "
            "rerun hem stochastic drift hem de scope expansion riski yaratırdı."
        ),
        "etki": (
            "Evidence preservation policy korundu. Protected CSV hashes unchanged "
            "kaldı; canonical run directories ve protected PNG/result artifacts "
            "değiştirilmedi."
        ),
        "note": (
            "Diagnostics read-only policy: frozen CSV'ler okunabilir; yeni CSV/PNG/MD "
            "outputs sadece docs/internal/posthoc_signal_analysis/ altında üretilir."
        ),
    },
    {
        "number": 54,
        "title": "No mechanistic overclaim",
        "wp": "Interpretation",
        "karar": (
            "Post-hoc diagnostics causal PPO internal mechanism proof olarak "
            "yorumlanmayacak. Labels'ın zero information içerdiği veya PPO actions'ın "
            "tamamen sigma_hat tarafından açıklandığı iddia edilmeyecek."
        ),
        "secenekler": (
            "(a) Diagnostics'i mechanistic proof gibi sunmak  |  "
            "(b) Diagnostics'i limited incremental value evidence olarak sunmak  |  "
            "(c) Action diagnostics'teki mixed result'ı gizlemek"
        ),
        "neden": (
            "Classifier ve incremental-prediction sonuçları güçlü supporting evidence "
            "verirken action-level diagnostics mixed'tir: sigma-only modeller özellikle "
            "skew'i tam açıklamaz. Bu nedenle action regressions sadece explicit label'ın "
            "limited incremental explanatory gain sağladığını gösterir; causal internal "
            "PPO mechanism proof değildir."
        ),
        "etki": (
            "thesis_29 wording'i defense-safe kaldı: 'further support', 'consistent with', "
            "'limited incremental contribution' gibi ifadeler kullanıldı. 'proved', "
            "'zero information', 'categorical labels are useless' ve benzeri ifadelerden "
            "kaçınıldı."
        ),
        "note": (
            "Ana bounded claim: tested synthetic HFMM environment içinde explicit "
            "categorical labels, continuous volatility signal beyond robust incremental "
            "value sağlamıyor."
        ),
    },
    {
        "number": 55,
        "title": "No further architecture / representation probing",
        "wp": "Scope",
        "karar": (
            "Bu tez versiyonunda transformer, attention fusion, hidden-layer probing, "
            "latent mutual information, new RL architecture veya representation probing "
            "deneyleri yapılmayacak; bunlar future work olarak bırakılacak."
        ),
        "secenekler": (
            "(a) Yeni architecture/probing deneyleri açmak  |  "
            "(b) Post-hoc lightweight diagnostics ile kapanmak  |  "
            "(c) WP6 sonrasında yeni thesis direction başlatmak"
        ),
        "neden": (
            "Yeni architecture veya representation probing deneyleri thesis scope'unu "
            "genişletir, yeni research questions doğurur ve existing evidence freeze "
            "politikasını zayıflatır. Mevcut tez için ihtiyaç, yeni RL direction değil, "
            "signal redundancy interpretation'ın defense-safe desteklenmesidir."
        ),
        "etki": (
            "Tez kapsamı kontrol altında tutuldu. Future work bölümü daha gelişmiş "
            "representation/architecture analizlerine açık kapı bırakır, ancak thesis_29 "
            "mevcut WP5/WP6 + post-hoc diagnostics kanıt hattıyla kapanır."
        ),
        "note": (
            "Bu karar, post-hoc diagnostics'i final supporting layer olarak sabitler; "
            "ek deney açma eşiği yükseltilmiştir."
        ),
    },
]


SUMMARY_ROWS = [
    (
        32,
        "Post-hoc diagnostics added as supporting main-text evidence, not appendix-only",
        "Post-hoc",
    ),
    (
        33,
        "Frozen artifact policy preserved: no PPO retraining, no WP5/WP6 rerun, no protected evidence modification",
        "Audit",
    ),
    (
        34,
        "No mechanistic overclaim: action diagnostics mixed; claim limited incremental value only",
        "Interpretation",
    ),
    (
        35,
        "No further architecture/probing expansion; transformer/attention/latent MI tests reserved for future work",
        "Scope",
    ),
]


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    shutil.copy2(SRC, DST)
    doc = Document(DST)

    template = doc.tables[-2]
    karar_tables = [_build_karar_table_v2(doc, template, **decision) for decision in DECISIONS]
    ozet_heading = _find_ozet_heading(doc)
    for tbl_xml in karar_tables:
        ozet_heading.addprevious(_blank_p())
        ozet_heading.addprevious(tbl_xml)

    for para in doc.paragraphs:
        _replace_in_paragraph(para, "Nisan 2026 — Sürüm 28", "Mayıs 2026 — Sürüm 29")
        _replace_in_paragraph(para, "En Kritik 31 Karar", "En Kritik 35 Karar")
        _replace_in_paragraph(para, "31 karar", "35 karar")

    summary = doc.tables[-1]
    template_row = summary.rows[1]._tr

    def _add_summary_row(num, text, wp):
        new_row = copy.deepcopy(template_row)
        summary._tbl.append(new_row)
        from docx.table import _Row

        row_obj = _Row(new_row, summary)
        _set_cell_text(row_obj.cells[0], str(num))
        _set_cell_text(row_obj.cells[1], text)
        _set_cell_text(row_obj.cells[2], wp)

    for row in SUMMARY_ROWS:
        _add_summary_row(*row)

    doc.save(DST)

    check = Document(DST)
    print(f"decisions_log_13.docx saved ({DST.stat().st_size:,} bytes)")
    print(f"Total tables: {len(check.tables)}")
    print(f"Summary rows: {len(check.tables[-1].rows)} (expected 36 = header + 35)")
    print("New karar table headers:")
    for idx in range(-5, -1):
        print(f"  {check.tables[idx].rows[0].cells[1].text}")


if __name__ == "__main__":
    main()
