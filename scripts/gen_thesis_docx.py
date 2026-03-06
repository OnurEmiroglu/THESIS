"""Generate structured thesis DOCX from scratch."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)


def set_font(run, size=11, bold=False, italic=False, name="Times New Roman"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic


def add_heading_custom(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    if level == 1:
        set_font(run, size=14, bold=True)
    elif level == 2:
        set_font(run, size=12, bold=True)
    else:
        set_font(run, size=11, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_para(doc, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_table_with_data(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        set_font(run, size=10, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = str(val)
            run = cell.paragraphs[0].runs[0]
            set_font(run, size=10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_font(run)


# ── KAPAK ──────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Karlsruhe Institute of Technology")
set_font(run, size=14, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Financial Engineering MSc Program\u0131")
set_font(run, size=12)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Y\u00fcksek Lisans Tez Tasla\u011f\u0131")
set_font(run, size=12, italic=True)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "Farkl\u0131 Volatilite Rejimleri Alt\u0131nda Peki\u015ftirmeli \u00d6\u011frenme ile\n"
    "Y\u00fcksek Frekansl\u0131 Piyasa Yap\u0131c\u0131l\u0131\u011f\u0131"
)
set_font(run, size=16, bold=True)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Onur Emiro\u011flu")
set_font(run, size=12, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Financial Engineering MSc Program\u0131")
set_font(run, size=12)

doc.add_page_break()

# ── 1. \u00d6ZET ─────────────────────────────────────────────────────────────
add_heading_custom(doc, "1. \u00d6ZET", level=1)

ozet = (
    "Bu \u00e7al\u0131\u015fmada, volatilite rejim bilgisinin (D\u00fc\u015f\u00fck/Orta/Y\u00fcksek) peki\u015ftirmeli \u00f6\u011frenme tabanl\u0131 "
    "piyasa yap\u0131c\u0131 ajanlar\u0131n performans\u0131na katk\u0131s\u0131 ara\u015ft\u0131r\u0131lmaktad\u0131r. Proximal Policy Optimization "
    "(PPO) algoritmas\u0131 kullan\u0131larak iki ajan e\u011fitilmi\u015ftir: rejim bilgisine sahip (ppo_aware) ve rejim "
    "bilgisinden yoksun (ppo_blind). Ajanlar, Poisson var\u0131\u015f s\u00fcre\u00e7li yo\u011funluk tabanl\u0131 dolum modeli "
    "ve Markov zinciri rejim ge\u00e7i\u015fleri i\u00e7eren tamamen sentetik bir piyasa ortam\u0131nda de\u011ferlendirilmi\u015ftir. "
    "Y\u00fcr\u00fct\u00fclen out-of-sample deney (20 ba\u011f\u0131ms\u0131z tohum, 1 milyon zaman ad\u0131m\u0131 e\u011fitim, %70/%30 "
    "walk-forward b\u00f6l\u00fcnmesi), her iki PPO varyant\u0131n\u0131n da Sharpe oran\u0131 a\u00e7\u0131s\u0131ndan klasik referans "
    "stratejilerini (Avellaneda-Stoikov ve naif sabit-spread) yakla\u015f\u0131k 6\u20137 kat ge\u00e7ti\u011fini ortaya "
    "koymaktad\u0131r. Buna kar\u015f\u0131n, rejim bilgisinin PPO performans\u0131na net bir katk\u0131 sa\u011flad\u0131\u011f\u0131 hipotezi "
    "istatistiksel olarak desteklenmemi\u015ftir: ppo_aware, yaln\u0131zca 20 tohumun 9\u2019unda ppo_blind\u2019\u0131 "
    "geride b\u0131rakm\u0131\u015ft\u0131r. Davran\u0131\u015fsal analiz, ppo_aware ajan\u0131n\u0131n rejime \u00f6zg\u00fc farkl\u0131 eylem da\u011f\u0131l\u0131mlar\u0131 "
    "geli\u015ftirdi\u011fini (\u00f6zellikle d\u00fc\u015f\u00fck volatilite rejiminde daha geni\u015f spread ve farkl\u0131 skew \u00f6r\u00fcnt\u00fcleri) "
    "g\u00f6stermektedir; ancak bu davran\u0131\u015fsal farkl\u0131la\u015fma performans avantaj\u0131na d\u00f6n\u00fc\u015fmemektedir. "
    "\u00c7al\u0131\u015fma, sentetik veri ile kontroll\u00fc ortam tasar\u0131m\u0131n\u0131n de\u011ferini, PPO\u2019nun envanter risk "
    "y\u00f6netimindeki \u00fcst\u00fcnl\u00fc\u011f\u00fcn\u00fc ve rejim fark\u0131ndal\u0131\u011f\u0131 hipotezi i\u00e7in \u201cnull result\u201d\u0131n yay\u0131nlanabilir "
    "bir akademik bulgu olu\u015fturdu\u011funu g\u00f6stermektedir."
)
add_para(doc, ozet)

add_para(
    doc,
    "Anahtar kelimeler: piyasa yap\u0131c\u0131l\u0131\u011f\u0131, peki\u015ftirmeli \u00f6\u011frenme, PPO, volatilite rejimleri, "
    "Avellaneda-Stoikov, Poisson dolum modeli, out-of-sample de\u011ferlendirme",
    italic=True,
)

doc.add_page_break()

# ── SEMBOLLER VE KISALTMALAR ─────────────────────────────────────────────
add_heading_custom(doc, "SEMBOLLER VE KISALTMALAR", level=1)
add_para(
    doc,
    "Bu b\u00f6l\u00fcmde \u00e7al\u0131\u015fma boyunca kullan\u0131lan matematiksel semboller, Yunan harfleri ve k\u0131saltmalar tan\u0131mlanmaktad\u0131r.",
)

add_heading_custom(doc, "Matematiksel Semboller", level=2)

semboller = [
    ("M_t", "t an\u0131ndaki mid-price (orta fiyat) [para birimi]"),
    ("\u03c3", "Volatilite parametresi \u2014 mid-price\u2019\u0131n birim zamandaki standart sapmas\u0131 [tick/\u221asaniye]"),
    ("\u03c3_base", "Temel volatilite parametresi; Orta (M) rejim i\u00e7in referans de\u011fer (0.8 tick/\u221asaniye)"),
    ("\u03c3_L, \u03c3_M, \u03c3_H", "D\u00fc\u015f\u00fck, Orta ve Y\u00fcksek rejimlere kar\u015f\u0131l\u0131k gelen volatilite de\u011ferleri"),
    ("\u0394t", "Sim\u00fclasyon zaman ad\u0131m\u0131 b\u00fcy\u00fckl\u00fc\u011f\u00fc (0.2 saniye)"),
    ("z", "Standart normal rassal de\u011fi\u015fken; z ~ N(0,1)"),
    ("\u03bb(\u03b4)", "\u03b4 tick uzakl\u0131\u011f\u0131ndaki kotasyon i\u00e7in Poisson dolum yo\u011funlu\u011fu [dolum/saniye]"),
    ("A", "S\u0131f\u0131r uzakl\u0131kta (\u03b4=0) temel dolum yo\u011funlu\u011fu (5.0 dolum/saniye)"),
    ("k", "Dolum yo\u011funlu\u011funun uzakl\u0131kla \u00fcssel azalma h\u0131z\u0131 (1.5 tick ba\u015f\u0131na)"),
    ("\u03b4", "Kotasyon uzakl\u0131\u011f\u0131 \u2014 bid veya ask fiyat\u0131n\u0131n mid-price\u2019tan tick cinsinden sapmas\u0131"),
    ("\u03b4_bid", "Bid (al\u0131\u015f) kotasyonunun mid-price\u2019tan uzakl\u0131\u011f\u0131 [tick]"),
    ("\u03b4_ask", "Ask (sat\u0131\u015f) kotasyonunun mid-price\u2019tan uzakl\u0131\u011f\u0131 [tick]"),
    ("P_fill", "Bir zaman ad\u0131m\u0131nda en az bir dolumun ger\u00e7ekle\u015fme olas\u0131l\u0131\u011f\u0131"),
    ("h", "Half-spread \u2014 kotasyon yar\u0131-aral\u0131\u011f\u0131 [tick]; h \u2208 {1, 2, 3, 4, 5}"),
    ("m", "Skew \u2014 asimetrik kotasyon kaymas\u0131 [tick]; m \u2208 {\u22122, \u22121, 0, +1, +2}"),
    ("q", "Envanter (inventory) \u2014 anda tutulan net pozisyon miktar\u0131 [lot]"),
    ("q_norm", "Normalize edilmi\u015f envanter; q_norm = q / inv_max_clip \u2208 [\u22121, +1]"),
    ("inv_max_clip", "Envanter k\u0131rpma s\u0131n\u0131r\u0131 (50 lot)"),
    ("\u03c4", "Kalan zaman fraksiyonu; \u03c4 = (T \u2212 t) / T \u2208 [0, 1]"),
    ("T", "Episode toplam ad\u0131m say\u0131s\u0131 (8000 ad\u0131m)"),
    ("R_t", "t ad\u0131m\u0131ndaki \u00f6d\u00fcl (reward)"),
    ("\u03b7", "Envanter ceza katsay\u0131s\u0131; \u03b7 = 0.001"),
    ("\u03b3", "Avellaneda-Stoikov risk aversion parametresi; \u03b3 = 0.01"),
    ("r", "Avellaneda-Stoikov rezervasyon fiyat\u0131"),
    ("d", "Avellaneda-Stoikov optimal yar\u0131-spread"),
    ("k_price", "Fiyat cinsinden ifade edilen yo\u011funluk azalma parametresi; k_price = k / tick_size"),
    ("P", "Markov zinciri ge\u00e7i\u015f matrisi (3\u00d73)"),
    ("p_ii", "i. rejimin kendi kendine ge\u00e7i\u015f olas\u0131l\u0131\u011f\u0131 (k\u00f6\u015fegen eleman)"),
    ("Sharpe", "Sharpe oran\u0131 \u2014 risk-d\u00fczeltmeli getiri \u00f6l\u00e7\u00fct\u00fc; Sharpe = (\u03bc/\u03c3) \u00d7 \u221a(1/\u0394t)"),
    ("inv_p99", "Mutlak envanter de\u011ferinin 99. persantili [lot]"),
    ("equity", "Toplam varl\u0131k de\u011feri; equity = cash + q \u00d7 M_t"),
]
add_table_with_data(doc, ["Sembol", "A\u00e7\u0131klama"], semboller, col_widths=[1.2, 4.8])

add_heading_custom(doc, "Yunan Harfleri", level=2)
yunan = [
    ("\u03b3 (gamma)", "Risk aversion parametresi (Avellaneda-Stoikov modelinde)"),
    ("\u03b7 (eta)", "Envanter ceza katsay\u0131s\u0131 (\u00f6d\u00fcl fonksiyonunda)"),
    ("\u03c3 (sigma)", "Volatilite \u2014 fiyat de\u011fi\u015fiminin standart sapmas\u0131"),
    ("\u03c4 (tau)", "Kalan zaman fraksiyonu"),
    ("\u03bb (lambda)", "Poisson dolum yo\u011funlu\u011fu"),
    ("\u03b4 (delta)", "Kotasyon uzakl\u0131\u011f\u0131 (mid-price\u2019tan tick cinsinden sapma)"),
    ("\u03bc (mu)", "Ortalama ad\u0131m ba\u015f\u0131na PnL (Sharpe hesab\u0131nda)"),
]
add_table_with_data(doc, ["Sembol", "A\u00e7\u0131klama"], yunan, col_widths=[1.2, 4.8])

add_heading_custom(doc, "K\u0131saltmalar", level=2)
kisaltmalar = [
    ("PPO", "Proximal Policy Optimization \u2014 yak\u0131nsak politika optimizasyonu algoritmas\u0131"),
    ("RL", "Reinforcement Learning \u2014 peki\u015ftirmeli \u00f6\u011frenme"),
    ("AS", "Avellaneda-Stoikov \u2014 klasik stokastik kontrol tabanl\u0131 piyasa yap\u0131c\u0131l\u0131\u011f\u0131 modeli"),
    ("HFT", "High-Frequency Trading \u2014 y\u00fcksek frekansl\u0131 al\u0131m-sat\u0131m"),
    ("OOS", "Out-of-Sample \u2014 modelin e\u011fitilmedi\u011fi veri \u00fczerinde de\u011ferlendirme"),
    ("WP", "Work Package \u2014 i\u015f paketi (proje a\u015famas\u0131)"),
    ("ABM", "Arithmetic Brownian Motion \u2014 aritmetik Brownian hareketi"),
    ("RV", "Realized Volatility \u2014 ger\u00e7ekle\u015fmi\u015f volatilite (kayan pencere ile hesaplanan)"),
    ("PnL", "Profit and Loss \u2014 kar ve zarar"),
    ("CUDA", "Compute Unified Device Architecture \u2014 NVIDIA GPU hesaplama mimarisi"),
    ("SB3", "Stable-Baselines3 \u2014 PPO implementasyonu i\u00e7in kullan\u0131lan Python k\u00fct\u00fcphanesi"),
    ("MlpPolicy", "Multi-Layer Perceptron Policy \u2014 \u00e7ok katmanl\u0131 alg\u0131lay\u0131c\u0131 politika a\u011f\u0131"),
    ("GAE", "Generalized Advantage Estimation \u2014 genelle\u015ftirilmi\u015f avantaj tahmini (PPO bile\u015feni)"),
    ("L / M / H", "Low / Medium / High \u2014 D\u00fc\u015f\u00fck / Orta / Y\u00fcksek volatilite rejimleri"),
]
add_table_with_data(doc, ["K\u0131saltma", "A\u00e7\u0131klama"], kisaltmalar, col_widths=[1.2, 4.8])

doc.add_page_break()

# ── TER\u0130MLER S\u00d6ZL\u00dc\u011e\u00dc ─────────────────────────────────────────────────────
add_heading_custom(doc, "TER\u0130MLER S\u00d6ZL\u00dc\u011e\u00dc", level=1)
add_para(
    doc,
    "Bu s\u00f6zl\u00fck, \u00e7al\u0131\u015fmada kullan\u0131lan teknik terimleri konuya a\u015fina olmayan okuyucular i\u00e7in tan\u0131mlamaktad\u0131r. "
    "Terimler alfabetik s\u0131raya g\u00f6re d\u00fczenlenmi\u015ftir.",
)

terimler = [
    ("Ablation (Ablasyon)", "Bir modelin belirli bir bile\u015feninin \u00e7\u0131kar\u0131larak ya da devre d\u0131\u015f\u0131 b\u0131rak\u0131larak etkisinin \u00f6l\u00e7\u00fclmesi amac\u0131yla yap\u0131lan kontroll\u00fc deney. Bu \u00e7al\u0131\u015fmada, rejim bilgisinin (use_regime=False) ve envanter ceza katsay\u0131s\u0131n\u0131n (\u03b7) etkisi ablasyon deneyleriyle \u00f6l\u00e7\u00fclm\u00fc\u015ft\u00fcr."),
    ("Action Space (Eylem Uzay\u0131)", "Bir peki\u015ftirmeli \u00f6\u011frenme ajan\u0131n\u0131n her ad\u0131mda se\u00e7ebilece\u011fi t\u00fcm olas\u0131 eylemlerin k\u00fcmesi. Bu \u00e7al\u0131\u015fmada eylem uzay\u0131, half-spread (h) ve skew (m) de\u011ferlerinden olu\u015fan ayr\u0131k bir yap\u0131dad\u0131r."),
    ("Adverse Selection (Ters Se\u00e7im)", "Piyasa yap\u0131c\u0131n\u0131n bilgili yat\u0131r\u0131mc\u0131lara kar\u015f\u0131 dezavantajl\u0131 konuma d\u00fc\u015fmesi durumu."),
    ("Agent (Ajan)", "Peki\u015ftirmeli \u00f6\u011frenmede, bir ortamda g\u00f6zlem yaparak karar veren ve \u00f6d\u00fcl sinyaline g\u00f6re davran\u0131\u015f\u0131n\u0131 g\u00fcncelleyen \u00f6\u011frenen sistem."),
    ("Ask (Sat\u0131\u015f Fiyat\u0131)", "Piyasa yap\u0131c\u0131n\u0131n satmaya haz\u0131r oldu\u011fu fiyat. Her zaman bid fiyat\u0131n\u0131n \u00fczerindedir."),
    ("Bid (Al\u0131\u015f Fiyat\u0131)", "Piyasa yap\u0131c\u0131n\u0131n almaya haz\u0131r oldu\u011fu fiyat. Her zaman ask fiyat\u0131n\u0131n alt\u0131ndad\u0131r."),
    ("Bid-Ask Spread", "Al\u0131\u015f ve sat\u0131\u015f fiyat\u0131 aras\u0131ndaki fark. Piyasa yap\u0131c\u0131n\u0131n temel gelir kayna\u011f\u0131d\u0131r."),
    ("Brownian Motion (Brownian Hareketi)", "Rastlant\u0131sal y\u00fcr\u00fcy\u00fc\u015f modellerinin s\u00fcrekli-zaman limiti. ABM\u2019de fiyat de\u011fi\u015fimleri normal da\u011f\u0131l\u0131mdan \u00e7ekilmektedir."),
    ("Clip (K\u0131rpma)", "Bir de\u011ferin belirli bir aral\u0131kla s\u0131n\u0131rland\u0131r\u0131lmas\u0131 i\u015flemi. Bu \u00e7al\u0131\u015fmada envanter \u00b150 lot ile k\u0131rp\u0131lmaktad\u0131r."),
    ("Convergence (Yak\u0131nsama)", "Peki\u015ftirmeli \u00f6\u011frenme e\u011fitiminde politikan\u0131n optimal ya da kararl\u0131 bir \u00e7\u00f6z\u00fcme ula\u015fmas\u0131 s\u00fcreci."),
    ("Entropy Coefficient (Entropi Katsay\u0131s\u0131)", "PPO\u2019da ke\u015fif-s\u00f6m\u00fcr\u00fc dengesini ayarlayan hiperparametre. ent_coef=0.01 kullan\u0131lm\u0131\u015ft\u0131r."),
    ("Episode", "Peki\u015ftirmeli \u00f6\u011frenmede ba\u015flang\u0131\u00e7tan biti\u015f ko\u015fuluna kadar s\u00fcren bir tam sim\u00fclasyon ko\u015fusu. Bu \u00e7al\u0131\u015fmada 8000 zaman ad\u0131m\u0131."),
    ("Exogenous Series (D\u0131\u015fsal Seri)", "Modelin d\u0131\u015f\u0131nda \u00f6nceden \u00fcretilmi\u015f ve t\u00fcm stratejilere ayn\u0131 bi\u00e7imde sunulan fiyat/volatilite serisi."),
    ("Fill (Dolum)", "Piyasa yap\u0131c\u0131n\u0131n kotasyon verdi\u011fi fiyattan ger\u00e7ekle\u015fen emir e\u015fle\u015fmesi."),
    ("Fill Rate (Dolum Oran\u0131)", "Toplam ad\u0131m say\u0131s\u0131na b\u00f6l\u00fcnen dolum say\u0131s\u0131."),
    ("Gymnasium", "OpenAI taraf\u0131ndan geli\u015ftirilen, peki\u015ftirmeli \u00f6\u011frenme ortamlar\u0131 i\u00e7in standart Python aray\u00fcz\u00fc."),
    ("Half-Spread (Yar\u0131-Aral\u0131k)", "Bid veya ask fiyat\u0131n\u0131n mid-price\u2019tan uzakl\u0131\u011f\u0131. Tam spread\u2019in yar\u0131s\u0131."),
    ("Hyperparameter (Hiperparametre)", "E\u011fitim s\u00fcrecinden \u00f6nce kullan\u0131c\u0131 taraf\u0131ndan belirlenen parametre."),
    ("Inventory (Envanter)", "Piyasa yap\u0131c\u0131n\u0131n anda elinde bulundurdu\u011fu net pozisyon miktar\u0131."),
    ("Inventory Risk (Envanter Riski)", "B\u00fcy\u00fck envanter pozisyonu tutarken fiyat hareketleri nedeniyle maruz kal\u0131nan zarar riski."),
    ("Latency (Gecikme)", "Kotasyon g\u00fcncellemesi ile piyasan\u0131n i\u015flemesi aras\u0131ndaki zaman fark\u0131. latency_steps=1."),
    ("Limit Order Book", "Beklemedeki t\u00fcm al\u0131\u015f ve sat\u0131\u015f emirlerinin fiyat-miktar listesi."),
    ("Lot", "Al\u0131m-sat\u0131m i\u015flemlerinde standart birim miktar."),
    ("Markov Chain (Markov Zinciri)", "Gelecekteki durumun yaln\u0131zca mevcut duruma ba\u011fl\u0131 oldu\u011fu olas\u0131l\u0131ksal s\u00fcre\u00e7."),
    ("Market Making (Piyasa Yap\u0131c\u0131l\u0131\u011f\u0131)", "Finansal piyasalarda s\u00fcrekli al\u0131\u015f ve sat\u0131\u015f kotasyonu vererek likidite sa\u011flama faaliyeti."),
    ("Mid-Price", "Anl\u0131k al\u0131\u015f ve sat\u0131\u015f fiyatlar\u0131n\u0131n ortalamas\u0131."),
    ("Model Misspecification", "Kullan\u0131lan modelin ger\u00e7ek veri \u00fcretim s\u00fcrecini tam olarak yans\u0131tmamas\u0131 durumu."),
    ("Null Result", "Ara\u015ft\u0131rma hipotezinin verilerle desteklenmedi\u011fi bulgu."),
    ("Observation Space (G\u00f6zlem Uzay\u0131)", "Ajan\u0131n her ad\u0131mda \u00e7evreden ald\u0131\u011f\u0131 bilginin vekt\u00f6r temsili. 6 boyutlu: [q_norm, \u03c3\u0302, \u03c4, r_L, r_M, r_H]."),
    ("One-Hot Encoding", "Kategorik bir de\u011fi\u015fkeni ikili vekt\u00f6rle temsil etme y\u00f6ntemi."),
    ("Out-of-Sample (OOS)", "Modelin e\u011fitilmedi\u011fi veri k\u00fcmesi \u00fczerinde yap\u0131lan de\u011ferlendirme."),
    ("Overfitting", "Modelin e\u011fitim verisine a\u015f\u0131r\u0131 uyum sa\u011flayarak yeni veride d\u00fc\u015f\u00fck performans g\u00f6stermesi."),
    ("Percentile (Persantil)", "Veri setinin belirli bir y\u00fczdesinin alt\u0131nda kald\u0131\u011f\u0131 de\u011fer."),
    ("Policy (Politika)", "Peki\u015ftirmeli \u00f6\u011frenmede ajan\u0131n g\u00f6zleme g\u00f6re eylem se\u00e7me stratejisi."),
    ("Poisson Process (Poisson S\u00fcreci)", "Olaylar\u0131n ba\u011f\u0131ms\u0131z ve sabit ortalama h\u0131zda rastlant\u0131sal ger\u00e7ekle\u015fti\u011fi s\u00fcre\u00e7."),
    ("PPO", "John Schulman ve ark. (2017) taraf\u0131ndan geli\u015ftirilen politika gradyan\u0131 tabanl\u0131 RL algoritmas\u0131."),
    ("Regime (Rejim)", "Piyasan\u0131n belirli bir volatilite seviyesinde bulundu\u011fu d\u00f6nem. L/M/H."),
    ("Reservation Price (Rezervasyon Fiyat\u0131)", "AS modelinde envanter riskini i\u00e7selle\u015ftiren d\u00fczeltilmi\u015f mid-price tahmini."),
    ("Reward (\u00d6d\u00fcl)", "RL\u2019de ajan\u0131n her ad\u0131mda \u00e7evreden ald\u0131\u011f\u0131 say\u0131sal sinyal. R_t = \u0394Equity \u2212 \u03b7\u00b7q\u00b2."),
    ("Risk-Adjusted Return", "Getirinin, \u00fcstlenilen risk miktar\u0131na b\u00f6l\u00fcnmesiyle elde edilen performans \u00f6l\u00e7\u00fct\u00fc."),
    ("Seed (Tohum)", "Rastlant\u0131sal say\u0131 \u00fcretecinin ba\u015flang\u0131\u00e7 de\u011feri. Tekrarlanabilirlik i\u00e7in kullan\u0131l\u0131r."),
    ("Sharpe Ratio (Sharpe Oran\u0131)", "Birim risk ba\u015f\u0131na elde edilen fazla getiriyi \u00f6l\u00e7en performans metrigi."),
    ("Skew (Asimetrik Kotasyon)", "Bid ve ask kotasyonlar\u0131n\u0131n mid-price\u2019a g\u00f6re e\u015fit olmayan bi\u00e7imde kayd\u0131r\u0131lmas\u0131."),
    ("Stable-Baselines3 (SB3)", "PPO dahil \u00e7e\u015fitli RL algoritmalar\u0131n\u0131n g\u00fcvenilir PyTorch implementasyonlar\u0131."),
    ("State (Durum)", "Peki\u015ftirmeli \u00f6\u011frenmede \u00e7evrenin belirli bir andaki tam g\u00f6zlemlenebilir temsili."),
    ("Sticky Transition Matrix", "K\u00f6\u015fegen elemanlar\u0131 y\u00fcksek Markov ge\u00e7i\u015f matrisi. Rejimlerin uzun s\u00fcre devam etmesini sa\u011flar."),
    ("Stochastic Control", "Rassal s\u00fcre\u00e7ler i\u00e7eren sistemlerde optimal karar verme teorisi."),
    ("Tick", "Fiyat\u0131n alabilece\u011fi en k\u00fc\u00e7\u00fck art\u0131\u015f birimi. tick_size = 0.01."),
    ("Timestep (Zaman Ad\u0131m\u0131)", "Sim\u00fclasyonun bir sonraki duruma ge\u00e7ti\u011fi en k\u00fc\u00e7\u00fck zaman birimi. \u0394t = 0.2 saniye."),
    ("Training (E\u011fitim)", "RL\u2019de ajan\u0131n politikas\u0131n\u0131 \u00f6d\u00fcl sinyallerine g\u00f6re iteratif bi\u00e7imde g\u00fcncellemesi s\u00fcreci."),
    ("Variance (Varyans)", "Bir de\u011fi\u015fkenin ortalama etraf\u0131ndaki yay\u0131l\u0131m\u0131n\u0131n karesi."),
    ("Walk-Forward Validation", "Zaman serisi verisini ge\u00e7mi\u015ften gelece\u011fe do\u011fru b\u00f6len de\u011ferlendirme y\u00f6ntemi."),
]
add_table_with_data(doc, ["Terim", "Tan\u0131m"], terimler, col_widths=[1.8, 4.2])

doc.add_page_break()

# ── 2. G\u0130R\u0130\u015e ─────────────────────────────────────────────────────────────
add_heading_custom(doc, "2. G\u0130R\u0130\u015e", level=1)

giris_paragraflar = [
    "Piyasa yap\u0131c\u0131l\u0131\u011f\u0131 (market making), finansal piyasalarda likidite sa\u011flayan ve s\u00fcrekli olarak hem al\u0131\u015f (bid) hem de sat\u0131\u015f (ask) kotasyonu veren ajanlar\u0131n stratejik davran\u0131\u015f\u0131n\u0131 inceleyen bir aland\u0131r. Piyasa yap\u0131c\u0131lar, spread geliri elde etmeye \u00e7al\u0131\u015f\u0131rken olumsuz fiyat hareketlerine maruz kalan envanter riskini de y\u00f6netmek zorundad\u0131r. Bu denge, \u00f6zellikle volatilitenin de\u011fi\u015fken oldu\u011fu y\u00fcksek frekansl\u0131 piyasalarda kritik \u00f6nem ta\u015f\u0131maktad\u0131r.",
    "Klasik yakla\u015f\u0131mlar aras\u0131nda Avellaneda ve Stoikov (2008) taraf\u0131ndan geli\u015ftirilen stokastik kontrol \u00e7er\u00e7evesi \u00f6ne \u00e7\u0131kmaktad\u0131r. Bu model, envanter riskini rezervasyon fiyat\u0131 kavram\u0131yla i\u00e7selle\u015ftirerek optimal bid-ask spread\u2019ini analitik olarak t\u00fcretmektedir. Ancak bu yakla\u015f\u0131m, sabit bir volatilite varsay\u0131m\u0131na dayanmakta ve piyasa rejimlerindeki de\u011fi\u015fimlere uyum sa\u011flayamamaktad\u0131r.",
    "Son y\u0131llarda peki\u015ftirmeli \u00f6\u011frenme (RL) y\u00f6ntemleri, dinamik piyasa ko\u015fullar\u0131na uyum sa\u011flayabilen ajan tasar\u0131m\u0131 i\u00e7in umut verici bir alternatif olarak \u00f6ne \u00e7\u0131km\u0131\u015ft\u0131r. PPO gibi politika gradyan\u0131 algoritmalar\u0131, \u00f6zellikle s\u00fcrekli durum uzaylar\u0131 ve karma\u015f\u0131k \u00f6d\u00fcl yap\u0131lar\u0131na sahip finansal ortamlarda ba\u015far\u0131yla uygulanm\u0131\u015ft\u0131r.",
    "Bu \u00e7al\u0131\u015fman\u0131n temel ara\u015ft\u0131rma sorusu \u015fudur: Volatilite rejim bilgisine (D\u00fc\u015f\u00fck/Orta/Y\u00fcksek) eri\u015fimi olan bir PPO ajan\u0131, bu bilgiden yoksun olan e\u015fde\u011ferine k\u0131yasla daha iyi performans sergiler mi?",
    "Bu soruyu yan\u0131tlamak i\u00e7in a\u015fa\u011f\u0131daki katk\u0131lar sunulmaktad\u0131r:",
]
for p_text in giris_paragraflar:
    add_para(doc, p_text)

katkilar = [
    "(1) Markov zinciri rejim ge\u00e7i\u015fleri ve Poisson tabanl\u0131 dolum modeli i\u00e7eren kontroll\u00fc bir sentetik piyasa ortam\u0131 tasarlanm\u0131\u015ft\u0131r.",
    "(2) Rejim fark\u0131ndal\u0131\u011f\u0131 ablasyonu ger\u00e7ekle\u015ftirilmi\u015f; ppo_aware ve ppo_blind ajanlar\u0131 20 ba\u011f\u0131ms\u0131z tohum ve walk-forward out-of-sample protokol\u00fc ile kar\u015f\u0131la\u015ft\u0131r\u0131lm\u0131\u015ft\u0131r.",
    "(3) Eylem da\u011f\u0131l\u0131m\u0131 analizi ile ajanlar\u0131n rejime \u00f6zg\u00fc davran\u0131\u015fsal farkl\u0131la\u015fmas\u0131 incelenmi\u015ftir.",
]
for k in katkilar:
    add_bullet(doc, k)

doc.add_page_break()

# ── 3. TEOR\u0130 VE METODOLOJ\u0130 ───────────────────────────────────────────────
add_heading_custom(doc, "3. TEOR\u0130 VE METODOLOJ\u0130", level=1)

add_heading_custom(doc, "3.1 Piyasa Modeli", level=2)
add_para(
    doc,
    "Piyasa fiyat\u0131, ayr\u0131k zamanl\u0131 aritmetik Brownian hareketi (ABM) ile modellenmektedir:",
)
add_para(doc, "M_{t+1} = M_t + \u03c3_r \u00b7 \u221a(\u0394t) \u00b7 z,    z ~ N(0,1)", align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(
    doc,
    "Burada M_t mid-price, \u03c3_r rejime ba\u011fl\u0131 volatilite parametresi (ticks/\u221asaniye cinsinden), "
    "\u0394t = 0.2 saniye zaman ad\u0131m\u0131 b\u00fcy\u00fckl\u00fc\u011f\u00fcd\u00fcr. Ba\u015flang\u0131\u00e7 fiyat\u0131 M_0 = 100.0, tick b\u00fcy\u00fckl\u00fc\u011f\u00fc 0.01\u2019dir.",
)
add_para(
    doc,
    "Temel volatilite parametresi \u03c3_base = 0.8 tick olarak belirlenmi\u015f; \u00fc\u00e7 rejim i\u00e7in \u00e7arpanlar "
    "[0.6, 1.0, 1.8] olarak tan\u0131mlanm\u0131\u015ft\u0131r. Buna g\u00f6re \u03c3_L = 0.48, \u03c3_M = 0.80, \u03c3_H = 1.44 "
    "tick/\u221asaniye de\u011ferleri elde edilmektedir.",
)

add_heading_custom(doc, "3.2 Dolum Modeli", level=2)
add_para(
    doc,
    "Emir dolumlar\u0131, yo\u011funluk tabanl\u0131 Poisson s\u00fcreci ile modellenmektedir. Delta tick cinsinden "
    "ifade edildi\u011finde, dolum yo\u011funlu\u011fu:",
)
add_para(doc, "\u03bb(\u03b4) = A \u00b7 exp(\u2212k \u00b7 \u03b4)", align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(
    doc,
    "form\u00fcl\u00fcyle hesaplanmaktad\u0131r. Burada A = 5.0 (delta=0\u2019daki temel yo\u011funluk, dolum/saniye) ve "
    "k = 1.5 (tick ba\u015f\u0131na \u00fcssel azalma) parametreleridir. Bir zaman ad\u0131m\u0131nda en az bir dolumun "
    "ger\u00e7ekle\u015fme olas\u0131l\u0131\u011f\u0131:",
)
add_para(doc, "P_fill = 1 \u2212 exp(\u2212\u03bb \u00b7 \u0394t)", align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(
    doc,
    "olarak verilmektedir. Hem bid hem ask taraf\u0131 i\u00e7in ba\u011f\u0131ms\u0131z Bernoulli denemeleri "
    "ger\u00e7ekle\u015ftirilmektedir.",
)
add_para(
    doc,
    "Gecikme modeli (latency_steps = 1), ajan\u0131n kotasyon hesaplamas\u0131nda bir \u00f6nceki ad\u0131m\u0131n "
    "mid-price de\u011ferini kullanmas\u0131na yol a\u00e7makta; bu durum ters se\u00e7im (adverse selection) riskini "
    "i\u00e7selle\u015ftirmektedir. Komisyon \u00fccreti her i\u015flem i\u00e7in 0.2 baz puan olarak uygulanmaktad\u0131r.",
)

add_heading_custom(doc, "3.3 Rejim Modeli", level=2)
add_para(
    doc,
    "Volatilite rejimleri, \u00fc\u00e7 durumlu (L: D\u00fc\u015f\u00fck, M: Orta, H: Y\u00fcksek) birinci derece Markov zinciri "
    "ile \u00fcretilmektedir. Ge\u00e7i\u015f matrisi yap\u0131\u015fkan (sticky) olacak \u015fekilde tasarlanm\u0131\u015ft\u0131r:",
)
add_para(
    doc,
    "P = [[0.9967, 0.0023, 0.0010], [0.0042, 0.9917, 0.0041], [0.0010, 0.0030, 0.9960]]",
    align=WD_ALIGN_PARAGRAPH.CENTER,
)
add_para(
    doc,
    "Beklenen rejim s\u00fcreleri s\u0131ras\u0131yla L: ~300, M: ~120, H: ~250 zaman ad\u0131m\u0131d\u0131r. Bu yap\u0131\u015fkanl\u0131k, "
    "kayan ger\u00e7ekle\u015fmi\u015f volatilite (rolling realized volatility \u2014 RV) ile rejim tespitinin m\u00fcmk\u00fcn "
    "olmas\u0131n\u0131 sa\u011flamaktad\u0131r.",
)
add_para(
    doc,
    "Rejim tespiti i\u00e7in RV pencere boyutu 50 ad\u0131m, \u0131s\u0131nma s\u00fcresi (warmup) 1000 ad\u0131m olarak "
    "belirlenmi\u015ftir. E\u015fik de\u011ferleri \u0131s\u0131nma d\u00f6neminin 33. ve 66. persentillerinden kalibre "
    "edilmektedir. Ger\u00e7ek zamanl\u0131 tespitte look-ahead bias engellenmi\u015ftir; regime_hat yaln\u0131zca "
    "ge\u00e7mi\u015f veriden hesaplanmaktad\u0131r. Elde edilen tespit do\u011frulu\u011fu %60.7\u2019dir (rastgele s\u0131n\u0131r: %33.3).",
)

add_heading_custom(doc, "3.4 Gymnasium Ortam\u0131", level=2)
add_para(
    doc,
    "OpenAI Gymnasium uyumlu MMEnv ortam\u0131 olu\u015fturulmu\u015ftur. G\u00f6zlem uzay\u0131 6 boyutludur:",
)
add_para(doc, "obs = [q_norm, \u03c3\u0302_t, \u03c4, r_L, r_M, r_H]", align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(
    doc,
    "Burada q_norm = inv / inv_max_clip envanter normalizasyonu, \u03c3\u0302_t kayan ger\u00e7ekle\u015fmi\u015f "
    "volatilite, \u03c4 = (T\u2212t)/T kalan zaman fraksiyonu, r_L/r_M/r_H ise rejim one-hot "
    "kodlamas\u0131d\u0131r (ppo_blind i\u00e7in hepsi s\u0131f\u0131r).",
)
add_para(doc, "Eylem uzay\u0131 MultiDiscrete([5, 5]) \u015feklinde tan\u0131mlanm\u0131\u015ft\u0131r:")
for line in [
    "h_idx \u2208 {0,1,2,3,4}  \u2192  h = h_idx + 1 tick  (half-spread)",
    "m_idx \u2208 {0,1,2,3,4}  \u2192  m = m_idx \u2212 2 tick  (skew)",
    "\u03b4_bid = max(1, h + m)",
    "\u03b4_ask = max(1, h \u2212 m)",
]:
    add_para(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "\u00d6d\u00fcl fonksiyonu \u015fu \u015fekilde tan\u0131mlanm\u0131\u015ft\u0131r:")
add_para(
    doc,
    "R_t = (equity_t \u2212 equity_{t\u22121}) \u2212 \u03b7 \u00b7 inv_t\u00b2",
    align=WD_ALIGN_PARAGRAPH.CENTER,
)
add_para(
    doc,
    "Burada \u03b7 = 0.001 envanter ceza katsay\u0131s\u0131d\u0131r (\u03b7 ablasyonu ile optimize edilmi\u015ftir).",
)

add_heading_custom(doc, "3.5 Referans Stratejiler", level=2)
add_heading_custom(doc, "3.5.1 Naif Sabit-Spread Stratejisi", level=3)
add_para(
    doc,
    "Her zaman ad\u0131m\u0131nda simetrik sabit half-spread uygulanmaktad\u0131r: \u03b4_bid = \u03b4_ask = h = 2 tick. "
    "Envanter fark\u0131ndal\u0131\u011f\u0131 veya skew mekanizmas\u0131 i\u00e7ermemektedir.",
)

add_heading_custom(doc, "3.5.2 Avellaneda-Stoikov Stratejisi", level=3)
add_para(doc, "Rezervasyon fiyat\u0131 ve yar\u0131-spread \u015fu form\u00fcllerle hesaplanmaktad\u0131r:")
add_para(doc, "r = mid \u2212 q \u00b7 \u03b3 \u00b7 \u03c3\u00b2 \u00b7 \u03c4", align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(
    doc,
    "d = \u00bd \u00b7 \u03b3 \u00b7 \u03c3\u00b2 \u00b7 \u03c4 + (1/\u03b3) \u00b7 ln(1 + \u03b3/k_price)",
    align=WD_ALIGN_PARAGRAPH.CENTER,
)
add_para(
    doc,
    "Burada \u03b3 = 0.01 risk aversion parametresi, k_price = k_ticks / tick_size fiyat cinsinden "
    "yo\u011funluk azalma parametresidir. Delta de\u011ferleri [1, 25] tick aral\u0131\u011f\u0131nda k\u0131rp\u0131lmaktad\u0131r.",
)

add_heading_custom(doc, "3.6 PPO E\u011fitimi ve De\u011ferlendirme Protokol\u00fc", level=2)
add_para(
    doc,
    "Stable-Baselines3 k\u00fct\u00fcphanesi kullan\u0131larak PPO e\u011fitimi ger\u00e7ekle\u015ftirilmi\u015ftir. Hiperparametreler:",
)

hp_rows = [
    ("total_timesteps", "1.000.000"),
    ("learning_rate", "0.0003"),
    ("n_steps", "2048"),
    ("batch_size", "256"),
    ("n_epochs", "10"),
    ("gamma", "0.999"),
    ("gae_lambda", "0.95"),
    ("clip_range", "0.2"),
    ("ent_coef", "0.01"),
]
add_table_with_data(doc, ["Parametre", "De\u011fer"], hp_rows, col_widths=[2.5, 3.5])
add_para(
    doc,
    "Her seed i\u00e7in episode uzunlu\u011fu 8000 ad\u0131m olup walk-forward b\u00f6l\u00fcnmesi %70/%30 "
    "(train: 5600, test: 2400 ad\u0131m) \u015feklinde uygulanm\u0131\u015ft\u0131r. De\u011ferlendirme yaln\u0131zca test "
    "b\u00f6l\u00fcm\u00fcnden ger\u00e7ekle\u015ftirilmektedir. Toplam 20 ba\u011f\u0131ms\u0131z tohum kullan\u0131lm\u0131\u015ft\u0131r (seeds: 1\u201320). "
    "E\u011fitim GPU \u00fczerinde ger\u00e7ekle\u015ftirilmi\u015ftir (NVIDIA CUDA, PyTorch 2.6.0+cu124).",
)

doc.add_page_break()

# ── 4. SONU\u00c7LAR VE TARTI\u015eMA ───────────────────────────────────────────────
add_heading_custom(doc, "4. SONU\u00c7LAR VE TARTI\u015eMA", level=1)

add_heading_custom(doc, "4.1 Ana Deney Sonu\u00e7lar\u0131 (20 Tohum, OOS)", level=2)
add_para(
    doc,
    "Tablo 1. Out-of-Sample Performans \u00d6zeti (20 Tohum, Ortalama \u00b1 Std)",
    bold=True,
)

tablo1_rows = [
    ("AS", "5.05", "4.72", "0.105", "29.95", "0.444"),
    ("Naif", "4.49", "3.49", "0.126", "21.20", "0.119"),
    ("ppo_aware", "4.10", "0.78", "0.715", "2.00", "0.236"),
    ("ppo_blind", "4.42", "0.71", "0.740", "2.05", "0.232"),
]
add_table_with_data(
    doc,
    ["Strateji", "Ort. Equity", "Std", "Ort. Sharpe", "inv_p99", "Fill Rate"],
    tablo1_rows,
    col_widths=[1.2, 1.1, 0.8, 1.2, 1.0, 1.0],
)

bulgular = [
    (
        "Birinci bulgu \u2014 PPO\u2019nun Sharpe \u00dcst\u00fcnl\u00fc\u011f\u00fc",
        "Her iki PPO varyant\u0131 da Sharpe oran\u0131 a\u00e7\u0131s\u0131ndan klasik referans stratejileri kar\u015f\u0131s\u0131nda "
        "belirgin \u00fcst\u00fcnl\u00fck sergilemi\u015ftir. ppo_aware ortalama Sharpe = 0.715, ppo_blind ise 0.740 "
        "de\u011feri elde ederken, AS i\u00e7in bu de\u011fer 0.105, naif strateji i\u00e7in ise 0.126 olarak "
        "\u00f6l\u00e7\u00fclm\u00fc\u015ft\u00fcr. Bu fark yakla\u015f\u0131k 6\u20137 kat b\u00fcy\u00fckl\u00fc\u011f\u00fcndedir ve temel mekanizmas\u0131 envanter "
        "kontrol\u00fcd\u00fcr: PPO varyantlar\u0131 inv_p99 \u2248 2 tick d\u00fczeyinde \u00e7al\u0131\u015f\u0131rken, AS ve naif stratejiler "
        "s\u0131ras\u0131yla 29.95 ve 21.20 de\u011ferlerine ula\u015fmaktad\u0131r.",
    ),
    (
        "\u0130kinci bulgu \u2014 Null Sonu\u00e7: Rejim Fark\u0131ndal\u0131\u011f\u0131 Hipotezi Desteklenmedi",
        "ppo_aware, 20 tohumun yaln\u0131zca 9\u2019unda ppo_blind\u2019\u0131 geride b\u0131rakm\u0131\u015ft\u0131r. Bu oran "
        "istatistiksel olarak anlaml\u0131 bir fark olu\u015fturmamaktad\u0131r ve rejim bilgisinin performans "
        "\u00fczerinde net bir katk\u0131 sa\u011flad\u0131\u011f\u0131 hipotezini desteklememektedir. Ortalama equity "
        "de\u011ferlerindeki fark da ihmal edilebilir d\u00fczeydedir: ppo_aware = 4.10 \u00b1 0.78, "
        "ppo_blind = 4.42 \u00b1 0.71.",
    ),
    (
        "\u00dc\u00e7\u00fcnc\u00fc bulgu \u2014 PPO Varyans Avantaj\u0131",
        "PPO varyantlar\u0131n\u0131n equity std de\u011ferleri (0.71\u20130.78) AS (4.72) ve naif (3.49) "
        "stratejilerine k\u0131yasla belirgin bi\u00e7imde d\u00fc\u015f\u00fckt\u00fcr. Bu durum, PPO\u2019nun yaln\u0131zca daha "
        "y\u00fcksek Sharpe oran\u0131 de\u011fil, ayn\u0131 zamanda daha tutarl\u0131 ve g\u00fcvenilir bir performans "
        "profili sundu\u011funa i\u015faret etmektedir.",
    ),
]
for baslik, metin in bulgular:
    p = doc.add_paragraph()
    run_b = p.add_run(baslik + ": ")
    set_font(run_b, bold=True)
    run_m = p.add_run(metin)
    set_font(run_m)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)

add_para(doc, "Bu null sonu\u00e7 i\u00e7in \u00fc\u00e7 a\u00e7\u0131klay\u0131c\u0131 hipotez \u00f6ne s\u00fcr\u00fclmektedir:")
null_hipotezler = [
    "(a) G\u00f6zlem uzay\u0131nda \u03c3\u0302_t de\u011feri halihaz\u0131rda mevcut oldu\u011fundan, one-hot rejim kodlamas\u0131 \u03c3\u0302_t\u2019den t\u00fcretilmi\u015f bilgiyi yedekli bi\u00e7imde sunmakta ve marjinal katk\u0131 sa\u011flamamaktad\u0131r.",
    "(b) \u00d6d\u00fcl fonksiyonu rejime \u00f6zg\u00fc davran\u0131\u015f\u0131 do\u011frudan te\u015fvik etmedi\u011finden, rejim bilgisini performansa d\u00f6n\u00fc\u015ft\u00fcrmek i\u00e7in sinyal yetersiz kalmaktad\u0131r.",
    "(c) %60.7 d\u00fczeyindeki rejim tespit do\u011frulu\u011fu, yanl\u0131\u015f s\u0131n\u0131fland\u0131rmalar\u0131n avantaj yerine g\u00fcr\u00fclt\u00fc \u00fcretmesine neden olmaktad\u0131r.",
]
for h in null_hipotezler:
    add_bullet(doc, h)

add_heading_custom(doc, "4.2 Eylem Analizi (Davran\u0131\u015fsal Farkl\u0131la\u015fma)", level=2)
add_para(
    doc,
    "Tablo 2. Rejim Baz\u0131nda Eylem Da\u011f\u0131l\u0131m\u0131 (Ortalama \u00b1 Std, 20 Tohum)",
    bold=True,
)

tablo2_rows = [
    ("AS", "L", "1.05", "0.03", "0.00", "0.06", "0.000"),
    ("AS", "M", "1.02", "0.02", "\u22120.03", "0.09", "0.000"),
    ("AS", "H", "1.02", "0.02", "\u22120.03", "0.09", "0.000"),
    ("Naif", "L", "2.00", "0.00", "0.00", "0.00", "0.000"),
    ("Naif", "M", "2.00", "0.00", "0.00", "0.00", "0.000"),
    ("Naif", "H", "2.00", "0.00", "0.00", "0.00", "0.000"),
    ("ppo_aware", "L", "3.30", "1.16", "\u22120.61", "0.54", "0.204"),
    ("ppo_aware", "M", "2.69", "2.07", "0.25", "1.91", "0.333"),
    ("ppo_aware", "H", "3.00", "2.00", "0.34", "1.52", "0.333"),
    ("ppo_blind", "L", "2.89", "1.00", "\u22120.21", "0.95", "0.000"),
    ("ppo_blind", "M", "2.86", "1.15", "\u22120.37", "0.71", "0.000"),
    ("ppo_blind", "H", "2.69", "1.31", "\u22120.41", "0.58", "0.000"),
]
add_table_with_data(
    doc,
    ["Strateji", "Rejim", "Ort. h", "Std h", "Ort. m", "Std m", "P(h=5)"],
    tablo2_rows,
    col_widths=[1.0, 0.7, 0.8, 0.7, 0.8, 0.7, 0.8],
)

add_para(
    doc,
    "ppo_aware ajan\u0131, rejimler aras\u0131nda istatistiksel olarak anlaml\u0131 davran\u0131\u015fsal farkl\u0131la\u015fma "
    "sergilemektedir. D\u00fc\u015f\u00fck volatilite rejiminde (L) daha geni\u015f spread (h = 3.30) ve negatif "
    "skew (m = \u22120.61) g\u00f6zlemlenirken, y\u00fcksek volatilite rejiminde (H) daha dar spread "
    "(h = 3.00) ve pozitif skew (m = +0.34) tercih edilmektedir. Bu asimetri, ajan\u0131n rejim "
    "bilgisini eylem politikas\u0131na entegre etti\u011finin davran\u0131\u015fsal kan\u0131t\u0131n\u0131 olu\u015fturmaktad\u0131r.",
)
add_para(
    doc,
    "Buna kar\u015f\u0131n ppo_blind, t\u00fcm rejimlerde tutarl\u0131 bi\u00e7imde m \u2248 \u22120.2 ile \u22120.4 aral\u0131\u011f\u0131nda "
    "seyreden tekd\u00fcze bir politika sergilemektedir. Bu durum, rejim bilgisinin eylem d\u00fczeyinde "
    "ger\u00e7ekten kullan\u0131ld\u0131\u011f\u0131n\u0131, ancak kullan\u0131m\u0131n performans fark\u0131na yans\u0131mad\u0131\u011f\u0131n\u0131 ortaya "
    "koymaktad\u0131r.",
)

doc.add_page_break()

# ── 5. SONU\u00c7 ─────────────────────────────────────────────────────────────
add_heading_custom(doc, "5. SONU\u00c7", level=1)

sonuc_paragraflar = [
    "Bu \u00e7al\u0131\u015fma, volatilite rejimleri alt\u0131nda peki\u015ftirmeli \u00f6\u011frenme tabanl\u0131 piyasa yap\u0131c\u0131l\u0131\u011f\u0131n\u0131 ara\u015ft\u0131rmaktad\u0131r. Temel bulgular \u015fu \u015fekilde \u00f6zetlenebilir:",
    "Birincisi, her iki PPO varyant\u0131 da Sharpe oran\u0131 a\u00e7\u0131s\u0131ndan Avellaneda-Stoikov ve naif referans stratejilerini yakla\u015f\u0131k 6\u20137 kat ge\u00e7mi\u015ftir. Bu \u00fcst\u00fcnl\u00fc\u011f\u00fcn temel kayna\u011f\u0131, PPO\u2019nun envanter riskini etkin bi\u00e7imde y\u00f6netmesi ve inv_p99 \u2248 2 tick d\u00fczeyinde tutmas\u0131d\u0131r.",
    "\u0130kincisi, rejim fark\u0131ndal\u0131\u011f\u0131 hipotezi istatistiksel olarak desteklenmemi\u015ftir. ppo_aware, 20 tohumun yaln\u0131zca 9\u2019unda ppo_blind\u2019\u0131 geride b\u0131rakm\u0131\u015f; ortalama performans fark\u0131 ihmal edilebilir d\u00fczeyde kalm\u0131\u015ft\u0131r. Bu null sonu\u00e7 i\u00e7in en g\u00fc\u00e7l\u00fc a\u00e7\u0131klama, g\u00f6zlem uzay\u0131nda halihaz\u0131rda bulunan \u03c3\u0302_t de\u011ferinin rejim bilgisini \u00f6rt\u00fck olarak i\u00e7ermesi ve one-hot rejim kodlamas\u0131n\u0131n yedekli bilgi sunmas\u0131d\u0131r.",
    "\u00dc\u00e7\u00fcnc\u00fcs\u00fc, ppo_aware davran\u0131\u015fsal d\u00fczeyde rejim farkl\u0131la\u015fmas\u0131 \u00f6\u011frenmesine kar\u015f\u0131n bu farkl\u0131la\u015fma performans avantaj\u0131na d\u00f6n\u00fc\u015fmemi\u015ftir. Bu bulgu, davran\u0131\u015fsal analizi performans metriklerinden ba\u011f\u0131ms\u0131z de\u011ferlendirmenin \u00f6nemini vurgulamaktad\u0131r.",
    "Gelecek \u00e7al\u0131\u015fmalar kapsam\u0131nda \u015fu y\u00f6nler \u00f6nerilmektedir:",
]
for p_text in sonuc_paragraflar:
    add_para(doc, p_text)

gelecek = [
    "(1) Rejim parametrelerini g\u00f6zlem uzay\u0131ndan \u00e7\u0131kararak bilgi fazlal\u0131\u011f\u0131n\u0131 giderme ve pure ablation tasar\u0131m\u0131.",
    "(2) Rejime ko\u015fullu \u00f6d\u00fcl tasar\u0131m\u0131 \u2014 y\u00fcksek volatilite rejiminde envanter cezas\u0131n\u0131n art\u0131r\u0131lmas\u0131 gibi mekanizmalar.",
    "(3) Skew penalty ablasyonu (c=0 kontrol grubu dahil) ile eylem d\u00fczeyinde d\u00fczenlile\u015ftirmenin sistematik de\u011ferlendirmesi.",
    "(4) Model yanl\u0131\u015f belirleme (model misspecification) benchmark\u2019\u0131 \u2014 AS parametrelerinin rejime ba\u011f\u0131ml\u0131 k\u0131l\u0131nmas\u0131.",
]
for g in gelecek:
    add_bullet(doc, g)

doc.add_page_break()

# ── KAYNAKLAR ─────────────────────────────────────────────────────────────
add_heading_custom(doc, "KAYNAKLAR", level=1)

kaynaklar = [
    "[1] Avellaneda, M., & Stoikov, S. (2008). High-frequency trading in a limit order book. Quantitative Finance, 8(3), 217-224.",
    "[2] Schulman, J., Wolski, F., Dhariwal, P., Radford, A., & Klimov, O. (2017). Proximal policy optimization algorithms. arXiv preprint arXiv:1707.06347.",
    "[3] Spooner, T., Fearnley, J., Savani, R., & Koukorinis, A. (2018). Market making via reinforcement learning. In Proceedings of the 17th International Conference on Autonomous Agents and MultiAgent Systems (pp. 434-442).",
    "[4] Raffin, A., Hill, A., Gleave, A., Kanervisto, A., Ernestus, M., & Dormann, N. (2021). Stable-baselines3: Reliable reinforcement learning implementations. Journal of Machine Learning Research, 22(268), 1-8.",
    "[5] Cont, R. (2001). Empirical characteristics of asset returns: Stylized facts. Quantitative Finance, 1(2), 223-236.",
    "[6] Hamilton, J. D. (1989). A new approach to the economic analysis of nonstationary time series and the business cycle. Econometrica, 57(2), 357-384.",
]
for k in kaynaklar:
    p = doc.add_paragraph()
    run = p.add_run(k)
    set_font(run)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)

doc.save("manuscript/thesis_final.docx")
print("Kaydedildi: manuscript/thesis_final.docx")
