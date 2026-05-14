"""Generate thesis_29.docx from thesis_28.docx.

Adds post-hoc signal redundancy diagnostics as supporting main-text evidence.
No PPO training, WP5/WP6 reruns, or protected evidence regeneration occurs.
"""

from __future__ import annotations

import io
import shutil
import sys
from pathlib import Path

import pandas as pd
from docx import Document

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

SRC = Path("manuscript/thesis_28.docx")
DST = Path("manuscript/thesis_29.docx")
POSTHOC = Path("docs/internal/posthoc_signal_analysis")


def _replace_in_paragraph(para, old_text: str, new_text: str) -> bool:
    full = "".join(run.text for run in para.runs)
    if old_text not in full:
        return False
    new_full = full.replace(old_text, new_text)
    if para.runs:
        para.runs[0].text = new_full
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(new_full)
    return True


def _find_paragraph(doc: Document, text: str):
    for para in doc.paragraphs:
        if para.text.strip() == text:
            return para
    raise RuntimeError(f"Paragraph not found: {text}")


def _add_paragraph_before(doc: Document, ref_para, text: str, style: str | None = None):
    para = doc.add_paragraph(text, style=style)
    ref_para._p.addprevious(para._p)
    return para


def _add_heading_before(doc: Document, ref_para, text: str, level: int):
    para = doc.add_heading(text, level=level)
    ref_para._p.addprevious(para._p)
    return para


def _add_table_before(doc: Document, ref_para, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style="Table Grid")
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for i, row in enumerate(rows, start=1):
        for j, value in enumerate(row):
            table.rows[i].cells[j].text = str(value)
    ref_para._p.addprevious(table._tbl)
    return table


def _read_posthoc_numbers() -> dict[str, float | int]:
    pred = pd.read_csv(POSTHOC / "predictability_metrics.csv")
    best_pred = pred[
        (pred["target_label"] == "regime_hat")
        & (pred["evaluation_mode"] == "within_source_calibrated")
        & (pred["model"] == "random_forest")
    ].iloc[0]

    inc = pd.read_csv(POSTHOC / "model_comparison_table.csv").iloc[0]
    action = pd.read_csv(POSTHOC / "action_model_metrics.csv")
    manifest = pd.read_csv(POSTHOC / "source_manifest.csv")
    alignment = pd.read_csv(POSTHOC / "action_curve_alignment.csv")

    deltas = []
    for keys, group in action.groupby(["run_group", "strategy", "target"], sort=True):
        a = group[group["model"] == "A_sigma_hat"].iloc[0]
        b = group[group["model"] == "B_sigma_hat_plus_regime"].iloc[0]
        deltas.append(
            {
                "target": keys[2],
                "r2_sigma": float(a["r2"]),
                "delta_r2": float(b["r2"] - a["r2"]),
            }
        )
    delta_df = pd.DataFrame(deltas)

    return {
        "source_count": int(len(manifest)),
        "accepted_action_curves": int((alignment["status"] == "accepted").sum()),
        "pred_bal_acc": float(best_pred["balanced_accuracy"]),
        "pred_nmi": float(best_pred["normalized_mutual_information"]),
        "pred_macro_f1": float(best_pred["macro_f1"]),
        "inc_delta_r2": float(inc["delta_oos_r2"]),
        "inc_delta_mae": float(inc["delta_mae"]),
        "inc_delta_rmse": float(inc["delta_rmse"]),
        "action_h_r2": float(delta_df[delta_df["target"] == "h"]["r2_sigma"].mean()),
        "action_m_r2": float(delta_df[delta_df["target"] == "m"]["r2_sigma"].mean()),
        "action_h_delta_r2": float(delta_df[delta_df["target"] == "h"]["delta_r2"].mean()),
        "action_m_delta_r2": float(delta_df[delta_df["target"] == "m"]["delta_r2"].mean()),
    }


def _insert_methodology(doc: Document, numbers: dict[str, float | int]) -> None:
    ref = _find_paragraph(doc, "4. SONUÇLAR VE TARTIŞMA")
    _add_heading_before(doc, ref, "3.9 Post-hoc Signal Redundancy Diagnostics", level=2)
    _add_paragraph_before(
        doc,
        ref,
        (
            "WP5 ve WP6 ana deneylerinden sonra, sinyal yedekliliği yorumunu "
            "desteklemek için hafif post-hoc tanı analizleri yürütülmüştür. Bu "
            "analizler yeni bir RL projesi veya yeni bir eğitim aşaması değildir: "
            "yalnızca donmuş WP2/WP5/WP6 artifact'ları ve "
            "docs/internal/posthoc_signal_analysis/ altındaki yeni özet çıktıları "
            "kullanılmıştır. PPO yeniden eğitilmemiş, WP5/WP6 yeniden koşturulmamış "
            "ve korumalı evidence CSV/figure artifact'ları değiştirilmemiştir."
        ),
    )
    _add_paragraph_before(
        doc,
        ref,
        (
            "Amaç, 'sigma_hat explicit kategorik rejim etiketlerinin ekonomik olarak "
            "ilgili bilgisinin çoğunu zaten taşıyor olabilir' yorumuna mekanik destek "
            "sağlamaktır; bu analizler birincil keşif deneyi veya causal PPO iç-mekanizma "
            "kanıtı olarak tasarlanmamıştır. source_manifest.csv, kullanılan "
            f"{numbers['source_count']} benzersiz WP2-style sentetik kaynağı listeler; "
            "action_curve_alignment.csv ise yalnızca donmuş curve/snapshot hizalaması "
            "doğrulanan action curve'lerinin kullanıldığını kayda geçirir."
        ),
    )
    _add_paragraph_before(
        doc,
        ref,
        (
            "Üç tanı bloğu uygulanmıştır: (A) regime_hat etiketinin yalnızca sigma_hat "
            "ile tahmin edilebilirliği; (B) sigma_hat gözlendikten sonra regime_hat'in "
            "gelecek mutlak mid-return için ek tahmin değeri; (C) PPO'nun seçtiği "
            "half-spread h ve skew m eylemlerinde, sigma_hat'e eklenen explicit label "
            "değişkenlerinin açıklayıcı katkısı. Tüm split'ler kronolojik train/test "
            "mantığını korur ve geleceğe bilgi sızıntısı yaratmayacak şekilde kurulmuştur."
        ),
    )
    _add_table_before(
        doc,
        ref,
        ["Blok", "Girdi / Hedef", "Yorum"],
        [
            [
                "A",
                "sigma_hat -> regime_hat",
                "Observed kategorik etiketin continuous volatility signal'dan ne ölçüde recover edilebildiği.",
            ],
            [
                "B",
                "target ~ sigma_hat vs target ~ sigma_hat + regime_hat",
                "Explicit label'ın sigma_hat sonrası held-out predictive contribution ölçümü.",
            ],
            [
                "C",
                "action ~ sigma_hat vs action ~ sigma_hat + regime_hat",
                "Eylem düzeyinde incremental explanatory gain; causal PPO mechanism proof değildir.",
            ],
        ],
    )


def _insert_results(doc: Document, numbers: dict[str, float | int]) -> None:
    ref = _find_paragraph(doc, "6. SONUÇ")
    _add_heading_before(doc, ref, "5.4 Post-hoc Signal Redundancy Diagnostics", level=2)
    _add_paragraph_before(
        doc,
        ref,
        (
            "WP6 sonrasında yürütülen post-hoc diagnostics, sinyal-yedekliliği yorumunu "
            "ana deneyleri yeniden koşturmadan sınayan üç yardımcı test sunar. Kaynak "
            "dosyalar predictability_summary.md / predictability_metrics.csv, "
            "incremental_value_summary.md / incremental_metrics.csv / "
            "model_comparison_table.csv, action_explanation_summary.md / "
            "action_model_metrics.csv ve final_signal_redundancy_assessment.md içinde "
            "saklanmaktadır. Bu bölüm, bu çıktıları ana metin içinde destekleyici evidence "
            "olarak raporlar."
        ),
    )
    _add_paragraph_before(
        doc,
        ref,
        (
            "Regime recoverability sonucu çok güçlüdür: source-calibrated modelde "
            "regime_hat yalnızca sigma_hat kullanılarak neredeyse tamamen recover "
            f"edilebilmektedir (balanced accuracy = {numbers['pred_bal_acc']:.4f}, "
            f"NMI = {numbers['pred_nmi']:.4f}, macro F1 = {numbers['pred_macro_f1']:.4f}). "
            "Bu bulgu beklenen yöndedir: regime_hat zaten rolling-volatility temelli "
            "thresholding ile üretildiği için, kaynak-içi kalibrasyonla sigma_hat'ten "
            "yüksek doğrulukla tahmin edilebilir."
        ),
    )
    _add_paragraph_before(
        doc,
        ref,
        (
            "İkinci tanı, sigma_hat gözlendikten sonra explicit regime_hat etiketinin "
            "gelecek adım mutlak mid-return tahminine ne kattığını ölçer. Model B "
            "(sigma_hat + regime_hat), Model A'ya (sigma_hat) göre held-out R²'yi yalnızca "
            f"{numbers['inc_delta_r2']:+.6f} değiştirmiştir; MAE değişimi "
            f"{numbers['inc_delta_mae']:+.6f}, RMSE değişimi "
            f"{numbers['inc_delta_rmse']:+.6f}'dır. Eğitim kümesinde ek step-function "
            "terimler istatistiksel olarak seçilebilir olsa da, held-out predictive gain "
            "pratik olarak çok küçüktür."
        ),
    )
    _add_paragraph_before(
        doc,
        ref,
        (
            "Üçüncü tanı daha sınırlı ve daha temkinli okunmalıdır. Donmuş WP5 curve CSV'leri "
            "yalnızca run-level WP2 snapshot ile regime_hat dizisi tam hizalanabildiğinde "
            f"kullanılmıştır; bu nedenle action analizi {numbers['accepted_action_curves']} "
            "hizalanmış curve ile sınırlıdır. Basit sigma-only action modelleri h için "
            f"ortalama R² = {numbers['action_h_r2']:.3f}, m için ortalama R² = "
            f"{numbers['action_m_r2']:.3f} üretmiştir; bu, PPO action'larının özellikle "
            "skew tarafında sigma_hat tarafından tamamen açıklanmadığını gösterir. Buna "
            "karşın explicit label eklemek incremental R²'yi h için "
            f"{numbers['action_h_delta_r2']:+.3f}, m için "
            f"{numbers['action_m_delta_r2']:+.3f} değiştirmiştir; yani label'ın ek açıklayıcı "
            "kazancı sınırlıdır."
        ),
    )
    _add_table_before(
        doc,
        ref,
        ["Tanı", "Ana sayı", "Okuma"],
        [
            [
                "A: regime_hat recoverability",
                f"Balanced accuracy {numbers['pred_bal_acc']:.4f}; NMI {numbers['pred_nmi']:.4f}",
                "Observed regime_hat source-calibrated sigma_hat'ten neredeyse tamamen recover edilebilir.",
            ],
            [
                "B: incremental prediction",
                (
                    f"ΔR² {numbers['inc_delta_r2']:+.6f}; "
                    f"ΔMAE {numbers['inc_delta_mae']:+.6f}; "
                    f"ΔRMSE {numbers['inc_delta_rmse']:+.6f}"
                ),
                "sigma_hat sonrası explicit label'ın held-out predictive contribution'ı çok küçüktür.",
            ],
            [
                "C: action explanation",
                (
                    f"ΔR² h {numbers['action_h_delta_r2']:+.3f}; "
                    f"ΔR² m {numbers['action_m_delta_r2']:+.3f}"
                ),
                "Action-level evidence mixed; limited incremental label gain, causal mechanism proof değil.",
            ],
        ],
    )
    _add_paragraph_before(
        doc,
        ref,
        (
            "Toplu yorum: Post-hoc diagnostics further support the signal-redundancy "
            "interpretation. In the tested synthetic setting, observed categorical regime "
            "labels are highly recoverable from source-calibrated sigma_hat and add only "
            "limited held-out predictive value once sigma_hat is observed. Action-level "
            "regressions show limited incremental explanatory gain from explicit labels, "
            "but they do not establish a causal internal PPO mechanism."
        ),
    )


def _update_abstract_and_conclusion(doc: Document) -> None:
    for para in doc.paragraphs:
        _replace_in_paragraph(para, "Nisan 2026 — Sürüm 28", "Mayıs 2026 — Sürüm 29")
        _replace_in_paragraph(
            para,
            "WP5, ppo_aware ile ppo_blind arasındaki pratik eşitliği kanıtlamak amacıyla TOST'u eşitlik yönünde kullanmıştır.",
            "WP5, ppo_aware ile ppo_blind arasındaki pratik eşitliği değerlendirmek amacıyla TOST'u eşitlik yönünde kullanmıştır.",
        )

    old_wp6 = (
        "Sinyal informativeness sweep'i (WP6), kategorik regime label'larının "
        "sigma_hat sürekli sinyali zayıfladığında dahi performansı iyileştirmediğini, "
        "hatta birlikte kullanıldığında yönsel olarak düşürdüğünü göstermiştir."
    )
    new_wp6 = (
        "Sinyal informativeness sweep'i (WP6), kategorik regime label'larının "
        "sigma_hat sürekli sinyali zayıfladığında dahi performansı iyileştirmediğini, "
        "hatta birlikte kullanıldığında yönsel olarak düşürdüğünü göstermiştir. "
        "Buna eklenen post-hoc tanı analizleri, observed categorical labels'ın "
        "source-calibrated sigma_hat'ten yüksek doğrulukla recover edilebildiğini ve "
        "sigma_hat gözlendikten sonra sınırlı incremental predictive value sunduğunu "
        "göstererek sinyal-yedekliliği yorumunu daha da desteklemektedir."
    )
    for para in doc.paragraphs:
        if old_wp6 in para.text:
            _replace_in_paragraph(para, old_wp6, new_wp6)
            break

    for para in doc.paragraphs:
        _replace_in_paragraph(
            para,
            "Anahtar kelimeler: piyasa yapıcılığı, pekiştirmeli öğrenme, PPO, volatilite rejimleri, Avellaneda-Stoikov, ablasyon, oracle agent, sinyal yedekliliği, model yanlış belirleme",
            "Anahtar kelimeler: piyasa yapıcılığı, pekiştirmeli öğrenme, PPO, volatilite rejimleri, Avellaneda-Stoikov, ablasyon, oracle agent, sinyal yedekliliği, post-hoc tanı analizleri, model yanlış belirleme",
        )

    future_ref = _find_paragraph(doc, "Gelecek Çalışmalar")
    _add_paragraph_before(
        doc,
        future_ref,
        (
            "Bulgu 8 — Post-hoc Signal Redundancy Diagnostics: Donmuş artifact'lar üzerinde "
            "yürütülen tanı analizleri, observed regime_hat etiketinin source-calibrated "
            "sigma_hat'ten neredeyse tamamen recover edilebildiğini ve sigma_hat sonrası "
            "held-out predictive contribution'ın çok sınırlı kaldığını göstermiştir. "
            "Action-level regressions explicit label'ın incremental explanatory gain'inin "
            "sınırlı olduğunu gösterse de, PPO'nun içsel causal mekanizmasını kanıtlamaz."
        ),
    )
    _add_paragraph_before(
        doc,
        future_ref,
        (
            "Sonuç olarak, test edilen sentetik HFMM ortamında explicit categorical regime "
            "labels, policy'ye zaten sunulan continuous volatility signal'ın ötesinde "
            "robust incremental value sağlamamaktadır. Bu ifade ortam ve kalibrasyon bandı "
            "ile sınırlıdır; labels'ın sıfır bilgi içerdiği veya tüm piyasa ortamlarında "
            "işlevsiz olduğu iddia edilmemektedir."
        ),
    )


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    for required in [
        "predictability_metrics.csv",
        "model_comparison_table.csv",
        "action_model_metrics.csv",
        "source_manifest.csv",
        "action_curve_alignment.csv",
        "predictability_summary.md",
        "incremental_value_summary.md",
        "action_explanation_summary.md",
        "final_signal_redundancy_assessment.md",
    ]:
        path = POSTHOC / required
        if not path.exists():
            raise FileNotFoundError(path)

    shutil.copy2(SRC, DST)
    doc = Document(DST)
    numbers = _read_posthoc_numbers()
    _update_abstract_and_conclusion(doc)
    _insert_methodology(doc, numbers)
    _insert_results(doc, numbers)
    doc.save(DST)

    print(f"thesis_29.docx saved ({DST.stat().st_size:,} bytes)")
    print("Post-hoc numbers inserted:")
    for key, value in numbers.items():
        print(f"  {key}: {value}")


if __name__ == "__main__":
    main()
