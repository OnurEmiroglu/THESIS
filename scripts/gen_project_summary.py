"""Danışman için proje özeti DOCX üretir."""

import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

dst = Document()

# ── helpers ────────────────────────────────────────────────
def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_normal(doc, text):
    return doc.add_paragraph(text)

def add_bold_normal(doc, bold_part, rest):
    p = doc.add_paragraph()
    r = p.add_run(bold_part)
    r.bold = True
    p.add_run(rest)
    return p

def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Table Grid')
    for j, h in enumerate(headers):
        t.rows[0].cells[j].text = h
        for p in t.rows[0].cells[j].paragraphs:
            for r in p.runs:
                r.bold = True
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            t.rows[i + 1].cells[j].text = str(val)
    return t

def add_bullet(doc, bold_part, rest):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(bold_part)
    r.bold = True
    p.add_run(rest)
    return p

# ── Kapak ──────────────────────────────────────────────────
p = dst.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Proje Özeti")
r.bold = True
r.font.size = Pt(18)

p = dst.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Yüksek Frekanslı Piyasa Yapıcılığı ile Pekiştirmeli Öğrenme")
r.font.size = Pt(14)

p = dst.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Onur Emiroğlu — Mart 2026")

dst.add_paragraph()  # boşluk

# ================================================================
# 1. PROJE GENEL BAKIŞ
# ================================================================
add_heading(dst, "1. Proje Genel Bakış", level=1)

p = dst.add_paragraph()
r = p.add_run("Araştırma sorusu: ")
r.bold = True
p.add_run(
    "Volatilite rejim bilgisine erişimi olan bir PPO ajanı, bu bilgiden yoksun olan "
    "eşdeğerine kıyasla daha iyi performans sergiler mi?"
)

add_normal(dst, (
    "Bu çalışma, pekiştirmeli öğrenme (PPO) tabanlı piyasa yapıcı ajanların volatilite rejim "
    "bilgisinden fayda sağlayıp sağlamadığını araştırmaktadır. Tamamen sentetik bir piyasa ortamında "
    "(Markov zinciri rejim geçişleri + Poisson dolum modeli) 5 farklı PPO varyantı 20 bağımsız tohum "
    "ile karşılaştırılmıştır. Temel bulgu: rejim bilgisi eklemek anlamlı bir performans artışı "
    "sağlamamaktadır (null result). Bu null result, gözlem uzayında zaten mevcut olan sigma_hat "
    "sinyalinin rejim bilgisini örtük olarak içermesiyle açıklanmaktadır."
))

# ================================================================
# 2. METODOLOJİ
# ================================================================
add_heading(dst, "2. Metodoloji", level=1)

# 2a. Piyasa Simülasyonu
add_heading(dst, "2a. Piyasa Simülasyonu", level=2)
add_table(dst,
    ["Bileşen", "Detay"],
    [
        ["Mid-price modeli", "Aritmetik Brownian Hareketi (ABM)"],
        ["Zaman adımı (Δt)", "0.2 saniye"],
        ["Başlangıç fiyatı", "100.0"],
        ["Tick büyüklüğü", "0.01"],
        ["Dolum modeli", "Poisson süreci: λ(δ) = A·exp(−k·δ), A=5.0, k=1.5"],
        ["Komisyon", "0.2 baz puan / işlem"],
        ["Latency", "1 adım (adverse selection içselleştirilmiş)"],
    ],
)

# 2b. Rejim Modeli
add_heading(dst, "2b. Rejim Modeli", level=2)
add_table(dst,
    ["Parametre", "Değer"],
    [
        ["Rejimler", "L (Düşük), M (Orta), H (Yüksek) volatilite"],
        ["σ_L / σ_M / σ_H", "0.48 / 0.80 / 1.44 tick/√saniye"],
        ["Geçiş matrisi", "Sticky Markov zinciri (köşegen ≈ 0.99)"],
        ["Beklenen rejim süresi", "L: ~300, M: ~120, H: ~250 adım"],
        ["Rejim tespiti (rv_baseline)", "%60.7 doğruluk (rastgele sınır: %33.3)"],
    ],
)

add_normal(dst, (
    "Volatilite rejimleri piyasada doğal olarak oluşan farklı oynaklık dönemlerini temsil eder — "
    "tıpkı sakin, orta ve fırtınalı deniz koşulları gibi. Markov zinciri yapısı sayesinde rejimler "
    "ani değil, yavaş geçişli olur; bu da kayan gerçekleşmiş volatilite (RV) ile rejim tespitini "
    "mümkün kılar."
))

# 2c. Gymnasium Ortamı ve PPO
add_heading(dst, "2c. Gymnasium Ortamı ve PPO", level=2)

add_normal(dst, "Gözlem Uzayı (6 boyut):")
add_table(dst,
    ["Eleman", "Açıklama"],
    [
        ["q_norm", "Normalize edilmiş envanter ∈ [-1, +1]"],
        ["σ̂_t", "Kayan gerçekleşmiş volatilite (50 adım pencere)"],
        ["τ", "Kalan zaman fraksiyonu ∈ [0, 1]"],
        ["r_L, r_M, r_H", "Rejim one-hot kodlaması (ppo_blind için hepsi 0)"],
    ],
)

add_normal(dst, "PPO Hiperparametreleri:")
add_table(dst,
    ["Parametre", "Değer"],
    [
        ["total_timesteps", "1.000.000"],
        ["learning_rate", "0.0003"],
        ["n_steps", "2048"],
        ["batch_size", "256"],
        ["n_epochs", "10"],
        ["gamma", "0.999"],
        ["ent_coef", "0.01"],
    ],
)

add_normal(dst, "5 PPO Varyantı:")
add_table(dst,
    ["Varyant", "σ̂ kullanır mı?", "Rejim kaynağı", "Ne gözlemler?"],
    [
        ["ppo_sigma_only", "Evet", "Yok", "Yalnızca σ̂ + sıfır rejim"],
        ["ppo_regime_only", "Hayır", "Tahmini (hat)", "Yalnızca rejim_hat"],
        ["ppo_combined", "Evet", "Tahmini (hat)", "σ̂ + rejim_hat"],
        ["ppo_oracle_pure", "Hayır", "Gerçek (true)", "Yalnızca rejim_true"],
        ["ppo_oracle_full", "Evet", "Gerçek (true)", "σ̂ + rejim_true"],
    ],
)

add_normal(dst, "Ödül fonksiyonu:")
add_normal(dst, "R_t = (equity_t − equity_{t−1}) − η · inv_t²")
add_normal(dst, "η = 0.001 (envanter ceza katsayısı, ablasyon ile optimize edilmiştir)")

# ================================================================
# 3. DENEY SONUÇLARI
# ================================================================
add_heading(dst, "3. Deney Sonuçları", level=1)

# 3a
add_heading(dst, "3a. Ana Deney — ppo_aware vs ppo_blind (20 Tohum, OOS)", level=2)
add_table(dst,
    ["Strateji", "Ort. Equity", "Std", "Ort. Sharpe", "inv_p99", "Fill Rate"],
    [
        ["AS", "5.05", "4.72", "0.105", "29.95", "0.444"],
        ["Naif", "4.49", "3.49", "0.126", "21.20", "0.119"],
        ["ppo_aware", "4.10", "0.78", "0.715", "2.00", "0.236"],
        ["ppo_blind", "4.42", "0.71", "0.740", "2.05", "0.232"],
    ],
)

add_normal(dst, (
    "PPO ajanları Sharpe oranı açısından klasik stratejileri yaklaşık 6-7 kat geçmektedir. Temel "
    "mekanizma envanter kontrolüdür: PPO ajanları pozisyonlarını sıkı tutarken (inv_p99 ≈ 2 lot), "
    "AS ve naif stratejiler 20-30 lot envanter taşımaktadır. Mutlak equity'de AS daha yüksek görünse "
    "de bu, yüksek envanter riski pahasına elde edilmektedir."
))

# 3b
add_heading(dst, "3b. Ablasyon Deneyi — 5 PPO Varyantı (20 Tohum)", level=2)
add_table(dst,
    ["Varyant", "Sharpe (ort. ± std)", "Final Equity (ort. ± std)", "inv_p99"],
    [
        ["ppo_sigma_only", "0.753 ± 0.111", "4.55 ± 0.65", "1.9 ± 0.7"],
        ["ppo_oracle_full", "0.722 ± 0.130", "4.07 ± 0.61", "2.0 ± 0.5"],
        ["ppo_regime_only", "0.698 ± 0.145", "4.01 ± 0.68", "2.1 ± 1.5"],
        ["ppo_combined", "0.696 ± 0.134", "3.91 ± 0.72", "1.8 ± 0.5"],
        ["ppo_oracle_pure", "0.684 ± 0.111", "3.93 ± 0.60", "1.9 ± 0.6"],
        ["naive (referans)", "0.127 ± 0.092", "4.49 ± 3.49", "21.2"],
        ["AS (referans)", "0.105 ± 0.082", "5.05 ± 4.72", "29.9"],
    ],
)

p = dst.add_paragraph()
r = p.add_run("Kritik bulgu — Oracle paradoksu: ")
r.bold = True
p.add_run(
    "Mükemmel rejim bilgisine sahip ppo_oracle_full (Sharpe: 0.722), yalnızca sigma'yı gözlemleyen "
    "ppo_sigma_only'yi (Sharpe: 0.753) geçememektedir. Bu, rejim etiketinin sigma_hat'in taşıdığı "
    "bilgiyi yedekli biçimde sunduğunu kanıtlamaktadır."
)

# 3c
add_heading(dst, "3c. Rejim Bazlı Performans (Ex-Post Attribution)", level=2)
add_table(dst,
    ["Strateji", "Rejim", "Sharpe", "inv_p99", "Fill Rate"],
    [
        ["AS", "L", "0.330", "22.26", "0.404"],
        ["AS", "M", "0.197", "27.49", "0.438"],
        ["AS", "H", "0.038", "27.66", "0.497"],
        ["naive", "L", "0.269", "16.35", "0.105"],
        ["naive", "M", "0.150", "18.87", "0.115"],
        ["naive", "H", "0.143", "19.20", "0.143"],
        ["PPO-aware", "L", "0.927", "1.85", "0.226"],
        ["PPO-aware", "M", "0.799", "1.60", "0.225"],
        ["PPO-aware", "H", "0.464", "1.79", "0.250"],
        ["PPO-blind", "L", "0.924", "2.00", "0.220"],
        ["PPO-blind", "M", "0.814", "1.88", "0.229"],
        ["PPO-blind", "H", "0.526", "1.91", "0.252"],
    ],
)

add_normal(dst, (
    "Not: Bu tablo gerçek rejim etiketleri (regime_true) üzerinden ex-post hesaplanmıştır; ajan "
    "karar anında yalnızca tahmini rejimi gözlemlemektedir."
))

# 3d
add_heading(dst, "3d. Detector Robustness (3 Dedektör × 20 Tohum)", level=2)
add_table(dst,
    ["Dedektör", "Doğruluk", "ppo_aware Sharpe", "ppo_blind Sharpe", "p-değeri (Sharpe)"],
    [
        ["rv_baseline", "%60.7", "0.718", "0.753", "0.114"],
        ["rv_dwell", "%60.4", "0.716", "0.753", "0.110"],
        ["HMM", "%81.8", "0.719", "0.753", "0.082"],
        ["ANOVA", "—", "—", "—", "F=0.003, p=0.997"],
    ],
)

add_normal(dst, (
    "%81.8 doğruluklu HMM dedektörü bile null result'u değiştirememektedir. Bu, sorunun dedektör "
    "kalitesinde değil, σ̂'nin zaten rejim bilgisini taşımasında yattığını kanıtlar."
))

# ================================================================
# 4. İSTATİSTİKSEL TESTLER
# ================================================================
add_heading(dst, "4. İstatistiksel Testler", level=1)

add_normal(dst, "Eşleştirilmiş t-testi sonuçları (n=20 seed, Bonferroni düzeltmesi: α = 0.01/10 = 0.001):")
add_table(dst,
    ["Karşılaştırma", "Metrik", "t", "p-değeri", "Cohen's d", "Anlamlı?"],
    [
        ["sigma_only vs combined", "Sharpe", "2.014", "0.058", "0.450", "Hayır"],
        ["sigma_only vs combined", "Equity", "3.334", "0.003", "0.746", "Hayır"],
        ["sigma_only vs oracle_full", "Sharpe", "1.651", "0.115", "0.369", "Hayır"],
        ["sigma_only vs oracle_full", "Equity", "3.789", "0.001", "0.847", "Hayır"],
        ["oracle_full vs combined", "Sharpe", "1.062", "0.301", "0.238", "Hayır"],
        ["oracle_full vs combined", "Equity", "0.939", "0.360", "0.210", "Hayır"],
        ["sigma_only vs regime_only", "Sharpe", "2.283", "0.034", "0.510", "Hayır"],
        ["sigma_only vs regime_only", "Equity", "4.510", "<0.001", "1.008", "EVET"],
        ["combined vs AS", "Sharpe", "16.615", "<0.001", "3.715", "EVET"],
        ["combined vs naive", "Sharpe", "17.878", "<0.001", "3.998", "EVET"],
    ],
)

add_normal(dst, (
    "Tablo iki kritik sonucu ortaya koymaktadır. Birincisi, PPO varyantları arasındaki farklılıklar "
    "(üst blok) Bonferroni düzeltmeli eşiği geçememektedir — null result istatistiksel olarak "
    "sağlamdır. İkincisi, PPO ile klasik stratejiler arasındaki fark son derece büyük ve anlamlıdır "
    "(Cohen's d > 3.7 — devasa etki büyüklüğü). σ̂ ile rejim etiketinin birbirini ikame ettiğini "
    "gösteren en güçlü kanıt oracle paradoksudur: gerçek rejim bilgisine sahip oracle_full ile "
    "tahmini rejim kullanan combined arasındaki fark istatistiksel olarak sıfırdan ayırt "
    "edilememektedir (p=0.301 Sharpe, p=0.360 equity)."
))

# ================================================================
# 5. TEMEL ÇIKARIMLAR
# ================================================================
add_heading(dst, "5. Temel Çıkarımlar", level=1)

add_bullet(dst, "PPO üstünlüğü: ",
    "Tüm PPO varyantları Sharpe açısından AS ve naif stratejiyi ~6-7× geçmektedir. "
    "Mekanizma: envanter disiplini (inv_p99 ≈ 2 vs 20-30 lot).")

add_bullet(dst, "Null result: ",
    "Rejim bilgisi eklemek anlamlı fayda sağlamamaktadır. Sharpe bazlı Wilcoxon p=0.261 "
    "(anlamlı değil); equity bazlı eşleştirilmiş t-test p=0.023 (ppo_blind lehine).")

add_bullet(dst, "Oracle paradoksu: ",
    "Mükemmel rejim bilgisi bile sigma_only'yi geçememektedir. Bu, rejim etiketinin gözlem "
    "uzayında zaten mevcut sigma_hat sinyalini yedekli biçimde sunduğunu kanıtlar.")

add_bullet(dst, "Dedektör bağımsızlığı: ",
    "Null result %60'tan %82'ye kadar üç farklı dedektörde tutarlıdır. ANOVA F=0.003, p=0.997.")

add_bullet(dst, "Sinyal yedekliliği tezi: ",
    "\"Rejim etiketleri, yalnızca rejim sinyali gözlem uzayında örtük olarak mevcut olmadığında "
    "faydalıdır.\"")

# ================================================================
# 6. GELECEK ÇALIŞMALAR
# ================================================================
add_heading(dst, "6. Gelecek Çalışmalar", level=1)

dst.add_paragraph(
    "Rejime koşullu ödül tasarımı: Yüksek volatilite döneminde envanter cezasını artıran adaptif η.",
    style='List Number',
)
dst.add_paragraph(
    "Model yanlış belirleme (misspecification) benchmark: Dolum yoğunluğu parametrelerinin (A, k) "
    "rejime bağlı kılınması — bu durumda sigma_hat'in taşıyamadığı ek bilgi oluşur ve rejim "
    "etiketinin faydalı olup olmadığı yeniden test edilebilir.",
    style='List Number',
)
dst.add_paragraph(
    "Yayın: İyi nitelendirilmiş null result ile kontrollü RL vs. klasik piyasa yapıcılığı karşılaştırması.",
    style='List Number',
)

# ================================================================
# KAYDET
# ================================================================
dst.save("manuscript/project_summary.docx")
print("project_summary.docx oluşturuldu.")
