"""decisions_log_5.docx -> decisions_log_6.docx:
Guncelleme 1: Karar #29 Bonferroni N=10 -> N=12, alpha=0.01/12=0.00083
Guncelleme 2: Karar #27 oracle paradox p=0.301 -> sigma_only vs oracle_full p=0.115
              (Karar #35 zaten dogru p=0.115 degerini iceriyor; asil hata #27'de)
Guncelleme 3: #35 Etki ve ozet satir #9: kanitlamistir -> desteklemektedir"""

from __future__ import annotations
import sys, io, shutil
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from pathlib import Path
from docx import Document
from docx2pdf import convert

SRC = Path("manuscript/decisions_log_5.docx")
DST = Path("manuscript/decisions_log_6.docx")
PDF = Path("manuscript/decisions_log_6.pdf")


def _replace_in_paragraph(para, old_text, new_text):
    """Replace old_text with new_text across all runs in a paragraph.
    Handles text split across multiple runs by joining, replacing,
    then putting result into first run and clearing the rest.
    Returns True if a replacement was made."""
    full = "".join(r.text for r in para.runs)
    if old_text not in full:
        return False
    new_full = full.replace(old_text, new_text)
    if para.runs:
        para.runs[0].text = new_full
        for r in para.runs[1:]:
            r.text = ""
    return True


def _replace_in_cell(cell, old_text, new_text):
    """Replace text in all paragraphs of a cell."""
    replaced = False
    for para in cell.paragraphs:
        if _replace_in_paragraph(para, old_text, new_text):
            replaced = True
    return replaced


def main():
    shutil.copy2(SRC, DST)
    doc = Document(DST)

    changes = []

    for tbl in doc.tables:
        hdr = tbl.rows[0].cells[1].text if len(tbl.rows[0].cells) > 1 else ""

        # ── GÜNCELLEME 1: Karar #29  Bonferroni N=10 → N=12 ──
        if "#29" in hdr:
            # R3 (Karar): B → A, 10 → 12, 0.001 → 0.00083
            if _replace_in_cell(
                tbl.rows[3].cells[1],
                "B \u2014 \u03b1=0.01/10=0.001 olarak d\u00fczeltildi.",
                "A \u2014 \u03b1=0.01/12=0.00083 olarak d\u00fczeltildi.",
            ):
                changes.append("#29 Karar: B/N=10/0.001 \u2192 A/N=12/0.00083")

            # R4 (Neden): 10 → 12 karşılaştırma
            if _replace_in_cell(
                tbl.rows[4].cells[1],
                "Tabloda fiilen 10 kar\u015f\u0131la\u015ft\u0131rma var; "
                "payda ger\u00e7ek test say\u0131s\u0131yla e\u015fle\u015fmeli.",
                "Tabloda fiilen 12 kar\u015f\u0131la\u015ft\u0131rma var "
                "(6 kar\u015f\u0131la\u015ft\u0131rma \u00d7 2 metrik); "
                "payda ger\u00e7ek test say\u0131s\u0131yla e\u015fle\u015fmeli.",
            ):
                changes.append("#29 Neden: 10 \u2192 12 kar\u015f\u0131la\u015ft\u0131rma")

        # ── GÜNCELLEME 3: Karar #35  "kanıtlamıştır" → "desteklemektedir" ──
        if "#35" in hdr:
            if _replace_in_cell(
                tbl.rows[5].cells[1],
                "kanıtlamıştır",
                "desteklemektedir",
            ):
                changes.append("#35 Etki: kanıtlamıştır → desteklemektedir")
            if _replace_in_cell(
                tbl.rows[5].cells[1],
                "Dedektör kalitesi argümanı kapatılmıştır. Oracle deneyi, "
                "herhangi bir dedektör doğruluğunun yeterli olacağını desteklemektedir.",
                "Dedektör kalitesi itirazı güçlü bir üst-sınır testiyle yanıtlanmıştır. "
                "Oracle deneyi, mükemmel rejim bilgisinin bile sigma_only'yi "
                "geçemediğini göstermektedir.",
            ):
                changes.append("#35 Etki: detector sufficiency overclaim softened")

        # ── GÜNCELLEME 2: Karar #27  oracle paradox p-değeri ──
        # Kullanıcı #35 dedi ancak #35 zaten doğru (p=0.115).
        # Asıl hata #27'de: p=0.301 oracle_full vs combined'a ait,
        # doğru karşılaştırma sigma_only vs oracle_full p=0.115.
        if "#27" in hdr:
            if _replace_in_cell(
                tbl.rows[3].cells[1],
                "paired t-test ile p=0.301 (Sharpe), p=0.360 (equity) "
                "bulundu \u2014 istatistiksel olarak anlaml\u0131 fark yok.",
                "paired t-test ile sigma_only vs oracle_full: "
                "p=0.115 (Sharpe) bulundu \u2014 istatistiksel olarak "
                "anlaml\u0131 fark yok.",
            ):
                changes.append(
                    "#27 Karar: p=0.301/0.360 \u2192 "
                    "sigma_only vs oracle_full p=0.115"
                )

    # ── GÜNCELLEME 3b: Özet tablosu satır #9 "kanıtlamak" → "desteklemek" ──
    sum_tbl = doc.tables[-1]
    if _replace_in_cell(
        sum_tbl.rows[9].cells[1],
        "kanıtlamak",
        "desteklemek",
    ):
        changes.append("Özet R9: kanıtlamak → desteklemek")

    doc.save(DST)
    print(f"decisions_log_6.docx saved ({DST.stat().st_size:,} bytes)")

    # ── PDF ──
    convert(str(DST), str(PDF))
    print(f"decisions_log_6.pdf saved ({PDF.stat().st_size:,} bytes)")

    # ── Doğrulama ──
    print(f"\n=== Changes ({len(changes)}) ===")
    for c in changes:
        print(f"  \u2713 {c}")

    doc2 = Document(DST)
    print("\n=== Verification ===")
    for tbl in doc2.tables:
        hdr = tbl.rows[0].cells[1].text if len(tbl.rows[0].cells) > 1 else ""
        if "#27" in hdr:
            print(f"#27 Karar: {tbl.rows[3].cells[1].text[:120]}")
        if "#29" in hdr:
            print(f"#29 Karar: {tbl.rows[3].cells[1].text[:120]}")
            print(f"#29 Neden: {tbl.rows[4].cells[1].text[:120]}")
        if "#35" in hdr:
            print(f"#35 Neden: {tbl.rows[4].cells[1].text[:120]}")
            print(f"#35 Etki:  {tbl.rows[5].cells[1].text[:120]}")
    sum2 = doc2.tables[-1]
    print(f"Özet R9:   {sum2.rows[9].cells[1].text[:120]}")


if __name__ == "__main__":
    main()
