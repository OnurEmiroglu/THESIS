"""decisions_log_8.docx -> decisions_log_9.docx
Appends four new decisions (#42 regime parameter retrospective, #43 canonical
sanity check, #44 audit drift correction, #45 misspec divergence) capturing
the WP5.5 canonical-realignment outcomes. Updates title-page marker to
"Sürüm 26" and summary from 21 karar -> 25 karar."""

from __future__ import annotations
import sys, io, shutil, copy
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement

SRC = Path("manuscript/decisions_log_8.docx")
DST = Path("manuscript/decisions_log_9.docx")

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _set_cell_text(cell, text, *, bold=None):
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].text = text
        for r in para.runs[1:]:
            r.text = ""
        if bold is not None:
            para.runs[0].bold = bold
    else:
        r = para.add_run(text)
        if bold is not None:
            r.bold = bold


def _build_karar_table_from_template(doc, template_tbl, number, title, wp, ikilem, secenekler, karar, neden, etki):
    new_xml = copy.deepcopy(template_tbl._tbl)
    from docx.table import Table
    new_tbl = Table(new_xml, doc.part)
    _set_cell_text(new_tbl.rows[0].cells[0], wp, bold=True)
    _set_cell_text(new_tbl.rows[0].cells[1], f"#{number}  {title}", bold=True)
    _set_cell_text(new_tbl.rows[1].cells[0], "İkilem", bold=True)
    _set_cell_text(new_tbl.rows[1].cells[1], ikilem)
    _set_cell_text(new_tbl.rows[2].cells[0], "Seçenekler", bold=True)
    _set_cell_text(new_tbl.rows[2].cells[1], secenekler)
    _set_cell_text(new_tbl.rows[3].cells[0], "Karar", bold=True)
    _set_cell_text(new_tbl.rows[3].cells[1], karar)
    _set_cell_text(new_tbl.rows[4].cells[0], "Neden", bold=True)
    _set_cell_text(new_tbl.rows[4].cells[1], neden)
    _set_cell_text(new_tbl.rows[5].cells[0], "Etki", bold=True)
    _set_cell_text(new_tbl.rows[5].cells[1], etki)
    return new_xml


def _replace_in_paragraph(para, old_text, new_text):
    full = "".join(r.text for r in para.runs)
    if old_text not in full:
        return False
    new_full = full.replace(old_text, new_text)
    if para.runs:
        para.runs[0].text = new_full
        for r in para.runs[1:]:
            r.text = ""
    return True


def main():
    shutil.copy2(SRC, DST)
    doc = Document(DST)

    template = doc.tables[-2]  # last individual Karar table (#41) in decisions_log_8

    k42 = _build_karar_table_from_template(
        doc, template,
        number=42,
        title="Rejim parametrelerinin retrospektif gözden geçirilmesi (sigma_mult, warmup)",
        wp="WP2/WP5",
        ikilem=(
            "Farklı iş paketlerinde kullanılan rejim parametreleri arasında tarihsel bir sapma tespit "
            "edildi: thesis §3.1 [0.6, 1.0, 1.8] / warmup=1000 olarak yazılı iken misspec-mild deneyi "
            "[0.5, 1.0, 2.0] / warmup=400 ile yürütülmüş. WP5.5 sinyal denetim konfigürasyonunda ise "
            "dt=1.0 / sigma_mid_ticks=0.5 drift'i yakalandı. Bu sapmaların ya geriye dönük "
            "homogenize edilmesi ya da açıkça belgelenmesi gerekiyordu."
        ),
        secenekler=(
            "A) Tüm geçmiş koşuları yeni (kanonik) parametrelerle yeniden eğit; büyük hesaplama "
            "maliyeti  |  B) Canonical operating point'i sabitle, yeni denetim/denemeler buna uysun, "
            "geçmiş deneyler olduğu gibi kalsın ve sapmalar tezde/log'da açıkça belirtilsin"
        ),
        karar=(
            "B — Kanonik nokta: dt=0.2, sigma_mid_ticks=0.8, sigma_mult=[0.6,1.0,1.8], warmup_steps=1000, "
            "rv_window=50, n_steps=8000, exec={A=5.0, k=1.5, fee_bps=0.2, latency_steps=1}, sticky "
            "3×3 trans_matrix. Misspec-mild varyantının [0.5,1.0,2.0]/warmup=400 ile yürütüldüğü tezde "
            "açıkça not edildi; geçmiş koşular yeniden üretilmedi."
        ),
        neden=(
            "Yeniden eğitim maliyeti (5 varyant × 20 tohum × 1M timestep) ana bulguya ek bilgi "
            "sağlamadan zaman ve enerji harcar. Kanonik noktanın yazılı hâle getirilmesi, ileride her "
            "yeni deneyin aynı referansı paylaşmasını garanti ediyor. Şeffaflık > retrospektif temizlik."
        ),
        etki=(
            "Yeni eklenen denemeler (WP5.5 sinyal denetimi dâhil) kanonik noktayı kullanıyor. Misspec "
            "sonuçlarının bu farkla birlikte yorumlanması için tezin §4.8'ine uyarı paragrafı eklendi."
        ),
    )

    k43 = _build_karar_table_from_template(
        doc, template,
        number=43,
        title="WP5.5 öncesi canonical operating point sanity-check zorunlu kılındı",
        wp="WP5.5",
        ikilem=(
            "WP5.5 (sinyal manipülasyon denetimi) doğrudan mevcut konfigürasyon dosyaları üzerinden "
            "koşturulacaktı. Ancak config dosyaları zaman içinde farklı iş paketleri için "
            "özelleştirilmiş; kanonik referansla birebir örtüşme kontrol edilmemişti. Sonuçların anlamlı "
            "kabul edilebilmesi için kullanılan parametrelerin tez metninde açıklanan ortamla uyumlu "
            "olması gerekliydi."
        ),
        secenekler=(
            "A) Config'i olduğu gibi kullan, sonuçları yorumlarken uyar  |  B) Önce yalnızca-okuma "
            "bir sanity-check yaparak kanonik noktayı resmen belirle, sapmaları raporla, sonra WP5.5 "
            "konfigürasyonunu buna göre hizala"
        ),
        karar=(
            "B — WP5.5 audit koşulmadan önce tez metni, config dosyaları ve önceki koşu snapshot'ları "
            "taranarak kanonik operating point resmen tespit edildi (Karar #42). İki mevcut sapma "
            "(WP5.5 audit config, misspec-mild config) açıkça dokümante edildi. WP5.5 koşusu ancak "
            "bundan sonra yürütüldü."
        ),
        neden=(
            "Sanity-check, sonuçların yanlış baz çizgisine göre raporlanıp geri çekilmesinden çok daha "
            "ucuz. Kanonik noktanın yazılı hâle getirilmesi sonraki tüm çalışmalar için tek referans "
            "noktası sağlıyor."
        ),
        etki=(
            "WP5.5 audit'i kanonik parametrelerle yeniden koşuldu; önceki drift'li iki koşu "
            "(20260419-214718, 20260419-221930) arşiv notuyla aynen saklandı. Karar izlenebilirliği "
            "korundu."
        ),
    )

    k44 = _build_karar_table_from_template(
        doc, template,
        number=44,
        title="WP5.5 audit drift düzeltildi; 'none' rejiminde yapısal metrikler ungated",
        wp="WP5.5",
        ikilem=(
            "İlk WP5.5 koşusu dt=1.0 ve sigma_mid_ticks=0.5 ile yürütülmüş; clean classification_accuracy ≈ "
            "0.40 ve L-ağırlıklı rejim dağılımı üretmiş, direction/monotonicity gate'lerini yanıltıcı şekilde "
            "FAIL göstermişti. Ek olarak 'none' (sabit fill_value=0.0) koşulunda classification_accuracy ve "
            "threshold_overlap metrikleri, eşik kalibrasyonu ile fill değeri arasındaki ilişkiyi yansıtıyor "
            "— bilgi içeriğine dair sinyal taşımıyor."
        ),
        secenekler=(
            "A) Drift'li koşuyu gerçek sonuç olarak raporla  |  B) Kanonik noktayla yeniden koştur, "
            "'none' koşulunda yapısal olarak tanımsız metrikleri PASS/FAIL değerlendirmesinden çıkar "
            "(reported-but-not-gated) ve bu politikayı audit özetinde belgele"
        ),
        karar=(
            "B — config/w55_audit.json kanonik noktaya hizalandı (dt=0.2, sigma_mid_ticks=0.8, "
            "sigma_mult=[0.6,1.0,1.8], warmup=1000, trans_matrix tam, exec bloğu açıkça yazılı); "
            "noise_std için 'auto' modu eklendi (0.5 × clean_sigma_std). Audit job'ı, 'none' koşulunda "
            "classification_accuracy ve threshold_overlap metriklerini raporlamaya devam ediyor ama "
            "direction/monotonicity gate'lerinde dışlıyor."
        ),
        neden=(
            "Sabit input altında bu iki metrik bilgi içeriğini değil fill/threshold ilişkisini ölçer; "
            "gating mantıken yanlış. Yeni koşu (20260422-170037): direction PASS, separation PASS, "
            "coarsen_safety PASS, monotonicity FAIL — genel öneri REVIEW. Monotonicity FAIL, ölçüt "
            "tasarımına dair bir gözlem olarak ayrıca tartışılıyor; üç degradasyon (noisy/lagged/"
            "coarsened) tek bir şiddet eksenine yerleştirilemiyor."
        ),
        etki=(
            "Tez §4'e 'Manipülasyon Geçerlilik Denetimi' alt bölümü eklendi; eski iki drift'li koşu "
            "'results/runs/wp55_audit_archive_note.md' ile kalıcı şekilde süpersed edildi. Audit "
            "reproducibility için commit'e alındı; yeni baseline açık."
        ),
    )

    k45 = _build_karar_table_from_template(
        doc, template,
        number=45,
        title="Misspec-mild ortam farkı tezde açık uyarı paragrafı olarak not edildi",
        wp="WP5",
        ikilem=(
            "Misspec-mild koşuları [0.5, 1.0, 2.0] / warmup=400 ile yürütülmüştü; §3.1'de tanımlanan "
            "kanonik [0.6, 1.0, 1.8] / warmup=1000 ortamından farklıydı. Aynı 20-tohumluk koşuyu "
            "kanonik parametrelerle yeniden üretmek ≈4-5 saat GPU maliyetiydi. Ya bu yeniden üretim "
            "yapılmalı ya da fark tez metninde açıkça belgelenmeliydi."
        ),
        secenekler=(
            "A) Kanonik parametrelerle 100 modeli baştan eğit, eski sonuçları değiştir  |  B) Mevcut "
            "misspec sonuçlarını koru, §4.8'e tasarım-farkı uyarı paragrafı ekle; sonucu 'kanonik "
            "ortamda yeniden üretilmemiştir' notuyla sun"
        ),
        karar=(
            "B — Misspec-mild tablosu ve yorumu olduğu gibi kalıyor; §4.8 Tasarım paragrafının başına "
            "ortam sapmasını açıkça belirten bir uyarı eklendi. Gelecek çalışmalar listesine kanonik "
            "parametrelerle yeniden üretim maddesi yazılı olarak bulunuyor."
        ),
        neden=(
            "Ana null bulgusu ([sigma_only ≈ oracle_full, p=0.881]) misspec-mild altında bile "
            "güçleniyor; ortam farkına rağmen argümanın yönü değişmiyor. Yeniden eğitim maliyeti, "
            "bulgunun yeniden üretilebilirlik gücüne ek bilgi sağlamıyor. Şeffaflık > saf temizlik."
        ),
        etki=(
            "§4.8 Tasarım paragrafı, okuyucunun sonucu doğru bağlamda yorumlamasını sağlayacak şekilde "
            "kanonik ortam farkını açıkça not ediyor. Log #42'deki kanonik nokta kararıyla bu karar "
            "çapraz referans veriyor."
        ),
    )

    body = doc.element.body
    ozet_heading = None
    for child in body.iterchildren():
        if child.tag == W_NS + "p":
            pStyle = child.find(W_NS + "pPr/" + W_NS + "pStyle")
            if pStyle is not None and pStyle.get(W_NS + "val") == "Heading1":
                text = "".join(t.text or "" for t in child.iter(W_NS + "t"))
                if "Özet" in text:
                    ozet_heading = child
                    break
    if ozet_heading is None:
        raise RuntimeError("Özet heading not found")

    def _blank_p():
        return OxmlElement("w:p")

    for tbl in (k42, k43, k44, k45):
        ozet_heading.addprevious(_blank_p())
        ozet_heading.addprevious(tbl)

    # Summary heading: 21 -> 25
    for child in body.iterchildren():
        if child.tag == W_NS + "p":
            text = "".join(t.text or "" for t in child.iter(W_NS + "t"))
            if "En Kritik 21 Karar" in text:
                from docx.text.paragraph import Paragraph
                p_obj = Paragraph(child, doc.part)
                _replace_in_paragraph(p_obj, "En Kritik 21 Karar", "En Kritik 25 Karar")

    for para in doc.paragraphs:
        if "ChatGPT ve Claude" in para.text and "21 karar" in para.text:
            _replace_in_paragraph(para, "21 karar", "25 karar")

    # Title-page: Sürüm 25 -> Sürüm 26
    for para in doc.paragraphs:
        if "Sürüm 25" in para.text:
            _replace_in_paragraph(para, "Sürüm 25", "Sürüm 26")

    # Summary rows #22-25
    sum_tbl = doc.tables[-1]
    template_row = sum_tbl.rows[1]._tr

    def _add_summary_row(num, text, wp):
        new_row = copy.deepcopy(template_row)
        sum_tbl._tbl.append(new_row)
        from docx.table import _Row
        row_obj = _Row(new_row, sum_tbl)
        _set_cell_text(row_obj.cells[0], str(num))
        _set_cell_text(row_obj.cells[1], text)
        _set_cell_text(row_obj.cells[2], wp)

    _add_summary_row(
        22,
        "Rejim parametre retrospektifi — kanonik operating point yazılı hâle getirildi; "
        "misspec-mild sapması belgelendi",
        "WP2/WP5",
    )
    _add_summary_row(
        23,
        "WP5.5 öncesi canonical operating point sanity-check zorunlu — sapmalar WP5.5 öncesi "
        "raporlandı",
        "WP5.5",
    )
    _add_summary_row(
        24,
        "WP5.5 audit drift düzeltildi ve 'none' rejiminde yapısal metrikler ungated — "
        "reported-but-not-gated politikası",
        "WP5.5",
    )
    _add_summary_row(
        25,
        "Misspec-mild ortam farkı §4.8'de açık uyarı paragrafıyla not edildi — yeniden üretim "
        "yapılmadı, şeffaflık tercih edildi",
        "WP5",
    )

    doc.save(DST)
    print(f"decisions_log_9.docx saved ({DST.stat().st_size:,} bytes)")

    doc2 = Document(DST)
    print(f"\n=== Verification ===")
    print(f"Total tables: {len(doc2.tables)}")
    for idx in (-5, -4, -3, -2):
        t = doc2.tables[idx]
        print(f"Karar table [{idx}] header: {t.rows[0].cells[1].text[:80]}")
    sum2 = doc2.tables[-1]
    print(f"Summary rows: {len(sum2.rows)} (should be 26 = header + 25)")
    print(f"  Last row: {[c.text[:60] for c in sum2.rows[-1].cells]}")
    for p in doc2.paragraphs[:20]:
        if "Sürüm" in p.text or "En Kritik" in p.text:
            print(f"Marker: {p.text}")


if __name__ == "__main__":
    main()
