"""decisions_log_2.docx -> decisions_log_3.docx: İki düzeltme uygula."""

from __future__ import annotations
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from lxml import etree

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
SRC = Path("manuscript/decisions_log_2.docx")
DST = Path("manuscript/decisions_log_3.docx")


def _styled_run(paragraph, text, bold=False, size_pt=10, font_name="Arial", color_hex=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = font_name
    if color_hex:
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor.from_string(color_hex)
    return run


def _set_cell_props(cell, width_dxa, fill_color):
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = etree.SubElement(tc, qn("w:tcPr"))
        tc.insert(0, tcPr)
    for old in tcPr.findall(qn("w:tcW")):
        tcPr.remove(old)
    tcW = etree.SubElement(tcPr, qn("w:tcW"))
    tcW.set(qn("w:w"), width_dxa)
    tcW.set(qn("w:type"), "dxa")
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    borders = etree.SubElement(tcPr, qn("w:tcBorders"))
    for side in ("top", "left", "bottom", "right"):
        b = etree.SubElement(borders, qn(f"w:{side}"))
        b.set(qn("w:val"), "single")
        b.set(qn("w:color"), "BBCFE0")
        b.set(qn("w:sz"), "1")
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = etree.SubElement(tcPr, qn("w:shd"))
    shd.set(qn("w:fill"), fill_color)
    shd.set(qn("w:val"), "clear")
    for old in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(old)
    tcMar = etree.SubElement(tcPr, qn("w:tcMar"))
    for side, val in [("top", "80"), ("left", "160"), ("bottom", "80"), ("right", "160")]:
        m = etree.SubElement(tcMar, qn(f"w:{side}"))
        m.set(qn("w:type"), "dxa")
        m.set(qn("w:w"), val)


def main():
    shutil.copy2(SRC, DST)
    doc = Document(DST)
    body = doc.element.body

    # ── DÜZELTME 1: Özet tablosu ──
    sum_tbl = doc.tables[-1]

    # (a) #9 metnini güncelle
    row9 = sum_tbl.rows[9]
    c1 = row9.cells[1]
    c1.paragraphs[0].clear()
    _styled_run(
        c1.paragraphs[0],
        "Detector robustness full run (120 model) — signal redundancy argümanını kanıtlamak ve null result'ı istatistiksel olarak teyit etmek",
        bold=False, size_pt=10, color_hex="2C2C2C",
    )
    _set_cell_props(c1, "6960", "E8F4FD")

    # (b) #11 satırını sil (row index 11)
    row11_elem = sum_tbl.rows[11]._tr
    sum_tbl._tbl.remove(row11_elem)

    # (c) Eski #12'yi #11 yap
    # Now row index 11 is the old #12
    new_row11 = sum_tbl.rows[11]
    c0 = new_row11.cells[0]
    c0.paragraphs[0].clear()
    _styled_run(c0.paragraphs[0], "11", bold=True, size_pt=10, color_hex="1F4E79")
    _set_cell_props(c0, "600", "E8F4FD")

    # ── DÜZELTME 2: "Genel ... Tez Yazım" bölümüne not ekle ──
    genel_heading = None
    for p in doc.paragraphs:
        if "Genel Proje Kararları — Tez Yazım Süreci" in p.text:
            genel_heading = p
            break
    assert genel_heading is not None, "Heading not found"

    genel_idx = list(body).index(genel_heading._element)
    # body[genel_idx+1] is blank paragraph, body[genel_idx+2] is first table (#29)
    # Insert note paragraph after the blank paragraph (before first table)
    insert_before = list(body)[genel_idx + 2]

    # Build note paragraph: italic, 11pt, Arial, color dark gray
    from docx.oxml import OxmlElement
    p_elem = OxmlElement("w:p")
    r = etree.SubElement(p_elem, qn("w:r"))
    rPr = etree.SubElement(r, qn("w:rPr"))
    rFonts = etree.SubElement(rPr, qn("w:rFonts"))
    for attr in ("ascii", "cs", "eastAsia", "hAnsi"):
        rFonts.set(qn(f"w:{attr}"), "Arial")
    etree.SubElement(rPr, qn("w:i"))
    etree.SubElement(rPr, qn("w:iCs"))
    sz = etree.SubElement(rPr, qn("w:sz"))
    sz.set(qn("w:val"), "22")  # 11pt
    szCs = etree.SubElement(rPr, qn("w:szCs"))
    szCs.set(qn("w:val"), "22")
    color = etree.SubElement(rPr, qn("w:color"))
    color.set(qn("w:val"), "555555")
    t = etree.SubElement(r, qn("w:t"))
    t.text = "Not: Aşağıdaki kararlar thesis_16 yazım sürecine aittir."

    body.insert(list(body).index(insert_before), p_elem)

    # Also add a blank spacer after the note
    spacer = OxmlElement("w:p")
    body.insert(list(body).index(insert_before), spacer)

    # ── Update summary heading: 12 -> 11 ──
    for p in doc.paragraphs:
        if "Kritik 12 Karar" in p.text:
            for run in p.runs:
                if "12" in run.text:
                    run.text = run.text.replace("En Kritik 12 Karar", "En Kritik 11 Karar")

    doc.save(DST)
    print(f"decisions_log_3.docx saved ({DST.stat().st_size:,} bytes)")

    # Verify
    doc2 = Document(DST)
    sum2 = doc2.tables[-1]
    print(f"\n=== SUMMARY TABLE ({len(sum2.rows)} rows) ===")
    for ri, row in enumerate(sum2.rows):
        cells = [c.text.strip() for c in row.cells]
        print(f"  row[{ri:2d}]: {cells}")

    print()
    for p in doc2.paragraphs:
        if "Kritik" in p.text:
            print(f"Heading: '{p.text}'")
        if "thesis_16" in p.text:
            print(f"Note found: '{p.text}'")


if __name__ == "__main__":
    main()
