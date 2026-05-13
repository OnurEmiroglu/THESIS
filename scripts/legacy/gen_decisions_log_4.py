"""decisions_log_3.docx -> decisions_log_4.docx:
Karar #32 (rejime koşullu envanter cezası) ekle, özet tabloyu güncelle."""

from __future__ import annotations
import sys, io, shutil
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from lxml import etree

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
SRC = Path("manuscript/decisions_log_3.docx")
DST = Path("manuscript/decisions_log_4.docx")


def _styled_run(paragraph, text, bold=False, size_pt=10, font_name="Arial", color_hex=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = font_name
    if color_hex:
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor.from_string(color_hex)
    return run


def _set_cell_props(cell, width_dxa, fill_color):
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = etree.SubElement(tc, qn("w:tcPr"))
        tc.insert(0, tcPr)
    for old in tcPr.findall(qn("w:tcW")):
        tcPr.remove(old)
    tcW = etree.SubElement(tcPr, qn("w:tcW"))
    tcW.set(qn("w:w"), width_dxa)
    tcW.set(qn("w:type"), "dxa")
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    borders = etree.SubElement(tcPr, qn("w:tcBorders"))
    for side in ("top", "left", "bottom", "right"):
        b = etree.SubElement(borders, qn(f"w:{side}"))
        b.set(qn("w:val"), "single")
        b.set(qn("w:color"), "BBCFE0")
        b.set(qn("w:sz"), "1")
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = etree.SubElement(tcPr, qn("w:shd"))
    shd.set(qn("w:fill"), fill_color)
    shd.set(qn("w:val"), "clear")
    for old in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(old)
    tcMar = etree.SubElement(tcPr, qn("w:tcMar"))
    for side, val in [("top", "80"), ("left", "160"), ("bottom", "80"), ("right", "160")]:
        m = etree.SubElement(tcMar, qn(f"w:{side}"))
        m.set(qn("w:type"), "dxa")
        m.set(qn("w:w"), val)


def _add_decision_table(doc, body, insert_before_elem, wp_tag, number, title, rows_data):
    """Insert a 2-column, 6-row decision table matching the style of #26-#31."""
    from docx.oxml import OxmlElement

    # Add a blank spacer paragraph before the table
    spacer = OxmlElement("w:p")
    body.insert(list(body).index(insert_before_elem), spacer)

    # Create table element
    tbl = doc.add_table(rows=6, cols=2, style='Normal Table')

    # Row 0: [WP tag, #N title]
    labels = [wp_tag, f"#{number}  {title}"]
    row_labels = ["", "İkilem", "Seçenekler", "Karar", "Neden", "Etki"]
    all_data = [labels] + [[row_labels[i], rows_data[i-1]] for i in range(1, 6)]

    for ri, (c0_text, c1_text) in enumerate(all_data):
        c0 = tbl.rows[ri].cells[0]
        c1 = tbl.rows[ri].cells[1]

        c0.paragraphs[0].clear()
        c1.paragraphs[0].clear()

        is_header = (ri == 0)
        fill_0 = "1F4E79" if is_header else "E8F4FD"
        fill_1 = "1F4E79" if is_header else "FFFFFF"
        color_0 = "FFFFFF" if is_header else "1F4E79"
        color_1 = "FFFFFF" if is_header else "2C2C2C"

        _styled_run(c0.paragraphs[0], c0_text,
                     bold=True, size_pt=10, color_hex=color_0)
        _styled_run(c1.paragraphs[0], c1_text,
                     bold=is_header, size_pt=10 if not is_header else 11,
                     color_hex=color_1)

        _set_cell_props(c0, "1400", fill_0)
        _set_cell_props(c1, "6960", fill_1)

    # Move table from end of document to before insert_before_elem
    tbl_elem = tbl._tbl
    body.remove(tbl_elem)
    body.insert(list(body).index(insert_before_elem), tbl_elem)

    return tbl


def main():
    shutil.copy2(SRC, DST)
    doc = Document(DST)
    body = doc.element.body

    # ── DEĞİŞİKLİK 1: Karar #32 ekle ──
    # Find the summary heading to insert before it
    summary_heading = None
    for p in doc.paragraphs:
        if "Kritik 11 Karar" in p.text:
            summary_heading = p
            break
    assert summary_heading is not None, "Summary heading not found"

    _add_decision_table(
        doc, body,
        insert_before_elem=summary_heading._element,
        wp_tag="WP5",
        number=32,
        title="Rejime koşullu envanter cezası: deneyelim mi?",
        rows_data=[
            (
                "Danışman rejime koşullu ödül tasarımını önerdi. ηH = 5×ηL ile "
                "ppo_combined, sigma_only'yi geçer mi?"
            ),
            (
                "A) Sabit η=0.001 ile devam et  |  B) ηL=0.0005, ηM=0.001, "
                "ηH=0.0025 konfigürasyonu ile 20 seed full run"
            ),
            (
                "B — Full run tamamlandı. Config: w5_eta_regime.json."
            ),
            (
                "Danışmanın dört önerisinden sonuncusuydu. Pilot (3 seed) null "
                "result korudu. Full run ile istatistiksel teyit gerekiyordu."
            ),
            (
                "ppo_combined vs ppo_sigma_only — Sharpe p=0.0016, Equity p=0.008, "
                "her iki metrikte sigma_only lehine anlamlı fark. Rejime koşullu η "
                "performansı artırmak yerine düşürdü. Signal redundancy argümanı ödül "
                "tasarımı boyutunda da teyit edildi. thesis_17'ye Section 4.7 olarak eklendi."
            ),
        ],
    )

    # ── DEĞİŞİKLİK 2: Özet tablosu güncelle ──
    # Update heading: 11 -> 12
    for p in doc.paragraphs:
        if "Kritik 11 Karar" in p.text:
            for run in p.runs:
                if "11" in run.text:
                    run.text = run.text.replace("En Kritik 11 Karar", "En Kritik 12 Karar")

    # Add row #12 to summary table (last table)
    sum_tbl = doc.tables[-1]
    new_row = sum_tbl.add_row()
    cells = new_row.cells

    cells[0].paragraphs[0].clear()
    _styled_run(cells[0].paragraphs[0], "12", bold=True, size_pt=10, color_hex="1F4E79")
    _set_cell_props(cells[0], "600", "E8F4FD")

    cells[1].paragraphs[0].clear()
    _styled_run(
        cells[1].paragraphs[0],
        "Rejime koşullu envanter cezası deneyi — signal redundancy'yi ödül tasarımı boyutunda da teyit etmek (p=0.0016)",
        bold=False, size_pt=10, color_hex="2C2C2C",
    )
    _set_cell_props(cells[1], "6960", "E8F4FD")

    cells[2].paragraphs[0].clear()
    _styled_run(cells[2].paragraphs[0], "WP5", bold=False, size_pt=10, color_hex="2C2C2C")
    _set_cell_props(cells[2], "1000", "E8F4FD")

    doc.save(DST)
    print(f"decisions_log_4.docx saved ({DST.stat().st_size:,} bytes)")

    # Verify
    doc2 = Document(DST)
    sum2 = doc2.tables[-1]
    print(f"\n=== SUMMARY TABLE ({len(sum2.rows)} rows) ===")
    for ri, row in enumerate(sum2.rows):
        cells = [c.text.strip()[:70] for c in row.cells]
        print(f"  row[{ri:2d}]: {cells}")

    print()
    for p in doc2.paragraphs:
        if "Kritik" in p.text:
            print(f"Heading: '{p.text}'")

    # Verify #32 exists
    for ti, tbl in enumerate(doc2.tables):
        r0 = tbl.rows[0].cells[1].text.strip()[:50] if len(tbl.rows[0].cells) > 1 else ""
        if "#32" in r0:
            print(f"\n#32 table found at table[{ti}]")
            for ri, row in enumerate(tbl.rows):
                c0 = row.cells[0].text.strip()[:15]
                c1 = row.cells[1].text.strip()[:70]
                print(f"  row[{ri}]: [{c0}] {c1}")


if __name__ == "__main__":
    main()
