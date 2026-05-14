"""decisions_log_11.docx -> decisions_log_12.docx

Appends Decisions #48-#51 documenting the audit-remediation sequence:
Lane A audit response, Lane B engineering fixes, Lane C remediation, and
procedural lessons.

Updates summary heading from "En Kritik 27 Karar" -> "En Kritik 31 Karar".
Updates the title-page marker from "Sürüm 27" -> "Sürüm 28" for the new
decision log only. decisions_log_11 remains unchanged.

Decision tables use the same v2 sub-section structure as #47:
  Karar / Seçenekler / Neden / Etki / Not  (no İkilem; trailing Not row)
The 6-row table template is reused; row labels are overwritten.
"""

from __future__ import annotations

import copy
import io
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

SRC = Path("manuscript/decisions_log_11.docx")
DST = Path("manuscript/decisions_log_12.docx")

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _set_cell_text(cell, text, *, bold=None):
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].text = text
        for r in para.runs[1:]:
            r.text = ""
        if bold is not None:
            para.runs[0].bold = bold
    else:
        r = para.add_run(text)
        if bold is not None:
            r.bold = bold


def _build_karar_table_v2(
    doc,
    template_tbl,
    number,
    title,
    wp,
    karar,
    secenekler,
    neden,
    etki,
    note,
):
    """Build a Karar table with Karar / Seçenekler / Neden / Etki / Not rows."""
    new_xml = copy.deepcopy(template_tbl._tbl)
    from docx.table import Table

    new_tbl = Table(new_xml, doc.part)
    _set_cell_text(new_tbl.rows[0].cells[0], wp, bold=True)
    _set_cell_text(new_tbl.rows[0].cells[1], f"#{number}  {title}", bold=True)
    _set_cell_text(new_tbl.rows[1].cells[0], "Karar", bold=True)
    _set_cell_text(new_tbl.rows[1].cells[1], karar)
    _set_cell_text(new_tbl.rows[2].cells[0], "Seçenekler", bold=True)
    _set_cell_text(new_tbl.rows[2].cells[1], secenekler)
    _set_cell_text(new_tbl.rows[3].cells[0], "Neden", bold=True)
    _set_cell_text(new_tbl.rows[3].cells[1], neden)
    _set_cell_text(new_tbl.rows[4].cells[0], "Etki", bold=True)
    _set_cell_text(new_tbl.rows[4].cells[1], etki)
    _set_cell_text(new_tbl.rows[5].cells[0], "Not", bold=True)
    _set_cell_text(new_tbl.rows[5].cells[1], note)
    return new_xml


def _replace_in_paragraph(para, old_text, new_text):
    full = "".join(r.text for r in para.runs)
    if old_text not in full:
        return False
    new_full = full.replace(old_text, new_text)
    if para.runs:
        para.runs[0].text = new_full
        for r in para.runs[1:]:
            r.text = ""
    return True


def _find_ozet_heading(doc):
    for child in doc.element.body.iterchildren():
        if child.tag != W_NS + "p":
            continue
        p_style = child.find(W_NS + "pPr/" + W_NS + "pStyle")
        if p_style is None or p_style.get(W_NS + "val") != "Heading1":
            continue
        text = "".join(t.text or "" for t in child.iter(W_NS + "t"))
        if "Özet" in text:
            return child
    raise RuntimeError("Özet heading not found")


def _blank_p():
    return OxmlElement("w:p")


DECISIONS = [
    {
        "number": 48,
        "title": "Lane A audit response scope",
        "wp": "Audit",
        "karar": (
            "Lane A bulguları thesis_28 içinde metodolojik caveat olarak "
            "yanıtlandı: C1 dwell filter look-ahead ana sonuç boru hattını "
            "etkilemeyen yardımcı detector comparison kapsamına alındı; C2 WP4 "
            "same-series train/eval yalnızca pilot/infrastructure olarak "
            "çerçevelendi; C3 WP6 noisy sigma calibration küçük noisy-only "
            "ölçek sorunu olarak caveat edildi."
        ),
        "secenekler": (
            "(a) Tüm Lane A bulguları için deneyleri yeniden koşturmak  |  "
            "(b) Boru hattı kapsamını ayırıp ana numerik iddiaları WP5 70/30 "
            "OOS ve WP6 sweep çıktıları üzerinden korumak  |  "
            "(c) Bulguları görmezden gelmek"
        ),
        "neden": (
            "C1 için main WP4/WP5/WP6 pipeline'ları causal rv_baseline kullanıyor; "
            "dwell filter offline auxiliary robustness comparison ile sınırlı. "
            "C2 için raporlanan PPO performans sayıları WP4'ten değil WP5'in "
            "70/30 OOS split sonuçlarından geliyor; WP4 eğitim/infrastructure "
            "pilot statüsünde. C3 için WP6 noisy sigma calibration etkisi "
            "noisy-only küçük ölçek kayması olarak sınıflandı; directional "
            "bulgular değişmediği için rerun yapılmadı."
        ),
        "etki": (
            "Thesis_28 sayısal iddiaları korunurken metodolojik sınırlar açık "
            "hale getirildi. Ana sonuçlar yeniden üretilmedi; caveat yaklaşımı "
            "kanıt hiyerarşisini netleştirdi: auxiliary detector analizi ayrı, "
            "reported PPO OOS sonuçları ayrı, noisy calibration sınırlaması ayrı."
        ),
        "note": (
            "Prosedürel ders: audit bulgusunun hangi evidence layer'a ait olduğu "
            "belirlenmeden apply yapılmadı. Bu ayrım yanlış audit mapping ve "
            "gereksiz rerun riskini düşürdü."
        ),
    },
    {
        "number": 49,
        "title": "Lane B engineering guardrails",
        "wp": "Infrastructure",
        "karar": (
            "Lane B kapsamında üç mühendislik güvence düzeltmesi kabul edildi: "
            "CSVMetricLogger schema consistency enforcement, resume sırasında "
            "config snapshot validation, ve CSV/TXT audit stability için "
            ".gitattributes + LF materialization."
        ),
        "secenekler": (
            "(a) Logger ve resume davranışını olduğu gibi bırakmak  |  "
            "(b) Sadece dokümantasyonla riski not etmek  |  "
            "(c) Küçük, local guardrail düzeltmeleriyle sessiz veri/konfigürasyon "
            "drift riskini azaltmak"
        ),
        "neden": (
            "CSVMetricLogger daha önce sessiz kolon uyumsuzluğu üretebilecek "
            "bir append pattern'ine açıktı; 26beecd ile explicit fieldnames "
            "schema check eklendi. Resume path'i yanlış config ile devam etme "
            "riskini taşıyordu; a63e640 ile snapshot mismatch default olarak "
            "reddedildi. Cross-platform hash/audit stabilitesi için ffe8d90 ile "
            ".gitattributes eklendi ve CSV/TXT line-ending davranışı sabitlendi."
        ),
        "etki": (
            "Bu değişiklikler mevcut numerik evidence artifact'larını yeniden "
            "üretmedi; gelecekteki koşuların izlenebilirliğini ve audit "
            "tekrarlanabilirliğini güçlendirdi. Hata yüzeyi runtime veya I/O "
            "guardrail seviyesinde daraltıldı."
        ),
        "note": (
            "Commit zinciri: 26beecd (B8 CSVMetricLogger), a63e640 (B9 resume "
            "validation), ffe8d90 (.gitattributes/LF audit stability)."
        ),
    },
    {
        "number": 50,
        "title": "Lane C active-code remediation",
        "wp": "Codebase",
        "karar": (
            "Lane C kapsamındaki aktif kod ve provenance düzeltmeleri ayrı "
            "küçük commit'lere bölündü: ANOVA framing netleştirildi; thesis "
            "figure ownership header'ları ayrıştırıldı; tarihsel generator'lar "
            "legacy alanına taşındı; aktif scriptlerde hardcoded run path'ler "
            "argparse default'larıyla parametrize edildi; WP2 per-run provenance "
            "artifact adı standardize edildi."
        ),
        "secenekler": (
            "(a) Tek büyük cleanup commit'i yapmak  |  "
            "(b) Sadece doküman notu eklemek  |  "
            "(c) Davranış yüzeyini koruyan, audit edilebilir küçük remediation "
            "commit'leri yapmak"
        ),
        "neden": (
            "0ed125e ile stats_detector_robustness.py ANOVA design framing'i "
            "açıklandı. 43ad288 ile figure_thesis.py ve figure_thesis_23.py "
            "ownership scope'u Fig 1-5 / Fig 6-9 olarak ayrıldı. 5ac5bcb ile "
            "historical thesis ve decisions-log generators scripts/legacy altına "
            "README ile taşındı. dca37b6 ile src/wp5/figure_thesis.py, "
            "src/wp5/figure_thesis_23.py ve scripts/eval_only_seed1to7.py "
            "hardcoded run path'leri byte-identical argparse defaults ile "
            "parametrize edildi. e0c47c5 ile src/wp2/synth_regime.py run-dir "
            "provenance artifact'ı ctx.run_dir/wp2_synth.csv olarak "
            "standardize edildi."
        ),
        "etki": (
            "No-arg davranışlar korunarak bakım yapılabilirlik ve provenance "
            "netliği artırıldı. thesis_28, CSV evidence ve PNG result artifacts "
            "değiştirilmedi; aktif downstream sayısal iddialar etkilenmedi."
        ),
        "note": (
            "C3 parametrize edilen aktif scriptler: src/wp5/figure_thesis.py, "
            "src/wp5/figure_thesis_23.py, scripts/eval_only_seed1to7.py. "
            "C4 WP2 global data/processed/wp2_synth.csv'yi latest convenience "
            "snapshot olarak korudu; run_dir/wp2_synth.csv specific-run "
            "provenance artifact oldu."
        ),
    },
    {
        "number": 51,
        "title": "Audit procedure and protected evidence",
        "wp": "Process",
        "karar": (
            "Audit-remediation sürecinde discovery-before-apply protokolü "
            "benimsendi; deney rerun yapılmadı; canonical Lane C protected "
            "SHA check her commit sonrası 4/4 MATCH kaldı; protected CSV "
            "evidence artifact'ları değiştirilmedi."
        ),
        "secenekler": (
            "(a) Hızlı apply ile tahmini düzeltmeler yapmak  |  "
            "(b) Discovery raporu, scoped apply, static verification ve protected "
            "hash check sırasını izlemek  |  "
            "(c) Rerun ile bütün evidence setini yeniden üretmek"
        ),
        "neden": (
            "Discovery adımı, ilk yanlış protected-file mapping riskini görünür "
            "kıldı ve canonical Lane C Rule dosyalarının doğru setini sabitledi: "
            "results/metrics_detector_compare.csv, "
            "docs/internal/wp6_sweep_full/summary_condition_variant.csv, "
            "docs/internal/wp6_sweep_full/summary_paired_combined_vs_sigma.csv, "
            "docs/internal/wp6_sweep_full/summary_paired_combined_vs_regime.csv. "
            "Her apply adımı bu set üzerinde post-commit 4/4 MATCH doğrulaması "
            "ile kapatıldı."
        ),
        "etki": (
            "Düzeltmeler evidence-preserving olarak kaldı: protected CSV "
            "artifact'ları, PNG result artifact'ları ve thesis_28 dokümanı "
            "değişmedi. Rerun yapılmadığı için sayısal iddialarda yeni "
            "stochastic variance veya artifact drift oluşmadı."
        ),
        "note": (
            "Bu karar, audit sürecinin kendisini reproducibility artifact olarak "
            "kayda geçirir. Ana ders: önce kapsam ve consumer/writer ilişkisi "
            "kanıtlanır, sonra küçük scoped apply yapılır, ardından canonical "
            "hash check ile evidence dokunulmazlığı doğrulanır."
        ),
    },
]


SUMMARY_ROWS = [
    (
        28,
        "Lane A audit response — C1 dwell auxiliary-only, C2 WP4 pilot-only, "
        "C3 WP6 noisy sigma calibration caveat; no rerun",
        "Audit",
    ),
    (
        29,
        "Lane B engineering guardrails — CSVMetricLogger schema, resume config "
        "validation, .gitattributes/LF audit stability",
        "Infrastructure",
    ),
    (
        30,
        "Lane C remediation — ANOVA framing, figure ownership, legacy generator "
        "archive, argparse path defaults, WP2 provenance artifact naming",
        "Codebase",
    ),
    (
        31,
        "Audit procedure — discovery before apply, no reruns, canonical Lane C "
        "protected SHA stayed 4/4 MATCH",
        "Process",
    ),
]


def main():
    shutil.copy2(SRC, DST)
    doc = Document(DST)

    # Last individual Karar table in log_11 is #47 — second-to-last table overall.
    template = doc.tables[-2]

    karar_tables = [
        _build_karar_table_v2(doc, template, **decision)
        for decision in DECISIONS
    ]

    ozet_heading = _find_ozet_heading(doc)
    for tbl_xml in karar_tables:
        ozet_heading.addprevious(_blank_p())
        ozet_heading.addprevious(tbl_xml)

    body = doc.element.body
    for child in body.iterchildren():
        if child.tag != W_NS + "p":
            continue
        text = "".join(t.text or "" for t in child.iter(W_NS + "t"))
        if any(s in text for s in ["En Kritik 27 Karar", "Sürüm 27"]):
            from docx.text.paragraph import Paragraph

            p_obj = Paragraph(child, doc.part)
            _replace_in_paragraph(p_obj, "En Kritik 27 Karar", "En Kritik 31 Karar")
            _replace_in_paragraph(p_obj, "Sürüm 27", "Sürüm 28")

    for para in doc.paragraphs:
        if "ChatGPT ve Claude" in para.text and "27 karar" in para.text:
            _replace_in_paragraph(para, "27 karar", "31 karar")

    sum_tbl = doc.tables[-1]
    template_row = sum_tbl.rows[1]._tr

    def _add_summary_row(num, text, wp):
        new_row = copy.deepcopy(template_row)
        sum_tbl._tbl.append(new_row)
        from docx.table import _Row

        row_obj = _Row(new_row, sum_tbl)
        _set_cell_text(row_obj.cells[0], str(num))
        _set_cell_text(row_obj.cells[1], text)
        _set_cell_text(row_obj.cells[2], wp)

    for row in SUMMARY_ROWS:
        _add_summary_row(*row)

    doc.save(DST)
    print(f"decisions_log_12.docx saved ({DST.stat().st_size:,} bytes)")

    doc2 = Document(DST)
    print("\n=== Verification ===")
    print(f"Total tables: {len(doc2.tables)}")
    print("New karar table headers:")
    for idx in range(-5, -1):
        print(f"  {doc2.tables[idx].rows[0].cells[1].text[:90]}")
    sum2 = doc2.tables[-1]
    print(f"Summary rows: {len(sum2.rows)} (should be 32 = header + 31)")
    print(f"  Last row: {[c.text[:70] for c in sum2.rows[-1].cells]}")
    for p in doc2.paragraphs[:40]:
        if "Sürüm" in p.text or "En Kritik" in p.text:
            print(f"Marker: {p.text}")


if __name__ == "__main__":
    main()
