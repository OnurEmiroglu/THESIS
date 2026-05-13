"""decisions_log_9.docx -> decisions_log_10.docx

Appends Decision #46 (signal informativeness sweep parameter calibration:
alpha=0.40, k=20). Updates title-page marker to "Sürüm 27" and summary
heading from "En Kritik 25 Karar" -> "En Kritik 26 Karar".

Decision #46 uses a different sub-section structure than #1-#45:
  Karar / Seçenekler / Neden / Etki / Not  (no İkilem; trailing Not row)
The 6-row table template is reused; row labels are overwritten.
"""

from __future__ import annotations
import sys, io, shutil, copy
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement

SRC = Path("manuscript/decisions_log_9.docx")
DST = Path("manuscript/decisions_log_10.docx")

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _set_cell_text(cell, text, *, bold=None):
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].text = text
        for r in para.runs[1:]:
            r.text = ""
        if bold is not None:
            para.runs[0].bold = bold
    else:
        r = para.add_run(text)
        if bold is not None:
            r.bold = bold


def _build_karar_table_v2(doc, template_tbl, number, title, wp,
                          karar, secenekler, neden, etki, note):
    """Build a Karar table with Karar / Seçenekler / Neden / Etki / Not rows."""
    new_xml = copy.deepcopy(template_tbl._tbl)
    from docx.table import Table
    new_tbl = Table(new_xml, doc.part)
    _set_cell_text(new_tbl.rows[0].cells[0], wp, bold=True)
    _set_cell_text(new_tbl.rows[0].cells[1], f"#{number}  {title}", bold=True)
    _set_cell_text(new_tbl.rows[1].cells[0], "Karar", bold=True)
    _set_cell_text(new_tbl.rows[1].cells[1], karar)
    _set_cell_text(new_tbl.rows[2].cells[0], "Seçenekler", bold=True)
    _set_cell_text(new_tbl.rows[2].cells[1], secenekler)
    _set_cell_text(new_tbl.rows[3].cells[0], "Neden", bold=True)
    _set_cell_text(new_tbl.rows[3].cells[1], neden)
    _set_cell_text(new_tbl.rows[4].cells[0], "Etki", bold=True)
    _set_cell_text(new_tbl.rows[4].cells[1], etki)
    _set_cell_text(new_tbl.rows[5].cells[0], "Not", bold=True)
    _set_cell_text(new_tbl.rows[5].cells[1], note)
    return new_xml


def _replace_in_paragraph(para, old_text, new_text):
    full = "".join(r.text for r in para.runs)
    if old_text not in full:
        return False
    new_full = full.replace(old_text, new_text)
    if para.runs:
        para.runs[0].text = new_full
        for r in para.runs[1:]:
            r.text = ""
    return True


def main():
    shutil.copy2(SRC, DST)
    doc = Document(DST)

    # Last individual Karar table in log_9 is #45 — second-to-last table overall
    template = doc.tables[-2]

    k46 = _build_karar_table_v2(
        doc, template,
        number=46,
        title=(
            "Signal Informativeness Sweep için noisy ve lagged koşullarının "
            "parametre kalibrasyonu (α = 0.40, k = 20)"
        ),
        wp="WP5.5",
        karar=(
            "Yaklaşan Signal Informativeness Sweep'inde noisy koşulu için noise "
            "standard deviation çarpanı α = 0.40 ve lagged koşulu için lag "
            "k = 20 step olarak sabitlendi."
        ),
        secenekler=(
            "(a) α = 0.20, k = 10 — konservatif, hafif degradation  |  "
            "(b) α = 0.40, k = 20 — orta-üst degradation, NRMSE bakımından "
            "yaklaşık eşleşmiş  |  "
            "(c) α = 0.80, k = 50 — agresif, none koşuluna yakın"
        ),
        neden=(
            "Offline calibration audit'i (run "
            "20260425-184732_seed42_wp55-calibration_1e806ff, 7 alpha × 5 seed + "
            "7 k × 5 seed = 70 ölçüm) seçimi yönlendirdi. Seçim kriterleri: "
            "(i) noisy ve lagged koşulları clean'den anlamlı biçimde ayrılmalı "
            "ama none'a çökmemeli, (ii) iki koşul birbirine yakın informational "
            "degradation şiddetinde olmalı ki sweep'te \"noise mı lag mı daha "
            "bozucu\" sorusu temiz cevaplanabilsin. α = 0.40 (Pearson 0.929, "
            "NRMSE 0.40, accuracy_drop 0.053) ve k = 20 (Pearson 0.937, "
            "NRMSE 0.35, accuracy_drop 0.042) bu iki kriteri karşılayan "
            "eşleştirilmiş çift olarak öne çıktı. Konservatif (a) seçeneği "
            "\"degradation var ama neredeyse yok\" bölgesinde kalarak sweep'i "
            "düzleştirme riski taşıyordu; agresif (c) seçeneği none koşuluna "
            "fazla yaklaşıyordu."
        ),
        etki=(
            "Sweep ana çalışması α = 0.40 ve k = 20 ile yürütülecek. Sensitivity "
            "analizi olarak α = 0.20 ve k = 10 appendix'te raporlanacak. "
            "Calibration audit CSV'leri (metrics_calibration_per_seed.csv, "
            "metrics_calibration_aggregated.csv) docs/internal/calibration_audit/ "
            "altında saklı; thesis Methods bölümünde bu kalibrasyona atıf "
            "yapılacak."
        ),
        note=(
            "Kalibrasyon detayı olarak konumlandırıldığı için hocaya sweep "
            "öncesi ayrı bir mail atılmadı; sweep tamamlandığında ilerleme "
            "güncellemesinde raporlanacak."
        ),
    )

    body = doc.element.body
    ozet_heading = None
    for child in body.iterchildren():
        if child.tag == W_NS + "p":
            pStyle = child.find(W_NS + "pPr/" + W_NS + "pStyle")
            if pStyle is not None and pStyle.get(W_NS + "val") == "Heading1":
                text = "".join(t.text or "" for t in child.iter(W_NS + "t"))
                if "Özet" in text:
                    ozet_heading = child
                    break
    if ozet_heading is None:
        raise RuntimeError("Özet heading not found")

    def _blank_p():
        return OxmlElement("w:p")

    ozet_heading.addprevious(_blank_p())
    ozet_heading.addprevious(k46)

    # Summary heading: 25 -> 26
    for child in body.iterchildren():
        if child.tag == W_NS + "p":
            text = "".join(t.text or "" for t in child.iter(W_NS + "t"))
            if "En Kritik 25 Karar" in text:
                from docx.text.paragraph import Paragraph
                p_obj = Paragraph(child, doc.part)
                _replace_in_paragraph(p_obj, "En Kritik 25 Karar", "En Kritik 26 Karar")

    for para in doc.paragraphs:
        if "ChatGPT ve Claude" in para.text and "25 karar" in para.text:
            _replace_in_paragraph(para, "25 karar", "26 karar")

    # Title-page: Sürüm 26 -> Sürüm 27
    for para in doc.paragraphs:
        if "Sürüm 26" in para.text:
            _replace_in_paragraph(para, "Sürüm 26", "Sürüm 27")

    # Summary row #26
    sum_tbl = doc.tables[-1]
    template_row = sum_tbl.rows[1]._tr

    def _add_summary_row(num, text, wp):
        new_row = copy.deepcopy(template_row)
        sum_tbl._tbl.append(new_row)
        from docx.table import _Row
        row_obj = _Row(new_row, sum_tbl)
        _set_cell_text(row_obj.cells[0], str(num))
        _set_cell_text(row_obj.cells[1], text)
        _set_cell_text(row_obj.cells[2], wp)

    _add_summary_row(
        26,
        "Signal Informativeness Sweep noisy/lagged parametre kalibrasyonu — "
        "α=0.40, k=20 sabitlendi (offline audit dayanağı)",
        "WP5.5",
    )

    doc.save(DST)
    print(f"decisions_log_10.docx saved ({DST.stat().st_size:,} bytes)")

    doc2 = Document(DST)
    print("\n=== Verification ===")
    print(f"Total tables: {len(doc2.tables)}")
    t = doc2.tables[-2]
    print(f"Last karar table header: {t.rows[0].cells[1].text[:80]}")
    sum2 = doc2.tables[-1]
    print(f"Summary rows: {len(sum2.rows)} (should be 27 = header + 26)")
    print(f"  Last row: {[c.text[:60] for c in sum2.rows[-1].cells]}")
    for p in doc2.paragraphs[:30]:
        if "Sürüm" in p.text or "En Kritik" in p.text:
            print(f"Marker: {p.text}")


if __name__ == "__main__":
    main()
