"""Generate thesis_24.docx.
Based on thesis_23 with added TOST equivalence test results for the ablation
(sigma_only vs oracle_full) and misspecification experiments."""

import copy, sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

dst = Document()

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_normal(doc, text):
    p = doc.add_paragraph(text)
    return p

def add_table_2col(doc, title, rows):
    """2-column table with header row."""
    t = doc.add_table(rows=1 + len(rows), cols=2, style='Table Grid')
    t.rows[0].cells[0].text = rows[0][0] if isinstance(title, list) else title.split("|")[0]
    t.rows[0].cells[1].text = rows[0][1] if isinstance(title, list) else title.split("|")[1]
    # header bold
    for cell in t.rows[0].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for i, (c1, c2) in enumerate(rows[1:] if isinstance(title, list) else rows, start=1):
        t.rows[i].cells[0].text = c1
        t.rows[i].cells[1].text = c2
    return t

def add_table(doc, headers, rows):
    """General n-column table."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Table Grid')
    for j, h in enumerate(headers):
        t.rows[0].cells[j].text = h
        for p in t.rows[0].cells[j].paragraphs:
            for r in p.runs:
                r.bold = True
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            t.rows[i + 1].cells[j].text = str(val)
    return t

# ================================================================
# COVER PAGE  — CHANGE A: "Nisan 2026 — Surum 18"
# ================================================================
p = dst.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Karlsruhe Institute of Technology").font.size = Pt(14)
p = dst.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Financial Engineering MSc Programı")
dst.add_paragraph()
p = dst.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Yüksek Lisans Tez Taslağı").font.size = Pt(12)
dst.add_paragraph()
p = dst.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Farklı Volatilite Rejimleri Altında Pekiştirmeli Öğrenme ile Yüksek Frekanslı Piyasa Yapıcılığı")
r.bold = True; r.font.size = Pt(14)
dst.add_paragraph()
p = dst.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Onur Emiroğlu")
p = dst.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Financial Engineering MSc Programı")
p = dst.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Nisan 2026 — Sürüm 23")

# ================================================================
# 1. OZET — CHANGE B: son paragraf degistirildi, CHANGE C: anahtar kelimeler
# ================================================================
add_heading(dst, "1. ÖZET", level=1)

add_normal(dst, (
    "Bu çalışmada, volatilite rejim bilgisinin (Düşük/Orta/Yüksek) pekiştirmeli öğrenme tabanlı piyasa "
    "yapıcı ajanların performansına katkısı araştırılmaktadır. Proximal Policy Optimization (PPO) "
    "algoritması kullanılarak beş ajan varyantı eğitilmiştir: yalnızca sigma (ppo_sigma_only), yalnızca "
    "rejim (ppo_regime_only), sigma + tahmini rejim (ppo_combined), yalnızca gerçek rejim (ppo_oracle_pure) "
    "ve sigma + gerçek rejim (ppo_oracle_full). Ajanlar, Poisson varış süreçli yoğunluk tabanlı dolum modeli "
    "ve Markov zinciri rejim geçişleri içeren tamamen sentetik bir piyasa ortamında değerlendirilmiştir."
))

add_normal(dst, (
    "Yürütülen out-of-sample deney (20 bağımsız tohum, 1 milyon zaman adımı eğitim, %70/%30 "
    "kronolojik eğitim/test bölünmesi), tüm PPO varyantlarının Sharpe oranı açısından klasik referans "
    "stratejilerini (Avellaneda-Stoikov ve naif sabit-spread) yaklaşık 6-7 kat geçtiğini ortaya koymaktadır. "
    "Buna karşın, rejim bilgisinin PPO performansına net bir katkı sağladığı hipotezi yalnızca istatistiksel "
    "olarak desteklenmemekle kalmamış; ppo_sigma_only ile ppo_oracle_full'un Sharpe performansları "
    "TOST eşdeğerlik testi ile pratik olarak eşdeğer bulunmuştur (TOST p<0.05, ±0.10 sınır)."
))

add_normal(dst, (
    "Ana Bulgu (Ablasyon): oracle_full ajanı (sigma + regime_true), mükemmel rejim etiketlerine sahip olmasına "
    "karşın ppo_sigma_only'yi istatistiksel olarak anlamlı biçimde geçememiştir (Sharpe: 0.753 vs 0.722, "
    "p = 0.115). Bu sonuç, rejim etiketinin — gerçek ya da tahmini — sigma_hat'in zaten taşıdığı bilginin "
    "üzerine anlamlı katkı sağlamadığı hipoteziyle güçlü biçimde tutarlıdır."
))

# CHANGE B: two new paragraphs replacing the old "Detector robustness deneyi" paragraph
add_normal(dst, (
    "Detector robustness deneyi — rv_baseline (%60.7), rv_dwell (%60.4) ve HMM (%81.8) dedektörleri ile "
    "20 tohumda yinelenen analiz — null sonucun dedektör seçiminden bağımsız olduğunu doğrulamaktadır "
    "(tüm dedektörler için Sharpe bazlı p > 0.08). Oracle deneyi ise %100 doğrulukla eşdeğer bir test "
    "sunmaktadır: ppo_oracle_full mükemmel rejim bilgisiyle dahi sigma_only'yi geçememiştir (p=0.115)."
))

add_normal(dst, (
    "Model misspecification robustness deneyi kapsamında fill model parametreleri A ve k rejime bağlı "
    "kılınmış (L: A=4.0/k=1.8, M: A=5.0/k=1.5, H: A=6.0/k=1.2) ve 5-varyant değerlendirmesi bu "
    "zorlaştırılmış ortamda yinelenmiştir. ppo_sigma_only yine en yüksek Sharpe değerini elde etmiş "
    "(0.686); oracle_full ile aradaki fark istatistiksel olarak anlamsız kalmıştır (p=0.881). Bu bulgu, "
    "signal redundancy argümanının çevresel koşullara karşı da sağlam olduğunu göstermektedir."
))

# CHANGE C: updated keywords
add_normal(dst, (
    "Anahtar kelimeler: piyasa yapıcılığı, pekiştirmeli öğrenme, PPO, volatilite rejimleri, "
    "Avellaneda-Stoikov, ablasyon, oracle agent, sinyal yedekliliği, model yanlış belirleme"
))

# ================================================================
# SEMBOLLER VE KISALTMALAR
# ================================================================
add_heading(dst, "SEMBOLLER VE KISALTMALAR", level=1)

add_normal(dst, (
    "Bu bölümde çalışma boyunca kullanılan matematiksel semboller, Yunan harfleri ve kısaltmalar "
    "tanımlanmaktadır."
))

# Matematiksel Semboller
add_heading(dst, "Matematiksel Semboller", level=2)
math_symbols = [
    ["Sembol", "Açıklama"],
    ["M_t", "t anındaki mid-price (orta fiyat) [para birimi]"],
    ["σ", "Volatilite parametresi — mid-price'ın birim zamandaki standart sapması [tick/√saniye]"],
    ["σ_base", "Temel volatilite parametresi; Orta (M) rejim için referans değer (0.8 tick/√saniye)"],
    ["σ_L, σ_M, σ_H", "Düşük, Orta ve Yüksek rejimlere karşılık gelen volatilite değerleri"],
    ["Δt", "Simülasyon zaman adımı büyüklüğü (0.2 saniye)"],
    ["z", "Standart normal rassal değişken; z ~ N(0,1)"],
    ["λ(δ)", "δ tick uzaklığındaki kotasyon için Poisson dolum yoğunluğu [dolum/saniye]"],
    ["A", "Sıfır uzaklıkta (δ=0) temel dolum yoğunluğu (5.0 dolum/saniye)"],
    ["k", "Dolum yoğunluğunun uzaklıkla üstel azalma hızı (1.5 tick başına)"],
    ["δ", "Kotasyon uzaklığı — bid veya ask fiyatının mid-price'tan tick cinsinden sapması"],
    ["δ_bid", "Bid (alış) kotasyonunun mid-price'tan uzaklığı [tick]"],
    ["δ_ask", "Ask (satış) kotasyonunun mid-price'tan uzaklığı [tick]"],
    ["P_fill", "Bir zaman adımında en az bir dolumun gerçekleşme olasılığı"],
    ["h", "Half-spread — kotasyon yarı-aralığı [tick]; h ∈ {1, 2, 3, 4, 5}"],
    ["m", "Skew — asimetrik kotasyon kayması [tick]; m ∈ {-2, -1, 0, +1, +2}"],
    ["q", "Envanter (inventory) — anda tutulan net pozisyon miktarı [lot]"],
    ["q_norm", "Normalize edilmiş envanter; q_norm = q / inv_max_clip ∈ [-1, +1]"],
    ["inv_max_clip", "Envanter kırpma sınırı (50 lot)"],
    ["τ", "Kalan zaman fraksiyonu; τ = (T - t) / T ∈ [0, 1]"],
    ["T", "Episode toplam adım sayısı (8000 adım)"],
    ["R_t", "t adımındaki ödül (reward)"],
    ["η", "Envanter ceza katsayısı; η = 0.001"],
    ["γ", "Avellaneda-Stoikov risk aversion parametresi; γ = 0.01"],
    ["r", "Avellaneda-Stoikov rezervasyon fiyatı"],
    ["d", "Avellaneda-Stoikov optimal yarı-spread"],
    ["k_price", "Fiyat cinsinden ifade edilen yoğunluk azalma parametresi; k_price = k / tick_size"],
    ["P", "Markov zinciri geçiş matrisi (3×3)"],
    ["Sharpe", "Sharpe oranı — risk-düzeltmeli getiri ölçütü; Sharpe = (μ/σ) × √(1/Δt)"],
    ["inv_p99", "Mutlak envanter değerinin 99. persantili [lot]"],
    ["equity", "Toplam varlık değeri; equity = cash + q × M_t"],
]
add_table(dst, math_symbols[0], math_symbols[1:])

# Yunan Harfleri
add_heading(dst, "Yunan Harfleri", level=2)
greek = [
    ["Sembol", "Açıklama"],
    ["γ (gamma)", "Risk aversion parametresi (Avellaneda-Stoikov modelinde)"],
    ["η (eta)", "Envanter ceza katsayısı (ödül fonksiyonunda)"],
    ["σ (sigma)", "Volatilite — fiyat değişiminin standart sapması"],
    ["τ (tau)", "Kalan zaman fraksiyonu"],
    ["λ (lambda)", "Poisson dolum yoğunluğu"],
    ["δ (delta)", "Kotasyon uzaklığı (mid-price'tan tick cinsinden sapma)"],
    ["μ (mu)", "Ortalama adım başına PnL (Sharpe hesabında)"],
]
add_table(dst, greek[0], greek[1:])

# Kısaltmalar
add_heading(dst, "Kısaltmalar", level=2)
abbrevs = [
    ["Kısaltma", "Açıklama"],
    ["PPO", "Proximal Policy Optimization — yakınsak politika optimizasyonu algoritması"],
    ["RL", "Reinforcement Learning — pekiştirmeli öğrenme"],
    ["AS", "Avellaneda-Stoikov — klasik stokastik kontrol tabanlı piyasa yapıcılığı modeli"],
    ["HFT", "High-Frequency Trading — yüksek frekanslı alım-satım"],
    ["OOS", "Out-of-Sample — modelin eğitilmediği veri üzerinde değerlendirme"],
    ["WP", "Work Package — iş paketi (proje aşaması)"],
    ["ABM", "Arithmetic Brownian Motion — aritmetik Brownian hareketi"],
    ["RV", "Realized Volatility — gerçekleşmiş volatilite (kayan pencere ile hesaplanan)"],
    ["PnL", "Profit and Loss — kar ve zarar"],
    ["CUDA", "Compute Unified Device Architecture — NVIDIA GPU hesaplama mimarisi"],
    ["SB3", "Stable-Baselines3 — PPO implementasyonu için kullanılan Python kütüphanesi"],
    ["MlpPolicy", "Multi-Layer Perceptron Policy — çok katmanlı algılayıcı politika ağı"],
    ["GAE", "Generalized Advantage Estimation — genelleştirilmiş avantaj tahmini (PPO bileşeni)"],
    ["L / M / H", "Low / Medium / High — Düşük / Orta / Yüksek volatilite rejimleri"],
]
add_table(dst, abbrevs[0], abbrevs[1:])

# ================================================================
# TERİMLER SÖZLÜĞÜ
# ================================================================
add_heading(dst, "TERİMLER SÖZLÜĞÜ", level=1)
add_normal(dst, (
    "Bu sözlük, çalışmada kullanılan teknik terimleri konuya aşina olmayan okuyucular için "
    "tanımlamaktadır. Terimler alfabetik sıraya göre düzenlenmiştir."
))

glossary = [
    ["Terim", "Tanım"],
    ["Ablation (Ablasyon)", "Bir modelin belirli bir bileşeninin çıkarılarak ya da devre dışı bırakılarak etkisinin ölçülmesi amacıyla yapılan kontrollü deney."],
    ["Action Space (Eylem Uzayı)", "Bir pekiştirmeli öğrenme ajanının her adımda seçebileceği tüm olası eylemlerin kümesi. Bu çalışmada half-spread (h) ve skew (m) değerlerinden oluşan ayrık bir yapıdadır."],
    ["Adverse Selection (Ters Seçim)", "Piyasa yapıcının bilgili yatırımcılara karşı dezavantajlı konuma düşmesi durumu."],
    ["Agent (Ajan)", "Pekiştirmeli öğrenmede, bir ortamda gözlem yaparak karar veren ve ödül sinyaline göre davranışını güncelleyen öğrenen sistem."],
    ["Ask (Satış Fiyatı)", "Piyasa yapıcının satmaya hazır olduğu fiyat. Her zaman bid fiyatının üzerindedir."],
    ["Bid (Alış Fiyatı)", "Piyasa yapıcının almaya hazır olduğu fiyat. Her zaman ask fiyatının altındadır."],
    ["Bid-Ask Spread", "Alış ve satış fiyatı arasındaki fark. Piyasa yapıcının temel gelir kaynağıdır."],
    ["Brownian Motion", "Rastlantısal yürüyüş modellerinin sürekli-zaman limiti. ABM'de fiyat değişimleri normal dağılımdan çekilmektedir."],
    ["Clip (Kırpma)", "Bir değerin belirli bir aralıkla sınırlandırılması işlemi. Bu çalışmada envanter ±50 lot ile kırpılmaktadır."],
    ["Convergence (Yakınsama)", "Pekiştirmeli öğrenme eğitiminde politikanın optimal ya da kararlı bir çözüme ulaşması süreci."],
    ["Entropy Coefficient", "PPO'da keşif-sömürü dengesini ayarlayan hiperparametre. Bu çalışmada ent_coef=0.01 kullanılmıştır."],
    ["Episode", "Pekiştirmeli öğrenmede başlangıçtan bitiş koşuluna kadar süren bir tam simülasyon koşusu. Bu çalışmada 8000 zaman adımından oluşmaktadır."],
    ["Exogenous Series (Dışsal Seri)", "Modelin dışında önceden üretilmiş ve tüm stratejilere aynı biçimde sunulan fiyat/volatilite serisi."],
    ["Fill (Dolum)", "Piyasa yapıcının kotasyon verdiği fiyattan gerçekleşen emir eşleşmesi."],
    ["Fill Rate (Dolum Oranı)", "Toplam adım sayısına bölünen dolum sayısı."],
    ["Gymnasium", "OpenAI tarafından geliştirilen, pekiştirmeli öğrenme ortamları için standart Python arayüzü."],
    ["Half-Spread (Yarı-Aralık)", "Bid veya ask fiyatının mid-price'tan uzaklığı. Bu çalışmada h sembolüyle gösterilir."],
    ["Inventory (Envanter)", "Piyasa yapıcının anda elinde bulundurduğu net pozisyon miktarı."],
    ["Inventory Risk (Envanter Riski)", "Piyasa yapıcının büyük envanter pozisyonu tutarken fiyat hareketleri nedeniyle maruz kaldığı zarar riski."],
    ["Latency (Gecikme)", "Piyasa yapıcının kotasyon güncellemesi ile piyasanın bu güncellemeyi işlemesi arasındaki zaman farkı."],
    ["Limit Order Book", "Bekleme sırasındaki tüm alış ve satış emirlerinin fiyat-miktar listesi."],
    ["Markov Chain", "Gelecekteki durumun yalnızca mevcut duruma bağlı olduğu olasılıksal süreç."],
    ["Market Making (Piyasa Yapıcılığı)", "Finansal piyasalarda sürekli alış ve satış kotasyonu vererek likidite sağlama faaliyeti."],
    ["Mid-Price", "Anlık alış (bid) ve satış (ask) fiyatlarının ortalaması."],
    ["Model Misspecification (Model Yanlış Belirleme)", "Modelin varsayımlarının gerçek veri üretim sürecinden sapması durumu. Bu çalışmada fill parametrelerinin rejime bağlı kılınmasıyla test edilmiştir."],
    ["Null Result", "Araştırma hipotezinin verilerce desteklenmediği bulgu."],
    ["Observation Space (Gözlem Uzayı)", "Ajanın her adımda çevreden aldığı bilginin vektör temsili. Bu çalışmada 6 boyutludur: [q_norm, σ̂, τ, r_L, r_M, r_H]."],
    ["One-Hot Encoding", "Kategorik bir değişkeni ikili vektörle temsil etme yöntemi."],
    ["Out-of-Sample (OOS)", "Modelin eğitilmediği, bağımsız bir veri kümesi üzerinde yapılan değerlendirme."],
    ["Percentile (Persantil)", "Veri setinin belirli bir yüzdesinin altında kaldığı değer."],
    ["Policy (Politika)", "Pekiştirmeli öğrenmede ajanın gözleme göre eylem seçme stratejisi."],
    ["Poisson Process", "Olayların bağımsız ve sabit ortalama hızda rastlantısal gerçekleştiği olasılıksal süreç."],
    ["PPO", "John Schulman ve ark. (2017) tarafından geliştirilen politika gradyanı tabanlı RL algoritması."],
    ["Regime (Rejim)", "Piyasanın belirli bir volatilite seviyesinde bulunduğu dönem. Bu çalışmada L/M/H."],
    ["Risk-Adjusted Return", "Getirinin, üstlenilen risk miktarına bölünmesiyle elde edilen performans ölçütü."],
    ["Seed (Tohum)", "Rastlantısal sayı üretecinin başlangıç değeri. Tekrarlanabilirlik için kullanılır."],
    ["Sharpe Ratio", "Birim risk başına elde edilen fazla getiriyi ölçen performans metriği."],
    ["Skew (Asimetrik Kotasyon)", "Bid ve ask kotasyonlarının mid-price'a göre eşit olmayan biçimde kaydırılması."],
    ["Stable-Baselines3 (SB3)", "PPO dahil çeşitli RL algoritmalarının güvenilir PyTorch implementasyonlarını içeren Python kütüphanesi."],
    ["Sticky Transition Matrix", "Köşegen elemanları yüksek Markov geçiş matrisi. Rejimlerin uzun süre devam etmesini sağlar."],
    ["Tick", "Finansal piyasalarda fiyatın alabileceği en küçük artış birimi. Bu çalışmada tick_size = 0.01."],
    ["Timestep (Zaman Adımı)", "Simülasyonun bir sonraki duruma geçtiği en küçük zaman birimi. Δt = 0.2 saniye."],
    ["Chronological Train/Test Split", "Zaman serisi verisinin kronolojik sıraya göre eğitim (%70) ve test (%30) kümelerine bölünmesi."],
]
add_table(dst, glossary[0], glossary[1:])

# ================================================================
# 2. GİRİŞ
# ================================================================
add_heading(dst, "2. GİRİŞ", level=1)

add_normal(dst, (
    "Piyasa yapıcılığı (market making), finansal piyasalarda likidite sağlayan ve sürekli olarak hem alış (bid) "
    "hem de satış (ask) kotasyonu veren ajanların stratejik davranışını inceleyen bir alandır. Piyasa yapıcılar, "
    "spread geliri elde etmeye çalışırken olumsuz fiyat hareketlerine maruz kalan envanter riskini de "
    "yönetmek zorundadır. Bu denge, özellikle volatilitenin değişken olduğu yüksek frekanslı piyasalarda "
    "kritik önem taşımaktadır."
))

add_normal(dst, (
    "Klasik yaklaşımlar arasında Avellaneda ve Stoikov (2008) tarafından geliştirilen stokastik kontrol "
    "çerçevesi öne çıkmaktadır. Bu model, envanter riskini rezervasyon fiyatı kavramıyla içselleştirerek "
    "optimal bid-ask spread'ini analitik olarak türetmektedir. Ancak bu yaklaşım, sabit bir volatilite "
    "varsayımına dayanmakta ve piyasa rejimlerindeki değişimlere uyum sağlayamamaktadır."
))

add_normal(dst, (
    "Bu çalışmanın temel araştırma sorusu şudur: Volatilite rejim bilgisine erişimi olan bir PPO ajanı, "
    "bu bilgiden yoksun olan eşdeğerine kıyasla daha iyi performans sergiler mi? Bu soruyu yanıtlamak için "
    "aşağıdaki katkılar sunulmaktadır:"
))

p = dst.add_paragraph("(1) Markov zinciri rejim geçişleri ve Poisson tabanlı dolum modeli içeren kontrollü sentetik bir piyasa ortamı tasarlanmıştır.", style='List Paragraph')
p = dst.add_paragraph("(2) Rejim farkındalığı ablasyonu gerçekleştirilmiş; 5 PPO varyantı 20 bağımsız tohum ile OOS protokolüyle karşılaştırılmıştır.", style='List Paragraph')
p = dst.add_paragraph("(3) Oracle agent deneyi ile mükemmel rejim bilgisinin üst sınır değeri ölçülmüştür.", style='List Paragraph')
p = dst.add_paragraph("(4) Detector robustness analizi gerçekleştirilmiş; null sonucun dedektör seçiminden bağımsızlığı doğrulanmıştır.", style='List Paragraph')

# ================================================================
# 3. TEORİ VE METODOLOJİ
# ================================================================
add_heading(dst, "3. TEORİ VE METODOLOJİ", level=1)

# 3.1 Piyasa Modeli
add_heading(dst, "3.1 Piyasa Modeli", level=2)
add_normal(dst, (
    "Piyasa fiyatı, ayrık zamanlı aritmetik Brownian hareketi (ABM) ile modellenmektedir:"
))
add_normal(dst, "M_{t+1} = M_t + σ_r · √(Δt) · z,    z ~ N(0,1)")
add_normal(dst, (
    "Burada M_t mid-price, σ_r rejime bağlı volatilite parametresi (ticks/√saniye cinsinden), "
    "Δt = 0.2 saniye zaman adımı büyüklüğüdür. Başlangıç fiyatı M_0 = 100.0, tick büyüklüğü 0.01'dir."
))
add_normal(dst, (
    "Temel volatilite parametresi σ_base = 0.8 tick olarak belirlenmiş; üç rejim için çarpanlar "
    "[0.6, 1.0, 1.8] olarak tanımlanmıştır. Buna göre σ_L = 0.48, σ_M = 0.80, σ_H = 1.44 tick/√saniye "
    "değerleri elde edilmektedir."
))

# 3.2 Dolum Modeli
add_heading(dst, "3.2 Dolum Modeli", level=2)
add_normal(dst, (
    "Emir dolumları, yoğunluk tabanlı Poisson süreci ile modellenmektedir. Delta tick cinsinden ifade "
    "edildiğinde, dolum yoğunluğu:"
))
add_normal(dst, "λ(δ) = A · exp(-k · δ)")
add_normal(dst, (
    "formülüyle hesaplanmaktadır. Burada A = 5.0 (delta=0'daki temel yoğunluk, dolum/saniye) ve k = 1.5 "
    "(tick başına üstel azalma) parametreleridir. Bir zaman adımında en az bir dolumun gerçekleşme olasılığı:"
))
add_normal(dst, "P_fill = 1 - exp(-λ · Δt)")
add_normal(dst, (
    "olarak verilmektedir. Hem bid hem ask tarafı için bağımsız Bernoulli denemeleri gerçekleştirilmektedir."
))
add_normal(dst, (
    "Gecikme modeli (latency_steps = 1), ajanın kotasyon hesaplamasında bir önceki adımın mid-price "
    "değerini kullanmasına yol açmaktadır; bu durum ters seçim (adverse selection) riskini "
    "içselleştirmektedir. Komisyon ücreti her işlem için 0.2 baz puan olarak uygulanmaktadır."
))

# 3.3 Rejim Modeli
add_heading(dst, "3.3 Rejim Modeli", level=2)
add_normal(dst, (
    "Volatilite rejimleri, üç durumlu (L: Düşük, M: Orta, H: Yüksek) birinci derece Markov zinciri ile "
    "üretilmektedir. Geçiş matrisi yapışkan (sticky) olacak şekilde tasarlanmıştır:"
))
add_normal(dst, "P = [[0.9967, 0.0023, 0.0010], [0.0042, 0.9917, 0.0041], [0.0010, 0.0030, 0.9960]]")
add_normal(dst, (
    "Beklenen rejim süreleri sırasıyla L: ~300, M: ~120, H: ~250 zaman adımıdır. Bu yapışkanlık, kayan "
    "gerçekleşmiş volatilite (rolling realized volatility — RV) ile rejim tespitinin mümkün olmasını "
    "sağlamaktadır."
))
add_normal(dst, (
    "Rejim tespiti için RV pencere boyutu 50 adım, ısınma süresi (warmup) 1000 adım olarak belirlenmiştir. "
    "Eşik değerleri ısınma döneminin 33. ve 66. persentillerinden kalibre edilmektedir. Gerçek zamanlı "
    "tespitte look-ahead bias engellenmiştir; regime_hat yalnızca geçmiş veriden hesaplanmaktadır. Elde "
    "edilen tespit doğruluğu %60.7'dir (rastgele sınır: %33.3)."
))
add_normal(dst, (
    "Detector robustness analizi kapsamında iki ek tespit yöntemi karşılaştırılmıştır: "
    "(i) dwell filtreli RV (rv_dwell): kısa süreli (< 5 adım) rejim geçişlerini düzelten filtre uygulanmış, "
    "tespit doğruluğu %60.4 olarak elde edilmiştir; "
    "(ii) Gaussian HMM (hmm): sigma_hat serisi üzerinde eğitilen üç durumlu Gaussian gizli Markov modeli, "
    "%81.8 doğruluk ile en yüksek tespit kalitesini sağlamıştır."
))

# 3.4 Gymnasium Ortamı
add_heading(dst, "3.4 Gymnasium Ortamı", level=2)
add_normal(dst, (
    "OpenAI Gymnasium uyumlu MMEnv ortamı oluşturulmuştur. Gözlem uzayı 6 boyutludur:"
))
add_normal(dst, "obs = [q_norm, σ̂_t, τ, r_L, r_M, r_H]")
add_normal(dst, (
    "Burada q_norm = clip(inv, -50, 50) / inv_max_clip envanter normalizasyonu, σ̂_t kayan gerçekleşmiş "
    "volatilite, τ = (T-t)/T kalan zaman fraksiyonu, r_L/r_M/r_H ise rejim one-hot kodlamasıdır "
    "(ppo_blind için hepsi sıfır)."
))
add_normal(dst, "Eylem uzayı MultiDiscrete([5, 5]) olarak tanımlanmıştır:")
add_normal(dst, "h_idx ∈ {0,1,2,3,4} → h = h_idx + 1 tick (half-spread)")
add_normal(dst, "m_idx ∈ {0,1,2,3,4} → m = m_idx - 2 tick (skew)")
add_normal(dst, "δ_bid = max(1, h + m)")
add_normal(dst, "δ_ask = max(1, h - m)")

# 3.5 Referans Stratejiler
add_heading(dst, "3.5 Referans Stratejiler", level=2)

add_heading(dst, "3.5.1 Naif Sabit-Spread Stratejisi", level=3)
add_normal(dst, (
    "Her zaman adımında simetrik sabit half-spread uygulanmaktadır: δ_bid = δ_ask = h = 2 tick. "
    "Envanter farkındalığı veya skew mekanizması içermemektedir."
))

add_heading(dst, "3.5.2 Avellaneda-Stoikov Stratejisi", level=3)
add_normal(dst, "Rezervasyon fiyatı ve yarı-spread şu formüllerle hesaplanmaktadır:")
add_normal(dst, "r = mid - q · γ · σ² · τ")
add_normal(dst, "d = ½ · γ · σ² · τ + (1/γ) · ln(1 + γ/k_price)")
add_normal(dst, (
    "Burada γ = 0.01 risk aversion parametresi, k_price = k_ticks / tick_size fiyat cinsinden yoğunluk "
    "azalma parametresidir. Delta değerleri [1, 25] tick aralığında kırpılmaktadır."
))

# 3.6 Ödül Fonksiyonu
add_heading(dst, "3.6 Ödül Fonksiyonu", level=2)
add_normal(dst, "R_t = (equity_t - equity_{t-1}) - η · inv_t²")
add_normal(dst, (
    "Burada η = 0.001 envanter ceza katsayısıdır (η ablasyonu ile optimize edilmiştir). "
    "Kodda ek bir skew cezası terimi (c·|m|) bulunmakla birlikte, ana deneylerde c = 0 olarak sabitlenmiş "
    "ve bu terim etkisiz bırakılmıştır."
))

add_normal(dst, "")
add_normal(dst, (
    "Ek olarak, η'nın rejime bağlı kılındığı bir varyant denenmiştir: "
    "ηL = 0.0005, ηM = 0.001, ηH = 0.0025 (ηH = 5×ηL). Bu konfigürasyonda "
    "ajan, yüksek volatilite rejimlerinde daha ağır envanter cezasıyla karşılaşmakta; "
    "böylece rejime özgü davranış ödül kanalı aracılığıyla doğrudan teşvik edilmektedir. "
    "Sonuçlar Section 4.7'de sunulmaktadır."
))

# 3.7 Ablasyon Tasarımı
add_heading(dst, "3.7 Gözlem Uzayı ve Ablasyon Tasarımı", level=2)
add_normal(dst, (
    "Gözlem uzayı 6 boyutlu ve sabit mimariyle tasarlanmıştır: obs = [q_norm, σ̂, τ, r_L, r_M, r_H]. "
    "Beş PPO varyantı, use_sigma ve regime_source parametreleriyle kontrol edilmektedir:"
))
variant_headers = ["Varyant", "use_sigma", "regime_source", "Gözlemlediği Bilgi"]
variant_rows = [
    ["ppo_sigma_only", "True", "none", "σ̂ + sıfır rejim"],
    ["ppo_regime_only", "False", "hat", "rejim_hat + sıfır σ̂"],
    ["ppo_combined", "True", "hat", "σ̂ + rejim_hat"],
    ["ppo_oracle_pure", "False", "true", "rejim_true + sıfır σ̂"],
    ["ppo_oracle_full", "True", "true", "σ̂ + rejim_true"],
]
add_table(dst, variant_headers, variant_rows)
add_normal(dst, "Tüm varyantlar 20 seed, 1M timestep ile eğitilmiş ve OOS test seti üzerinde değerlendirilmiştir.")

# 3.8 PPO Eğitimi
add_heading(dst, "3.8 PPO Eğitimi ve Değerlendirme Protokolü", level=2)
add_normal(dst, "Stable-Baselines3 kütüphanesi kullanılarak PPO eğitimi gerçekleştirilmiştir. Hiperparametreler:")

hp_headers = ["Parametre", "Değer"]
hp_rows = [
    ["total_timesteps", "1.000.000"],
    ["learning_rate", "0.0003"],
    ["n_steps", "2048"],
    ["batch_size", "256"],
    ["n_epochs", "10"],
    ["gamma", "0.999"],
    ["gae_lambda", "0.95"],
    ["clip_range", "0.2"],
    ["ent_coef", "0.01"],
]
add_table(dst, hp_headers, hp_rows)
add_normal(dst, (
    "Her seed için episode uzunluğu 8000 adım olup kronolojik eğitim/test bölünmesi %70/%30 "
    "(train: 5600, test: 2400 adım) şeklinde uygulanmıştır. Değerlendirme yalnızca test bölümünden "
    "gerçekleştirilmektedir. Toplam 20 bağımsız tohum kullanılmıştır (seeds: 1-20)."
))

# ================================================================
# 4. SONUÇLAR VE TARTIŞMA
# ================================================================
add_heading(dst, "4. SONUÇLAR VE TARTIŞMA", level=1)

# 4.1 Ana Deney Sonuçları
add_heading(dst, "4.1 Ana Deney Sonuçları (20 Tohum, OOS)", level=2)
add_normal(dst, "Tablo 1, orijinal WP5 deneyinin (ppo_aware vs ppo_blind) out-of-sample sonuçlarını özetlemektedir.")

tbl1_h = ["Strateji", "Ort. Equity", "Std", "Ort. Sharpe", "inv_p99", "Fill Rate"]
tbl1_r = [
    ["AS", "5.05", "4.72", "0.105", "29.95", "0.444"],
    ["Naif", "4.49", "3.49", "0.126", "21.20", "0.119"],
    ["ppo_aware", "4.10", "0.78", "0.715", "2.00", "0.236"],
    ["ppo_blind", "4.42", "0.71", "0.740", "2.05", "0.232"],
]
add_table(dst, tbl1_h, tbl1_r)

add_normal(dst, (
    "Birinci bulgu — PPO'nun Sharpe Üstünlüğü: Her iki PPO varyantı da Sharpe oranı açısından klasik referans "
    "stratejilerini belirgin biçimde aşmaktadır (~6-7×). Bu üstünlüğün temel mekanizması envanter kontroldür: "
    "PPO varyantları inv_p99 ≈ 2 lot düzeyinde çalışırken, AS ve naif stratejiler sırasıyla 29.95 ve 21.20 "
    "değerlerine ulaşmaktadır. Mutlak equity açısından AS (5.05) ve naif (4.49) stratejilerin PPO varyantlarından "
    "(4.10-4.42) yüksek görünmesi, bu stratejilerin geniş spreadlerden yüksek brüt gelir elde etmesinden "
    "kaynaklanmaktadır; ancak söz konusu gelir yüksek envanter riski pahasına elde edilmektedir."
))

add_normal(dst, (
    "İkinci bulgu — Null Sonuç: ppo_aware, Sharpe bazında 20 tohumun 11'inde ppo_blind'ı geride bırakmıştır. "
    "Sharpe bazında istatistiksel olarak anlamlı fark gözlemlenmemiştir (paired t-test p = 0.261); ancak "
    "equity metriğinde ppo_blind lehine anlamlı fark mevcuttur (p = 0.023). Bu çalışmada risk-adjusted "
    "performans ölçütü olarak Sharpe oranı esas alınmaktadır."
))

add_normal(dst, (
    "Üçüncü bulgu — PPO Varyans Avantajı: PPO varyantlarının equity std değerleri (0.71-0.78) AS (4.72) ve "
    "naif (3.49) stratejilerine kıyasla belirgin biçimde düşüktür."
))

# Şekil 1
dst.add_picture(r"C:\Users\onure\Desktop\THESIS\results\plots\thesis\fig1_sharpe_inv.png", width=Inches(5.5))
add_normal(dst, (
    "Şekil 1. Out-of-sample performans özeti (20 tohum, ortalama ± std). Sol: Sharpe oranı. Sağ: envanter riski "
    "(inv_p99). Şekil, PPO'nun üstünlüğünün kaynağını görsel olarak ortaya koymaktadır. Sol panelde her iki PPO "
    "varyantının Sharpe oranı AS ve Naive stratejileri belirgin biçimde aşmaktadır. Sağ panelde PPO ajanları "
    "inv_p99 ≈ 2 lot düzeyinde çalışırken klasik stratejiler 20-30 lot envanter taşımaktadır."
))

# Üç hipotez
add_normal(dst, "Bu null sonuç için üç açıklayıcı hipotez öne sürülmektedir:")
dst.add_paragraph(
    "(a) Gözlem uzayında σ̂_t değeri hali hazırda mevcut olduğundan, one-hot rejim kodlaması "
    "σ̂_t'den türetilmiş bilgiyi yedekli biçimde sunmakta ve marjinal katkı sağlamamaktadır.",
    style='List Paragraph'
)
dst.add_paragraph(
    "(b) Ödül fonksiyonu rejime özgü davranışı doğrudan teşvik etmediğinden, rejim bilgisini "
    "performansa dönüştürmek için sinyal yetersiz kalmaktadır.",
    style='List Paragraph'
)
dst.add_paragraph(
    "(c) Detector robustness deneyi bu hipotezi geçersiz kılmaktadır: %81.8 doğruluklu HMM "
    "dedektörü ile dahi null sonuç korunmaktadır (p = 0.082). Dolayısıyla tespit kalitesi "
    "belirleyici değildir; açıklayıcı güç (a) hipotezinde yoğunlaşmaktadır.",
    style='List Paragraph'
)

# Şekil 2
dst.add_picture(r"C:\Users\onure\Desktop\THESIS\results\plots\thesis\fig2_paired_seed.png", width=Inches(5.5))
add_normal(dst, (
    "Şekil 2. Seed bazında PPO-aware ve PPO-blind karşılaştırması (her nokta bir seed, n=20). Kesikli çizgi y=x "
    "referansıdır. Sol: Sharpe (paired t-test p=0.261). Sağ: Final equity (paired t-test p=0.023). "
    "Sol panelde noktaların y=x çizgisi etrafında simetrik dağılımı Sharpe bazında istatistiksel anlamlılık "
    "eşiğine ulaşılmadığını göstermektedir. Sağ panelde ise noktaların çoğunluğunun y=x çizgisinin altında "
    "kalması, equity metriğinde PPO-blind'ın sistematik avantajına işaret etmektedir."
))

# 4.2 Pure Ablation
add_heading(dst, "4.2 Pure Ablation Deneyi (Sinyal Yedekliliği Testi)", level=2)
add_normal(dst, (
    "Motivasyon: Danışmanın temel eleştirisi, σ̂'nin her iki ajanda da gözlem uzayında mevcut olması "
    "nedeniyle aware-blind karşılaştırmasının confounded olduğuydu. Bu deneyle, σ̂'nin çıkarıldığı ve "
    "rejim etiketinin tek başına test edildiği pure ablation gerçekleştirilmiştir."
))

add_normal(dst, "Tablo 2, 5 varyantın 20 seed ortalamasını göstermektedir.")

tbl2_h = ["Varyant", "Sharpe (ort. ± std)", "Final Equity (ort. ± std)", "inv_p99 (ort. ± std)"]
tbl2_r = [
    ["ppo_sigma_only", "0.753 ± 0.111", "4.55 ± 0.65", "1.9 ± 0.7"],
    ["ppo_oracle_full", "0.722 ± 0.130", "4.07 ± 0.61", "2.0 ± 0.5"],
    ["ppo_regime_only", "0.698 ± 0.145", "4.01 ± 0.68", "2.1 ± 1.5"],
    ["ppo_combined", "0.696 ± 0.134", "3.91 ± 0.72", "1.8 ± 0.5"],
    ["ppo_oracle_pure", "0.684 ± 0.111", "3.93 ± 0.60", "1.9 ± 0.6"],
    ["naive", "0.127 ± 0.092", "4.49 ± 3.49", "21.2 ± 12.2"],
    ["AS", "0.105 ± 0.082", "5.05 ± 4.72", "29.9 ± 12.6"],
]
add_table(dst, tbl2_h, tbl2_r)

# thesis_23 extension: 5-variant visual summary
dst.add_page_break()
dst.add_picture(r"C:\Users\onure\Desktop\THESIS\results\plots\thesis_23\fig6_ablation_summary.png", width=Inches(5.5))
add_normal(dst, (
    "Şekil 3. Beş PPO varyantı için pure ablation özeti (20 tohum, ortalama ± std). Sol panel Sharpe, "
    "orta panel final equity, sağ panel inv_p99 değerlerini göstermektedir. AS ve naive referansları "
    "bu görsele dahil edilmemiştir; amaç explicit rejim etiketi, oracle rejim etiketi ve sigma_hat "
    "sinyalinin marjinal katkısını PPO varyantları içinde doğrudan karşılaştırmaktır."
))

add_normal(dst, (
    "Bu görsel, Tablo 2'deki ana örüntüyü özetlemektedir: sigma_only en yüksek Sharpe "
    "ortalamasına sahiptir ve oracle_full ya da combined varyantlarının üzerinde anlamlı bir iyileşme "
    "gözlenmemektedir. Bu desen, explicit rejim etiketinin sigma_hat üzerinde ek risk-adjusted bilgi "
    "taşımadığı yönündeki signal redundancy yorumuyla uyumludur."
))

add_normal(dst, (
    "Kritik Bulgu: ppo_sigma_only en yüksek Sharpe değerini (0.753) elde etmiştir — oracle_full'ı "
    "(σ̂ + regime_true, Sharpe = 0.722) ve combined'ı (σ̂ + regime_hat, Sharpe = 0.696) geçmektedir. "
    "Mükemmel rejim bilgisi bile σ̂'nin sağladığı bilginin üzerine anlamlı katkı sağlamamaktadır."
))

# 4.3 İstatistiksel Testler
add_heading(dst, "4.3 İstatistiksel Testler (Ablasyon)", level=2)
add_normal(dst, (
    "Tablo 3, 12 karşılaştırmanın eşleştirilmiş t-testi sonuçlarını göstermektedir "
    "(n=20 seed, Bonferroni düzeltmesi: α = 0.01/12 = 0.00083)."
))

tbl3_h = ["Karşılaştırma", "Metrik", "t", "p-değeri", "Cohen's d", "Anlamlı?"]
tbl3_r = [
    ["sigma_only vs combined", "sharpe", "2.014", "5.83e-02", "0.450", "Hayır"],
    ["sigma_only vs combined", "equity", "3.334", "3.49e-03", "0.746", "Hayır"],
    ["sigma_only vs oracle_full", "sharpe", "1.651", "1.15e-01", "0.369", "Hayır"],
    ["sigma_only vs oracle_full", "equity", "3.789", "1.24e-03", "0.847", "Hayır"],
    ["oracle_full vs combined", "sharpe", "1.062", "3.01e-01", "0.238", "Hayır"],
    ["oracle_full vs combined", "equity", "0.939", "3.60e-01", "0.210", "Hayır"],
    ["sigma_only vs regime_only", "sharpe", "2.283", "3.41e-02", "0.510", "Hayır"],
    ["sigma_only vs regime_only", "equity", "4.510", "2.40e-04", "1.008", "EVET"],
    ["combined vs AS", "sharpe", "16.615", "8.99e-13", "3.715", "EVET"],
    ["combined vs AS", "final_equity", "-1.064", "3.01e-01", "-0.238", "Hayır"],
    ["combined vs naive", "sharpe", "17.878", "2.42e-13", "3.998", "EVET"],
    ["combined vs naive", "final_equity", "-0.728", "4.76e-01", "-0.163", "Hayır"],
]
add_table(dst, tbl3_h, tbl3_r)

add_normal(dst, (
    "Oracle Paradoksu Özeti: oracle_full vs combined farkı Sharpe'ta p = 0.301, equity'de p = 0.360 — "
    "istatistiksel olarak hiçbir anlamlılık yok. Mükemmel etiket, tahmini etiketten ayırt edilemez."
))

add_normal(dst, (
    "TOST equivalence test (bounds ±0.10) confirms practical equivalence between ppo_sigma_only and "
    "ppo_oracle_full (p=0.0005). The 95% confidence interval for the Sharpe difference is "
    "[−0.007, +0.068], which falls within the equivalence bound. This provides positive evidence that "
    "the performance difference is practically negligible, not merely statistically inconclusive."
))

dst.add_picture(r"C:\Users\onure\Desktop\THESIS\results\plots\thesis_23\fig7_oracle_paired_seed.png", width=Inches(5.5))
add_normal(dst, (
    "Şekil 4. Oracle paradoksunun seed bazında görselleştirilmesi. Sol panel doğrudan karşılaştırmayı "
    "gösterir: ppo_sigma_only ile ppo_oracle_full Sharpe değerleri aynı test seed'leri üzerinde "
    "eşleştirilmiştir (paired t-test p = 0.115). Sağ panel oracle_full ile combined varyantlarını "
    "karşılaştırır (p = 0.301). Kesikli çizgi y=x referansıdır."
))

add_normal(dst, (
    "Bu şeklin birincil yorumu sol paneldir: mükemmel rejim etiketlerine sahip oracle_full, sigma_only "
    "varyantını Sharpe bazında istatistiksel olarak anlamlı biçimde geçememektedir. Dolayısıyla oracle "
    "varyantı, dedektör gürültüsü eleştirisini sınırlayan güçlü bir üst-sınır testi sunar; ancak p=0.115 "
    "sonucu, kanıtı deterministik bir eşdeğerlik iddiası olarak değil, anlamlı üstünlük bulunmaması "
    "olarak yorumlanmalıdır."
))

# 4.4 Eylem Analizi
add_heading(dst, "4.4 Eylem Analizi (Davranışsal Farklılaşma)", level=2)

add_normal(dst, "Tablo 4. Rejim Bazında Eylem Dağılımı (Ortalama ± Std, 20 Tohum)")
tbl_action_h = ["Strateji", "Rejim", "Ort. h", "Std h", "Ort. m", "Std m", "P(h=5)"]
tbl_action_r = [
    ["AS", "L", "1.00", "0.00", "-0.01", "0.02", "0.000"],
    ["AS", "M", "1.00", "0.00", "-0.02", "0.03", "0.000"],
    ["AS", "H", "1.00", "0.00", "-0.02", "0.03", "0.000"],
    ["Naif", "L", "2.00", "0.00", "0.00", "0.00", "0.000"],
    ["Naif", "M", "2.00", "0.00", "0.00", "0.00", "0.000"],
    ["Naif", "H", "2.00", "0.00", "0.00", "0.00", "0.000"],
    ["ppo_aware", "L", "1.43", "0.63", "-0.05", "0.22", "0.000"],
    ["ppo_aware", "M", "1.68", "0.72", "-0.01", "0.25", "0.008"],
    ["ppo_aware", "H", "1.74", "0.71", "+0.00", "0.22", "0.000"],
    ["ppo_blind", "L", "1.39", "0.47", "-0.03", "0.24", "0.000"],
    ["ppo_blind", "M", "1.42", "0.41", "-0.00", "0.23", "0.000"],
    ["ppo_blind", "H", "1.60", "0.42", "-0.00", "0.17", "0.000"],
]
add_table(dst, tbl_action_h, tbl_action_r)

add_normal(dst, (
    "ppo_aware ajanı, rejimler arasında sınırlı davranışsal farklılaşma sergilemektedir. Tüm rejimlerde "
    "ortalama half-spread h ≈ 1.4-1.8 tick aralığında seyretmekte, skew değerleri ise m ≈ 0 civarında "
    "kalmaktadır. ppo_blind ile karşılaştırıldığında anlamlı bir politika ayrışması gözlemlenmemektedir."
))

# Şekil 5
dst.add_picture(r"C:\Users\onure\Desktop\THESIS\results\plots\thesis\fig5_action_analysis.png", width=Inches(5.5))
add_normal(dst, (
    "Şekil 5. Rejim bazında eylem dağılımı (20 tohum, ortalama ± std). Üst: half-spread h. Alt: skew m. "
    "Sol sütun: PPO-aware. Sağ sütun: PPO-blind. Half-spread h değerleri L'den H'ye hafif artış göstermekte; "
    "ancak yüksek varyans nedeniyle bu artış istatistiksel olarak güçlü değildir."
))

# 4.5 Rejim Bazlı Performans Analizi
add_heading(dst, "4.5 Rejim Bazlı Performans Analizi", level=2)

add_normal(dst, (
    "Tablo 5, test döneminde üç volatilite rejiminde (L, M, H) elde edilen strateji performanslarını "
    "özetlemektedir."
))

add_normal(dst, (
    "Not: Rejim bazlı kırılım, gerçek rejim etiketleri (regime_true) üzerinden hesaplanmıştır; ajan ise "
    "yalnızca tahmini rejimi (regime_hat) gözlemlemektedir. Bu tablo, ex-post rejim attribution'ı "
    "sunmaktadır; ajan karar anında yalnızca regime_hat'i gözlemlemekte olup gerçek rejim bilgisine "
    "erişimi bulunmamaktadır."
))

tbl_regime_h = ["Strateji", "Rejim", "Sharpe", "inv_p99", "Fill Rate"]
tbl_regime_r = [
    ["AS", "L", "0.330", "22.26", "0.404"],
    ["AS", "M", "0.197", "27.49", "0.438"],
    ["AS", "H", "0.038", "27.66", "0.497"],
    ["naive", "L", "0.269", "16.35", "0.105"],
    ["naive", "M", "0.150", "18.87", "0.115"],
    ["naive", "H", "0.143", "19.20", "0.143"],
    ["PPO-aware", "L", "0.927", "1.85", "0.226"],
    ["PPO-aware", "M", "0.799", "1.60", "0.225"],
    ["PPO-aware", "H", "0.464", "1.79", "0.250"],
    ["PPO-blind", "L", "0.924", "2.00", "0.220"],
    ["PPO-blind", "M", "0.814", "1.88", "0.229"],
    ["PPO-blind", "H", "0.526", "1.91", "0.252"],
]
add_table(dst, tbl_regime_h, tbl_regime_r)

add_normal(dst, (
    "PPO ajanlarının her üç rejimde de AS ve naive baseline'larına kıyasla çok daha düşük envanter riski "
    "(inv_p99 ≈ 2 vs. 20-28) sergilediği görülmektedir. Her iki PPO ajanı da düşük volatilite (L) "
    "rejiminde en yüksek Sharpe değerlerini kaydetmiştir."
))

# Şekil 6
dst.add_picture(r"C:\Users\onure\Desktop\THESIS\results\plots\thesis\fig3_regime_sharpe.png", width=Inches(5.5))
add_normal(dst, (
    "Şekil 6. Volatilite rejimine göre Sharpe oranı (20 tohum, ortalama ± std). PPO'nun üstünlüğünün "
    "tüm volatilite rejimlerinde geçerli olduğu görülmektedir. Yüksek volatilite (H) rejiminde tüm "
    "stratejilerin Sharpe değerleri düşmekte; ancak PPO varyantları bu rejimde dahi klasik stratejileri "
    "belirgin biçimde geride bırakmaktadır."
))

# 4.6 Detector Robustness
add_heading(dst, "4.6 Detector Robustness Analizi", level=2)

tbl_det_h = ["Dedektör", "Doğruluk", "ppo_aware Sharpe", "ppo_blind Sharpe", "Sharpe p-değeri"]
tbl_det_r = [
    ["rv_baseline", "%60.7", "0.718", "0.753", "0.114"],
    ["rv_dwell", "%60.4", "0.716", "0.753", "0.110"],
    ["HMM", "%81.8", "0.719", "0.753", "0.082"],
]
add_table(dst, tbl_det_h, tbl_det_r)

add_normal(dst, (
    "Not: ppo_blind bu tabloda 0.753 Sharpe değeriyle görünmektedir; bu değer, farklı dedektörlerle "
    "yeniden eğitilen modellerin (wp5-detector-full run'ı) ortalamasını yansıtmaktadır. Section 4.1'deki "
    "0.740 değeri ise orijinal wp5-eval-main run'ından elde edilmiştir."
))

add_normal(dst, (
    "ANOVA testi: F = 0.003, p = 0.997 — ppo_aware performansı dedektör seçimine göre anlamlı biçimde "
    "değişmemektedir. Null sonuç, zayıf tespit kalitesinden kaynaklanmamaktadır."
))

# Şekil 7
dst.add_picture(r"C:\Users\onure\Desktop\THESIS\results\plots\thesis\fig4_detector_robustness.png", width=Inches(5.5))
add_normal(dst, (
    "Şekil 7. Dedektör bazında ortalama Sharpe (20 tohum). PPO-blind Sharpe değeri tüm dedektörlerde "
    "0.753 olarak sabit kalmakta; PPO-aware ise dedektör seçiminden bağımsız olarak ~0.717-0.719 "
    "düzeyinde seyretmektedir."
))

# 4.7 Rejime Koşullu Envanter Cezası Deneyi (moved BEFORE Section 5)
add_heading(dst, "4.7 Rejime Koşullu Envanter Cezası Deneyi", level=2)
add_normal(dst, (
    "Danışmanın önerisi doğrultusunda, null sonucun ödül tasarımı boyutunda da sınanması "
    "amacıyla rejime koşullu envanter cezası deneyi gerçekleştirilmiştir. Sabit η = 0.001 "
    "yerine ηL = 0.0005, ηM = 0.001, ηH = 0.0025 konfigürasyonu kullanılmış; beş varyant "
    "20 tohum, 1M timestep ile eğitilmiştir."
))
add_normal(dst, "Tablo 6, sonuçları özetlemektedir.")
add_table(dst,
    ["Varyant", "Ort. Sharpe", "Std Sharpe", "Ort. Equity"],
    [
        ["ppo_sigma_only", "0.714", "0.119", "3.91"],
        ["ppo_combined",   "0.629", "0.147", "3.50"],
        ["ppo_oracle_full","0.638", "0.158", "3.38"],
        ["ppo_oracle_pure","0.578", "0.247", "3.45"],
        ["ppo_regime_only","0.513", "0.238", "3.07"],
    ]
)

dst.add_picture(r"C:\Users\onure\Desktop\THESIS\results\plots\thesis_23\fig8_eta_regime_summary.png", width=Inches(5.5))
add_normal(dst, (
    "Şekil 8. Rejime koşullu envanter cezası deneyinde PPO varyantlarının Sharpe ve final equity "
    "sonuçları (20 tohum, ortalama ± std). Görsel, sigma_only ile combined varyantları arasındaki "
    "eşleştirilmiş karşılaştırmayı özellikle vurgular: Sharpe için p = 0.0016, final equity için "
    "p = 0.008."
))

add_normal(dst, (
    "Bu görsel, rejime koşullu ödül tasarımının explicit rejim etiketini faydalı hale getirmediğini "
    "göstermektedir. combined varyantı hem sigma_hat hem tahmini rejim etiketini gözlemlerken, "
    "sigma_only Sharpe bazında istatistiksel olarak daha güçlü kalmaktadır. Bu sonuç ödül kanalı "
    "üzerinden yapılan rejim hassaslaştırmasının signal redundancy yorumunu ortadan kaldırmadığını "
    "destekler."
))

add_normal(dst, (
    "Eşleştirilmiş t-testi (ppo_combined vs ppo_sigma_only): Sharpe p = 0.0016, "
    "Equity p = 0.008 — her iki metrikte ppo_sigma_only lehine istatistiksel olarak "
    "anlamlı fark. Rejime koşullu η, ppo_combined'ın performansını artırmak yerine "
    "anlamlı biçimde düşürmüştür. ppo_combined gözlem uzayında σ̂ aracılığıyla zaten "
    "rejim bilgisine sahipken, rejime koşullu η bu örtük sinyali gereksiz biçimde "
    "pekiştirmiş ve ajan optimizasyonunu zorlaştırmıştır. Bu bulgu, signal redundancy "
    "argümanını ödül tasarımı boyutunda da desteklemektedir."
))

# ================================================================
# CHANGE D — Section 4.8: Model Misspecification Robustness Deneyi
# ================================================================
add_heading(dst, "4.8 Model Misspecification Robustness Deneyi", level=2)

add_normal(dst, (
    "Motivasyon: Mevcut sentetik ortamda fill model parametreleri A ve k tüm rejimler için sabit "
    "tutulmaktadır. Bir eleştiri olarak, gerçek piyasalarda yüksek volatilite rejimlerinde likidite "
    "dinamiklerinin farklılaşabileceği öne sürülebilir. Bu deneyle, fill parametrelerinin rejime bağlı "
    "olduğu daha zorlu bir ortamda null result'ın geçerliliği test edilmektedir."
))

add_normal(dst, (
    "Tasarım: MMSimulator generic kalmış; env.py içinde her adımda regime_true'ya göre ExecParams "
    "override edilmiştir. Mild misspecification parametreleri: L rejimi A=4.0, k=1.8; M rejimi A=5.0, "
    "k=1.5 (baz değer); H rejimi A=6.0, k=1.2. Bu parametre konfigürasyonu, yüksek volatilite "
    "rejiminde gerçekçi bir likidite daralmasını değil, PPO ajanlarının değişen dolum dinamiklerine "
    "karşı dayanıklılığını test etmeye yönelik adversarial bir parametre şokunu temsil etmektedir. "
    "Aynı 5-varyant (sigma_only, regime_only, combined, "
    "oracle_pure, oracle_full), 20 tohum, 1M timestep eğitim protokolü uygulanmıştır."
))

add_normal(dst, "Tablo 7. Model Misspecification (Mild) OOS Sonuçları (20 Tohum)")
add_table(dst,
    ["Varyant", "Ort. Sharpe"],
    [
        ["ppo_sigma_only",  "0.685"],
        ["ppo_oracle_full", "0.682"],
        ["ppo_combined",    "0.651"],
        ["ppo_regime_only", "0.634"],
        ["ppo_oracle_pure", "0.602"],
    ]
)

dst.add_picture(r"C:\Users\onure\Desktop\THESIS\results\plots\thesis_23\fig9_misspec_summary.png", width=Inches(5.5))
add_normal(dst, (
    "Şekil 9. Mild model misspecification ortamında PPO varyantlarının Sharpe karşılaştırması "
    "(20 tohum, ortalama ± std). Fill parametreleri A ve k rejime bağlı olacak şekilde değiştirilmiştir; "
    "sigma_only ile oracle_full arasındaki eşleştirilmiş Sharpe karşılaştırması p = 0.881'dir."
))

add_normal(dst, (
    "Bu görsel, model misspecification sonucunu özetlemektedir: rejime bağlı A/k altında "
    "sigma_only ve oracle_full neredeyse aynı Sharpe düzeyinde kalmaktadır. Bu bulgu yalnızca bu mild "
    "misspecification tasarımı için bir robustness sonucu olarak yorumlanmalıdır; tüm piyasa model "
    "yanlış belirlemeleri altında evrensel eşdeğerlik iddiası değildir."
))

add_normal(dst, (
    "Eşleştirilmiş t-testi sonuçları: ppo_sigma_only vs ppo_oracle_full: p=0.881; "
    "ppo_sigma_only vs ppo_combined: p=0.217; ppo_sigma_only vs ppo_oracle_pure: p=0.098. "
    "Hiçbir karşılaştırma istatistiksel anlamlılık eşiğine ulaşmamıştır."
))

add_normal(dst, (
    "TOST equivalence test (bounds ±0.05) further confirms practical equivalence under model "
    "misspecification (p=0.039). The 95% confidence interval [−0.048, +0.056] lies entirely within "
    "the ±0.10 bound. Cohen's d = 0.035 indicates negligible effect size, strengthening the signal "
    "redundancy interpretation across both standard and adversarial environments."
))

add_normal(dst, (
    "Bu bulgular, signal redundancy argümanının çevresel model parametrelerindeki değişimlere karşı "
    "sağlam olduğunu ortaya koymaktadır. Fill dinamiklerinin rejime bağlı olduğu daha zorlu bir ortamda "
    "dahi ppo_sigma_only, mükemmel rejim bilgisine sahip oracle_full'dan istatistiksel olarak ayırt "
    "edilemez performans göstermiştir (p=0.881)."
))

# ================================================================
# 5. SONUÇ — CHANGE E: Bulgu 6 eklendi
# ================================================================
add_heading(dst, "5. SONUÇ", level=1)

add_normal(dst, "Bu çalışma, volatilite rejimleri altında pekiştirmeli öğrenme tabanlı piyasa yapıcılığını araştırmaktadır. Temel bulgular:")

add_normal(dst, (
    "Bulgu 1 — PPO üstünlüğü: Tüm PPO varyantları Sharpe oranı açısından Avellaneda-Stoikov ve naif "
    "referans stratejilerini yaklaşık 6-7 kat geçmiştir. Bu üstünlüğün temel kaynağı envanter yönetimi "
    "verimliliğidir: PPO inv_p99 ≈ 2 lot, baseline'lar 20-30 lot."
))

add_normal(dst, (
    "Bulgu 2 — Null Sonuç: Rejim bilgisinin performansı artırdığı hipotezi istatistiksel olarak "
    "desteklenmemiştir. Sharpe bazında istatistiksel olarak anlamlı fark gözlemlenmemiştir "
    "(paired t-test p = 0.261); equity metriğinde ppo_blind lehine anlamlı fark mevcuttur "
    "(p = 0.023). Bu null sonuç, "
    "gözlem uzayında hali hazırda bulunan σ̂_t'nin rejim bilgisini örtük olarak içermesiyle açıklanmaktadır."
))

add_normal(dst, (
    "Bulgu 3 — Pure Ablation (Sinyal Yedekliliği Testi): ppo_sigma_only en yüksek performansı elde "
    "etmiştir — mükemmel rejim etiketlerine sahip oracle_full ajanı dahi sigma_only'yi Sharpe bazında "
    "istatistiksel olarak anlamlı biçimde geçememiştir (p = 0.115). Bonferroni-düzeltilmiş eşik altında "
    "equity farkı da istatistiksel olarak anlamlı kabul edilmemiştir."
))

add_normal(dst, (
    "Bulgu 4 — Dedektör Bağımsızlığı: Null sonuç, %60.4'ten %81.8'e kadar üç farklı dedektör "
    "performansında tutarlı biçimde korunmuştur (ANOVA: F = 0.003, p = 0.997)."
))

add_normal(dst, (
    "Bulgu 5 — Rejime Koşullu Envanter Cezası: Basit rejime koşullu envanter cezası "
    "(ηH = 5×ηL), explicit rejim bilgisini faydalı hale getirmemiştir; sigma_only "
    "varyantı istatistiksel olarak anlamlı biçimde güçlü kalmıştır (p = 0.0016)."
))

# CHANGE E: Bulgu 6
add_normal(dst, (
    "Bulgu 6 — Model Misspecification Robustness: Fill model parametrelerinin (A, k) rejime bağlı "
    "kılındığı zorlaştırılmış ortamda da null sonuç geçerliliğini korumuştur. ppo_sigma_only yine en "
    "yüksek performansı göstermiş; oracle_full ile aradaki fark istatistiksel olarak anlamsız kalmıştır "
    "(p=0.881). Bu bulgu, signal redundancy argümanının çevresel koşullara karşı sağlamlığını "
    "(robustness) doğrulamaktadır."
))

# CHANGE F: Gelecek Çalışmalar — item (2) updated
add_heading(dst, "Gelecek Çalışmalar", level=2)
dst.add_paragraph(
    "(1) Basit rejime koşullu envanter cezası denendi ve fayda sağlamadı; "
    "daha zengin rejime koşullu ödül tasarımları gelecekte incelenebilir.",
    style='List Paragraph'
)
dst.add_paragraph(
    "(2) Model misspecification strong variant — daha sert parametre farklılaşması "
    "(L: A=3.0/k=2.0, H: A=8.0/k=1.0) ile robustness testinin genişletilmesi.",
    style='List Paragraph'
)
dst.add_paragraph("(3) Paper yayını: iyi nitelendirilmiş null result ile kontrollü RL vs. klasik piyasa yapıcılığı kıyaslaması.", style='List Paragraph')

# ================================================================
# KAYNAKLAR
# ================================================================
add_heading(dst, "KAYNAKLAR", level=1)
add_normal(dst, "[1] Avellaneda, M., & Stoikov, S. (2008). High-frequency trading in a limit order book. Quantitative Finance, 8(3), 217-224.")
add_normal(dst, "[2] Schulman, J., Wolski, F., Dhariwal, P., Radford, A., & Klimov, O. (2017). Proximal policy optimization algorithms. arXiv preprint arXiv:1707.06347.")
add_normal(dst, "[3] Spooner, T., Fearnley, J., Savani, R., & Koukorinis, A. (2018). Market making via reinforcement learning. In Proceedings of AAMAS (pp. 434-442).")
add_normal(dst, "[4] Raffin, A., Hill, A., Gleave, A., Kanervisto, A., Ernestus, M., & Dormann, N. (2021). Stable-baselines3: Reliable reinforcement learning implementations. JMLR, 22(268), 1-8.")
add_normal(dst, "[5] Cont, R. (2001). Empirical characteristics of asset returns: Stylized facts. Quantitative Finance, 1(2), 223-236.")
add_normal(dst, "[6] Hamilton, J. D. (1989). A new approach to the economic analysis of nonstationary time series and the business cycle. Econometrica, 57(2), 357-384.")

# ================================================================
# APPENDIX B — Source Code File Index
# ================================================================
add_heading(dst, "Appendix B: Source Code File Index", level=1)
add_normal(dst, "This appendix lists all active source code files in the project repository. Files are organized by module.")

app_h = ["Dosya Yolu", "Açıklama"]
app_r = [
    ["run.py", "Ana iş paketi dispatcher'ı"],
    ["src/run_context.py", "RunContext, logger, seed yönetimi, config snapshot"],
    ["src/w0_smoke.py", "WP0 smoke test"],
    ["src/w1_as_baseline.py", "Avellaneda-Stoikov baseline stratejisi"],
    ["src/w1_compare.py", "WP1 strateji karşılaştırması"],
    ["src/w1_naive_sweep.py", "Naive sabit-spread sweep deneyi"],
    ["src/w3_sanity.py", "Gymnasium ortamı sanity check"],
    ["src/wp1/sim.py", "Piyasa simülatörü (ABM + Poisson dolum modeli)"],
    ["src/wp2/synth_regime.py", "Sentetik rejim üretimi (Markov zinciri)"],
    ["src/wp2/job_w2_synth.py", "WP2 job entry"],
    ["src/wp2/compare_detectors.py", "Rejim dedektörü karşılaştırması"],
    ["src/wp3/env.py", "Gymnasium MMEnv ortamı"],
    ["src/wp4/job_w4_ppo.py", "PPO eğitimi (Stable-Baselines3)"],
    ["src/wp5/job_w5_eval.py", "Out-of-sample değerlendirme (20 tohum)"],
    ["src/wp5/job_w5_ablation_eta.py", "η envanter ceza katsayısı ablasyon deneyi"],
    ["src/wp5/job_w5_ablation_skew.py", "Skew penalty ablasyon deneyi"],
    ["src/wp5/job_w5_detector_compare.py", "Detector robustness deneyi (120 model)"],
    ["src/wp5/analyze_actions.py", "Eylem dağılımı analizi"],
    ["src/wp5/figure_thesis.py", "Tez figürlerinin üretimi"],
    ["src/wp5/figure_thesis_23.py", "thesis_23 ek figürlerinin üretimi"],
    ["src/wp5/stats_detector_robustness.py", "İstatistiksel testler (paired t-test ve ANOVA)"],
    ["scripts/gen_thesis_docx.py", "Tez DOCX üretim scripti"],
    ["scripts/gen_thesis_draft.js", "Tez taslak üretim scripti (Node.js)"],
]
add_table(dst, app_h, app_r)

# ================================================================
# SAVE
# ================================================================
dst.save("manuscript/thesis_24.docx")
print("thesis_24.docx saved successfully.")
