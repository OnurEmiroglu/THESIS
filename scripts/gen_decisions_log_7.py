"""decisions_log_6.docx -> decisions_log_7.docx
Iki yeni karar eklendi: #39 TOST equivalence test (null result -> positive evidence)
ve #40 TOST threshold choice (+/-0.10 primary bound).
Numaralar +3 kaydirildi cunku #36-#38 zaten thesis_19/20/21 revizyonlari icin kullanilmis.
Ozet tablosu 18 -> 20 karara guncellendi."""

from __future__ import annotations
import sys, io, shutil, copy
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

SRC = Path("manuscript/decisions_log_6.docx")
DST = Path("manuscript/decisions_log_7.docx")


def _set_cell_text(cell, text, *, bold=None):
    """Overwrite a cell's text in its first paragraph, preserving styling of the first run."""
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].text = text
        for r in para.runs[1:]:
            r.text = ""
        if bold is not None:
            para.runs[0].bold = bold
    else:
        r = para.add_run(text)
        if bold is not None:
            r.bold = bold


def _build_karar_table_from_template(doc, template_tbl, number, title, wp, ikilem, secenekler, karar, neden, etki):
    """Deep-copy template table XML, swap in new text. Returns new table element."""
    new_xml = copy.deepcopy(template_tbl._tbl)
    from docx.table import Table
    new_tbl = Table(new_xml, doc.part)
    # Row 0: WP tag | #N  Title
    _set_cell_text(new_tbl.rows[0].cells[0], wp, bold=True)
    _set_cell_text(new_tbl.rows[0].cells[1], f"#{number}  {title}", bold=True)
    # Row 1: Ikilem
    _set_cell_text(new_tbl.rows[1].cells[0], "İkilem", bold=True)
    _set_cell_text(new_tbl.rows[1].cells[1], ikilem)
    # Row 2: Secenekler
    _set_cell_text(new_tbl.rows[2].cells[0], "Seçenekler", bold=True)
    _set_cell_text(new_tbl.rows[2].cells[1], secenekler)
    # Row 3: Karar
    _set_cell_text(new_tbl.rows[3].cells[0], "Karar", bold=True)
    _set_cell_text(new_tbl.rows[3].cells[1], karar)
    # Row 4: Neden
    _set_cell_text(new_tbl.rows[4].cells[0], "Neden", bold=True)
    _set_cell_text(new_tbl.rows[4].cells[1], neden)
    # Row 5: Etki
    _set_cell_text(new_tbl.rows[5].cells[0], "Etki", bold=True)
    _set_cell_text(new_tbl.rows[5].cells[1], etki)
    return new_xml


def _replace_in_paragraph(para, old_text, new_text):
    full = "".join(r.text for r in para.runs)
    if old_text not in full:
        return False
    new_full = full.replace(old_text, new_text)
    if para.runs:
        para.runs[0].text = new_full
        for r in para.runs[1:]:
            r.text = ""
    return True


def main():
    shutil.copy2(SRC, DST)
    doc = Document(DST)

    # Template: Karar #38 table (last 6x2 decision table) to preserve styling/widths
    # doc.tables ordering: index 36 = Karar #38 (6x2, WP5 block)
    template = doc.tables[36]

    # Build #39 TOST equivalence choice
    k39 = _build_karar_table_from_template(
        doc, template,
        number=39,
        title="Null result: t-test \"fark yok\" mu, TOST \"fark ihmal edilebilir\" mi?",
        wp="WP5",
        ikilem=(
            "Null result \"fark yok\" mu yoksa \"fark ihmal edilebilir\" mi olarak sunulmalı?"
        ),
        secenekler=(
            "A) Klasik t-test ile \"p > 0.05, fark bulunamadı\" demek  |  "
            "B) TOST equivalence test ile \"fark pratik olarak ihmal edilebilir aralıkta\" kanıtlamak"
        ),
        karar="B — TOST eklendi.",
        neden=(
            "\"Absence of evidence is not evidence of absence\" eleştirisini önlemek için. "
            "TOST, null result'ı pasif bir bulgu olmaktan çıkarıp aktif bir kanıta dönüştürür."
        ),
        etki=(
            "Ana ablasyon: TOST ±0.10 eşdeğer (p=0.0005), 95% CI [−0.007, +0.068]. "
            "thesis_24'te ilgili bölümlere eklendi."
        ),
    )

    # Build #40 TOST threshold
    k40 = _build_karar_table_from_template(
        doc, template,
        number=40,
        title="TOST eşiği (equivalence bound) nasıl belirlenmeli?",
        wp="WP5",
        ikilem="TOST eşiği (equivalence bound) nasıl belirlenmeli?",
        secenekler=(
            "A) ±0.05 (çok sıkı)  |  B) ±0.10 (orta, pratik anlamlılık için makul)  |  "
            "C) ±0.15 (çok gevşek)"
        ),
        karar="B — ±0.10 birincil eşik olarak seçildi.",
        neden=(
            "Ortalama Sharpe ~0.75 seviyesinde; 0.10 fark %13'lük göreli sapma demek. "
            "Bu, market making bağlamında pratik olarak ihmal edilebilir bir fark. "
            "Misspec ortamında ±0.05 bile eşdeğerlik sağlandı (p=0.039)."
        ),
        etki=(
            "İki deney için tutarlı eşik. Paper'da savunulabilir metodolojik seçim."
        ),
    )

    # Insert new tables BEFORE "Özet: En Kritik 18 Karar" heading
    body = doc.element.body
    W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    ozet_heading = None
    for child in body.iterchildren():
        if child.tag == W_NS + "p":
            pStyle = child.find(W_NS + "pPr/" + W_NS + "pStyle")
            if pStyle is not None and pStyle.get(W_NS + "val") == "Heading1":
                text = "".join(t.text or "" for t in child.iter(W_NS + "t"))
                if "Özet" in text:
                    ozet_heading = child
                    break
    if ozet_heading is None:
        raise RuntimeError("Özet heading not found")

    # Build a blank paragraph element (copy from doc's existing blank paragraph)
    from docx.oxml import OxmlElement
    def _blank_p():
        p = OxmlElement("w:p")
        return p

    # Insert: [blank-p, k39, blank-p, k40] before ozet_heading
    ozet_heading.addprevious(_blank_p())
    ozet_heading.addprevious(k39)
    ozet_heading.addprevious(_blank_p())
    ozet_heading.addprevious(k40)

    # Update Özet heading: "18 Karar" -> "20 Karar"
    for child in body.iterchildren():
        if child.tag == W_NS + "p":
            text = "".join(t.text or "" for t in child.iter(W_NS + "t"))
            if "Özet: En Kritik 18 Karar" in text:
                # find the paragraph wrapper and replace text
                from docx.text.paragraph import Paragraph
                p_obj = Paragraph(child, doc.part)
                _replace_in_paragraph(p_obj, "En Kritik 18 Karar", "En Kritik 20 Karar")

    # Update description paragraph that follows the heading if it mentions "18"
    for para in doc.paragraphs:
        if "ChatGPT ve Claude" in para.text and "18" in para.text:
            _replace_in_paragraph(para, "18 karar", "20 karar")

    # Add rows to summary table (last table) for #19 TOST choice, #20 TOST threshold
    sum_tbl = doc.tables[-1]
    # Template row is R1 (first data row) — copy its XML to preserve styling
    template_row = sum_tbl.rows[1]._tr
    for num, title, wp in [
        ("19", "TOST equivalence test — null result'ı pasif bulgudan aktif kanıta çevirmek", "WP5"),
        ("20", "TOST eşiği ±0.10 seçimi — pratik anlamlılık için makul bound", "WP5"),
    ]:
        new_row = copy.deepcopy(template_row)
        sum_tbl._tbl.append(new_row)
        from docx.table import _Row
        row_obj = _Row(new_row, sum_tbl)
        _set_cell_text(row_obj.cells[0], num)
        _set_cell_text(row_obj.cells[1], title)
        _set_cell_text(row_obj.cells[2], wp)

    doc.save(DST)
    print(f"decisions_log_7.docx saved ({DST.stat().st_size:,} bytes)")

    # Verification
    doc2 = Document(DST)
    print(f"\n=== Verification ===")
    print(f"Total tables: {len(doc2.tables)}")
    for i, tbl in enumerate(doc2.tables[-4:-1], start=len(doc2.tables)-4):
        hdr = tbl.rows[0].cells[1].text[:70] if len(tbl.rows[0].cells) > 1 else ''
        print(f"  Table {i}: {hdr}")
    sum2 = doc2.tables[-1]
    print(f"Summary rows: {len(sum2.rows)}")
    print(f"  R{len(sum2.rows)-2}: {[c.text[:50] for c in sum2.rows[-2].cells]}")
    print(f"  R{len(sum2.rows)-1}: {[c.text[:50] for c in sum2.rows[-1].cells]}")


if __name__ == "__main__":
    main()
