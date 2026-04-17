"""decisions_log.docx -> decisions_log_2.docx: Yeni kararları ekle (#26-#31)."""

from __future__ import annotations
import copy as copy_mod
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, Emu
from docx.oxml.ns import qn
from lxml import etree

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

SRC = Path("manuscript/decisions_log.docx")
DST = Path("manuscript/decisions_log_2.docx")

# ── Styling constants (matched from original doc) ──────────────────
# Header table: col0 = WP tag (teal bg), col1 = title (blue bg)
HDR_COL0_W = "900"
HDR_COL0_FILL = "1ABC9C"
HDR_COL1_W = "8460"
HDR_COL1_FILL = "2E75B6"

# Detail table: col0 = label (light blue bg), col1 = content (lighter blue)
DET_COL0_W = "2000"
DET_COL0_FILL = "E8F0F8"
DET_COL1_W = "7360"
DET_COL1_FILL = "E8F4FD"

# Cell borders
BORDER_COLOR = "BBCFE0"

# Summary table
SUM_HDR_FILL = "1F4E79"
SUM_ROW_FILL = "E8F4FD"


def _set_cell_props(cell, width_dxa, fill_color):
    """Set cell width, borders, shading, margins to match original."""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = etree.SubElement(tc, qn("w:tcPr"))
        tc.insert(0, tcPr)

    # width
    for old in tcPr.findall(qn("w:tcW")):
        tcPr.remove(old)
    tcW = etree.SubElement(tcPr, qn("w:tcW"))
    tcW.set(qn("w:w"), width_dxa)
    tcW.set(qn("w:type"), "dxa")

    # borders
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    borders = etree.SubElement(tcPr, qn("w:tcBorders"))
    for side in ("top", "left", "bottom", "right"):
        b = etree.SubElement(borders, qn(f"w:{side}"))
        b.set(qn("w:val"), "single")
        b.set(qn("w:color"), BORDER_COLOR)
        b.set(qn("w:sz"), "1")

    # shading
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = etree.SubElement(tcPr, qn("w:shd"))
    shd.set(qn("w:fill"), fill_color)
    shd.set(qn("w:val"), "clear")

    # margins
    for old in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(old)
    tcMar = etree.SubElement(tcPr, qn("w:tcMar"))
    for side, val in [("top", "80"), ("left", "160"), ("bottom", "80"), ("right", "160")]:
        m = etree.SubElement(tcMar, qn(f"w:{side}"))
        m.set(qn("w:type"), "dxa")
        m.set(qn("w:w"), val)


def _styled_run(paragraph, text, bold=False, size_pt=10, font_name="Arial", color_hex=None):
    """Add a run with exact formatting."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = font_name
    if color_hex:
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor.from_string(color_hex)
    return run


def _set_table_width(tbl, width_dxa="9360"):
    """Set w:tblW on table."""
    tblPr = tbl._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = etree.SubElement(tbl._tbl, qn("w:tblPr"))
        tbl._tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tw = etree.SubElement(tblPr, qn("w:tblW"))
    tw.set(qn("w:type"), "dxa")
    tw.set(qn("w:w"), width_dxa)

    # borders
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    borders = etree.SubElement(tblPr, qn("w:tblBorders"))
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = etree.SubElement(borders, qn(f"w:{side}"))
        b.set(qn("w:val"), "single")
        b.set(qn("w:color"), "auto")
        b.set(qn("w:sz"), "4")


# ── Decision data ──────────────────────────────────────────────────

DECISIONS_PAGE9 = [
    {
        "wp": "WP5",
        "num": 26,
        "title": "Ablasyon tasarımı: kaç varyant yeterli?",
        "ikilem": "aware vs blind ikili karşılaştırması yeterliydi, ama danışman \"σ̂ her ikisinde de var, bu confounded\" dedi. Kaç varyant tasarlayalım?",
        "secenekler": "A) 2 varyant: aware vs blind  |  B) 5 varyant: sigma_only, regime_only, combined, oracle_pure, oracle_full",
        "karar": "B — 5 varyant tasarlandı ve 20 seed ile eğitildi.",
        "neden": "Danışmanın confounding eleştirisini doğrudan yanıtlamak için σ̂'yi tamamen çıkaran (regime_only, oracle_pure) ve mükemmel rejim bilgisi veren (oracle_pure, oracle_full) varyantlar gerekiyordu.",
        "etki": "Oracle paradoksu ortaya çıktı: oracle_full (Sharpe: 0.722), sigma_only'yi (Sharpe: 0.753) geçemedi. Bu, null result'ın en güçlü kanıtı oldu.",
    },
    {
        "wp": "WP5",
        "num": 27,
        "title": "Oracle paradoksu: anomali mi, bulgu mu?",
        "ikilem": "oracle_full sigma_only'den daha düşük Sharpe aldı. Bu bir bug mı yoksa gerçek bir akademik bulgu mu?",
        "secenekler": "A) Kodu kontrol et, bug ara  |  B) Bulgu olarak kabul et, istatistiksel olarak doğrula",
        "karar": "B — Önce kod doğrulandı (bug yok), sonra paired t-test ile p=0.301 (Sharpe), p=0.360 (equity) bulundu — istatistiksel olarak anlamlı fark yok.",
        "neden": "Mükemmel rejim bilgisi bile sigma_only'yi geçemiyorsa bu signal redundancy'nin doğrudan kanıtı. Anomali değil, tezin temel argümanının en sert testi.",
        "etki": "\"Oracle paradoksu\" tezin kritik bulgularından biri olarak konumlandırıldı. Danışmanın confounding eleştirisine nihai yanıt verildi.",
    },
    {
        "wp": "WP5",
        "num": 28,
        "title": "Detector robustness full run tamamlandı: sonuç ne?",
        "ikilem": "Full run (3 detector × 20 seed × 2 strateji = 120 model) tamamlandı. Sonuçlar pilot ile tutarlı mı?",
        "secenekler": "A) Pilot yeterliydi, full run gereksizdi  |  B) Full run null result'ı istatistiksel olarak güçlendirdi",
        "karar": "B — rv_baseline p=0.114, rv_dwell p=0.110, HMM p=0.082 (hepsi anlamlı değil). ANOVA: F=0.003, p=0.997.",
        "neden": "%81.8 doğruluklu HMM ile dahi null result korunuyor. \"Detection kalitesi düşük olduğu için null çıktı\" itirazı tamamen elendi.",
        "etki": "Null result üç bağımsız dedektörde tutarlı — signal redundancy argümanı istatistiksel olarak sağlam.",
    },
]

DECISIONS_PAGE10 = [
    {
        "wp": "Genel",
        "num": 29,
        "title": "Bonferroni düzeltmesi: kaç karşılaştırma sayacağız?",
        "ikilem": "Ablasyon tablosunda 10 satır var ama script başlangıçta α=0.01/12=0.00083 kullanıyordu. Doğru payda ne?",
        "secenekler": "A) 12 kullan (6 karşılaştırma × 2 metrik)  |  B) 10 kullan (gerçek satır sayısı)",
        "karar": "B — α=0.01/10=0.001 olarak düzeltildi.",
        "neden": "Tabloda fiilen 10 karşılaştırma var; payda gerçek test sayısıyla eşleşmeli. Pratik etki ihmal edilebilir (tüm \"EVET\" kararları aynı kaldı) ama iç tutarlılık zorunlu.",
        "etki": "thesis_15 → thesis_16 geçişinde düzeltildi. İstatistiksel doğruluk sağlandı.",
    },
    {
        "wp": "Genel",
        "num": 30,
        "title": "\"Kanıtlamaktadır\" dili: çok güçlü mü?",
        "ikilem": "Tezde \"rejim etiketinin bilgiyi yedekli sunduğunu kanıtlamaktadır\" gibi ifadeler vardı. Sentetik ortamda bu kadar kesin dil uygun mu?",
        "secenekler": "A) Güçlü dili koru, bulgunun ağırlığını vurgula  |  B) \"Tutarlıdır\", \"güçlü kanıt sunar\" gibi epistemik açıdan daha doğru ifadeler kullan",
        "karar": "B — Tüm \"kanıtlamaktadır\" ifadeleri yumuşatıldı.",
        "neden": "Ampirik çalışmalarda \"kanıtlar\" yerine \"destekler/ile tutarlıdır\" akademik standarttır. Danışman ve hakem yorumunda bu fark kritik olabilir.",
        "etki": "thesis_16'da dil tutarlı ve savunulabilir hale geldi.",
    },
    {
        "wp": "Genel",
        "num": 31,
        "title": "Şekil numaraları kaydı: nasıl önleriz?",
        "ikilem": "PNG dosyaları (fig3_regime_sharpe, fig4_detector_robustness, fig5_action_analysis) yanlış şekil açıklamalarına eşleştirilmişti.",
        "secenekler": "A) Manuel kontrol her sürümde  |  B) Script'te şekil-açıklama eşleşmesini yorumla belgele",
        "karar": "B — gen_thesis scripti her add_picture() çağrısına açıklayıcı yorum eklendi. thesis_16'da düzeltme yapıldı.",
        "neden": "Şekil kayması okuyucuya \"tez acele toparlanmış\" izlenimi verir. Akademik belgede görsel-metin tutarlılığı zorunlu.",
        "etki": "thesis_16 ile tüm 5 şekil doğru yerleşti. Versiyon kontrol yorumları eklendi.",
    },
]

SUMMARY_NEW_ROWS = [
    ("10", "5 varyantlı ablasyon + oracle paradoksu", "WP5"),
    ("11", "Detector robustness full run (120 model) null result teyidi", "WP5"),
    ("12", "Dil yumuşatma + istatistiksel iç tutarlılık (Bonferroni, şekil sırası)", "Yazım"),
]


def add_decision_block(doc, body, insert_before, d):
    """Build header + detail tables for one decision and insert before given element."""
    # ── Header table ──
    hdr_tbl = doc.add_table(rows=1, cols=2)
    _set_table_width(hdr_tbl)
    c0 = hdr_tbl.cell(0, 0)
    c0.paragraphs[0].clear()
    _styled_run(c0.paragraphs[0], d["wp"], bold=True, size_pt=10)
    _set_cell_props(c0, HDR_COL0_W, HDR_COL0_FILL)

    c1 = hdr_tbl.cell(0, 1)
    c1.paragraphs[0].clear()
    _styled_run(c1.paragraphs[0], f"#{d['num']}  {d['title']}", bold=True, size_pt=11)
    _set_cell_props(c1, HDR_COL1_W, HDR_COL1_FILL)

    # Move to correct position
    body.insert(list(body).index(insert_before), hdr_tbl._tbl)

    # ── Detail table ──
    labels = ["İkilem", "Seçenekler", "Karar", "Neden", "Etki"]
    values = [d["ikilem"], d["secenekler"], d["karar"], d["neden"], d["etki"]]
    det_tbl = doc.add_table(rows=5, cols=2)
    _set_table_width(det_tbl)
    for i, (lbl, val) in enumerate(zip(labels, values)):
        cl = det_tbl.cell(i, 0)
        cl.paragraphs[0].clear()
        _styled_run(cl.paragraphs[0], lbl, bold=True, size_pt=10)
        _set_cell_props(cl, DET_COL0_W, DET_COL0_FILL)

        cr = det_tbl.cell(i, 1)
        cr.paragraphs[0].clear()
        _styled_run(cr.paragraphs[0], val, bold=False, size_pt=10)
        _set_cell_props(cr, DET_COL1_W, DET_COL1_FILL)

    body.insert(list(body).index(insert_before), det_tbl._tbl)

    # ── Blank paragraph spacer ──
    spacer = doc.add_paragraph("")
    body.insert(list(body).index(insert_before), spacer._element)


def add_heading_before(doc, body, insert_before, text):
    """Add a Heading 1 paragraph before insert_before element, matching original XML."""
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    p_elem = OxmlElement("w:p")
    # pPr with style + spacing
    pPr = etree.SubElement(p_elem, qn("w:pPr"))
    pStyle = etree.SubElement(pPr, qn("w:pStyle"))
    pStyle.set(qn("w:val"), "Heading1")
    spacing = etree.SubElement(pPr, qn("w:spacing"))
    spacing.set(qn("w:before"), "360")
    spacing.set(qn("w:after"), "120")
    # Run
    r = etree.SubElement(p_elem, qn("w:r"))
    rPr = etree.SubElement(r, qn("w:rPr"))
    rFonts = etree.SubElement(rPr, qn("w:rFonts"))
    for attr in ("ascii", "cs", "eastAsia", "hAnsi"):
        rFonts.set(qn(f"w:{attr}"), "Arial")
    etree.SubElement(rPr, qn("w:b"))
    etree.SubElement(rPr, qn("w:bCs"))
    color = etree.SubElement(rPr, qn("w:color"))
    color.set(qn("w:val"), "1F4E79")
    sz = etree.SubElement(rPr, qn("w:sz"))
    sz.set(qn("w:val"), "32")
    szCs = etree.SubElement(rPr, qn("w:szCs"))
    szCs.set(qn("w:val"), "32")
    t = etree.SubElement(r, qn("w:t"))
    t.text = text

    body.insert(list(body).index(insert_before), p_elem)

    # Blank paragraph after heading
    spacer = doc.add_paragraph("")
    body.insert(list(body).index(insert_before), spacer._element)


def main():
    # Copy original -> new file
    shutil.copy2(SRC, DST)
    doc = Document(DST)
    body = doc.element.body

    # ── Find the "Özet" heading ──
    ozet_para = None
    for p in doc.paragraphs:
        if "Özet: En Kritik 9 Karar" in p.text:
            ozet_para = p
            break
    assert ozet_para is not None, "Could not find 'Özet: En Kritik 9 Karar'"
    ozet_elem = ozet_para._element

    # ── Page 10: Genel Proje Kararları — Tez Yazım Süreci ──
    # (Insert page 10 first so it ends up AFTER page 9 in final order)
    add_heading_before(doc, body, ozet_elem, "Genel Proje Kararları — Tez Yazım Süreci")
    for d in DECISIONS_PAGE10:
        add_decision_block(doc, body, ozet_elem, d)

    # ── Page 9: WP5 — Ablasyon ve İleri Analizler ──
    # Find the "Genel Proje Kararları — Tez Yazım Süreci" heading we just added
    genel_new_elem = None
    for p in doc.paragraphs:
        if "Genel Proje Kararları — Tez Yazım Süreci" in p.text:
            genel_new_elem = p._element
            break
    assert genel_new_elem is not None

    add_heading_before(doc, body, genel_new_elem, "WP5 — Ablasyon ve İleri Analizler")
    for d in DECISIONS_PAGE9:
        add_decision_block(doc, body, genel_new_elem, d)

    # ── Update summary heading: 9 -> 12 ──
    for run in ozet_para.runs:
        if "9" in run.text:
            run.text = run.text.replace("En Kritik 9 Karar", "En Kritik 12 Karar")

    # ── Add 3 rows to summary table (last table in doc) ──
    sum_tbl = doc.tables[-1]
    for num_str, karar_text, wp_text in SUMMARY_NEW_ROWS:
        row = sum_tbl.add_row()

        # Cell 0: number
        c0 = row.cells[0]
        c0.paragraphs[0].clear()
        _styled_run(c0.paragraphs[0], num_str, bold=True, size_pt=10, color_hex="1F4E79")
        _set_cell_props(c0, "600", SUM_ROW_FILL)

        # Cell 1: decision text
        c1 = row.cells[1]
        c1.paragraphs[0].clear()
        _styled_run(c1.paragraphs[0], karar_text, bold=False, size_pt=10, color_hex="2C2C2C")
        _set_cell_props(c1, "6960", SUM_ROW_FILL)

        # Cell 2: WP
        c2 = row.cells[2]
        c2.paragraphs[0].clear()
        _styled_run(c2.paragraphs[0], wp_text, bold=False, size_pt=10, color_hex="555555")
        _set_cell_props(c2, "1800", SUM_ROW_FILL)

    doc.save(DST)
    print(f"decisions_log_2.docx saved ({DST.stat().st_size:,} bytes)")

    # Verify
    doc2 = Document(DST)
    print(f"Paragraphs: {len(doc2.paragraphs)}")
    print(f"Tables: {len(doc2.tables)}")
    sum_tbl2 = doc2.tables[-1]
    print(f"Summary table rows: {len(sum_tbl2.rows)} (expected 13: header + 12)")
    for p in doc2.paragraphs:
        if "Kritik" in p.text:
            print(f"Summary heading: '{p.text}'")


if __name__ == "__main__":
    main()
