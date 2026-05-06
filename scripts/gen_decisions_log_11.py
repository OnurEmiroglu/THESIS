"""decisions_log_10.docx -> decisions_log_11.docx

Appends Decision #47 (WP6 signal informativeness sweep — chapter reframing
from informativeness-threshold to encoding-interference, Plot 4 not pursued).
Updates title-page marker to "Sürüm 28" and summary heading from
"En Kritik 26 Karar" -> "En Kritik 27 Karar".

Decision #47 uses the same v2 sub-section structure as #46:
  Karar / Seçenekler / Neden / Etki / Not  (no İkilem; trailing Not row)
The 6-row table template is reused; row labels are overwritten.
"""

from __future__ import annotations
import sys, io, shutil, copy
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement

SRC = Path("manuscript/decisions_log_10.docx")
DST = Path("manuscript/decisions_log_11.docx")

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _set_cell_text(cell, text, *, bold=None):
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].text = text
        for r in para.runs[1:]:
            r.text = ""
        if bold is not None:
            para.runs[0].bold = bold
    else:
        r = para.add_run(text)
        if bold is not None:
            r.bold = bold


def _build_karar_table_v2(doc, template_tbl, number, title, wp,
                          karar, secenekler, neden, etki, note):
    """Build a Karar table with Karar / Seçenekler / Neden / Etki / Not rows."""
    new_xml = copy.deepcopy(template_tbl._tbl)
    from docx.table import Table
    new_tbl = Table(new_xml, doc.part)
    _set_cell_text(new_tbl.rows[0].cells[0], wp, bold=True)
    _set_cell_text(new_tbl.rows[0].cells[1], f"#{number}  {title}", bold=True)
    _set_cell_text(new_tbl.rows[1].cells[0], "Karar", bold=True)
    _set_cell_text(new_tbl.rows[1].cells[1], karar)
    _set_cell_text(new_tbl.rows[2].cells[0], "Seçenekler", bold=True)
    _set_cell_text(new_tbl.rows[2].cells[1], secenekler)
    _set_cell_text(new_tbl.rows[3].cells[0], "Neden", bold=True)
    _set_cell_text(new_tbl.rows[3].cells[1], neden)
    _set_cell_text(new_tbl.rows[4].cells[0], "Etki", bold=True)
    _set_cell_text(new_tbl.rows[4].cells[1], etki)
    _set_cell_text(new_tbl.rows[5].cells[0], "Not", bold=True)
    _set_cell_text(new_tbl.rows[5].cells[1], note)
    return new_xml


def _replace_in_paragraph(para, old_text, new_text):
    full = "".join(r.text for r in para.runs)
    if old_text not in full:
        return False
    new_full = full.replace(old_text, new_text)
    if para.runs:
        para.runs[0].text = new_full
        for r in para.runs[1:]:
            r.text = ""
    return True


def main():
    shutil.copy2(SRC, DST)
    doc = Document(DST)

    # Last individual Karar table in log_10 is #46 — second-to-last table overall
    template = doc.tables[-2]

    k47 = _build_karar_table_v2(
        doc, template,
        number=47,
        title="Signal degradation hypothesis reversal",
        wp="WP6",
        karar=(
            "WP6 Signal Informativeness Sweep'in yorumu, Plot 1-3 sonuçları "
            "sonrası informativeness-threshold çerçevesinden encoding-interference "
            "çerçevesine geçirildi. Plot 4 (per-regime breakdown / action analysis) "
            "bu sürümde takip edilmedi."
        ),
        secenekler=(
            "(a) Orijinal informativeness-threshold çerçevesini korumak ve null "
            "result olarak raporlamak  |  "
            "(b) Encoding-interference çerçevesine geçmek ve daha güçlü ampirik "
            "deseni öne çıkarmak  |  "
            "(c) Plot 4'ü ekleyip per-regime / action analyses ile mekanizma "
            "tanımlamaya çalışmak"
        ),
        neden=(
            "Plot 1 (monotonic-gap): sigma_only condition-invariant kaldı (0.756-"
            "0.783); beklenen monotonik daralma gözlenmedi — orijinal hipotez "
            "tutmadı. Plot 2 (paired-seed combined vs sigma_only): combined < "
            "sigma_only direction her 4 informative koşulda (paired t p<0.05, "
            "Cohen's dz -0.57 ila -0.91); pratik (TOST δ=0.05) yalnızca lagged'de "
            "(mean_diff=-0.101, dz=-0.91). Plot 3 (paired-seed combined vs "
            "regime_only): mean diff null ama variance heterojen (std_diff > "
            "std_regime_only 3/4 koşulda); TOST equivalence n=20'de underpowered "
            "(observed std_diff ≈ 0.13-0.16, minimum detectable equivalence band "
            "≈ ±0.07-0.08, δ=0.05'in altında değil)."
        ),
        etki=(
            "Chapter 5 yeniden çerçevelendi: orijinal hipotez reddedildi, daha "
            "güçlü ampirik desen (encoding-interference: combined sistematik "
            "olarak sigma_only altında, regime_only seviyesine yakın) öne "
            "çıkarıldı, mekanizma (interference vs crowd-out vs optimization "
            "noise) bu deneyden tanımlanamaz olarak açıkça not edildi. Plot 4 "
            "takip edilmedi — üç plot beraber chapter spine'ı kuruyor; ek analiz "
            "core sweep sorusuna ortogonal yeni cepheler (per-regime breakdown, "
            "action analysis) açacaktı."
        ),
        note=(
            "Metodolojik simetri: WP5'te TOST equivalence'ı kanıtlamak için "
            "kullanılmıştı (ppo_aware ≈ ppo_blind, max-p kuralı). WP6 Plot 2'de "
            "ters yönde non-equivalence için kullanıldı (combined ≠ sigma_only, "
            "min-p kuralı); Plot 3 WP5 yönüne döndü. Bu simetri thesis chapter "
            "intro'sunda açıkça belirtilecek. \"encoding-interference\" mekanizma "
            "değil betimleyici etiket olarak kullanılıyor; mekanizma "
            "tanımlamasına gitmemek bilinçli bir kısıtlama."
        ),
    )

    body = doc.element.body
    ozet_heading = None
    for child in body.iterchildren():
        if child.tag == W_NS + "p":
            pStyle = child.find(W_NS + "pPr/" + W_NS + "pStyle")
            if pStyle is not None and pStyle.get(W_NS + "val") == "Heading1":
                text = "".join(t.text or "" for t in child.iter(W_NS + "t"))
                if "Özet" in text:
                    ozet_heading = child
                    break
    if ozet_heading is None:
        raise RuntimeError("Özet heading not found")

    def _blank_p():
        return OxmlElement("w:p")

    ozet_heading.addprevious(_blank_p())
    ozet_heading.addprevious(k47)

    # Summary heading: 26 -> 27
    for child in body.iterchildren():
        if child.tag == W_NS + "p":
            text = "".join(t.text or "" for t in child.iter(W_NS + "t"))
            if "En Kritik 26 Karar" in text:
                from docx.text.paragraph import Paragraph
                p_obj = Paragraph(child, doc.part)
                _replace_in_paragraph(p_obj, "En Kritik 26 Karar", "En Kritik 27 Karar")

    for para in doc.paragraphs:
        if "ChatGPT ve Claude" in para.text and "26 karar" in para.text:
            _replace_in_paragraph(para, "26 karar", "27 karar")

    # Title-page: Sürüm 27 -> Sürüm 28
    for para in doc.paragraphs:
        if "Sürüm 27" in para.text:
            _replace_in_paragraph(para, "Sürüm 27", "Sürüm 28")

    # Summary row #27
    sum_tbl = doc.tables[-1]
    template_row = sum_tbl.rows[1]._tr

    def _add_summary_row(num, text, wp):
        new_row = copy.deepcopy(template_row)
        sum_tbl._tbl.append(new_row)
        from docx.table import _Row
        row_obj = _Row(new_row, sum_tbl)
        _set_cell_text(row_obj.cells[0], str(num))
        _set_cell_text(row_obj.cells[1], text)
        _set_cell_text(row_obj.cells[2], wp)

    _add_summary_row(
        27,
        "WP6 Signal Informativeness Sweep — hipotez çevirimi (informativeness-"
        "threshold → encoding-interference); Plot 4 takip edilmedi",
        "WP6",
    )

    doc.save(DST)
    print(f"decisions_log_11.docx saved ({DST.stat().st_size:,} bytes)")

    doc2 = Document(DST)
    print("\n=== Verification ===")
    print(f"Total tables: {len(doc2.tables)}")
    t = doc2.tables[-2]
    print(f"Last karar table header: {t.rows[0].cells[1].text[:80]}")
    sum2 = doc2.tables[-1]
    print(f"Summary rows: {len(sum2.rows)} (should be 28 = header + 27)")
    print(f"  Last row: {[c.text[:60] for c in sum2.rows[-1].cells]}")
    for p in doc2.paragraphs[:30]:
        if "Sürüm" in p.text or "En Kritik" in p.text:
            print(f"Marker: {p.text}")


if __name__ == "__main__":
    main()
